#!/usr/bin/env python3
"""VANTORA P2 — Field-Workflow Convergence Complete -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def bl(t,size=9.3,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.0,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("P2 — Field-Workflow Convergence"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("All six operational approvals on the engine — flagged OFF, reversible, staging-validated"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Legacy paths preserved · pilot behaviour unchanged (all flags OFF) · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Summary")
para("All six operational field approvals are now convergeable onto the configurable engine, each behind its own feature "
     "flag (default OFF). With flags OFF the pilot behaves exactly as today — the legacy field-queue / screen paths are "
     "untouched. With a flag ON, the create-path starts an engine instance and the decision flows through the unified inbox "
     "to an OUTCOME HANDLER that reuses the existing, proven decision RPC. No RPC was rewritten; nothing was activated.",
     bold=True)
tbl(["Workflow","Flag","Entity","Approver permission","Branch-scoped","Reuses RPC (approve / reject)"],
[
 ["Load Request","KAKO_APPROVAL_LOADREQ","stock_request","stock_request.approve","Yes","erp_approve_stock_request / status=rejected"],
 ["Day-Close","KAKO_APPROVAL_DAYCLOSE","work_session","day.approve_close_exception","Yes","erp_approve_day_close / reopen (close_status=open)"],
 ["Out-of-Route Visit","KAKO_APPROVAL_VISIT","visit_compliance","visit.approve_out_of_route","Yes*","erp_decide_visit_compliance (approve+reject)"],
 ["Customer Transfer","KAKO_APPROVAL_CUSTTRANSFER","customer_transfer","customer.transfer","No (cross-branch)","erp_approve_customer_transfer / status=rejected"],
 ["Van Transfer","KAKO_APPROVAL_VANTRANSFER","van_transfer","stock.transfer.approve","Yes*","erp_approve_van_transfer / erp_reject_van_transfer"],
 ["Van Reconciliation","KAKO_APPROVAL_VANRECON","van_reconciliation","reconciliation.approve","Yes*","erp_settle_van_reconciliation / erp_reject_van_reconciliation"],
],widths=[1.3,1.5,1.2,1.4,0.9,2.0],size=7.6)
para("* branch_scoped flag set; when no branch is supplied in context the engine resolves the permission company-wide "
     "(safe degrade). Day-Close passes the session branch, so it fans out to in-branch approvers.", italic=True, color=GREY, size=8.5)

part("1 · Per-workflow: design · validation · risk · rollback · pilot impact")
def wf(title, design, validation, risk, rollback, pilot):
    h2(title)
    para("Design — "+design, size=9.2)
    para("Validation — "+validation, size=9.2, color=GREEN)
    para("Risk — "+risk, size=9.2, color=AMBER)
    para("Rollback — "+rollback, size=9.2)
    para("Pilot impact — "+pilot, size=9.2, bold=True)

wf("1 · Load Request",
   "createStockRequest starts stock_request_approval (branch-scoped) when flagged; handler runs the existing erp_approve_stock_request on approve, marks rejected on reject.",
   "Branch-scoped start created 4 'user' tasks for exactly the in-branch stock_request.approve holders (admin/branchmgr/supervisor/warehouse); requester excluded. tsc/suite/build green.",
   "Medium — stock-to-van movement; reuses the proven atomic RPC; flag-gated.",
   "Flag OFF (instant) → legacy /inventory/requests path. Delete definition by key. Verified definition is deletable, no data impact.",
   "None while OFF (default). Legacy approve/reject unchanged.")
wf("2 · Day-Close Exception",
   "closeDay starts day_close_approval (branch-scoped to the session branch) when the close goes pending_approval; handler runs erp_approve_day_close on approve, reopens the day (close_status=open) on reject — a NEW explicit reject the legacy path lacked.",
   "Branch-scoped start created 3 'user' tasks for the in-branch day.approve_close_exception holders (supervisor/branch_manager/admin). Governance flags on. Cleaned up.",
   "Medium — daily operational; reject reopens (recoverable); flag-gated.",
   "Flag OFF → legacy approveDayClose via the field queue. Delete definition. Reversible.",
   "None while OFF. (Behaviour note: when ON, a rejected day-close REOPENS the day so the rep can retry — new, documented behaviour.)")
wf("3 · Out-of-Route Visit",
   "checkInVisit starts visit_compliance_approval when a pending compliance row is logged; handler calls the existing erp_decide_visit_compliance for both approve and reject.",
   "Start created a single 'permission' task (visit.approve_out_of_route), decidable via the engine's permission authorization. Cleaned up.",
   "Medium — high volume; validate inbox load (erp_workflow_my_tasks is indexed) before wide activation.",
   "Flag OFF → legacy decideVisitCompliance via the field queue. Delete definition. Reversible.",
   "None while OFF.")
wf("4 · Customer Transfer",
   "transferCustomer (require approval) starts customer_transfer_approval (company-wide — inherently cross-branch) when status=pending; handler runs erp_approve_customer_transfer on approve, marks rejected on reject — a NEW explicit reject.",
   "Start created a single 'permission' task (customer.transfer). Governance flags on. Cleaned up.",
   "Medium — customer ownership change; flag-gated.",
   "Flag OFF → legacy approveCustomerTransfer. Delete definition. Reversible.",
   "None while OFF. (When ON, an explicit reject marks the transfer rejected.)")
wf("5 · Van Transfer",
   "requestVanTransfer starts van_transfer_approval when not auto-approved (status=pending); handler reuses erp_approve_van_transfer / erp_reject_van_transfer. The under-threshold auto-approve is preserved.",
   "Start created a single 'permission' task (stock.transfer.approve). Cleaned up.",
   "Medium-High — moves stock; the auto-approve threshold and the proven RPCs are reused unchanged.",
   "Flag OFF → legacy approve/reject + auto-approve. Delete definition. Reversible.",
   "None while OFF. Auto-approve below threshold unaffected even when ON.")
wf("6 · Van Reconciliation",
   "computeVanReconciliation starts van_reconciliation_approval when variance puts it pending_approval; handler reuses erp_settle_van_reconciliation / erp_reject_van_reconciliation. The under-variance auto-draft is preserved.",
   "Start created a single 'permission' task (reconciliation.approve). Cleaned up.",
   "Medium-High — financial variance; reuses the proven settle/reject RPCs.",
   "Flag OFF → legacy settle/reject. Delete definition. Reversible.",
   "None while OFF. Under-threshold auto-draft unaffected.")

part("2 · Anti-double-decision + guardrails")
bl("Field-queue caps now hide a legacy type once its P2 flag is ON, so the same approval never shows twice (legacy item + "
   "engine task). Flags OFF ⇒ caps unchanged.")
bl("Every workflow: additive migration (idempotent), dormant until flagged, legacy path preserved, governance flags on "
   "(self-approval block + mandatory reject reason), audit via the engine, reversible by flag or by deleting the definition.")
bl("No RPC rewritten, no data migration, no breaking change. Pilot behaviour is identical (all flags OFF).")
para("Recommended activation order (post pilot sign-off, one at a time, validate each): Load Request → Day-Close → Visit → "
     "Customer Transfer → Van Transfer → Van Reconciliation. Turn on UNIFIED_INBOX so approvers see engine tasks in one place.",
     bold=True, color=NAVY)

out="docs/audits/VANTORA-P2-Field-Workflow-Convergence.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
