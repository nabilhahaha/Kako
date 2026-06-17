#!/usr/bin/env python3
"""VANTORA Pilot Tenant Access Package -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E); CYAN=RGBColor(0x0E,0x7C,0x86)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11.5); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def kv(label,val,vcolor=DARK):
    p=d.add_paragraph(); r=p.add_run(label+":  "); r.bold=True; r.font.size=Pt(9.5)
    r2=p.add_run(val); r2.font.size=Pt(9.5); r2.font.color.rgb=vcolor
def tbl(headers,rows,widths=None,size=8.6,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.6); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Pilot Tenant Access Package"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Logins · roles · landing pages · responsibilities · workflows to test · end-to-end order"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Live on vantora-staging · 7 users · all passwords test.123 · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Tenant")
tbl(["Field","Value"],[
 ["Tenant name","VANTORA Pilot FMCG (DEMO)"],
 ["Tenant ID","612af0bd-973c-4fed-8e76-80cf444ef9e0"],
 ["Active modules","21 (FMCG business-type set: sales, pos, inventory, warehousing, purchasing, returns, sales_orders, accounting, crm, workflow, analytics, field_ops, distribution + engines)"],
 ["Branches","1 — PILOT (Pilot Branch)"],
 ["Warehouses","2 — PILOT-WH (Main) · PILOT-VAN (Van)"],
 ["Demo route","PILOT-R1 — Cairo East Route (rep = Salesman, 5 customers in sequence)"],
 ["Demo customers","5 — see list below"],
 ["Demo products","8 — see list below"],
 ["Sales history","12 issued invoices, 7 collections, ~4,906 EGP open AR, 4 live pending approvals"],
],widths=[1.6,5.0])

part("Pilot users (7)")
para("All passwords: test.123 (temporary — change on first real use). Each account opens directly on its work screen "
     "(role-aware landing).", bold=True)

USERS=[
 ("Pilot Company Admin","Company Admin (admin)","admin@pilot.test","/dashboard",
  "Owns the tenant: branches, warehouses, users/roles, catalog, policies, oversight.",
  ["Review the dashboard KPIs + one-tap quick actions","Settings · Staff → add a user & assign a role","Open Reports (populated)","Switch language AR ↔ EN"]),
 ("Pilot Branch Manager","Branch Manager (branch_manager)","branchmgr@pilot.test","/manager",
  "Runs the branch: purchasing, local customers, approvals, team oversight (no company settings/billing).",
  ["Branch cockpit KPIs","Approvals → clear branch requests","Purchasing → raise/receive a PO","Customers → edit a local customer"]),
 ("Pilot Supervisor","Supervisor (supervisor)","supervisor@pilot.test","/approvals/queue",
  "Supervises the field team: approves day-close, out-of-route visits, customer/van transfers; monitors coverage.",
  ["Approval Queue → filter by type/status","Approve a day-close with a comment","Reject an out-of-route visit","Approve the pending van transfer (mobile tab too)"]),
 ("Pilot Salesman","Salesman / Van Rep (salesman)  [= Cash Van]","salesman@pilot.test","/today",
  "Works the route: sell, collect, request van load, close the day. Cannot post GL or change price/credit.",
  ["Today → work the 5 route customers","POS / invoice a sale (price locked)","Record a collection (no double-charge on double-tap)","Request a van load · Close the day"]),
 ("Pilot Accountant","Accountant (accountant)","accountant@pilot.test","/collections",
  "Finance: record collections, post the GL, supplier payments, AR aging & reports. Cannot sell.",
  ["Collections → record a payment","Accounting · Vouchers → post a voucher","AR aging / financial reports","Confirm 'issue invoice' is blocked"]),
 ("Pilot Warehouse Keeper","Warehouse Keeper (warehouse_keeper)","warehouse@pilot.test","/inventory/requests",
  "Stock: approve van load requests, receive POs, adjust/transfer/count. Cannot sell/collect.",
  ["Inventory · Requests → approve a load","Receive a purchase order","Adjust stock / run a count","Confirm 'record collection' is blocked"]),
 ("Pilot Viewer","Viewer (viewer)","viewer@pilot.test","/dashboard",
  "Read-only observer: reports, dashboards, inventory view. No write actions.",
  ["Open dashboards & reports","Browse inventory (read-only)","Confirm every write action is blocked"]),
]
for name,role,email,land,resp,flows in USERS:
    h2(name)
    kv("Role", role); kv("Email", email); kv("Temporary password", "test.123", AMBER); kv("Landing page", land, GREEN)
    kv("Main responsibilities", resp)
    para("Key workflows to test:", bold=True, size=9.5)
    for fwk in flows: b(fwk, size=9)

part("Demo data")
h2("Customers (5)")
tbl(["Code","Name","Credit limit (EGP)"],[
 ["PILOT-C01","Al Nour Grocery","20,000"],["PILOT-C02","El Salam Market","15,000"],
 ["PILOT-C03","City Mini Market","30,000"],["PILOT-C04","Family Supermarket","25,000"],
 ["PILOT-C05","Corner Shop","8,000"],
],widths=[1.4,3.0,1.6])
h2("Products (8)")
tbl(["Code","Name","Sell price (EGP)"],[
 ["PILOT-P01","Sunflower Oil 1L","70"],["PILOT-P02","White Sugar 1kg","28"],
 ["PILOT-P03","Egyptian Rice 1kg","25"],["PILOT-P04","Black Tea 250g","42"],
 ["PILOT-P05","Bar Soap 120g","9"],["PILOT-P06","Tomato Paste 380g","17"],
 ["PILOT-P07","Pasta 400g","13"],["PILOT-P08","Powder Detergent 1kg","55"],
],widths=[1.4,3.0,1.6])
para("All products carry 1,000 units of stock in the Main Warehouse. Route PILOT-R1 sequences all 5 customers for the salesman.", italic=True)

part("Recommended testing order  (complete end-to-end walkthrough)")
para("Run the accounts in this order so each step sets up the next — field activity creates the approvals that managers "
     "then clear, and finance closes the loop.", bold=True)
tbl(["#","Account","Why this order","What to produce"],[
 ["1","Company Admin (admin@pilot.test)","Confirm the system is set up & populated; orient on the dashboard","Verify data, add a test user, open reports"],
 ["2","Salesman (salesman@pilot.test)","Generate the day's field activity (this creates pending approvals)","Sell + collect; request a van load; close the day"],
 ["3","Supervisor (supervisor@pilot.test)","Clear what the rep just created","Approve day-close, visit, van transfer in the queue"],
 ["4","Warehouse Keeper (warehouse@pilot.test)","Fulfil stock + replenish","Approve load request; receive a PO; adjust stock"],
 ["5","Accountant (accountant@pilot.test)","Close the financial loop on the day's sales","Record collections; post a voucher; review AR aging"],
 ["6","Branch Manager (branchmgr@pilot.test)","Branch-level oversight & purchasing","Review branch KPIs; raise/receive a PO; clear approvals"],
 ["7","Viewer (viewer@pilot.test)","Confirm read-only safety","Browse reports; confirm no write action is possible"],
],widths=[0.3,2.0,2.2,2.1],
 fill={(i,1):BLUEBG for i in range(7)})
para("")
para("Tip: do steps 2–5 on a phone to validate the mobile field experience (bottom-nav for the rep, Approvals tab for the "
     "supervisor). Capture anything slow/confusing in the Pilot Operations Pack feedback table.", italic=True, color=NAVY)

out="docs/audits/VANTORA-Pilot-Access-Package.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
