#!/usr/bin/env python3
"""VANTORA Commercial Demo Presentation v2 (10-section outline) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); CYAN=RGBColor(0x0E,0x7C,0x86)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(11)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def slide(k,t):
    d.add_page_break()
    p=d.add_paragraph(); r=p.add_run(k.upper()); r.font.color.rgb=CYAN; r.font.size=Pt(10); r.bold=True
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(22); r.bold=True; d.add_paragraph()
def para(t,color=DARK,size=11,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=11,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=9.5,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)
for _ in range(6): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(52); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("FMCG Distribution, run on one platform"); r.font.size=Pt(16); r.font.color.rgb=CYAN; r.bold=True
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Commercial & Demo Presentation"); r.font.size=Pt(12); r.font.color.rgb=GREY

slide("1 · Overview","What is VANTORA")
para("VANTORA is a mobile-first, multi-tenant platform that runs an FMCG distributor end-to-end — the van on the street and "
     "the ledger in the office, in one system. It replaces the usual desktop-ERP + separate field-app + spreadsheets stack.", size=12, bold=True)
b("One record: the van sale IS the invoice IS the GL entry.")
b("Bilingual (Arabic/English, RTL-native), online + offline, cloud-hosted.")

slide("2 · ERP","ERP capabilities")
b("Sales: invoices, POS / quick-sale, returns, credit notes, ETA e-invoicing readiness.")
b("Inventory: multi-warehouse, batches/expiry (FEFO), transfers, counts, tenant-set valuation/costing.")
b("Purchasing: suppliers, POs, goods receipt, supplier returns & payments.")
b("Accounting: chart of accounts, automatic journals, vouchers, AR/AP, aging, reports, exports.")
b("Customers & pricing: master data, segments/channels, credit limits + approval, price rules — all governed.")

slide("3 · Distribution","Distribution capabilities")
b("Routes & territories with customer sequencing, working days, van assignment.")
b("Van sales: load → sell-off-van → collect → reconcile, stock in real time.")
b("Coverage & journey compliance: planned vs visited, out-of-route, day-close exceptions.")
b("Trade spend (promotions/accruals/claims/ROI), merchandising / perfect store, critical alerts.")

slide("4 · Field Force","Field Force capabilities")
b("Purpose-built 'Today' rep home — route, visit, sell, collect, close — designed for a phone.")
b("Offline-capable selling & collecting; sync on reconnect.")
b("GPS & visit compliance; van load requests & transfers.")
b("Role-aware landing: a rep opens on their day, not a back-office dashboard.")

slide("5 · Approvals","Approval workflows")
para("A single mobile Approval Queue replaces phone-call / WhatsApp sign-offs. Field managers act in one place.", bold=True)
tbl(["Workflow","Requester → Approver"],[
 ["Day-close exception","Salesman → Supervisor"],
 ["Out-of-route visit","Salesman → Supervisor"],
 ["Van load / stock transfer","Salesman → Supervisor / Warehouse"],
 ["Customer transfer (reassign)","Manager → Manager"],
 ["Credit-limit change","Rep → Manager"],
 ["Trade-spend promotion","Trade mktg → Manager"],
],widths=[3.0,3.4])
b("Filters (type/status), approve/reject with comments, full history + audit trail. Mobile bottom-nav tab for approvers.")

slide("6 · Architecture","Multi-tenant architecture")
b("True multi-tenant SaaS: every company isolated by Postgres row-level security (RLS) — verified live (a rep sees 0 of another tenant's 154 customers).")
b("New tenant + its roles + modules seed automatically on creation; ready in minutes.")
b("Per-company modules + feature flags + templates (Lite / Standard / Enterprise) — powerful backend, simple frontend.")
b("Cloud-native: Next.js on Vercel + Supabase Postgres; low IT footprint for the customer.")

slide("7 · Mobile","Mobile experience")
b("Mobile-first across the board: the 'More' drawer mirrors the full system, so nothing is desktop-only.")
b("The rep's core loop is ~4 taps (Home / Today / Customers / Sell / Inventory); approvers get a one-tap Approvals tab.")
b("Offline mode keeps the field working with no signal.")
b("Arabic RTL and English LTR render natively.")

slide("8 · Security","Security & permissions")
b("Granular permissions with enforced separation of duties: a rep sells & collects but cannot post the GL, change prices/credit, or approve loads (validated live).")
b("Three access tiers: Platform Owner (vendor) → Tenant Admin → scoped role.")
b("Idempotent money operations (a double-tap never double-charges) and atomic, branch-scoped document numbering.")
b("Every sensitive action is permission-gated and audit-logged; tenant data never crosses tenants.")

slide("9 · Proof","FMCG pilot results")
b("Full platform / role / coverage / workflow audits passed; 6 critical issues fixed and validated on live data.")
b("Tenant isolation, idempotent collections, sequential invoice numbering, separation of duties — all verified on the live pilot tenant.")
b("Every core FMCG workflow completes end-to-end by the intended role; 1300+ automated tests green; production build green.")
b("A populated pilot tenant (route, 12 invoices, collections, ~4,906 EGP AR, 4 live approvals) is ready to demo today.")

slide("10 · Positioning","Competitive positioning")
tbl(["VANTORA vs…","Edge"],[
 ["Desktop ERPs (SAP/Oracle/local)","Mobile-first field force + distribution built in, not bolted on; fast to onboard"],
 ["Van-sales / DMS point apps","Real ERP + GL underneath — no reconciliation gap; one vendor"],
 ["Spreadsheets / manual","Governed pricing, credit, approvals, audit; real-time visibility"],
 ["Generic SaaS","Arabic-first, FMCG-shaped, multi-industry core (also pharmacy/retail/clinic)"],
],widths=[2.2,4.2])
para("")
para("VANTORA is pilot-ready for FMCG distribution and extensible across industries — positioned to win the region's "
     "distributors and expand from the sell-collect loop into full finance and field-force engines.", size=12, bold=True, color=NAVY)
out="docs/audits/VANTORA-Commercial-Presentation-v2.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
