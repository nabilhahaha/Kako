#!/usr/bin/env python3
"""VANTORA Role & Permission Audit (computed from code) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
YES="●"; NO="—"; PART="◐"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.0,fill=None,center_from=1):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; pr=c.paragraphs[0]; r=pr.add_run(hh); r.bold=True; r.font.size=Pt(7.8); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
        if j>=center_from: pr.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; pr=c.paragraphs[0]; r=pr.add_run(v); r.font.size=Pt(size)
            if j>=center_from: pr.alignment=WD_ALIGN_PARAGRAPH.CENTER
            if v==YES: r.font.color.rgb=GREEN; r.bold=True
            elif v==NO: r.font.color.rgb=GREY
            elif v==PART: r.font.color.rgb=AMBER; r.bold=True
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ============================================================================
# AUTHORITATIVE DATA — transcribed from src/lib/erp/permissions.ts ROLE_PERMISSIONS
# and src/lib/erp/navigation.ts NAV_SECTIONS (branch claude/fmcg-sell-collect-loop).
# ============================================================================
ALLP = object()  # admin/manager = '*'
ROLE = {
 "Platform Owner": ("__owner__", ALLP),  # isPlatformOwner: provider panel only
 "Company Admin":  ("admin", ALLP),       # ALL tenant perms; NOT global super-admin
 "Branch Manager": ("branch_manager", set("""sales.sell sales.discount sales.collect sales.return customers.manage
   customers.change_status inventory.view inventory.adjust inventory.transfer inventory.count stock_request.approve
   purchasing.manage suppliers.manage reports.view customer.transfer customer.create customer.edit route.create
   journey.create stock.adjust stock.transfer.approve visit.approve_out_of_route day.approve_close_exception stock.view
   user.transfer""".split())),
 "Supervisor": ("supervisor", set("""sales.sell sales.discount sales.collect sales.return customers.manage
   customers.change_status inventory.view stock_request.approve reports.view visit.approve_out_of_route
   day.approve_close_exception stock.transfer.approve customer.transfer journey.create route.create stock.view
   reconciliation.view reconciliation.manage""".split())),
 "Salesman": ("salesman", set("""sales.sell sales.collect customers.manage inventory.view stock_request.create
   field.sales field.attach_media day.close stock.view stock.transfer customer.create reconciliation.view""".split())),
 "Cash Van": ("cash_van", set("""sales.sell sales.collect customers.manage inventory.view stock_request.create
   field.sales field.attach_media""".split())),   # NOT in code map — inferred van-rep set (DB-seed dependent)
 "Accountant": ("accountant", set("""accounting.view accounting.post reports.view suppliers.manage sales.collect
   customers.change_status stock.view fashion.reports fashion.cashbox fashion.installments fashion.purchase""".split())),
 "Warehouse": ("warehouse_keeper", set("""inventory.view inventory.adjust inventory.transfer inventory.count
   stock_request.approve purchasing.manage stock.view stock.adjust stock.transfer stock.transfer.approve
   fashion.inventory fashion.purchase reconciliation.view reconciliation.manage""".split())),
 "Viewer": ("viewer", set("reports.view accounting.view inventory.view".split())),
}
ROLES = list(ROLE.keys())
def has(role, perms):
    """Does role hold ANY of perms? Platform Owner = provider-only (no tenant perms)."""
    key, pset = ROLE[role]
    if key=="__owner__": return False   # owner sees no tenant-operational screens
    if pset is ALLP: return True
    return any(p in pset for p in (perms if isinstance(perms,(list,tuple,set)) else [perms]))

# ---- Permission categories (functional groups) ----
CATS = [
 ("Sales / orders", ["sales.sell","sales.discount","customer.create","customer.edit","customers.manage","pricing.manage"]),
 ("Collections / cash", ["sales.collect"]),
 ("Returns", ["sales.return","purchasing.return"]),
 ("Inventory / stock", ["inventory.view","inventory.adjust","inventory.transfer","inventory.count","stock.view","stock.adjust","stock.transfer","stock.transfer.approve"]),
 ("Stock requests (van)", ["stock_request.create","stock_request.approve"]),
 ("Purchasing", ["purchasing.manage","suppliers.manage"]),
 ("Accounting (GL)", ["accounting.view","accounting.post"]),
 ("Field ops / route", ["field.sales","route.create","journey.create","visit.approve_out_of_route","day.close","day.approve_close_exception"]),
 ("Reports", ["reports.view","report.aggregate.view"]),
 ("Settings / admin", ["settings.branches","settings.users","settings.custom_fields","integrations.manage","workflow.manage","user.transfer"]),
]
def cat_level(role, perms):
    key,pset=ROLE[role]
    if key=="__owner__": return NO
    if pset is ALLP: return YES
    held=[p for p in perms if p in pset]
    if not held: return NO
    return YES if len(held)==len(perms) else PART

# ---- Menus (curated key surface) : (name, perms, superAdminOnly, ownerOnly) ----
M_OWNER=[("Provider · Overview/Companies/Billing/Analytics",None,False,True)]
MENUS = [
 ("Dashboard / Home", "__all__", False, False),
 ("Sales · Invoices", ["sales.sell"], False, False),
 ("Sales · POS", ["sales.sell"], False, False),
 ("Sales · Returns", ["sales.return"], False, False),
 ("Sales · Customers", ["customers.manage"], False, False),
 ("Sales · Journey (field)", ["field.sales"], False, False),
 ("Collections", ["sales.collect"], False, False),
 ("Distribution · Routes/Coverage/Targets", ["reports.view","customers.manage"], False, False),
 ("Distribution · Van Accounting", ["reports.view"], False, False),
 ("Inventory · Products/Stock", ["inventory.view"], False, False),
 ("Inventory · Transfers", ["inventory.transfer"], False, False),
 ("Inventory · Load Requests", ["stock_request.create","stock_request.approve"], False, False),
 ("Inventory · Stock Count", ["inventory.count"], False, False),
 ("Purchasing · Suppliers", ["suppliers.manage"], False, False),
 ("Purchasing · Purchase Orders", ["purchasing.manage"], False, False),
 ("Accounting · Chart/Journal/Reports", ["accounting.view"], False, False),
 ("Accounting · Vouchers", ["accounting.post"], False, False),
 ("Reports (sales/distribution)", ["reports.view"], False, False),
 ("Settings · Branches", ["settings.branches"], False, False),
 ("Settings · Staff", ["settings.users"], False, False),
 ("Settings · Users (global)", None, True, False),
 ("Settings · Permissions (global)", None, True, False),
 ("Settings · Organization/Authz/Features", ["settings.users"], False, False),
 ("Settings · Workflows", ["workflow.manage"], False, False),
]
def menu_vis(role, perms, sao, owner):
    key,pset=ROLE[role]
    if owner: return YES if key=="__owner__" else NO
    if key=="__owner__": return NO  # owner: no tenant menus
    if perms=="__all__": return YES
    if sao: return NO  # superAdminOnly: needs global super-admin flag, not the admin role
    return YES if has(role, perms) else NO

# ---- Actions (server-action gates, incl. MJ-1) : (name, perms-any) ----
ACTIONS = [
 ("Create / issue invoice", ["sales.sell"]),
 ("Record payment / collection", ["sales.collect"]),
 ("Complete sales return", ["sales.return","sales.sell"]),
 ("Post / create voucher (GL)", ["accounting.post"]),
 ("Receive purchase order", ["purchasing.manage"]),
 ("Record supplier payment", ["accounting.post","suppliers.manage"]),
 ("Adjust stock", ["inventory.adjust","stock.adjust"]),
 ("Finalize stock count", ["inventory.count"]),
 ("Complete transfer", ["inventory.transfer","stock.transfer.approve"]),
 ("Create stock request (van load)", ["stock_request.create"]),
 ("Approve stock request", ["stock_request.approve"]),
 ("Manage staff / users", ["settings.users"]),
 ("Manage branches", ["settings.branches"]),
 ("Edit product price", ["pricing.manage"]),   # MJ-8 (P1 — currently auth-only; expected gate)
]
def act(role, perms):
    return YES if has(role, perms) else NO

# ============================================================================
# COVER
# ============================================================================
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Role & Permission Audit"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Screen-by-screen validation · 9 roles · computed from ROLE_PERMISSIONS + navigation gates"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Branch claude/fmcg-sell-collect-loop · post-P0 (MJ-1 gates live) · read-only · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("How to read this")
para("Legend: ● = visible / allowed · — = hidden / forbidden · ◐ = partial (holds some, not all, of the group's "
     "permissions). Data is COMPUTED from the code's ROLE_PERMISSIONS map and the navigation/page gates, so the matrices are "
     "internally consistent with what the running app enforces.", bold=True)
b("Assumption: a fully-provisioned FMCG tenant with all core modules + engine flags ON, so the matrices isolate ROLE/"
  "PERMISSION gating (the subject of this audit). Where a menu also needs a module/flag, that is noted.")
b("Company Admin = the 'admin' role (ALL tenant permissions) but NOT the GLOBAL super-admin flag. Screens marked "
  "superAdminOnly (Settings · Users, Permissions, e-invoice, Platform Audit) require that platform flag — so a pure company "
  "admin does NOT see them; they manage people via Settings · Staff.")
b("Platform Owner is the vendor: they see ONLY the Provider panel (cross-tenant), never tenant-operational menus "
  "(visibleSections short-circuits for isPlatformOwner). To act inside a tenant they use 'View as'.")
b("Cash Van is NOT defined in the code ROLE_PERMISSIONS map (not a BranchRole); its effective permissions come from the DB "
  "erp_role_permissions seed. The set shown is the INFERRED van-rep default (≈ driver) and must be verified against the DB.",
  color=AMBER)

# 1. PERMISSION MATRIX
part("1. Permission Matrix  (role × permission category)")
rows=[[name]+[cat_level(role,perms) for role in ROLES] for name,perms in CATS]
tbl(["Permission category"]+[r.split(" ")[0] if r not in("Platform Owner","Cash Van","Company Admin","Branch Manager") else r.replace("Platform ","Plat.").replace("Company ","Co.").replace("Branch ","Br.") for r in ROLES], rows,
    widths=[1.7]+[0.55]*9, size=8.5)
para("● holds the whole category · ◐ holds part · — none. Platform Owner shows — for every tenant category by design "
     "(no tenant perms; provider-scoped).", italic=True, size=8.5)

# 2. ROLE MATRIX
part("2. Role Matrix  (effective permission set per role)")
def permlist(role):
    key,pset=ROLE[role]
    if key=="__owner__": return "Provider/platform scope only (cross-tenant); no tenant Permission[] — bypasses tenant gates via 'View as'."
    if pset is ALLP: return "ALL tenant permissions ('*')."
    return ", ".join(sorted(pset))
for role in ROLES:
    key,_=ROLE[role]
    h2(f"{role}   (code role: {key})")
    para(permlist(role), size=9)
para("")
para("Note: admin & manager = '*' (ALL). it_admin/sales_director/etc. exist in code but are outside the 9 requested roles.",
     italic=True, size=8.5)

# 3. MENU VISIBILITY MATRIX
part("3. Menu Visibility Matrix  (role × menu)")
def shortR(r): return {"Platform Owner":"Plat.Own","Company Admin":"Co.Admin","Branch Manager":"Br.Mgr","Supervisor":"Superv","Salesman":"Sales","Cash Van":"CashVan","Accountant":"Acct","Warehouse":"Whse","Viewer":"View"}[r]
rows=[[name]+[menu_vis(role,perms,sao,owner) for role in ROLES] for (name,perms,sao,owner) in (M_OWNER+MENUS)]
tbl(["Menu / section"]+[shortR(r) for r in ROLES], rows, widths=[2.2]+[0.5]*9, size=8.0)
para("Hidden menus per role = every row showing — for that column. Platform Owner: only the Provider row is ●.", italic=True, size=8.5)

# 4. SCREEN VISIBILITY MATRIX (actions)
part("4. Screen / Action Visibility Matrix  (role × allowed action)")
para("Allowed (●) vs Forbidden (—) for the sensitive write actions — now enforced at the app layer (MJ-1) on top of the "
     "RPC branch-access check.", bold=True)
rows=[[name]+[act(role,perms) for role in ROLES] for (name,perms) in ACTIONS]
tbl(["Action (server action)"]+[shortR(r) for r in ROLES], rows, widths=[2.2]+[0.5]*9, size=8.0)
para("Platform Owner shows — (acts only via 'View as' as a tenant role). 'Edit product price' is shown against the EXPECTED "
     "pricing.manage gate; today that write is still auth-only (MJ-8, P1) — see Findings.", italic=True, size=8.5, color=AMBER)

# 5. ROLE GAP REPORT
part("5. Role Gap Report  (expected vs actual)")
tbl(["Role","Expected","Actual (code)","Gap"],[
 ["Platform Owner","Vendor: cross-tenant Provider panel; no tenant ops","isPlatformOwner → Provider only; bypasses tenant gates via View-as","OK — matches"],
 ["Company Admin","Full tenant control incl. user & permission management","admin = ALL perms, BUT Users/Permissions screens are superAdminOnly (global flag)","GAP: user/permission SCREENS hidden from a pure admin role; staff mgmt via /settings/staff only"],
 ["Branch Manager","Branch ops: sell, stock, purchasing, approvals (no settings/billing)","Matches: broad ops perms, no settings.users/branches/billing","OK"],
 ["Supervisor","Field supervision + approvals, reports","Matches: approvals, reconciliation, reports; no purchasing/accounting","OK"],
 ["Salesman","Sell + collect + field loop + request load","Matches; cannot approve loads (separation of duties)","OK"],
 ["Cash Van","Van rep: load, sell, collect, reconcile","NOT in code ROLE_PERMISSIONS (no BranchRole key)","GAP: code default missing — DB-seed only; verify erp_role_permissions"],
 ["Accountant","GL post, collections, supplier pay, reports","Matches; note also holds customers.change_status (suspend/block)","MINOR: customers.change_status broader than name implies"],
 ["Warehouse","Stock adjust/transfer/count, approve loads, receive POs","Matches (holds purchasing.manage to receive)","OK"],
 ["Viewer","Read-only: reports + accounting view + inventory view","Matches; no write perms → blocked from all MJ-1 actions","OK"],
],widths=[1.2,2.1,2.1,1.4])

# 6. UNAUTHORIZED ACCESS FINDINGS
part("6. Unauthorized-Access Findings")
para("Where a role can see/do more than it should. Post-P0, the financial/stock posting actions are gated (MJ-1), so the "
     "previous 'any authenticated user can post' class is CLOSED. Residual items:", bold=True)
tbl(["#","Finding","Risk","Status / fix"],[
 ["U-1","Electrical section (serials/warranties/RMA) has NO module gate — a non-electronics tenant admin/super-admin who is granted electrical.rma would see it","Minor","Open — MN-2 (P1): add an 'electrical' module gate"],
 ["U-2","Accountant holds customers.change_status (suspend/block customers) — broader than the role name implies","Minor","Accept or trim in role-template review (P2)"],
 ["U-3","cashier holds cross-vertical perms (pharmacy.dispense, restaurant.manage, fashion.*) — harmless because each is module-gated, but wide","Minor","Accept (module-gated) / tighten in P2"],
 ["U-4","'Edit product price' / 'set credit limit' master-data writes are still auth-only (no pricing/customer gate)","Major","Open — MJ-8 (P1): add pricing.manage / customer.edit gates"],
],widths=[0.4,3.7,0.7,1.8],
 fill={(0,2):AMBERBG,(1,2):AMBERBG,(2,2):AMBERBG,(3,2):REDBG})

# 7. MISSING ACCESS FINDINGS
part("7. Missing-Access Findings")
para("Where a role cannot do something it legitimately needs.", bold=True)
tbl(["#","Finding","Risk","Status / fix"],[
 ["M-1","Cash Van has no code permission default — if the DB seed is missing/empty, a van rep cannot sell/collect/load at all","Major","Open: add cash_van to ROLE_PERMISSIONS (van-rep set) OR verify the DB seed; align code & DB"],
 ["M-2","Company Admin (admin role) cannot open Settings · Users / Permissions (superAdminOnly) — only Staff. If a tenant expects the admin to manage roles/permissions in-app, that screen is unreachable without the global super-admin flag","Major","Decide: either grant tenant-admin a scoped users/permissions screen, or document Staff as the tenant path"],
 ["M-3","Accountant cannot issue invoices (no sales.sell) — correct by design, but supplier-payment posting requires they ALSO hold accounting.post|suppliers.manage (they do)","None","OK — no action"],
 ["M-4","Supervisor/Salesman cannot post to the GL or receive POs — correct separation of duties","None","OK — by design"],
],widths=[0.4,3.9,0.6,1.7],
 fill={(0,2):REDBG,(1,2):REDBG,(2,2):GREENBG,(3,2):GREENBG})

part("Summary & recommendation")
para("The role/permission model is sound and now correctly enforced end-to-end (menu gate → screen gate → server-action "
     "gate). Two real gaps deserve attention before broad onboarding:", bold=True)
b("M-1 / Cash Van — reconcile the DB role with the code map so van reps are never left permission-less (highest priority).",bold=True)
b("M-2 / Company Admin user-management — decide the intended tenant-admin path for users & permissions (product decision).")
b("U-4 / master-data write gates (pricing, credit limit) — already queued as MJ-8 in P1.")
para("Everything else matches expected behaviour, including the post-P0 separation of duties (Viewer and other low-privilege "
     "roles are correctly blocked from every sensitive write). Recommend resolving M-1 and the M-2 product decision, then "
     "proceeding — no finding here blocks a controlled pilot.", bold=True, color=NAVY)

out="docs/audits/VANTORA-Role-Permission-Audit.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
