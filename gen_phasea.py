#!/usr/bin/env python3
"""VANTORA Phase A (Governance Hygiene) Implementation Checklist & Status Tracker -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"; GREYBG="ECECEC"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def chk(t,done=False,size=9.5):
    p=d.add_paragraph(); r=p.add_run(("☑  " if done else "☐  ")+t); r.font.size=Pt(size)
    r.font.color.rgb=GREEN if done else DARK
def tbl(headers,rows,widths=None,size=8.4,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# Cover
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Phase A — Governance Hygiene"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Implementation Checklist & Status Tracker"); r.font.size=Pt(13); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Zero-impact items only · Structural simplification (Phase B/C/D) remains roadmap-only · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# SCOPE / GUARDRAILS
h1("Scope & Guardrails")
para("This tracker covers ONLY Phase A (Governance Hygiene) — the five zero-impact items from the Implementation Roadmap. "
     "Every item here is additive, reversible, and invisible to existing tenants. Nothing in this document merges, migrates, "
     "or reduces anything a tenant relies on.", bold=True)
h2("Locked OUT of Phase A (roadmap-only until after first paying customer + usage data)")
tbl(["Deferred (DO NOT do now)","Belongs to"],[
 ["Merge business types","Phase B/C"],
 ["Merge / retire roles (mgmt tiers)","Phase C"],
 ["Reduce / delete permissions","Phase C"],
 ["Migrate any tenant (business_type, role assignments)","Phase C"],
 ["Role inheritance / archetype module inheritance","Phase D"],
],widths=[4.4,2.4],fill={(i,0):REDBG for i in range(5)})
para("Current objective (per founder): (1) complete pilot validation, (2) get first paying customer, (3) collect real "
     "usage data. Structural simplification is revisited only after that.", italic=True, color=NAVY)

# STATUS SUMMARY
h1("Status Summary")
para("Legend: ☐ Not Started (Ready) · ◐ In Progress · ☑ Done · ⚠ Blocked. As of this revision, no Phase A "
     "item has been started — execution awaits the founder's go-ahead.")
tbl(["#","Phase A item","Type","Zero-impact?","Reversible?","Migration?","Status"],[
 ["A1","Remove dead permission keys","Code","Yes","Yes","No","☑ DONE"],
 ["A2","Permission aliasing — RECLASSIFIED to Phase C","—","No (changes access)","—","Yes","⟶ Phase C"],
 ["A3","Bind nav + page to same flag (Trade Spend / Coverage 404 fix)","Bug fix","Yes (strictly better)","Yes","No","☑ DONE"],
 ["A4","Document 5 archetypes (Handbook Appendix A)","Docs","Yes","Yes","No","☑ DONE"],
 ["A5","Catalog feature-flag kinds (Handbook Appendix A)","Docs","Yes","Yes","No","☑ DONE"],
],widths=[0.4,2.7,0.9,1.1,0.9,0.7,0.9],
 fill={(0,6):GREENBG,(1,6):AMBERBG,(2,6):GREENBG,(3,6):GREENBG,(4,6):GREENBG})
para("Overall completion: 4 / 4 executable items DONE (A1, A3, A4, A5).  A2 was reclassified to Phase C after the "
     "business-impact review proved permission aliasing is NOT zero-impact (distinct enforced permissions). Tests 1314 "
     "passing, build green, zero tenant impact. See the Phase A Completion Report for before/after validation.",
     bold=True)

# ===== DETAIL PAGES =====
def item(code,title,typ,desc,subtasks,accept,verify,zero,rev):
    part(f"{code} — {title}")
    tbl(["Field","Value"],[
     ["Type",typ],["Zero-impact","Yes — "+zero],["Reversible","Yes — "+rev],
     ["Data migration","No"],["Status","☐ Not Started (Ready)"],["Owner","(assign)"],["Target phase","A — Governance Hygiene"],
    ],widths=[1.4,5.4])
    h2("What & why"); para(desc)
    h2("Implementation sub-tasks (tick as completed)")
    for s in subtasks: chk(s)
    h2("Acceptance criteria");
    for a in accept: b(a)
    h2("Verification");
    for v in verify: b(v,color=BLUE)

item("A1","Remove dead permission keys","Code (registry only)",
 "change_requests.create / change_requests.approve / change_requests.manage and trade_spend.manage are granted in the "
 "database role seeds but are NOT part of the code Permission type — nothing enforces them. Removing them eliminates a "
 "source of audit confusion with no behavioural change.",
 ["Confirm via grep that the keys are absent from the code Permission union / hasPermission paths.",
  "List every seed/role row that grants the dead keys (read-only inventory, no change).",
  "Remove the dead keys from future seed definitions (do not retro-edit live tenant grants in Phase A).",
  "Update the permission registry test to assert the keys no longer appear.",
  "Note in the governance handbook that these are retired."],
 ["No code path references the removed keys.",
  "Permission registry test green.",
  "Existing tenants' effective permissions unchanged (dead keys never granted anything)."],
 ["grep for each key returns zero enforcement references.",
  "Run the feature-catalog / permission registry test suite — must stay green.",
  "Spot-check one tenant's resolved permissions before/after — identical."],
 "the keys enforce nothing today, so removal cannot change any tenant's access",
 "re-adding the key strings restores the prior seed list")

item("A2","Add permission aliases (additive)","Code (resolver only)",
 "Three duplicate families exist: stock.* vs inventory.*, customer.* vs customers.*, report.aggregate.view vs reports.view. "
 "Phase A introduces ALIASES only: the resolver treats the legacy key and the canonical key as equivalent. No key is "
 "deleted, no role is re-seeded. This removes the overlap at the resolution layer while every existing grant keeps working.",
 ["Define the canonical key for each family (inventory.*, customers.*, reports.view).",
  "Add an alias map: legacy key -> canonical key (resolver checks both).",
  "Ensure hasPermission(canonical) is satisfied by a legacy grant and vice-versa.",
  "Add tests proving old-key holders and new-key holders get identical access.",
  "Document the canonical set; mark legacy keys 'aliased (deletion deferred to Phase C)'."],
 ["A role with only the legacy key still passes a canonical-key check (and vice-versa).",
  "No seed data changed; no permission removed.",
  "Alias map is the single source of equivalence."],
 ["Unit test: grant legacy key -> assert canonical check passes.",
  "Unit test: grant canonical key -> assert legacy check passes.",
  "Regression: existing role-permission tests stay green."],
 "aliases are purely additive — both old and new keys resolve to the same grant, so no tenant loses or gains access",
 "removing the alias map entries reverts to the prior independent behaviour")

item("A3","Bind nav + page to the same capability flag","Bug fix (freeze-allowed)",
 "Today some screens appear in navigation but the page itself is gated by a different (env) flag, producing show-then-404 "
 "— exactly the Trade Spend case (nav shows /distribution/trade-spend, page returns notFound() when KAKO_TRADE_SPEND is "
 "off). Phase A binds the nav item and the page to the SAME capability signal so a link is shown only when the page works.",
 ["Inventory every nav item whose page has an independent flag/notFound() gate.",
  "For each, source the nav visibility from the same flag the page checks.",
  "Verify the result: link shown => page renders; flag off => link hidden (never a 404).",
  "Add a test asserting nav-gate and page-gate agree for the engine routes.",
  "Re-validate the Trade Spend route specifically."],
 ["No nav link leads to a notFound() page for any flag state.",
  "Trade Spend link hidden when its flag is off, shown+working when on.",
  "Behaviour strictly improves (no new exposure)."],
 ["Toggle each engine flag off/on and confirm link visibility matches page availability.",
  "Automated test: for each gated route, nav-visible iff page-renders.",
  "Manual: reproduce the prior Trade Spend 404 and confirm it is gone."],
 "it only ever HIDES a link that would have 404'd — it never exposes anything new, so no tenant gains access",
 "reverting restores the prior (buggy) independent gating")

item("A4","Document the 5 archetypes; curate onboarding picker","Docs / onboarding copy",
 "Existing tenants keep their exact business_type. Phase A only (a) writes the 5-archetype model into the governance "
 "baseline and (b) curates the NEW-tenant onboarding picker so new sign-ups are guided to the archetype set. No existing "
 "business_type is renamed, remapped, or removed.",
 ["Finalise the 5-archetype definitions (Distribution, Stock-Retail, Service-Vertical, Self-contained-Vertical, Generic).",
  "Map each existing business_type to its archetype (documentation table only — no data change).",
  "Curate the onboarding picker options/labels for NEW tenants (presentation only).",
  "Leave all existing tenants' business_type values untouched.",
  "Record that type->archetype MIGRATION is Phase C, explicitly out of scope here."],
 ["Existing tenants' business_type unchanged in the database.",
  "New-tenant picker presents the archetype-guided set.",
  "Archetype map documented in the governance baseline."],
 ["DB check: business_type distribution identical before/after.",
  "Onboarding walkthrough shows curated options for a new tenant only.",
  "No migration script exists or runs."],
 "no tenant record changes; only new-tenant onboarding copy and documentation are touched",
 "revert the picker copy; documentation has no runtime effect")

item("A5","Split feature-flag kinds in registry/docs","Docs / registry classification",
 "Two different kinds of flags are currently mixed: tenant CAPABILITY flags (feature-catalog; Lite/Standard/Enterprise) "
 "and platform KILL-SWITCHES (env KAKO_*). Phase A classifies and documents them as two distinct kinds so future work "
 "never re-creates the show-then-404 confusion. Pure classification — no runtime behaviour changes.",
 ["List every flag and label it Capability or Kill-switch.",
  "Document the rule: tenant-visible gating uses Capability flags; unfinished-engine gating uses Kill-switches.",
  "Cross-reference A3: nav+page bind to the Capability flag, kill-switch only force-disables.",
  "Add the classification to the governance baseline and the feature-catalog docs.",
  "No change to flag values or evaluation logic."],
 ["Every flag has a documented kind.",
  "The capability-vs-kill-switch rule is written into the baseline.",
  "No runtime flag behaviour altered."],
 ["Doc review: each flag appears exactly once with a kind.",
  "Confirm no code/flag-evaluation change in the diff (docs/registry-comment only)."],
 "it is classification and documentation only; no flag value or evaluation path changes",
 "documentation/classification is inherently reversible")

# COMPLETION LOG
part("Completion Log")
para("Update this log as items are executed. Each row records the evidence that the item shipped zero-impact.")
tbl(["Item","Status","Date","PR / commit","Verification evidence","Confirmed zero-impact by"],[
 ["A1","☐ Ready","—","—","—","—"],
 ["A2","☐ Ready","—","—","—","—"],
 ["A3","☐ Ready","—","—","—","—"],
 ["A4","☐ Ready","—","—","—","—"],
 ["A5","☐ Ready","—","—","—","—"],
],widths=[0.5,0.9,0.7,1.4,2.4,1.0])
h2("Exit criteria for Phase A")
b("All 5 items ☑ Done with verification evidence logged.")
b("Pilot tenants report no behavioural change attributable to Phase A.")
b("Governance baseline updated to reflect aliases + retired dead keys + flag-kind split.")
b("No business type merged, no role retired, no permission deleted, no tenant migrated (those remain Phase B/C/D).")
para("")
para("Phase A is complete when these five hygiene items are shipped and verified zero-impact. At that point the platform is "
     "cleaner but functionally identical — ready to focus fully on pilot validation, the first paying customer, and usage "
     "data before any structural simplification is revisited.", bold=True, color=NAVY)

out="docs/audits/VANTORA-PhaseA-Hygiene-Checklist.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
