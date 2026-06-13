#!/usr/bin/env python3
"""VANTORA Phase A — Completion Report (before/after validation) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows,widths=None,size=8.4,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# Cover
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Phase A — Completion Report"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Governance Hygiene · Before/After validation · Tests + Build"); r.font.size=Pt(12); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Branch: claude/fmcg-sell-collect-loop · 4 of 5 items executed · A2 reclassified to Phase C · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# 1. EXEC
h1("1. Executive summary")
para("Phase A (Governance Hygiene) is complete. Four zero-impact items were executed and verified; one (A2, permission "
     "consolidation) was correctly reclassified to Phase C after the business-impact review proved it would change access. "
     "The full test suite (1314 passing) and the production build (green) confirm no behavioural change to any existing "
     "tenant. No role, business type, permission family, or tenant was modified.", bold=True)
tbl(["Item","What shipped","Type","Status"],[
 ["A1","Removed inert engine permission keys from the entitlement map","Code","✔ DONE"],
 ["A2","Permission consolidation — held; reclassified to Phase C","—","⟶ Phase C"],
 ["A3","Bound nav visibility to the page's flag (no show-then-404)","Code","✔ DONE"],
 ["A4","Ratified the 5-archetype model + business-type map (Handbook Appendix A)","Docs","✔ DONE"],
 ["A5","Cataloged the two feature-flag kinds + ~38 KAKO_* tokens (Handbook Appendix A)","Docs","✔ DONE"],
],widths=[0.5,3.7,0.7,0.9],
 fill={(0,3):GREENBG,(1,3):AMBERBG,(2,3):GREENBG,(3,3):GREENBG,(4,3):GREENBG})

# 2. CODE CHANGES
part("2. Exact code changes (5 files, +20 / −11)")
para("Minimal, additive diff. Two production files for behaviour (registry.ts, nav-flags.ts + navigation.ts), two test "
     "files updated to lock in the new contracts.")
tbl(["File","Change","Lines"],[
 ["src/lib/entitlements/registry.ts","Removed 4 inert map keys (change_requests.* ×3, trade_spend.manage); added explanatory note","−4 keys"],
 ["src/lib/entitlements/entitlements.test.ts","Assert the removed keys are now unmapped ([])","±5"],
 ["src/lib/erp/nav-flags.ts","Emit 'trade_spend' / 'distribution' tokens when their kill-switch is ON","+7"],
 ["src/lib/erp/navigation.ts","Added flag:'trade_spend' to Trade Spend item; flag:'distribution' to Coverage item","2 items"],
 ["src/lib/erp/navigation-engine-gates.test.ts","Added the two new tokens to the test flag set","+1"],
],widths=[2.6,3.4,0.8])

# 3. BEFORE / AFTER
part("3. Before / After validation (per item)")
h2("A1 — dead permission keys")
tbl(["","Before","After"],[
 ["Code","registry.ts mapped change_requests.create/approve/manage + trade_spend.manage → modules","Those 4 entries removed; keys unmapped"],
 ["Enforcement","None — keys absent from the Permission union, granted to no role, no hasPermission() call","None — unchanged"],
 ["Effective access","modulesForPermission('trade_spend.manage') = ['trade_spend'] but never invoked","= [] (never gated); behaviour identical"],
 ["Test","asserted the keys mapped","asserts the keys are unmapped (regression guard)"],
],widths=[1.1,2.9,2.8])
para("Result: zero behavioural change — the keys enforced nothing before or after.", italic=True, color=GREEN)
h2("A3 — nav/page flag binding (show-then-404 fix)")
tbl(["","Before","After"],[
 ["Trade Spend nav","Shown to any tenant with the trade_spend module + reports.view","Shown only when KAKO_TRADE_SPEND is ON (matches the page)"],
 ["Trade Spend page","notFound() when KAKO_TRADE_SPEND off → user hits 404 from a visible link","Unchanged — but the link is now hidden when it would 404"],
 ["Coverage nav","Shown; page notFound() when KAKO_DISTRIBUTION off → 404","Shown only when KAKO_DISTRIBUTION is ON"],
 ["Net UX","Show-then-404 (the pilot-reported Trade Spend bug)","Link shown ⇔ page renders. Strictly better; never exposes anything new"],
],widths=[1.1,2.9,2.8])
para("Result: the show-then-404 class is closed for both hard-404 routes. Redirect-style pages were out of scope (graceful, "
     "not a 404).", italic=True, color=GREEN)
h2("A4 — archetype documentation")
para("Before: the 5-archetype model existed only as narrative. After: a ratified business-type→archetype map table in the "
     "Governance Handbook (Appendix A). No business_type value changed; documentation only.", color=GREEN)
h2("A5 — feature-flag catalog")
para("Before: capability flags and kill-switches were conceptually mixed (the root of show-then-404). After: a catalog in the "
     "Handbook (Appendix A) defining the two kinds, the governance rule, and the ~38 KAKO_* tokens. No runtime change.", color=GREEN)
h2("A2 — held (reclassified)")
para("Evidence (fmcg/actions.ts, field/van-sales, permissions.ts) shows stock.*/inventory.*, customer.*/customers.*, "
     "report.aggregate.view/reports.view are distinct, separately-enforced permissions. Aliasing would widen access, so it "
     "was moved to Phase C with a migration plan. Roadmap + Handbook updated accordingly.", color=AMBER)

# 4. TEST + BUILD
part("4. Test & build evidence")
h2("Unit / integration tests — vitest run")
code("Test Files  180 passed | 69 skipped (249)")
code("     Tests  1314 passed | 181 skipped (1495)")
code("     0 failed")
para("Affected suites specifically re-run and green: entitlements, navigation, navigation-engine-gates, navigation-routes, "
     "permissions (66 passed).")
h2("Type check — tsc --noEmit")
code("tsc exit: 0  (no type errors)")
h2("Production build — next build")
code("✓ Compiled successfully")
code("✓ Route table generated (all app routes); exit code 0")
tbl(["Gate","Result"],[
 ["Full test suite","PASS (1314, 0 failed)"],["Type check","PASS (exit 0)"],["Production build","PASS (exit 0)"],
],widths=[2.5,3.0],fill={(0,1):GREENBG,(1,1):GREENBG,(2,1):GREENBG})

# 5. ARTIFACTS + GUARDRAILS
part("5. Updated artifacts & guardrails")
h2("Governance baseline updated")
b("Governance Handbook — added Appendix A (Phase A ratified): A4 archetype map, A5 flag catalog, A1/A3 status, A2 reclassification.")
b("Implementation Roadmap — A1/A3/A4/A5 marked DONE; A2 moved from Quick Wins (Immediate) to Phase C with the evidence note.")
b("Phase A checklist tracker — items A1/A3/A4/A5 complete; A2 marked reclassified.")
h2("Guardrails honoured (what was NOT done)")
tbl(["Constraint","Honoured?"],[
 ["No role consolidation","Yes — 27 roles untouched"],
 ["No business-type consolidation","Yes — all 21 types untouched; documentation only"],
 ["No permission-family consolidation","Yes — A2 explicitly held; no alias added"],
 ["No tenant migration","Yes — no data migration of any kind"],
 ["A2 not implemented","Yes — reclassified to Phase C"],
],widths=[3.4,2.1],fill={(i,1):GREENBG for i in range(5)})

# 6. SIGN-OFF
part("6. Sign-off & next steps")
para("Phase A is complete and verified zero-impact. The platform is cleaner (inert keys gone, nav/page flags aligned, "
     "governance baseline documented) but functionally identical for every existing tenant.", bold=True, color=NAVY)
h2("Recommended next steps (unchanged objective)")
b("Resume the stated objective: (1) complete pilot validation, (2) get the first paying customer, (3) collect real usage data.")
b("Operational note (not code): the A3 binding means Trade Spend / Coverage links appear only where KAKO_TRADE_SPEND / "
  "KAKO_DISTRIBUTION are set. To expose Trade Spend in the pilot, set KAKO_TRADE_SPEND=1 in the deployment env.")
b("Hold all structural simplification (Phase B/C/D), including A2, until after the first paying customer + usage data.")
para("")
para("Phase A: DONE. No structural change was made. The next governance step (Phase B) begins only on your go-ahead, after "
     "the pilot and first paying customer.", bold=True, color=NAVY)

out="docs/audits/VANTORA-PhaseA-Completion-Report.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
