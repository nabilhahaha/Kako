#!/usr/bin/env python3
"""VANTORA Governance Implementation Roadmap (founder decision report) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(19); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.4,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# Cover
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Governance Implementation Roadmap"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("A founder decision report — what to simplify, when, and what to never touch"); r.font.size=Pt(11); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Companion to the VANTORA Governance Architecture Handbook · Planning only — no code changes · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---------- 1. EXECUTIVE RECOMMENDATION ----------
h1("1. Executive Recommendation")
para("This roadmap turns the Governance Handbook's findings into a sequenced, risk-rated plan. It is a decision document, "
     "not a change. Nothing here modifies code, schema, roles, modules, business types, or permissions. It tells you what "
     "is safe to simplify now, what must wait for the pilot or the first paying customer, and what should never be touched.",
     bold=True)
para("The core message: VANTORA's complexity is taxonomy, not capability. That distinction makes the cleanup unusually safe. "
     "The biggest wins — deduplicating permission families and collapsing identical business types into labels — can be done "
     "as additive, backward-compatible aliasing with ZERO impact on any current tenant. The risky work (retiring management "
     "role tiers, re-scoping roles) is deferred until after the pilot and first paying customer, when you have real usage "
     "data to validate the merge.")
tbl(["Recommendation","Founder guidance"],[
 ["Adopt the handbook as baseline","Yes — freeze it as the official governance reference now."],
 ["Start simplifying during the pilot?","Only the ZERO-impact items (Quick Wins). Everything structural waits."],
 ["Touch roles/business types mid-pilot?","No. Pilot tenants are watching behavior; defer all merges."],
 ["What never gets touched","Platform↔tenant separation; engine flag-gating; core ERP entitlement model."],
 ["First structural change","Permission deduplication via aliases — additive, reversible, invisible to users."],
],widths=[2.4,4.6])
para("Bottom line: approve the Quick Wins (Section 3) now; schedule the structural simplification (Section 6) for the "
     "post-pilot window; treat Section 4 (High-Risk) as gated behind real customer data. Net target: business types 21→~5, "
     "roles 27→~8+scope, permissions 91→~60 — phased so no tenant ever sees a regression.", color=NAVY, bold=True)

# ---------- 2. PRIORITY MATRIX ----------
part("2. Priority Matrix")
para("Every handbook recommendation, classified by WHEN it should happen, with benefit / risk / effort / tenant impact / "
     "data-migration. Timing buckets: IMM=Immediate (safe now) · PP=Post-Pilot · APC=After First Paying Customer · "
     "LT=Long-Term Architecture.")
LEG={"IMM":GREENBG,"PP":BLUEBG,"APC":AMBERBG,"LT":REDBG}
def C(x): return LEG[x]
rows=[
 # rec, when, benefit, risk, effort, tenant impact, migration
 ["[A1 ✔ DONE] Drop dead permission keys (change_requests.*, trade_spend.manage)","IMM","Removes confusion; cleaner audit","Low","S","None (unenforced)","No"],
 ["[A2 ⟶ Phase C] Consolidate permission families — NOT aliasable: they are distinct enforced perms","APC","Real dedup, but needs migration","High","L","YES — changes effective access","YES"],
 ["[A3 ✔ DONE] Tie nav + page gating to the SAME flag (fix show-then-404)","IMM","Fixes real UX bug (Trade Spend 404)","Low","S","Positive (no broken links)","No"],
 ["[A4 ✔ DONE] Document the 5 archetypes (no type change)","IMM","Clarifies onboarding","Low","S","None (existing tenants keep type)","No"],
 ["[A5 ✔ DONE] Catalog the two feature-flag kinds (capability vs kill-switch)","IMM","Prevents future show-then-404","Low","S","None","No"],
 ["Don't seed optional/vertical roles by default in new tenants","PP","Simpler new-tenant setup","Low","S","New tenants only","No"],
 ["Introduce role 'scope' dimension (branch/area/region/national) additively","PP","Foundation to retire mgmt tiers","Medium","L","None until adopted","No"],
 ["Collapse identical-module business types into 'Retail' archetype (new tenants)","PP","Fewer types to reason about","Medium","M","New tenants only","No"],
 ["Fold sales_orders capability into sales (nav grouping, keep data)","PP","One fewer module concept","Medium","M","None (data unchanged)","No"],
 ["Merge change_requests/critical_alerts presentation under workflow/analytics","PP","Fewer top-level modules","Medium","M","Nav only","No"],
 ["Retire 5 management role tiers → Manager + scope (migrate existing assignments)","APC","27→~8 roles; huge clarity","High","L","Existing tenants w/ those roles","YES"],
 ["Migrate existing tenants' business_type to archetype + label","APC","Real type-count reduction","High","L","All multi-type tenants","YES"],
 ["Physically remove deduplicated permissions after alias soak","APC","Final permission cleanup","Medium","M","All tenants (post-soak)","YES (cleanup)"],
 ["Role inheritance from base templates (stop re-listing perms per role)","LT","Lower long-term maintenance","Medium","L","None (internal model)","Maybe"],
 ["Archetype-driven module sets (type inherits archetype, adds vertical)","LT","Cleanest taxonomy","High","L","Internal; tenant-invisible","Maybe"],
 ["First-class engine permissions in code (typed, not flag-only)","LT","Stronger authz model","Medium","M","None","No"],
 ["MFA / SSO / session permissions (currently a Gap)","LT","Enterprise security readiness","High","L","Opt-in per tenant","No"],
]
tbl(["Recommendation","When","Benefit","Risk","Eff.","Tenant impact","Migr."],
    [[r[0],r[1],r[2],r[3],r[4],r[5],r[6]] for r in rows],
    widths=[2.5,0.5,1.3,0.55,0.4,1.25,0.55],size=7.6,
    fill={(i,1):C(r[1]) for i,r in enumerate(rows)})
para("Effort: S=days · M=1–2 weeks · L=multi-week. Legend colours: green=Immediate, blue=Post-Pilot, amber=After Paying "
     "Customer, red=Long-Term.", italic=True, size=8.5)

# ---------- 3. QUICK WINS ----------
part("3. Quick Wins  —  ZERO impact on current functionality")
para("These are the recommendations that can be implemented NOW with zero behavioural change for any existing tenant. They "
     "are additive, reversible, and invisible to users. This is the single most important section: it is the safe set.",
     bold=True, color=GREEN)
para("UPDATE (post business-impact review): the original Quick-Win set listed permission ALIASING as zero-impact. The "
     "codebase review proved it is NOT — the 'duplicate' families are distinct enforced permissions, so aliasing changes "
     "access. It was removed from Quick Wins and reclassified to Phase C. The remaining FOUR were executed.", bold=True, color=AMBER)
tbl(["Quick win","Why it's zero-impact","Status"],[
 ["A1 · Remove dead permission keys (change_requests.*, trade_spend.manage)","Not in the code Permission type, granted to no role — removal changes no behaviour","✔ DONE"],
 ["A3 · Bind nav visibility to the same flag the page enforces","A screen either shows AND works, or is hidden — strictly better than today's show-then-404","✔ DONE"],
 ["A4 · Document the 5 archetypes (existing types unchanged)","Existing tenants keep their exact business_type; documentation only","✔ DONE"],
 ["A5 · Catalog the two feature-flag kinds (capability vs kill-switch)","Pure classification + documentation; no runtime change","✔ DONE"],
 ["A2 · Permission aliasing — REMOVED FROM QUICK WINS","NOT zero-impact: distinct enforced permissions; aliasing widens access","⟶ Phase C"],
],widths=[2.5,3.2,0.9],fill={**{(i,2):GREENBG for i in range(4)},(4,2):AMBERBG,(4,0):AMBERBG})
para("Founder decision (executed): A1, A3, A4, A5 shipped as the governance-hygiene batch — tests 1314 passing, build green, "
     "zero behavioural change to existing tenants. A2 held for Phase C with a migration plan. No role, business type, "
     "permission family, or tenant was changed.", bold=True)

# ---------- 4. HIGH-RISK CHANGES ----------
part("4. High-Risk Changes  —  gate behind real customer data")
para("These deliver the largest taxonomy reduction but carry real risk because they touch live role assignments or a "
     "tenant's business_type. They REQUIRE data migration and must not be attempted during the pilot.", bold=True, color=RED)
tbl(["Change","Risk","Why risky","Precondition before starting"],[
 ["Retire 5 management role tiers → Manager + scope","High","Existing users hold area/regional/national/sales_director/branch roles; remapping can change who can approve/see what","After first paying customer; snapshot + reversible migration + per-tenant validation"],
 ["Migrate business_type → archetype + vertical label","High","business_type drives module sets and home routing; a bad map can hide modules a tenant relies on","Archetype map proven on new tenants first; dual-read during transition"],
 ["Physically delete deduplicated permissions","Medium-High","If any custom role still references the old key post-alias, a grant could vanish","Alias soak period (≥1 release) + usage scan showing zero references"],
 ["Role inheritance refactor","Medium-High","Changes how every role resolves its permissions internally","Comprehensive role-permission regression suite green"],
],widths=[2.0,0.7,2.4,2.0],
 fill={(0,1):REDBG,(1,1):REDBG,(2,1):AMBERBG,(3,1):AMBERBG})
para("Founder guidance: none of these are urgent. They are the 'after we have paying customers and real usage telemetry' "
     "tier. Each must ship behind a reversible migration with a per-tenant validation gate. Do not bundle them — sequence "
     "one at a time.", bold=True)

# ---------- 5. PILOT-SAFE CHANGES ----------
part("5. Pilot-Safe Changes")
para("What may be touched while the pilot is live without putting pilot feedback at risk. The rule: anything that does not "
     "change what a pilot user sees or can do is pilot-safe.")
tbl(["Pilot-safe (OK during pilot)","NOT pilot-safe (defer)"],[
 ["All Quick Wins (Section 3)","Retiring management role tiers"],
 ["Dead-key removal","Re-scoping existing role assignments"],
 ["Permission aliasing (additive)","Migrating any tenant's business_type"],
 ["Nav/page flag binding (fixes a bug)","Physically deleting permissions"],
 ["Documentation of archetypes & flag kinds","Module nav reorganisation that moves where users click"],
 ["Curating the new-tenant onboarding picker","Anything requiring data migration"],
],widths=[3.4,3.4],
 fill={**{(i,0):GREENBG for i in range(6)}, **{(i,1):REDBG for i in range(6)}})
para("Why this split: the pilot's purpose is to gather behavioural feedback. Any change that moves buttons, hides modules, "
     "or alters who-can-do-what contaminates that signal. Additive/invisible changes do not.", italic=True)

# ---------- 6. POST-PILOT SIMPLIFICATION PLAN ----------
part("6. Post-Pilot Simplification Plan")
para("The sequenced structural work, to begin once the pilot closes and (for the High-Risk tier) the first paying customer "
     "is live. Each phase is independently shippable and non-breaking on its own.")
h2("Phase A — Governance Hygiene (Immediate / pilot-safe)")
b("Ship the Quick Wins batch: dead-key removal, permission aliases, nav/page flag binding, archetype + flag-kind docs.")
b("Outcome: permission overlap eliminated at the alias layer; the show-then-404 class of bug closed. Zero tenant impact.")
h2("Phase B — Additive Foundations (Post-Pilot)")
b("Introduce the role SCOPE dimension additively (no role retired yet) and stop seeding optional/vertical roles by default.")
b("Curate new-tenant onboarding to the 5 archetypes; new tenants get archetype types, existing tenants unchanged.")
b("Group sales_orders under sales and present change_requests/critical_alerts under workflow/analytics (nav only).")
b("Outcome: new tenants are born simple; existing tenants untouched; foundations laid for the merges.")
h2("Phase C — Structural Merge (After First Paying Customer)")
b("Retire the 5 management role tiers into Manager + scope, with a reversible per-tenant migration and validation gate.")
b("Migrate existing tenants' business_type to archetype + label using the map proven on new tenants in Phase B.")
b("After the alias soak, physically remove the deduplicated permissions once a usage scan shows zero references.")
b("Outcome: the real count reductions land — roles 27→~8+scope, business types 21→~5, permissions 91→~60.")
h2("Phase D — Long-Term Architecture")
b("Role inheritance from base templates; archetype-driven module inheritance; first-class engine permissions; MFA/SSO.")
b("Outcome: lowest long-term maintenance surface; enterprise security readiness.")
tbl(["Phase","Trigger","Touches tenants?","Migration?"],[
 ["A — Hygiene","Now / during pilot","No","No"],
 ["B — Foundations","Pilot closed","New tenants only","No"],
 ["C — Structural merge","First paying customer","Yes (reversible)","Yes"],
 ["D — Long-term","Post-scale","Internal only","Maybe"],
],widths=[1.8,2.0,1.6,1.4],
 fill={(0,2):GREENBG,(1,2):BLUEBG,(2,2):AMBERBG,(3,2):GREENBG})

# ---------- 7. FINAL TARGET ARCHITECTURE ROADMAP ----------
part("7. Final Target Architecture Roadmap")
para("Where each dimension lands, and the phase that gets it there. Every current capability is preserved.", bold=True)
tbl(["Dimension","Today","Target","Lands in phase","Migration"],[
 ["Business Types","21","~5 archetypes (+labels)","B (new) → C (existing)","C only"],
 ["Modules","31","~24 in 3 tiers","B (presentation)","No"],
 ["Roles","27","~8 templates + scope","B (scope) → C (retire tiers)","C only"],
 ["Permissions (DB)","91","~60 verb.resource","A (alias) → C (delete)","C only"],
 ["Permission groups","16","8 categories","A","No"],
 ["Feature-flag kinds","2 mixed","2 separated","A","No"],
 ["Security (MFA/SSO)","Gap","First-class","D","No"],
],widths=[1.6,0.7,2.0,1.7,0.9],
 fill={(0,2):GREENBG,(1,2):GREENBG,(2,2):GREENBG,(3,2):GREENBG,(4,2):GREENBG,(5,2):GREENBG,(6,2):BLUEBG})
h2("The one-line sequence")
para("A (now, zero-impact) → B (post-pilot, new-tenant only) → C (after paying customer, migrated + reversible) → "
     "D (long-term). Nothing structural touches a live tenant until Phase C, and even then behind a reversible migration "
     "with a per-tenant validation gate.", bold=True, color=NAVY)
h2("What must NEVER be touched")
b("Platform Owner ↔ Tenant Admin ↔ scoped-role separation (the Provider panel boundary). Correct and load-bearing.")
b("Engine modules as company/business-type driven + kill-switch flag (not plan-gated). Correct design.")
b("Core ERP modules as plan-licensable entitlements. Correct.")
b("Tenant data isolation / RLS. Never weaken for the sake of taxonomy simplification.")

para("")
para("Founder takeaway: approve Phase A now (zero risk, real cleanup), let the pilot run undisturbed, and hold every "
     "structural merge until you have a paying customer and a reversible migration. This roadmap simplifies VANTORA by "
     "roughly half its taxonomy surface without ever putting a live tenant at risk.", bold=True, color=NAVY)

out="docs/audits/VANTORA-Governance-Implementation-Roadmap.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
