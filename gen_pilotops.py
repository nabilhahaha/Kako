#!/usr/bin/env python3
"""VANTORA Pilot Operations Pack (manual UI checklist + feedback) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.2,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Pilot Operations Pack"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Role-by-role UI verification · daily-activity scripts · usability log · feedback templates"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Pilot tenant VANTORA Pilot FMCG (DEMO) on staging · operate, observe, capture · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("How to use this pack")
para("VANTORA is now in PILOT OPERATION. This pack is the operator's tool to verify the system through the UI, run the daily "
     "FMCG activities on the live demo data, and capture usability + issues in a consistent format. Tick each step, note "
     "anything that feels slow/confusing, and log issues using the categories at the back. Only issues found in real usage "
     "should drive fixes — no new modules, workflows or architecture.", bold=True)
h2("Logins (staging) — all passwords test.123")
tbl(["Role","Email","Lands on (verified)"],[
 ["Company Admin","admin@pilot.test","/dashboard"],
 ["Salesman / Van rep","salesman@pilot.test","/today"],
 ["Supervisor","supervisor@pilot.test","/approvals/queue"],
 ["Accountant","accountant@pilot.test","/collections"],
 ["Warehouse Keeper","warehouse@pilot.test","/inventory/requests"],
],widths=[1.8,2.4,2.0])

h1("Access layer — already verified on the live tenant")
para("Confirmed by direct database checks (login, landing, and capability per role). Use the UI checklist next to confirm "
     "the visual/click experience matches.", bold=True)
tbl(["User","Login","Lands on","Sell","Collect","Post GL","Approve","Purchasing"],[
 ["admin","OK","/dashboard","Yes","Yes","Yes","Yes","Yes"],
 ["salesman","OK","/today","Yes","Yes","No","No","No"],
 ["supervisor","OK","/approvals/queue","Yes","Yes","No","Yes","No"],
 ["accountant","OK","/collections","No","Yes","Yes","No","No"],
 ["warehouse","OK","/inventory/requests","No","No","No","Yes","Yes"],
],widths=[1.2,0.6,1.6,0.6,0.7,0.7,0.7,0.9])
para("Separation of duties holds: a salesman can sell/collect but not post the GL or approve; an accountant posts but can't "
     "sell; a warehouse keeper approves loads + receives but can't sell/collect. All five authenticate with test.123.",
     italic=True, color=GREEN)

# UI checklist per role
def checklist(role, rows):
    part(f"UI verification — {role}")
    para("Tick each step in the browser. Mark Pass/Fail and note anything slow or confusing.", italic=True)
    tbl(["#","Step","Expected","Pass/Fail","Notes"],
        [[str(i+1)]+r for i,r in enumerate(rows)],
        widths=[0.3,2.2,2.0,0.8,1.2])

checklist("Salesman (salesman@pilot.test)",[
 ["Log in","Lands on /today (your day)",""],
 ["Open route / customers","See 5 customers in sequence",""],
 ["Create an invoice (or POS sell)","Price is pre-filled and locked; invoice issues",""],
 ["Record a collection","Customer balance reduces; receipt created",""],
 ["Request a van load (/inventory/van-transfer)","Submits a pending request to the supervisor",""],
 ["Close the day","A coverage exception is sent for approval",""],
 ["Mobile: use bottom-nav Home/Today/Customers/Sell/Inventory","All reachable; 'More' shows the rest",""],
 ["Try to post a voucher / change a price","Blocked (no permission)",""],
])
checklist("Supervisor (supervisor@pilot.test)",[
 ["Log in","Lands on /approvals/queue with 4 pending items",""],
 ["Filter by type / status","List filters correctly",""],
 ["Approve a day-close with a comment","Item moves to Approved; audit recorded",""],
 ["Reject an out-of-route visit","Item moves to Rejected",""],
 ["Approve the pending van transfer","Stock moves; status completed",""],
 ["Mobile: Approvals bottom-nav tab","Opens the same queue in one tap",""],
])
checklist("Accountant (accountant@pilot.test)",[
 ["Log in","Lands on /collections",""],
 ["Review outstanding AR (~4,906 EGP)","Customer balances shown",""],
 ["Record a collection","Balance reduces; receipt created",""],
 ["Post a voucher (Accounting · Vouchers)","GL journal posts",""],
 ["Open AR aging / financial reports","Populated with pilot history",""],
 ["Try to issue an invoice","Blocked (no sales.sell)",""],
])
checklist("Warehouse Keeper (warehouse@pilot.test)",[
 ["Log in","Lands on /inventory/requests",""],
 ["Approve a pending stock/van load request","Stock moves to the van",""],
 ["Receive a purchase order","Stock + AP updated",""],
 ["Adjust stock / run a count","Movement recorded",""],
 ["Try to record a collection","Blocked (no sales.collect)",""],
])
checklist("Company Admin (admin@pilot.test)",[
 ["Log in","Lands on /dashboard with KPIs + quick actions",""],
 ["Use a quick action (New Invoice / Collect / New Customer)","Opens the right screen in one tap",""],
 ["Settings · Staff → add a user, assign a role","User created; role assigned",""],
 ["Open Reports","Populated with pilot sales/collections",""],
 ["Switch language AR ↔ EN","UI flips RTL/LTR correctly",""],
])

part("Daily FMCG activity script  (end-to-end on the demo data)")
tbl(["#","Activity","Roles involved","Expected outcome"],[
 ["1","Morning: rep starts day, works the route","Salesman","Visits logged, sells + collects"],
 ["2","Sell-to-collect loop","Salesman","Invoice issued, payment recorded, no duplicate on double-click"],
 ["3","Van replenishment","Salesman → Supervisor","Load request → approved → stock on van"],
 ["4","Out-of-route exception","Salesman → Supervisor","Visit flagged → approved/rejected in the queue"],
 ["5","End of day close","Salesman → Supervisor","Coverage exception → approved"],
 ["6","Finance run","Accountant","Post collections, review AR aging, post vouchers"],
 ["7","Replenish warehouse","Warehouse → (PO)","Receive PO, stock updated"],
 ["8","Oversight","Admin / Branch Mgr","Dashboard KPIs, reports, approvals cleared"],
],widths=[0.3,2.0,1.8,2.4])

part("Usability observation log  (fill during the pilot)")
tbl(["Date","Role","Screen","What felt slow / confusing","Suggestion","Severity"],[
 ["","","","","",""],["","","","","",""],["","","","","",""],
 ["","","","","",""],["","","","","",""],["","","","","",""],
],widths=[0.8,1.0,1.1,1.8,1.4,0.8])

part("Issue / feedback capture  (categories only — no feature requests)")
para("Log only real problems found in usage. Allowed categories: Bug · Permission · Mobile · Performance · Deployment. A "
     "feature request is NOT logged unless a real daily process is blocked.", bold=True)
tbl(["ID","Date","Role","Category","Screen / step","Description","Severity (Blocker/Major/Minor)","Status"],[
 ["F-001","","","","","","",""],["F-002","","","","","","",""],
 ["F-003","","","","","","",""],["F-004","","","","","","",""],
 ["F-005","","","","","","",""],
],widths=[0.5,0.7,0.8,0.9,1.2,1.6,1.0,0.8])
h2("Triage rule")
b("Blocker → fix immediately (stops a daily process). Major → fix this pilot iteration. Minor → backlog.")
b("Fix ONLY what real usage surfaces. No new modules, workflows, or architecture changes during the pilot.")
para("")
para("Status: the access/permission/data layer is verified live and passing. This pack drives the remaining human UI/"
     "usability confirmation and structured feedback capture for the operational pilot.", bold=True, color=NAVY)

out="docs/audits/VANTORA-Pilot-Operations-Pack.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
