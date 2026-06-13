#!/usr/bin/env python3
"""VANTORA Platform Coverage Audit -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.2,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Platform Coverage Audit"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Visible · Reachable · Usable · Protected — screen-by-screen, role-by-role, workflow-by-workflow"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Branch claude/fmcg-sell-collect-loop · post-P0 · M-1/M-2/U-4 resolved · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Executive summary")
para("Read-only coverage sweep of the whole FMCG surface (nav, desktop, mobile, pages, actions, RPCs, APIs) plus the three "
     "carried-over items (M-1 Cash Van, M-2 tenant-admin model, U-4 price/credit gating). The platform is largely "
     "complete & reachable: ZERO dead links, mobile == desktop visibility (the 'More' drawer renders the full sidebar), and "
     "all 15 API routes are by-design or already consumed. The real gap is a cluster of SUPERVISOR-APPROVAL workflows that "
     "are fully built in the backend (actions + RPCs + seeded permissions) but have NO screen — so those workflows cannot be "
     "completed end-to-end by the intended role.", bold=True)
tbl(["Result","Count","Severity"],[
 ["Dead links (nav → missing page)","0","—"],
 ["Implemented but Hidden (backend-not-exposed)","13 actions / 3 RPCs","Major (approval surfaces)"],
 ["Visible but Incomplete (stub tile)","1","Minor"],
 ["Missing Navigation","2","Minor"],
 ["Mobile Visibility Issues","1 (structural: none)","Minor"],
 ["Carried items resolved","M-1, M-2, U-4","Closed / proposed / fixed"],
],widths=[3.3,1.4,1.7],fill={(1,2):AMBERBG})

# ===== 10 CATEGORIES =====
part("1. Implemented & Visible  (the bulk of the platform)")
para("Every navigation href resolves to a real page (0 dead links). The core FMCG loop — sell, collect, inventory, "
     "purchasing, accounting, distribution dashboards, pharmacy pack — is implemented, in navigation, on desktop, and "
     "reachable on mobile via the 'More' drawer. Verified working (not stubs): integrations sub-screens, van sell/collect "
     "(server-authoritative pricing + settlement), customer/supplier/product CRUD, reports.")

part("2. Implemented but Hidden  (backend-not-exposed — the headline gap)")
para("Exported server actions (and their RPCs) that are fully implemented with seeded permissions but have NO UI caller — so "
     "the capability exists yet no role can reach it in-app.", bold=True)
tbl(["Capability","Actions (file)","Permission (seeded)","Risk"],[
 ["Van stock TRANSFER request→approve→reject","requestVanTransfer / approveVanTransfer / rejectVanTransfer (field/actions.ts) + RPCs erp_*_van_transfer","stock_request.* / approvals","Major"],
 ["Day-close EXCEPTION approval (supervisor)","approveDayClose (field/actions.ts) — reps closeDay, supervisors can't approve","day.approve_close_exception","Major"],
 ["Out-of-route VISIT compliance decision","decideVisitCompliance (field/actions.ts:167)","visit.approve_out_of_route","Major"],
 ["Customer / rep REASSIGNMENT approve","transferCustomer / approveCustomerTransfer / transferUser","customer.transfer / user.transfer","Major"],
 ["Trade-spend promo APPROVE / CANCEL","approveTradeSpend / cancelTradeSpend (page is view-only)","reports.view (+ flag)","Major (flag-gated)"],
 ["Reopen day","reopenDay (rep/actions.ts)","field.sales","Minor"],
 ["resolvePriceAction / uomToBaseAction","fmcg/actions.ts — likely dead exports (POS resolves inline)","—","Minor"],
],widths=[2.0,2.4,1.2,0.9],
 fill={(0,3):AMBERBG,(1,3):AMBERBG,(2,3):AMBERBG,(3,3):AMBERBG,(4,3):AMBERBG})
para("Plus 1 ORPHAN page: /forms/customer-data-update (FORM_BUILDER-gated) is reachable by URL only — no launcher links to "
     "it. Minor / likely flag-staged.", italic=True)

part("3. Visible but Incomplete")
b("'Confirm Load' tile on the Van-Sales 'My Day' hub (field/van-sales/page.tsx:28-39) shows a 'Coming soon' chip and is "
  "non-clickable — even though a working /field/van-sales/confirm route EXISTS. The tile simply isn't wired (stale chip). "
  "Minor; fix = wire the tile to /confirm or remove the chip.")
b("Flag-gated engine pages (alerts, change-requests, trade-spend, coverage, workflow-templates, forms) are disabled-by-"
  "default and correctly hidden by their nav flag (A3) — NOT stubs; full backends exist when enabled.")

part("4. Missing UI  (workflows with no screen for the intended role)")
para("These are the same as Section 2's Major items, framed as the missing SCREEN. The supervisor/manager approval surfaces "
     "are the priority — reps can act, but their approver cannot complete the loop in-app:", bold=True)
b("Supervisor approval queue for: day-close exceptions, out-of-route visits, customer/rep transfers, van transfers.")
b("Trade-spend: approve/cancel controls on the (currently view-only) trade-spend page.")
para("Fix type: WIRE existing actions to a screen (e.g. extend /supervisor or /approvals with an approval queue). No new "
     "backend — purely UI exposure of built capability.", italic=True)

part("5. Missing Navigation")
tbl(["Item","Detail","Risk","Fix"],[
 ["Van-Sales hub /field/van-sales/*","No nav entry; reachable only via the /today bottom tab's 'My Day' card when van_sales is active","Minor","Add a flag-gated nav item for the van-sales hub"],
 ["/forms/customer-data-update","No launcher / nav entry; deep-link only","Minor","Add a launch button on the customer/field screen, or confirm by-design"],
],widths=[1.9,3.0,0.7,1.6])

part("6. Missing Permissions  — RESOLVED (U-4)")
para("The only open permission gap (master-data price & credit-limit writes were auth-only) is now CLOSED in this pass — see "
     "Section 'U-4' below. No other tested role is missing a permission it needs.", bold=True, color=GREEN)

part("7. Role Mismatches")
tbl(["Finding","Detail","Risk","Status"],[
 ["Electrical section no module gate","Relies only on electrical.rma perm; a non-electronics admin granted it would see RMA screens","Minor","Open — MN-2 (P1)"],
 ["Accountant holds customers.change_status","Broader than the name implies (can suspend/block + now set credit)","Minor","Accept / trim in role-template review"],
 ["No other role mismatch","Menu gate → page guard → action gate are consistent (RBAC audit)","—","OK"],
],widths=[1.8,3.0,0.7,1.5])

part("8. Mobile Visibility")
para("Structural finding: the mobile 'More' drawer renders the SAME visibleSections list as the desktop sidebar "
     "(sidebar.tsx) — so mobile visibility EQUALS desktop visibility; only the 4 quick bottom tabs (Home, Today, Customers/"
     "Sell, Inventory) differ. There are no desktop-only screens hidden from mobile.", bold=True)
tbl(["Field-rep destination","Mobile reachable?"],[
 ["Today / journey","Yes — bottom tab"],
 ["Record collection","Yes — More drawer"],
 ["Van reconciliation","Yes — More drawer"],
 ["Offline mode","Yes — More drawer"],
 ["Van-sales hub (load/sell/collect)","Indirect — via /today 'My Day' card (no own nav entry; see §5)"],
],widths=[3.0,2.5])

part("9. Dead Links / Dead Screens")
b("Dead links: NONE — every nav/bottom-nav href resolves to an existing page.",color=GREEN)
b("Dead screen: the 'Confirm Load' coming-soon tile (§3) is the only visibly-inert element, and a working route for it "
  "exists. Everything else renders working content or a legitimate empty state.")

part("10. Backend Features Not Exposed in UI")
para("Consolidated from Section 2: 13 server actions + 3 RPCs implemented with no UI surface. The Major cluster is the "
     "field-supervisor approval set; the rest are minor/dead exports. API routes (15) are all by-design (cron/integration/"
     "health) or already consumed — 0 gaps.", bold=True)

# ===== CARRIED ITEMS =====
part("M-1 — Cash Van role alignment  (CLOSED)")
para("Investigated code AND database. Finding: 'cash_van' is NOT a role in either layer.", bold=True)
b("BranchRole union (types.ts) defines 21 roles — admin, manager, sales_director, national/regional/area managers, "
  "branch_manager, it_admin, supervisor, accountant, cashier, salesman, driver, technician, doctor, receptionist, stylist, "
  "housekeeping, warehouse_keeper, staff, viewer. There is NO cash_van.")
b("DB erp_role_permissions seed (migrations 0017/0102/0107 + vertical packs) seeds admin, manager, branch_manager, "
  "supervisor, salesman, accountant, cashier, warehouse_keeper, staff, viewer + vertical roles. There is NO cash_van — it "
  "appears only in a code COMMENT (0269).")
b("Therefore code and DB ALREADY AGREE — there is no mismatch to fix. The van cash-rep is the 'salesman' role, which carries "
  "the full van loop: field.sales, sales.sell, sales.collect, stock_request.create, stock.transfer, reconciliation.view.")
para("Resolution: M-1 is closed by confirmation. 'Cash Van' ≡ the salesman role; any informal 'cash van' assignment should "
     "map to salesman (or 'driver' for a lighter variant). Adding a distinct cash_van role would be a NEW feature (type "
     "union + exhaustive maps + DB seed + home routing + i18n) and is intentionally NOT done, per the 'no new features' "
     "priority. If a separate van-sales selling model is later desired, it is a deliberate roadmap item.", bold=True, color=NAVY)

part("M-2 — Tenant-admin user & permission management model  (PROPOSAL)")
para("The earlier finding ('company admin cannot manage users/permissions') was partly a naming issue. The admin role "
     "(settings.users) DOES have a full tenant-admin path; the super-admin-only screens are PLATFORM-level (cross-tenant).",
     bold=True)
tbl(["Capability","Tenant-admin path (admin role)","Platform path (global super-admin)"],[
 ["Create staff & assign roles to branches","Settings · Staff  (perm settings.users) ✓","Settings · Users (superAdminOnly)"],
 ["Customize role→permission / policies","Settings · Authz + Action Policies + Features (settings.users) ✓","Settings · Permissions (superAdminOnly)"],
 ["e-invoice / global audit","—","superAdminOnly"],
],widths=[2.0,2.6,1.9])
h2("Recommended model (keep the boundary; one clarity tweak)")
b("KEEP Settings · Users and Settings · Permissions super-admin-only — they are cross-tenant/platform screens; exposing "
  "them to tenant admins would breach the platform boundary.",bold=True)
b("The tenant admin's user-management home is Settings · Staff (+ Authz / Action Policies for governance). This is correct "
  "and sufficient for the pilot — a company admin can onboard staff, assign roles, and tune policies today.")
b("Optional clarity-only improvement: relabel 'Staff' → 'Users & Roles' so tenant admins discover it as the user-management "
  "screen. No structural/permission change. (Not done in this pass — it's a label/i18n decision.)")
para("Conclusion: no real capability gap; the model is sound. Recommendation: adopt the boundary above as the official "
     "tenant-admin model; the only optional change is the label.", bold=True, color=NAVY)

part("U-4 — Pricing & credit-limit gating  (FIXED in this pass)")
para("Master-data price and credit-limit writes were auth-only (any logged-in user could set arbitrary prices / credit "
     "limits). Now field-level permission-gated — additive, non-breaking (non-price/non-credit edits unaffected).", bold=True)
tbl(["Write","Gate added","Behaviour when unauthorized"],[
 ["Product cost_price / sell_price (upsertProduct)","pricing.manage OR product.create","Edit: preserve existing prices · Create: default 0 (other fields still editable)"],
 ["Customer credit_limit (upsertCustomer)","customers.status.change (via can())","Edit: preserve existing limit · Create: default 0 (real increases route via the credit-limit request workflow)"],
],widths=[2.1,1.9,2.5])
b("Of the 9 audited roles, only admin/manager (and sales_director/NSM for pricing) may now change prices; credit limits "
  "require credit/status authority (admin/manager/branch_manager/supervisor/accountant). salesman/cashier/viewer are "
  "blocked from both — verified by an added regression test.")
para("Verified: tsc clean · 1318 tests pass (incl. the new U-4 assertions) · build green.", bold=True, color=GREEN)

# ===== WORKFLOW E2E =====
part("Workflow-by-workflow: can the intended role complete it end-to-end?")
tbl(["Workflow","Intended role(s)","End-to-end in UI?","Gap"],[
 ["Sell → collect (invoice/POS → payment)","salesman / cashier","YES","—"],
 ["Van load: request → approve → load","salesman → warehouse/supervisor","YES (stock_request)","—"],
 ["Purchase: PO → receive → AP post","procurement/whse → accountant","YES","—"],
 ["Sales return: create → complete","cashier/salesman → mgr","YES","—"],
 ["Collections → reconciliation","salesman → supervisor","YES","—"],
 ["GL: voucher → post","accountant","YES","—"],
 ["Day close → exception APPROVE","salesman → supervisor","NO","approveDayClose has no screen (§2)"],
 ["Out-of-route visit → APPROVE","salesman → supervisor","NO","decideVisitCompliance has no screen"],
 ["Customer/rep transfer → APPROVE","mgr → mgr","NO","transfer approve actions have no screen"],
 ["Van stock transfer → APPROVE","rep → supervisor","NO","van-transfer actions have no screen"],
 ["Trade-spend promo → APPROVE","trade mktg/mgr","NO","approve/cancel not on the page (flag-gated)"],
],widths=[2.3,1.6,1.0,1.6],
 fill={(6,2):REDBG,(7,2):REDBG,(8,2):REDBG,(9,2):REDBG,(10,2):AMBERBG})
para("The core sell→collect→inventory→finance loop is complete end-to-end. The INCOMPLETE workflows are all SUPERVISOR/"
     "MANAGER APPROVAL surfaces — built in the backend, missing only the screen.", bold=True)

# ===== RECO =====
part("Recommendations & priority")
h2("Before the FMCG pilot (UI-exposure only — no new backend)")
b("P-A (Major): add a supervisor/manager APPROVAL QUEUE that wires the existing actions — approveDayClose, "
  "decideVisitCompliance, approveCustomerTransfer/transferUser, approve/reject van transfer. This is the single biggest "
  "usability gap; all logic + permissions already exist.",bold=True)
b("P-B (Minor): add a nav entry for the Van-Sales hub (/field/van-sales) so the rep loop is reachable independently of the "
  "/today card; wire or remove the 'Confirm Load' coming-soon tile.")
b("P-C (Minor): decide M-2 label ('Staff' → 'Users & Roles'); add electrical module gate (MN-2); launcher for "
  "/forms/customer-data-update or confirm by-design.")
h2("Already done in this pass")
b("U-4 price/credit gating — FIXED. M-1 — CLOSED (cash_van ≡ salesman; code/DB aligned). M-2 — model proposed (keep "
  "boundary; optional label).",color=GREEN)
para("")
para("Net: the platform is visible, reachable, and (post-P0 + U-4) protected. The remaining work before pilot is to EXPOSE "
     "already-built approval workflows in the UI (P-A) — no new features, exactly matching the stated priority.",
     bold=True, color=NAVY)

out="docs/audits/VANTORA-Platform-Coverage-Audit-v2.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
