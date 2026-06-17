#!/usr/bin/env python3
"""VANTORA Phase A — Business Impact Review & Go/No-Go Decision Report -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"; GREYBG="ECECEC"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows,widths=None,size=8.4,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)
def verdict(label,col):
    p=d.add_paragraph(); r=p.add_run("  "+label+"  "); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF)
    _sh(p._p.getparent() if False else None,"")  # noop
    return p

# Cover
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Phase A — Business Impact Review"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Per-item Go / No-Go decision report (A1–A5)"); r.font.size=Pt(13); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Evidence-based from the codebase · No code changes · Decision report only · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# EXEC + SUMMARY
h1("How to read this report")
para("Each Phase A item (A1–A5) is reviewed independently against the live codebase — exact files, exact lines, and the "
     "six impact dimensions you asked for. Every item gets its own rollback plan, risk score, and a standalone Go / No-Go. "
     "Items are NOT batched. No code was changed to produce this; it is a decision report.", bold=True)
h2("Headline recommendation")
para("Four of the five items are safe to execute. ONE — A2 (permission aliasing) — must remain roadmap-only. The codebase "
     "shows the so-called 'duplicate' permission families are in fact intentionally distinct, separately-enforced permissions "
     "(e.g. report.aggregate.view is a deliberate scale-safe boundary; stock.* gates van operations while inventory.* gates "
     "the warehouse). Aliasing them would change effective access for real roles — that is not a Phase A hygiene change, it "
     "is a structural permission redesign and belongs in Phase C.", bold=True, color=RED)
h2("Decision summary")
tbl(["Item","Change","Type","Affects existing tenants?","Risk","Decision"],[
 ["A1","Remove dead permission keys (change_requests.*, trade_spend.manage)","Code only","No","Zero","GO"],
 ["A2","Alias 'duplicate' permissions (stock/inventory, customer/customers, report.aggregate/reports)","Code only","YES — changes effective access","Medium–High","NO-GO (defer to Phase C)"],
 ["A3","Bind nav + page to the same flag (Trade Spend 404 fix)","Code only","Visual only (hides a broken link)","Low","GO"],
 ["A4","Document 5 archetypes (documentation only)","Docs only","No","Zero","GO"],
 ["A5","Split feature-flag kinds in docs/registry","Docs only","No","Zero","GO"],
],widths=[0.4,2.6,0.8,1.5,0.8,1.1],
 fill={(0,5):GREENBG,(1,5):REDBG,(2,5):GREENBG,(3,5):GREENBG,(4,5):GREENBG,
       (0,4):GREENBG,(1,4):REDBG,(2,4):AMBERBG,(3,4):GREENBG,(4,4):GREENBG})
para("Net: execute A1, A3, A4, A5. Hold A2. The Phase A goal (clean up with zero tenant impact) is best served by NOT doing "
     "A2 as an alias — doing so would silently widen access, the opposite of hygiene.", bold=True)

# ============ DETAIL TEMPLATE ============
def item(code_,title,files,lines_label,lines,affects,typ,rollback,risk,riskcol,decision,deccol,rationale,notes=None):
    part(f"{code_} — {title}")
    # 1. Files
    h2("1. Exact files that will change")
    for f in files: code(f)
    # 2. Lines/modules
    h2("2. Exact lines / modules affected")
    para(lines_label, italic=True, size=9)
    for ln in lines: code(ln)
    # 3. Impact dimensions
    h2("3. Does it affect…")
    tbl(["Dimension","Affected?","Detail"],
        [[k,v[0],v[1]] for k,v in affects.items()],
        widths=[1.7,0.8,4.3],
        fill={(i,1):(REDBG if affects[k][0].startswith("YES") else GREENBG) for i,k in enumerate(affects)})
    # 4. Type
    h2("4. Change type")
    para(typ, bold=True)
    # 5. Rollback
    h2("5. Rollback plan")
    for rb in rollback: b(rb)
    # 6. Risk
    h2("6. Risk score")
    t=d.add_table(rows=1,cols=1); t.style="Table Grid"
    c=t.rows[0].cells[0]; c.text=""; r=c.paragraphs[0].add_run("  "+risk+"  "); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=NAVY; _sh(c,riskcol)
    # Decision
    h2("Final recommendation")
    t=d.add_table(rows=1,cols=1); t.style="Table Grid"
    c=t.rows[0].cells[0]; c.text=""; r=c.paragraphs[0].add_run("  "+decision+"  "); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,deccol)
    para(rationale)
    if notes: para(notes, italic=True, color=AMBER)

# ---- A1 ----
item("A1","Remove dead permission keys",
 ["src/lib/entitlements/registry.ts   (PERMISSION_MODULES map)",
  "src/lib/entitlements/entitlements.test.ts   (assertions)"],
 "The keys exist ONLY as entries in the entitlement permission→module map. They are NOT in the code Permission union "
 "(src/lib/erp/permissions.ts), are granted to NO role, and are checked by NO hasPermission() call. They are inert.",
 ["registry.ts:58  'change_requests.create': ['change_requests'],",
  "registry.ts:59  'change_requests.approve': ['change_requests'],",
  "registry.ts:60  'change_requests.manage': ['change_requests'],",
  "registry.ts:65  'trade_spend.manage': ['trade_spend'],",
  "entitlements.test.ts:66/68  modulesForPermission('change_requests.approve' / 'trade_spend.manage')"],
 {"Existing tenants":["No","No tenant role grants these keys; nothing resolves them"],
  "Existing users":["No","No user holds them"],
  "Existing roles":["No","Not present in any role definition in permissions.ts"],
  "Existing permissions":["No","Not in the Permission union; only inert map entries"],
  "Existing business types":["No","Unrelated"],
  "Existing deployments":["No","modulesForPermission() for these strings is never called in any live path"]},
 "Code only (registry map + its unit test). No DB change, no documentation dependency.",
 ["Re-add the four map lines and the two test assertions (a 6-line revert).",
  "Fully covered by git revert of the single commit.",
  "No data to restore — nothing was ever stored against these keys."],
 "Zero Risk", GREENBG, "GO", "1B7F3B",
 "Safe to execute. The keys enforce nothing; removing them only tidies the entitlement map and its test. Recommend pairing "
 "with a one-line note in the governance baseline that these engine permissions are typed-and-enforced when their engine "
 "ships (Phase C), not via this inert map.",
 "Caution: do NOT also delete any erp_change_request_definition rows or DB seeds in Phase A — this item is limited to the "
 "inert code map. DB cleanup (if any) is Phase C.")

# ---- A2 ----
item("A2","Alias 'duplicate' permissions  —  NOT a duplicate set",
 ["src/lib/erp/permissions.ts   (Permission union + role grants)",
  "src/app/(app)/fmcg/actions.ts, src/app/(app)/field/**, src/lib/erp/fmcg-action-gate.ts (enforcement sites)"],
 "EVIDENCE CONTRADICTS THE PREMISE. These are not synonyms — they are distinct, separately-enforced permissions with "
 "different semantics, assigned to different roles:",
 ["permissions.ts:13-16  inventory.view/adjust/transfer/count   = WAREHOUSE stock management",
  "permissions.ts:51-54  stock.view/adjust/transfer.approve     = VAN/field stock (FMCG S5) — different operations",
  "field/actions.ts:314  hasPermission(ctx,'stock.transfer')    (van) — NOT inventory.transfer",
  "field/van-sales/warehouse/page.tsx:21  hasPermission(ctx,'stock.adjust')  gates the VAN warehouse",
  "permissions.ts:10-12 customers.manage/approve/change_status  = management;  :45-48 customer.create/edit/import/transfer = field onboarding",
  "fmcg/actions.ts:418  reports.view OR report.aggregate.view (deliberate pair)",
  "fmcg/actions.ts:546  report.aggregate.view ONLY — a narrower 'scale-safe' boundary (intentional)"],
 {"Existing tenants":["YES","Every tenant using FMCG/field roles relies on the stock.* vs inventory.* split"],
  "Existing users":["YES","A user with inventory.view would suddenly pass stock.adjust (van) checks, and vice-versa"],
  "Existing roles":["YES","warehouse_keeper, salesman, supervisor hold these as DISTINCT grants"],
  "Existing permissions":["YES","Collapses 3 intentional families; erases the report.aggregate scale-safe boundary"],
  "Existing business types":["Indirect","FMCG/distribution depend on the field-vs-warehouse separation"],
  "Existing deployments":["YES","Authorization behaviour changes on next deploy"]},
 "Code only — but it is an authorization (access-control) change, the highest-sensitivity kind of code change.",
 ["No clean rollback once roles act under merged semantics: removing the alias would re-split access and could strand "
  "users who relied on the widened grant.",
  "Would require a full role-permission regression suite + per-tenant access snapshot BEFORE and AFTER — i.e. Phase C rigour.",
  "Because it changes effective access, 'reversible' cannot be guaranteed without a migration plan."],
 "Medium–High Risk", REDBG, "NO-GO  (keep roadmap-only → Phase C)", "B31B1B",
 "The handbook labelled these 'duplicate families', but the code proves they are intentionally distinct, actively-enforced "
 "permissions. Aliasing them is NOT zero-impact hygiene — it silently WIDENS access (e.g. warehouse↔van, management↔field "
 "onboarding) and erases the deliberate report.aggregate.view scale-safe boundary. That is a structural permission "
 "redesign requiring real usage data and a migration plan. It belongs in Phase C, after the first paying customer.",
 "This is exactly why the business-impact review was worth doing before touching code: A2 looked like the flagship 'safe "
 "dedup' but is the one item that would change tenant behaviour. Recommend updating the Roadmap/Handbook to reclassify "
 "permission consolidation from Phase A to Phase C, with explicit per-family analysis.")

# ---- A3 ----
item("A3","Bind nav + page to the same capability flag (Trade Spend 404 fix)",
 ["src/lib/erp/navigation.ts   (nav item gate)",
  "src/app/(app)/distribution/trade-spend/page.tsx   (page gate) — reference only",
  "src/lib/trade-spend/flags.ts   (TRADE_SPEND_ENABLED) — reference only"],
 "The mismatch is exact and visible in code: the nav item is gated by module+perm but NOT by the env flag the page "
 "requires, so the link shows and the page 404s.",
 ["navigation.ts:390  { href:'/distribution/trade-spend', perm:'reports.view', module:['distribution','trade_spend'] }",
  "trade-spend/page.tsx:22  if (!TRADE_SPEND_ENABLED()) notFound();",
  "flags.ts:6  TRADE_SPEND_ENABLED = () => on(process.env.KAKO_TRADE_SPEND)",
  "Fix: make the nav item's visibility also require the same TRADE_SPEND_ENABLED()/flag signal (and apply the same pattern "
  "to other engine routes whose page has an independent flag gate)."],
 {"Existing tenants":["Visual only","A link that currently 404s becomes hidden when its flag is off"],
  "Existing users":["Visual only","Fewer dead links; no change to what they can actually do"],
  "Existing roles":["No","Permission gates unchanged"],
  "Existing permissions":["No","No permission added/removed"],
  "Existing business types":["No","Module gating unchanged"],
  "Existing deployments":["YES (intended)","Where KAKO_TRADE_SPEND is off, the link disappears instead of 404-ing"]},
 "Code only (navigation gating). No DB, no docs dependency. Strictly-better behaviour (never exposes anything new).",
 ["Revert the navigation.ts gate change (single commit) to restore prior link visibility.",
  "No data involved; instant rollback.",
  "Behaviour is monotonic — the fix can only hide a would-be-404 link, so rollback risk is nil."],
 "Low Risk", GREENBG, "GO", "1B7F3B",
 "This is a genuine bug fix and falls squarely inside the approved freeze categories (Bug / Deployment). It removes the "
 "show-then-404 confusion the pilot already hit on Trade Spend. The only visible effect is that a broken link is hidden "
 "when its engine flag is off. Recommend shipping it as the single highest-value Phase A item.",
 "Scope discipline: in Phase A, bind only the nav gate to the existing flag. Do NOT introduce the new capability-flag "
 "abstraction here — that design lands with A5's documentation and any refactor is deferred.")

# ---- A4 ----
item("A4","Document the 5 archetypes (documentation only)",
 ["docs/audits/ (governance baseline) — documentation",
  "src/lib/erp/setup-wizard.ts, src/app/setup/* — REFERENCE ONLY (not edited in Phase A)"],
 "Phase A is limited to WRITING the 5-archetype model and the business_type→archetype map into the governance baseline. "
 "The onboarding-picker curation (a UX/code change) is explicitly DEFERRED so this item stays documentation-only.",
 ["No source line changes. Documentation table maps each of the 21 existing business types to one of 5 archetypes.",
  "setup-wizard.ts / setup/page.tsx are referenced as the FUTURE (Phase B) place to curate the picker — not touched now."],
 {"Existing tenants":["No","No business_type value changes; pure documentation"],
  "Existing users":["No","Unaffected"],
  "Existing roles":["No","Unaffected"],
  "Existing permissions":["No","Unaffected"],
  "Existing business types":["No","All 21 remain exactly as defined; only documented as archetypes"],
  "Existing deployments":["No","No code/runtime change"]},
 "Documentation only (governance baseline). No code, no DB.",
 ["Revert the documentation edit. Nothing runtime to roll back.",
  "Zero data impact."],
 "Zero Risk", GREENBG, "GO  (documentation only)", "1B7F3B",
 "Safe. Keep A4 strictly to documentation in Phase A. The picker-curation sub-task (touching setup-wizard.ts) is a Phase B "
 "new-tenant-only change and must NOT be done now, to preserve the zero-impact guarantee.",
 "If the picker curation is wanted sooner, treat it as a separate, explicitly-scoped Low-Risk Phase B item — not folded "
 "into this documentation task.")

# ---- A5 ----
item("A5","Split feature-flag kinds in docs/registry (documentation only)",
 ["docs/audits/ (governance baseline) — documentation",
  "src/lib/erp/feature-catalog.ts, src/lib/**/flags.ts — REFERENCE ONLY (classified, not edited)"],
 "Phase A classifies the two flag kinds in documentation: tenant CAPABILITY flags (feature-catalog; Lite/Standard/"
 "Enterprise) vs platform KILL-SWITCHES (~38 env KAKO_* flags). No flag value or evaluation path changes.",
 ["~38 distinct KAKO_* kill-switches exist (e.g. KAKO_TRADE_SPEND, KAKO_ALERTS, KAKO_CHANGE_REQUESTS, KAKO_EINVOICE …).",
  "feature-catalog.ts provides the tenant capability system (the Pharmacy Lite/Standard/Enterprise templates).",
  "Deliverable = a documentation table labelling each flag Capability or Kill-switch + the governance rule."],
 {"Existing tenants":["No","Classification/documentation only"],
  "Existing users":["No","Unaffected"],
  "Existing roles":["No","Unaffected"],
  "Existing permissions":["No","Unaffected"],
  "Existing business types":["No","Unaffected"],
  "Existing deployments":["No","No flag value or evaluation logic changes"]},
 "Documentation only (optionally a non-functional code COMMENT in registry). No behaviour change.",
 ["Revert the documentation edit. Nothing runtime to roll back.",
  "Zero data impact."],
 "Zero Risk", GREENBG, "GO  (documentation only)", "1B7F3B",
 "Safe. This is the policy half of the A3 fix: it writes down the rule (tenant gating uses capability flags; unfinished "
 "engines use kill-switches) so the show-then-404 class of bug is prevented going forward. Keep it documentation-only in "
 "Phase A; any registry refactor is Phase D.")

# ============ FINAL ============
part("Final Go / No-Go — consolidated")
tbl(["Item","Decision","Risk","Type","One-line reason"],[
 ["A1","GO","Zero","Code","Inert keys; removal changes nothing"],
 ["A2","NO-GO","Med–High","Code","'Duplicates' are distinct enforced permissions — aliasing widens access"],
 ["A3","GO","Low","Code","Real bug fix; only hides a would-be-404 link"],
 ["A4","GO","Zero","Docs","Archetype documentation only; no tenant touched"],
 ["A5","GO","Zero","Docs","Flag-kind classification only; no behaviour change"],
],widths=[0.5,1.0,1.0,0.8,3.4],
 fill={(0,1):GREENBG,(1,1):REDBG,(2,1):GREENBG,(3,1):GREENBG,(4,1):GREENBG})
h2("Recommended action")
b("EXECUTE: A1, A3, A4, A5 — all Zero/Low risk, none affect existing tenant access. Suggested order: A4, A5 (docs), A3 "
  "(bug fix), A1 (inert cleanup).",bold=True)
b("HOLD: A2 — reclassify permission consolidation to Phase C (after first paying customer + usage data), with per-family "
  "analysis and a migration plan. Update the Roadmap/Handbook to reflect that A2 is NOT zero-impact.",bold=True,color=RED)
b("UNCHANGED: no role merged, no business type merged, no permission reduced, no tenant migrated. The four GO items are "
  "additive/documentation/visual-only.")
para("")
para("Bottom line: 4 of 5 Phase A items are cleared for execution with zero tenant-access impact; A2 is correctly caught as "
     "a structural change in disguise and held for Phase C. This is the decision — implementation still awaits your explicit "
     "go-ahead, item by item.", bold=True, color=NAVY)

out="docs/audits/VANTORA-PhaseA-Business-Impact-Review.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
