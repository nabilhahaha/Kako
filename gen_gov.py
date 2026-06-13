#!/usr/bin/env python3
"""VANTORA Governance Architecture Handbook (evidence-based) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(19); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.4,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# Cover
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Governance Architecture Handbook"); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Business Types · Modules · Roles · Permissions · Feature Flags — the complete platform structure for the founder & platform owner"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Evidence-based from vantora-staging (rsjvgehvastmawzwnqcs) + codebase · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# EXEC SUMMARY
h1("Executive Summary")
para("This handbook maps VANTORA's governance model exactly as it exists today. The figures are pulled from the live "
     "database and the codebase, not from memory.")
tbl(["Dimension","Count","Observation"],[
 ["Business Types","21","Many share identical module sets — collapsible into ~5 archetypes"],
 ["Modules","31","8 core ERP + 5 capability + 6 FMCG engines + 12 vertical/POS"],
 ["Roles","27","Deep FMCG management hierarchy + vertical-specialist roles → over-granular for SMEs"],
 ["Permissions (DB)","91","Code recognises 75 — a ~16-permission code↔DB gap (dead/orphan keys)"],
 ["Permission groups","16","Re-mappable to 8 governance categories"],
],widths=[1.6,0.7,4.7])
para("Headline finding: VANTORA is not over-built in capability — it is over-fragmented in taxonomy. The complexity you "
     "feel comes from (1) too many business types with near-identical module sets, (2) a deep role hierarchy copied from "
     "large-enterprise org charts, and (3) overlapping permission families (customer/customers, inventory/stock, "
     "report/reports) plus dead engine-permission keys. None require rebuilding — they require consolidation. "
     "Parts 7–8 give the simplification and the ideal target model.", bold=True)

# ============ PART 1 ============
part("Part 1 — Business Type Architecture")
para("21 business types are defined (erp_business_type_modules). Each maps to a curated module set. Status: the platform "
     "code/screens exist for all; maturity varies. Legend: PR=Pilot Ready, I=Implemented, P=Partially Implemented.")
tbl(["Business Type","Purpose / target customer","Modules (n)","Status"],[
 ["fmcg","FMCG distributors / wholesalers with field sales","21 (all engines)","Pilot Ready (engines flag-gated)"],
 ["pharmacy","Pharmacies & pharmacy chains","12","Pilot Ready (data-proven)"],
 ["supermarket / market","Grocery / mini-market retail (POS)","12","Implemented"],
 ["wholesale","Wholesale traders","14","Implemented"],
 ["electronics","Electronics/electrical retail+wholesale (RMA/serials)","13","Partially Implemented"],
 ["delivery","Distribution/delivery operations","13","Implemented"],
 ["general","Generic ERP (any business)","14","Implemented"],
 ["auto_parts / bakery / bookstore / herbalist","Specialty retail with stock+POS","11 each (identical set)","Implemented"],
 ["butchery / workshop","Retail/service with stock, no warehousing","9–10","Implemented"],
 ["clinic","Medical clinics (patients, visits, services)","6","Partially Implemented"],
 ["restaurant / cafe","Restaurants & cafés (tables, kitchen, orders)","7","Partially Implemented"],
 ["salon","Salons (appointments, tickets, services)","6","Partially Implemented"],
 ["hotel","Hotels (rooms, bookings)","6","Partially Implemented"],
 ["laundry","Laundries (orders, services)","6","Partially Implemented"],
 ["clothing","Fashion stores (self-contained pack)","1 (fashion)","Partially Implemented"],
 ["services","Service businesses (sales, no inventory)","6","Implemented"],
],widths=[1.9,3.0,1.2,1.4])
h2("Roles & workflows per business type (representative)")
para("Each business type's roles are seeded from the role catalog; workflows = the generic approval engine plus pack "
     "specifics (e.g. pharmacy dispense, FMCG change-requests). Vertical packs add vertical roles (clinic→doctor/"
     "receptionist; salon→stylist; hotel→housekeeping/receptionist; workshop→technician).")
h2("How business types differ — 5 real archetypes")
tbl(["Archetype","Business types","Defining modules"],[
 ["Distribution","fmcg, delivery, wholesale, general","+ distribution, field_ops, sales_orders (+ FMCG engines for fmcg)"],
 ["Stock + POS retail","supermarket, pharmacy, auto_parts, bakery, bookstore, herbalist, butchery, workshop, electronics","inventory, pos, sales, purchasing, returns (+ vertical: pharmacy/market)"],
 ["Service vertical","clinic, salon, hotel, laundry, restaurant, cafe","vertical module + crm/accounting/analytics (little/no inventory/sales)"],
 ["Self-contained vertical","clothing (fashion)","single 'fashion' module — its own sub-system"],
 ["Generic / light","services, general","core ERP only; no specialization"],
],widths=[1.5,3.0,2.6])
para("Insight: 9 of the 21 types ('auto_parts/bakery/bookstore/herbalist' share an IDENTICAL 11-module set; "
     "'restaurant/cafe' identical; several services identical). The business-type list is finer-grained than the actual "
     "module differentiation — see Part 7.", italic=True)

# ============ PART 2 ============
part("Part 2 — Module Architecture")
para("31 modules across 4 categories. 'Plan-gated' modules are licensed via subscription plans (ALL_MODULES); 'engine' "
     "modules are company/business-type driven (never plan-gated) and additionally require a KAKO_* env flag.")
h2("Core ERP modules")
tbl(["Module","Purpose","Used by (BTs)","Roles","Readiness"],[
 ["sales","Selling, invoices, orders","most","salesman, cashier, admin, manager","Pilot Ready"],
 ["pos","Point of sale","retail/distribution BTs","cashier, salesman","Pilot Ready"],
 ["inventory","Stock, movements, batches","stock BTs","warehouse_keeper, inventory_controller","Pilot Ready"],
 ["warehousing","Transfers, counts, requests","stock+distribution BTs","warehouse_keeper","Pilot Ready"],
 ["purchasing","POs, receipts, suppliers","most retail/distribution","procurement, branch_manager","Pilot Ready"],
 ["returns","Sales returns","retail/distribution","cashier, salesman, accountant","Pilot Ready"],
 ["sales_orders","Order capture","distribution BTs","salesman, supervisor","Implemented"],
 ["accounting","GL, journals, vouchers","all","accountant","Implemented"],
],widths=[1.2,1.8,1.6,1.7,0.9])
h2("Capability modules (plan-licensable)")
tbl(["Module","Purpose","Roles","Readiness"],[
 ["crm","Customers & relationships","all sales roles","Pilot Ready"],
 ["workflow","Approvals & governed processes","admin, approvers","Implemented"],
 ["analytics","Dashboards & reporting","manager, viewer, accountant","Implemented"],
 ["field_ops","Field execution (visits, journeys, surveys, targets)","salesman, merchandiser, supervisor","Implemented"],
 ["integrations","API / webhooks / sync","it_admin, admin","Implemented"],
],widths=[1.3,2.6,2.3,1.0])
h2("FMCG engine modules (business-type driven + KAKO_* flag)")
tbl(["Module","Purpose","Dependencies","Roles","Readiness"],[
 ["route_management","Routes, territories, journeys","field_ops, distribution","supervisor, area_manager","Implemented (flag)"],
 ["van_sales","Van load/sell/return/reconcile","field_ops, inventory","cash_van, salesman, supervisor","Implemented (flag)"],
 ["trade_spend","Promotions, accruals, claims, ROI","accounting","sales_director, accountant, admin","Implemented (Foundation)"],
 ["merchandising","Perfect store, MSL, grading, surveys","field_ops","merchandiser, supervisor","Implemented (flag)"],
 ["change_requests","Governed master-data changes","workflow","salesman, supervisor, credit_controller","Implemented (flag)"],
 ["critical_alerts","OOS/credit/collection/route/inventory alerts","analytics","supervisor, manager","Implemented (flag)"],
],widths=[1.4,1.9,1.3,1.6,1.0])
h2("Vertical modules")
tbl(["Module","Business type","Readiness"],[
 ["pharmacy","pharmacy","Pilot Ready"],["market","supermarket","Implemented"],["wholesale","wholesale/electronics","Implemented"],
 ["distribution","fmcg/delivery/wholesale/general","Implemented"],["fashion","clothing","Partially Implemented"],
 ["clinic","clinic","Partially"],["restaurant","restaurant/cafe","Partially"],["salon","salon","Partially"],
 ["hotel","hotel","Partially"],["laundry","laundry","Partially"],
],widths=[1.4,3.0,1.3])

# ============ PART 3 ============
part("Part 3 — Role Hierarchy")
para("27 roles exist (erp_role_permissions). Organised into 8 governance layers. 'Mandatory' = needed for the tenant to "
     "operate; 'Optional' = scale/vertical-specific.")
tbl(["Layer","Roles","Why it exists","Mandatory?"],[
 ["Platform","(is_platform_owner / platform staff)","Vendor operates the SaaS across tenants","Mandatory (vendor)"],
 ["Company","admin (super-admin), it_admin","Top tenant authority + IT/data admin","admin Mandatory; it_admin Optional"],
 ["Operations (mgmt)","manager, national_sales_manager, regional_manager, area_manager, sales_director, branch_manager","Sales/ops management hierarchy","manager Mandatory; rest Optional (scale)"],
 ["Field Sales","salesman, cash_van, driver","Execute the route/van journey","salesman Mandatory for distribution"],
 ["Merchandising","merchandiser, supervisor","In-store standards + field supervision","Optional (FMCG)"],
 ["Finance","accountant, credit_controller, collection_officer","Collections, AR/AP, credit control","accountant Mandatory; others Optional"],
 ["Warehouse","warehouse_keeper, inventory_controller","Stock operations","warehouse_keeper Mandatory for stock BTs"],
 ["Procurement","procurement, (branch_manager)","Supplier & purchasing","Optional"],
 ["Vertical specialists","doctor, receptionist, stylist, technician, housekeeping, cashier, staff, viewer","Pack-specific (clinic/salon/workshop/hotel/retail)","Only for that vertical"],
],widths=[1.4,2.4,2.2,1.2])
h2("Reporting lines (FMCG example)")
para("Platform Owner → (tenant) Admin → Sales Director / GM(manager) → National Sales Manager → Regional Manager → Area "
     "Manager → Supervisor → Salesman & Merchandiser. Finance, Warehouse and Procurement report to Admin/GM in parallel.")
para("Honest note: six management tiers (sales_director/national/regional/area/branch/manager) is an enterprise org chart. "
     "Most SME tenants need only 2–3 of them — see Part 7.", italic=True, color=AMBER)

# ============ PART 4 ============
part("Part 4 — Permission Architecture")
para("The code groups the 75 recognised permissions into 16 groups; the database grants 91. These re-map to 8 governance "
     "categories:")
tbl(["Governance category","Source groups / examples","~count"],[
 ["Administrative","settings (users, branches, custom_fields), user.import/transfer, integrations.manage","~9"],
 ["Operational (field)","field_ops (visit, day, survey, assortment, grade, target), route/journey","~13"],
 ["Financial","accounting (post/view), credit.request.*, return.reason","~6"],
 ["Inventory","inventory.*, stock.*, product.*, reconciliation.*, stock_request.*, uom","~17"],
 ["Sales","sales.*, customer.*/customers.*, pricing.*, field.sales","~16"],
 ["Workflow","workflow.manage, change_requests.*","~4"],
 ["Reporting","reports.view, report.aggregate.view","~2"],
 ["Security / Access","settings.users (authz), action policies, audit (governed via settings)","~3"],
 ["Vertical","pharmacy/clinic/salon/hotel/laundry/restaurant/market/wholesale/electrical/fashion","~22"],
],widths=[1.7,3.9,0.9])
h2("Findings")
tbl(["Type","Finding","Detail"],[
 ["Dead / orphan","change_requests.create/approve/manage, trade_spend.manage","Granted in DB but NOT in the code Permission type; only referenced by a registry test → never enforced"],
 ["Duplicate","inventory.* vs stock.*","Two families for the same thing (adjust/transfer/view) — pick one"],
 ["Duplicate","customer.* vs customers.*","customer.create/edit/import vs customers.manage/approve/change_status overlap"],
 ["Duplicate","report.aggregate.view vs reports.view","Near-identical reporting gates"],
 ["Confusing","field.sales vs field_ops perms","'field.sales' lives in the Sales group but gates field execution"],
 ["Confusing","pricing.manage vs pricing.view","Fine, but often conflated by seeders"],
 ["Missing","Explicit security perms (MFA/SSO/session)","No first-class security/identity permissions — Roadmap"],
 ["Missing","First-class engine perms in code","trade_spend/route/merch gated by module+flag, not by a typed permission"],
],widths=[1.2,2.2,3.4],
 fill={(0,0):REDBG,(1,0):AMBERBG,(2,0):AMBERBG,(3,0):AMBERBG,(4,0):AMBERBG,(5,0):AMBERBG,(6,0):BLUEBG,(7,0):BLUEBG})

# ============ PART 5 ============
part("Part 5 — Relationship Matrices")
h2("5.1 Business Type × Module category")
tbl(["Business Type","Core ERP","POS","Inventory","Distribution","Engines","Vertical"],[
 ["fmcg","Yes","Yes","Yes","Yes","Yes (all 6)","wholesale"],
 ["delivery / general","Yes","Yes*","Yes","Yes","—","—"],
 ["wholesale","Yes","—","Yes","Yes","—","wholesale"],
 ["electronics","Yes","Yes","Yes","—","—","wholesale (RMA)"],
 ["pharmacy","Yes","Yes","Yes","—","—","pharmacy"],
 ["supermarket","Yes","Yes","Yes","—","—","market"],
 ["auto_parts/bakery/bookstore/herbalist/butchery/workshop","Yes","Yes","Yes","—","—","—"],
 ["clinic/salon/hotel/laundry","—","—","—","—","—","vertical only"],
 ["restaurant/cafe","Yes","—","Yes","—","—","restaurant"],
 ["clothing","self-contained","—","—","—","—","fashion"],
 ["services","Yes(sales)","—","—","—","—","—"],
],widths=[2.4,0.9,0.6,0.9,1.1,1.0])
h2("5.2 Module category × Role layer")
tbl(["Module","Field Sales","Merch","Warehouse","Finance","Procurement","Mgmt/Admin"],[
 ["Sales/POS","Yes","—","—","collect","—","oversee"],
 ["Inventory/Warehousing","—","—","Yes","—","receive","oversee"],
 ["Purchasing","—","—","receive","—","Yes","oversee"],
 ["Accounting","—","—","—","Yes","—","oversee"],
 ["field_ops/route/van","Yes","Yes","—","—","—","supervise"],
 ["merchandising","—","Yes","—","—","—","supervise"],
 ["trade_spend","—","—","—","Yes","—","approve"],
 ["change_requests","raise","raise","—","credit","—","approve"],
 ["workflow/analytics","—","—","—","view","—","manage/view"],
],widths=[1.5,1.0,0.7,1.0,0.9,1.1])
h2("5.3 Role × Permission category (key FMCG roles)")
tbl(["Role","Admin","Oper.","Fin.","Inv.","Sales","Workflow","Report"],[
 ["admin","Yes","Yes","Yes","Yes","Yes","Yes","Yes"],
 ["manager","—","Yes","view","view","oversee","approve","Yes"],
 ["supervisor","—","Yes","—","—","—","approve","Yes"],
 ["salesman","—","Yes","collect","—","Yes","raise","—"],
 ["merchandiser","—","Yes","—","—","—","raise","—"],
 ["accountant","—","—","Yes","—","collect","—","Yes"],
 ["warehouse_keeper","—","—","—","Yes","—","—","—"],
 ["procurement","—","—","—","adjust","—","—","—"],
],widths=[1.6,0.8,0.8,0.7,0.8,0.8,0.9])
h2("5.4 Business Type × Role (who you need)")
tbl(["Business Type","Mandatory roles","Optional roles"],[
 ["fmcg","admin, manager, salesman, accountant, warehouse_keeper","supervisor, merchandiser, area/regional/national mgr, sales_director, cash_van, credit_controller, collection_officer, procurement"],
 ["pharmacy","admin, cashier/pharmacist, accountant","manager, warehouse_keeper, procurement"],
 ["supermarket/retail","admin, cashier, warehouse_keeper","manager, accountant, procurement"],
 ["service vertical","admin, vertical role (doctor/stylist/…)","accountant, receptionist"],
],widths=[1.5,2.4,3.0])

# ============ PART 6 ============
part("Part 6 — FMCG Organization Mapping")
tbl(["Real FMCG role","VANTORA role","Primary modules"],[
 ["General Manager","manager (or admin)","analytics, dashboards, approvals (all, read)"],
 ["National Sales Manager","national_sales_manager","analytics, distribution, targets (national)"],
 ["Regional Manager","regional_manager","distribution, route_management, targets (region)"],
 ["Area Manager","area_manager","route_management, coverage, targets (area)"],
 ["Supervisor","supervisor","field_ops, route_management, merchandising (approvals)"],
 ["Salesman","salesman","sales, pos, field_ops, van_sales, change_requests"],
 ["Merchandiser","merchandiser","merchandising, field_ops, surveys (no sell)"],
 ["Trade Marketing","sales_director","trade_spend (promotions/accruals/ROI)"],
 ["Finance","accountant / credit_controller / collection_officer","accounting, collections, cashbox, trade_spend (settle)"],
 ["Warehouse","warehouse_keeper / inventory_controller","inventory, warehousing, transfers, counts"],
 ["Procurement","procurement / branch_manager","purchasing, suppliers"],
],widths=[1.9,1.9,3.1])

# ============ PART 7 ============
part("Part 7 — Simplification Review (honest)")
tbl(["Question","Answer","Why"],[
 ["Too many roles?","YES","27 roles; 6 sales-mgmt tiers + 8 vertical-specialist roles. SMEs use ~6–8."],
 ["Too many modules?","BORDERLINE","31 is reasonable, but engine modules duplicate gating (module + env flag)."],
 ["Too many permissions?","YES","91 in DB vs 75 in code; duplicate families (inventory/stock, customer/customers)."],
 ["Too many business types?","YES","21 types, but only ~5 real archetypes; 9 share an identical module set."],
],widths=[1.7,1.1,4.0],
 fill={(0,1):REDBG,(1,1):AMBERBG,(2,1):REDBG,(3,1):REDBG})
h2("Merge")
b("Business types: collapse 'auto_parts/bakery/bookstore/herbalist/butchery/workshop' into one 'Retail (stock+POS)' archetype with a label, not a distinct module set.")
b("Roles: merge the sales-management tiers into 2 configurable levels (Manager, Area/Regional) with scope, not 6 fixed roles.")
b("Permissions: merge inventory.* and stock.* into one family; customer.* into customers.*; report.aggregate.view into reports.view.")
h2("Remove")
b("Dead permissions: change_requests.*, trade_spend.manage (not enforced by code) — remove or wire to real gates.")
b("Vertical roles/permissions for verticals you will not ship soon (Roadmap) — keep out of the seeded default set.")
h2("Inherit")
b("Make roles inherit from a base (e.g. all field roles inherit field.sales) rather than re-listing permissions per role.")
b("Make business type inherit an archetype's module set, then add only the vertical module.")
h2("Keep separate")
b("Platform vs tenant (Provider panel vs company) — correct and must stay.")
b("Engine modules as company/business-type driven (not plan-gated) — correct design.")
b("Core ERP modules as plan-licensable entitlements — correct.")

# ============ PART 8 ============
part("Part 8 — Recommended Future Architecture")
h2("Business Type Structure — archetype + label")
b("5 archetypes (Distribution, Stock-Retail, Service-Vertical, Self-contained-Vertical, Generic) each owning a module set; "
  "a 'business type' becomes a label + a vertical module on top of an archetype.")
h2("Module Structure — 3 clean tiers")
b("Tier 1 Core (always): sales, inventory, purchasing, accounting, crm.  Tier 2 Capability (licensed): workflow, analytics, "
  "field_ops, integrations.  Tier 3 Pack (vertical/engine): pharmacy, distribution engines, etc.")
h2("Role Structure — template roles + scope")
b("~8 template roles (Admin, Manager, Sales, Field-Sales, Merchandiser, Finance, Warehouse, Procurement) + a SCOPE "
  "dimension (branch/area/region/national) instead of 6 fixed management roles. Vertical roles load only with the pack.")
h2("Permission Structure — verb.resource, deduplicated")
b("One consistent verb.resource scheme (e.g. inventory.adjust, sales.sell), no parallel families; ~60 clean permissions "
  "instead of 91 with overlaps; permissions inherit via role templates.")
h2("Approval Structure")
b("One workflow engine; approval chains declared per action (discount, credit change, transfer, write-off, promotion) "
  "with scope-aware approvers — replacing ad-hoc per-feature gates.")
h2("Access Structure")
b("Three layers only: Platform Owner (vendor) → Tenant Admin → Scoped role. Temporary-access grants for elevation. "
  "Add MFA/SSO (currently a Gap).")
h2("Feature Flag Structure")
b("Two kinds, clearly separated: (a) tenant capability flags (feature-catalog, Lite/Standard/Enterprise templates) — the "
  "source of truth for what a tenant sees; (b) platform kill-switches (env KAKO_*) for unfinished engines. "
  "Tie nav + page gating to the SAME flag so a screen never shows-then-404s (the Trade Spend issue).")

# ============ PART 9 ============
part("Part 9 — Complexity Hotspot Classification (founder view)")
para("Every business type, module, permission group and role is classified on one scale so you can see at a glance what to "
     "keep, what to simplify, and what to retire. This is the decision layer on top of Parts 1–8.", bold=True)
h2("Classification scale")
tbl(["Class","Meaning","Action"],[
 ["Essential","Platform cannot operate without it","Keep — invest"],
 ["Useful","Adds real value for many tenants","Keep"],
 ["Optional","Valuable only at scale or for one vertical","Keep, but don't seed by default"],
 ["Redundant","Overlaps something else that already exists","Deduplicate"],
 ["Merge","Should fold into a sibling / parent","Consolidate"],
 ["Remove","Dead, unused, or not on the near roadmap","Retire / defer"],
],widths=[1.1,3.6,1.9],
 fill={(0,0):GREENBG,(1,0):BLUEBG,(2,0):AMBERBG,(3,0):REDBG,(4,0):AMBERBG,(5,0):REDBG})

h2("9.1 Business Types (21) classified")
_bt=[
 ["fmcg","Essential","Flagship distribution archetype",GREENBG],
 ["pharmacy","Essential","Active pilot, data-proven",GREENBG],
 ["supermarket / market","Essential","Core stock+POS retail",GREENBG],
 ["wholesale","Useful","Distinct (wholesale module)",BLUEBG],
 ["general","Useful","Generic ERP fallback",BLUEBG],
 ["delivery","Useful","Distribution variant of fmcg",BLUEBG],
 ["services","Useful","Sales-only light tenant",BLUEBG],
 ["electronics","Optional","RMA/serial variant of retail",AMBERBG],
 ["clinic","Optional","Service vertical (partial UI)",AMBERBG],
 ["restaurant / cafe","Merge","Identical module set → one 'Food Service'",AMBERBG],
 ["salon","Optional","Service vertical (partial)",AMBERBG],
 ["hotel","Optional","Service vertical (partial)",AMBERBG],
 ["laundry","Optional","Service vertical (partial)",AMBERBG],
 ["clothing (fashion)","Optional","Self-contained pack",AMBERBG],
 ["auto_parts","Merge","Identical 11-module set → 'Retail'",REDBG],
 ["bakery","Merge","Identical 11-module set → 'Retail'",REDBG],
 ["bookstore","Merge","Identical 11-module set → 'Retail'",REDBG],
 ["herbalist","Merge","Identical 11-module set → 'Retail'",REDBG],
 ["butchery","Merge","Retail-with-stock, near-identical",REDBG],
 ["workshop","Merge","Service+stock → 'Repair/Service'",REDBG],
 ["(label-only duplicates)","Remove","Keep as labels, not distinct types",REDBG],
]
tbl(["Business Type","Class","Rationale"],[[r[0],r[1],r[2]] for r in _bt],widths=[2.0,1.0,3.6],
    fill={(i,1):r[3] for i,r in enumerate(_bt)})
para("Summary: 3 Essential · 4 Useful · 6 Optional · 8 Merge/Remove → collapses to ~5 archetypes.", italic=True)

h2("9.2 Modules (31) classified")
_md=[
 ["sales","Essential",GREENBG],["inventory","Essential",GREENBG],["purchasing","Essential",GREENBG],
 ["accounting","Essential",GREENBG],["pos","Essential",GREENBG],["crm","Essential",GREENBG],
 ["warehousing","Useful",BLUEBG],["returns","Useful",BLUEBG],["sales_orders","Merge",AMBERBG],
 ["workflow","Useful",BLUEBG],["analytics","Useful",BLUEBG],["field_ops","Useful",BLUEBG],
 ["integrations","Optional",AMBERBG],
 ["distribution","Essential",GREENBG],["route_management","Useful",BLUEBG],["van_sales","Useful",BLUEBG],
 ["merchandising","Useful",BLUEBG],["trade_spend","Optional",AMBERBG],["change_requests","Merge",AMBERBG],
 ["critical_alerts","Merge",AMBERBG],
 ["pharmacy","Essential",GREENBG],["market","Useful",BLUEBG],["wholesale","Useful",BLUEBG],
 ["fashion","Optional",AMBERBG],["clinic","Optional",AMBERBG],["restaurant","Optional",AMBERBG],
 ["salon","Optional",AMBERBG],["hotel","Optional",AMBERBG],["laundry","Optional",AMBERBG],
 ["electrical","Optional",AMBERBG],["services","Merge",AMBERBG],
]
tbl(["Module","Class","Module","Class"],
    [[_md[i][0],_md[i][1],(_md[i+1][0] if i+1<len(_md) else ""),(_md[i+1][1] if i+1<len(_md) else "")] for i in range(0,len(_md),2)],
    widths=[1.8,1.0,1.8,1.0])
para("Summary: sales_orders→merge into sales; change_requests/critical_alerts→fold under workflow/analytics; "
     "services module→drop (sales covers it). Modules consolidate from 31 to ~24 distinct, organised as 3 tiers.", italic=True)

h2("9.3 Permission Groups (16) classified")
tbl(["Permission group","Class","Action"],[
 ["settings / users / branches","Essential","Keep"],
 ["sales","Essential","Keep"],
 ["inventory","Essential","Keep (absorb 'stock')"],
 ["stock","Redundant","Merge into inventory"],
 ["customer","Redundant","Merge into customers"],
 ["customers","Essential","Keep (canonical)"],
 ["accounting / credit","Essential","Keep"],
 ["field_ops","Useful","Keep"],
 ["reports","Essential","Keep (absorb report.aggregate)"],
 ["report.aggregate","Redundant","Merge into reports"],
 ["pricing","Useful","Keep"],
 ["workflow","Useful","Keep"],
 ["change_requests","Remove","Dead in code — wire or drop"],
 ["integrations","Optional","Keep (capability)"],
 ["pharmacy / vertical","Optional","Load with pack only"],
 ["misc (uom, reconciliation, user.transfer)","Merge","Fold into inventory/admin"],
],widths=[2.6,1.1,2.9],
 fill={(3,1):REDBG,(4,1):REDBG,(9,1):REDBG,(12,1):REDBG,(15,1):AMBERBG})
para("Summary: 3 redundant + 1 dead + 1 merge → 16 groups consolidate to ~8 governance categories (Part 4).", italic=True)

h2("9.4 Roles (27) classified")
_rl=[
 ["admin","Essential",GREENBG],["manager","Essential",GREENBG],["salesman","Essential",GREENBG],
 ["accountant","Essential",GREENBG],["warehouse_keeper","Essential",GREENBG],["cashier","Essential",GREENBG],
 ["supervisor","Useful",BLUEBG],["merchandiser","Useful",BLUEBG],["procurement","Useful",BLUEBG],
 ["credit_controller","Useful",BLUEBG],["it_admin","Useful",BLUEBG],["viewer","Useful",BLUEBG],
 ["area_manager","Merge",AMBERBG],["regional_manager","Merge",AMBERBG],["national_sales_manager","Merge",AMBERBG],
 ["sales_director","Merge",AMBERBG],["branch_manager","Merge",AMBERBG],
 ["collection_officer","Merge",AMBERBG],["inventory_controller","Merge",AMBERBG],["cash_van","Merge",AMBERBG],
 ["driver","Optional",AMBERBG],["staff","Optional",AMBERBG],
 ["doctor","Optional",AMBERBG],["receptionist","Optional",AMBERBG],["stylist","Optional",AMBERBG],
 ["technician","Optional",AMBERBG],["housekeeping","Optional",AMBERBG],
]
tbl(["Role","Class","Role","Class"],
    [[_rl[i][0],_rl[i][1],(_rl[i+1][0] if i+1<len(_rl) else ""),(_rl[i+1][1] if i+1<len(_rl) else "")] for i in range(0,len(_rl),2)],
    widths=[1.9,1.0,1.9,1.0])
para("Summary: 6 Essential + 6 Useful core template roles. The 5 management tiers (area/regional/national/sales_director/"
     "branch) collapse into 2 (Manager + scope). collection_officer/inventory_controller/cash_van fold into "
     "accountant/warehouse_keeper/salesman + scope. 7 vertical roles load only with their pack. Net: 27 → ~8 template "
     "roles + a scope dimension + pack roles.", italic=True)

# ---- Current vs Future ----
part("Part 10 — Current State vs Recommended Future State")
para("The explicit target model. Every current capability is preserved; the taxonomy shrinks.", bold=True)
tbl(["Dimension","Current","Recommended future","Reduction","How"],[
 ["Business Types","21","~5 archetypes (+ labels)","-76%","Collapse identical module sets; type = archetype + vertical label"],
 ["Modules","31","~24 (3 tiers: Core/Capability/Pack)","-23%","Merge sales_orders→sales, change_requests/critical_alerts→workflow/analytics, drop services module"],
 ["Roles","27","~8 template roles + scope","-70%","Template roles + branch/area/region/national scope; pack roles load on demand"],
 ["Permissions (DB)","91","~60 (verb.resource)","-34%","Dedupe inventory/stock, customer/customers, reports/aggregate; drop dead engine keys"],
 ["Permission groups","16","8 governance categories","-50%","Re-map per Part 4"],
 ["Feature-flag kinds","2 mixed","2 separated (capability vs kill-switch)","—","Tie nav + page to same capability flag (fixes show-then-404)"],
],widths=[1.5,0.8,2.1,0.8,2.3],
 fill={(0,3):GREENBG,(1,3):GREENBG,(2,3):GREENBG,(3,3):GREENBG,(4,3):GREENBG})
h2("Ideal future counts — the founder's one-line answer")
b("Business Types: ~5 archetypes (Distribution, Stock-Retail, Service-Vertical, Self-contained-Vertical, Generic).",bold=True)
b("Modules: ~24 effective, presented as 3 tiers (6 Core · 6 Capability · ~12 Pack/vertical, only the relevant pack loads).",bold=True)
b("Roles: ~8 template roles (Admin, Manager, Sales, Field-Sales, Merchandiser, Finance, Warehouse, Procurement) + a scope "
  "dimension; vertical roles load with the pack.",bold=True)
b("Permissions: ~60 deduplicated verb.resource permissions (from 91), grouped into 8 governance categories.",bold=True)
para("Net effect: the cognitive surface (types × roles × permission families a founder/admin must reason about) drops by "
     "roughly half, with zero loss of shipped capability. Sequence: (1) dedupe permissions + drop dead keys, (2) collapse "
     "business types to archetypes, (3) introduce role-scope to retire the management tiers, (4) split the feature-flag "
     "kinds. Each step is independent and non-breaking.", bold=True, color=NAVY)

para("")
para("This handbook is the official VANTORA governance reference. The model is sound; the path to clarity is consolidation "
     "(Parts 7–8), not rebuilding. Adopting the archetype + template-role + deduplicated-permission model would cut the "
     "cognitive surface roughly in half while preserving every current capability.", bold=True, color=NAVY)

out="docs/audits/VANTORA-Governance-Architecture-Handbook.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
