#!/usr/bin/env python3
"""VANTORA — RP Migration-Repair Dry-Run HANDOFF PACKAGE (deploy owner) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"; GREYBG="ECECEC"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)

def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows,widths=None,size=8.4,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ---------------- Cover ----------------
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Route Planner Migration-Repair — Dry-Run Handoff"); r.bold=True; r.font.size=Pt(21); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Reconcile tracking rows 0358–0364 · For the deploy / pipeline owner · Clone-only procedure · No production action"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Supabase rsjvgehvastmawzwnqcs (vantora-staging) · Companion script: docs/rp_migration_repair_dryrun.sql · June 22, 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---------------- Why handoff ----------------
h1("Why this is a handoff (not an executed dry-run)")
para("Per instruction, the in-session attempt to create an isolated branch was retried and FAILED again — the "
     "execution environment gates create_branch (“MCP tool call requires approval”) and the Supabase MCP cannot run "
     "the real Supabase CLI (migration list / repair / db diff). As instructed, the run was stopped and packaged for "
     "the deploy owner. No production query was issued in this step; nothing was simulated on production; no raw INSERT "
     "was run anywhere.", bold=True)
tbl(["Capability needed","Available here?"],
    [["Create Supabase branch/clone","NO — create_branch is harness-gated"],
     ["Run supabase migration list / repair / db diff (CLI)","NO — MCP has no CLI"],
     ["Read-only SQL on production","Yes (used earlier; not re-run in this step)"]],
    widths=[3.6,3.0], fill={(0,1):REDBG,(1,1):REDBG,(2,1):GREENBG})

# ---------------- 1. Clone target ----------------
h1("1 · Clone / branch target")
b("Target to clone: Supabase project rsjvgehvastmawzwnqcs (vantora-staging) — NEVER production.")
b("Create via:  supabase branches create rp-drift-dryrun --experimental   (or dump+restore into a scratch project).")
b("Delete when done:  supabase branches delete rp-drift-dryrun --experimental.")
para("Note: a fresh branch replays only RECORDED migrations (0353–0357), so 0358–0364 objects will be ABSENT initially. "
     "To reproduce production's drift faithfully, apply the 0358–0364 SQL to the clone WITHOUT recording it (psql the "
     "files; do not `db push`), per STEP 0 of the companion script.", italic=True, size=9, color=GREY)

# ---------------- 2. Baseline ----------------
h1("2 · Baseline migration state (from prior read-only evidence)")
para("Recorded today (note: version = apply-timestamp, name = file stem):")
tbl(["version","name","recorded?"],
    [["20260620142235","0353_route_planner_access","YES"],
     ["20260620142253","0354_rp_reporting_graph","YES"],
     ["20260620142324","0355_rp_integration","YES"],
     ["20260620142400","0356_rp_request_center","YES"],
     ["20260620142410","0357_rp_schema_health","YES"],
     ["—","0358_rp_connector_admin","NO (object exists)"],
     ["—","0359_rp_planning_persistence_a","NO (object exists)"],
     ["—","0360_rp_dataset_persistence","NO (object exists)"],
     ["—","0361_rp_saved_plans","NO (object exists)"],
     ["—","0362_rp_mission_perms","NO (object exists)"],
     ["—","0363_rp_missions","NO (object exists)"],
     ["—","0364_rp_plan_approvals","NO (object exists)"]],
    widths=[1.7,3.1,1.8],
    fill={(i,2):GREENBG for i in range(5)} | {(i,2):AMBERBG for i in range(5,12)})
b("RP objects: 15 / 15 tables present (verified read-only).", color=GREEN, bold=True)
b("Pilot data: 28 rows across 7 populated RP tables (verified read-only).", color=GREEN, bold=True)

# ---------------- 3. Repair steps ----------------
h1("3 · Repair command / steps used (CLONE ONLY)")
para("CRITICAL — confirm the deploy convention first. Existing rows use timestamp versions + file-stem names, which is "
     "NOT what `supabase migration repair 0358` would write by default (it keys on the 0XXX token). Match the existing "
     "convention. Recommended path (direct, controlled INSERT of tracking rows only — no DDL, no data):", bold=True, color=RED)
code("begin;")
code("insert into supabase_migrations.schema_migrations (version, name) values")
code("  ('20260620142420','0358_rp_connector_admin'),")
code("  ('20260620142430','0359_rp_planning_persistence_a'),")
code("  ('20260620142440','0360_rp_dataset_persistence'),")
code("  ('20260620142450','0361_rp_saved_plans'),")
code("  ('20260620142500','0362_rp_mission_perms'),")
code("  ('20260620142510','0363_rp_missions'),")
code("  ('20260620142520','0364_rp_plan_approvals');")
code("commit;   -- on the CLONE only")
para("Versions above are examples that sort AFTER 0357; substitute the real apply-timestamps your pipeline assigns, "
     "preserving 0358 < … < 0364 order. Alternative CLI path (only if your pipeline timestamp-renames files):", size=9, color=GREY)
code("supabase migration repair --status applied 0358 0359 0360 0361 0362 0363 0364")

# ---------------- 4-5 expected outputs ----------------
h1("4–5 · Post-repair migration list & DB diff (expected)")
tbl(["Step","Command","Expected result"],
    [["Post-repair list","supabase migration list","0353–0364 all applied; NO pending"],
     ["Schema diff","supabase db diff","No schema differences (objects already match files)"]],
    widths=[1.5,2.4,2.7], fill={(0,2):GREENBG,(1,2):GREENBG})

# ---------------- 6-8 integrity ----------------
h1("6–8 · Integrity confirmations (expected on clone)")
tbl(["Check","Before","After","Expected"],
    [["Data fingerprint (7 tables)","28 rows","28 rows","IDENTICAL"],
     ["RP objects present","15","15","IDENTICAL"],
     ["Schema drift (db diff)","—","none","No change"]],
    widths=[2.6,1.3,1.3,1.4],
    fill={(0,3):GREENBG,(1,3):GREENBG,(2,3):GREENBG})
para("Because repair only inserts tracking rows, data and schema are mathematically unchanged; the before/after "
     "fingerprints must match exactly. Any mismatch = stop and investigate.", size=9, italic=True, color=GREY)

# ---------------- 9 errors ----------------
h1("9 · Errors / warnings")
b("BLOCKER (environment): create_branch is harness-gated; Supabase MCP cannot run the CLI — hence this handoff.", color=RED)
b("RISK (convention): timestamp-version vs 0XXX-filename mismatch — confirm the pipeline before choosing Path A vs B.", color=AMBER)
b("No other errors. No production action taken. No data touched.", color=GREEN)

# ---------------- 10 recommendation ----------------
h1("10 · Recommendation — is production repair safe?")
para("LIKELY SAFE, but gated on a clean clone dry-run by the deploy owner. The change is metadata-only (insert 7 "
     "tracking rows); the objects already exist and the pilot data is untouched, so risk is low. Do NOT proceed on "
     "production until: (a) the clone dry-run shows no pending migrations and a clean db diff with identical "
     "before/after fingerprints, AND (b) the deploy owner confirms the correct version convention, AND (c) you give "
     "explicit approval after reviewing the clone results.", bold=True)

h2("Scope guardrails honored")
para("Clone/branch only (attempted). No production writes. No production query in this step. No migration repair on "
     "production. No simulation of production repair. No raw INSERT on production. No Route Planner product-code merge. "
     "No data/route-guard/permissions changes. No PR #325. No field-insights. Production repair awaits separate "
     "explicit approval.", italic=True, color=GREY, size=9)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"VANTORA_RP_Migration_Repair_Handoff.docx")
d.save(out)
print("WROTE", out, os.path.getsize(out), "bytes")
