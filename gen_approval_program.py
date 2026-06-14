#!/usr/bin/env python3
"""VANTORA Approval Architecture Convergence Program -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    for ln in t.split("\n"):
        p=d.add_paragraph(); r=p.add_run(ln); r.font.name="Consolas"; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0x22,0x22,0x22)
def tbl(headers,rows,widths=None,size=8.0,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ---- Cover ----
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Approval Architecture Convergence Program"); r.bold=True; r.font.size=Pt(21); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("One enterprise-grade, multi-tenant approval engine — end-to-end execution program"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Target architecture · gap analysis · roadmap · risk · backward compatibility · phases · effort · blueprint · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---- Exec ----
h1("Executive summary")
para("This program converges VANTORA's nine approval workflows onto ONE configurable, multi-tenant approval engine — the "
     "engine the platform already has — and retires the six hardcoded one-off approvals. The objective is a single "
     "enterprise-grade approval model that scales across FMCG, Pharmacy, Retail, Clinics and future verticals, where each "
     "tenant's governance is CONFIGURATION, not code.", bold=True)
para("Crucially, the existing Phase-2 engine is already far richer than its early migrations suggested. Inspecting the LIVE "
     "definitions revealed: per-tenant + versioned definitions with draft/publish and a visual builder; conditional, "
     "multi-step, parallel, quorum routing; permission-aware decision authorization (erp_workflow_user_can_act already "
     "honours approver_type='permission'); an authority-limit layer (approval_action + erp_within_limit + erp_role_limits); "
     "and task notifications. The convergence is therefore mostly WIRING + a few missing primitives, not a rebuild.",
     bold=True, color=GREEN)
h2("What was built in this step (pilot-safe foundation)")
b("Permission-based approver RESOLUTION (migration 0296): erp_workflow_resolve_users now fans out / notifies "
  "approver_type='permission', mirroring erp_user_has_permission exactly. approver_type can now be 'permission'.", color=GREEN)
b("Governance flags + enforcement (0296 + 0297): block_self_approval and require_reject_reason on a definition; the decide "
  "RPC enforces them — a faithful SUPERSET of the live function (all existing logic verbatim). Default OFF ⇒ zero "
  "behaviour change until a tenant opts in.", color=GREEN)
b("Inbox fix (TS): the Workflow Inbox now surfaces role- AND permission-assigned tasks (it previously hid them), mirroring "
  "the engine's authorization. This is the concrete gap that left credit-limit approvers with no UI path. Unit-tested.", color=GREEN)
b("Prepared (not applied): a per-tenant blueprint for Credit-Limit, Trade-Spend and Price-Change definitions.", color=GREEN)
para("Everything above is ADDITIVE and DORMANT: no existing workflow uses the new capabilities yet, so the live pilot is "
     "unaffected. Nothing was applied to the pilot database in this step — the migrations are prepared in the repo for the "
     "team to apply on the normal pipeline.", bold=True)

# ---- 1. Target architecture ----
part("1 · Final target architecture")
para("A single engine. A request creates a workflow INSTANCE; the engine routes it through one or more LEVELS by RULES; "
     "approvers decide in ONE inbox; the engine records audit + applies the outcome via a per-entity handler; SLAs escalate "
     "the undecided. Domain code never decides WHO or WHEN — only WHAT happens on approval.")
h2("Resolution model (request → approver)")
b("CAPABILITY GATE — the action always re-checks a permission server-side (unchanged security floor).")
b("RULE MATCH — request type + conditions (amount thresholds via condition gt/lt; branch scope; segment) choose the LEVEL(s).")
b("ASSIGNEE per level — permission-holder OR role, resolved within branch scope; named-user only as override/escalation.")
b("GUARDS — cannot-approve-own, mandatory reject reason, quorum (N-of-M) for parallel levels, authority limits (amount).")
b("ESCALATION — SLA timer routes undecided tasks to escalate_to (higher level / named user); status 'escalated'.")
h2("Data model (already present, extended)")
tbl(["Object","Role","Status"],
[
 ["erp_workflow_definitions","Per-tenant, versioned (draft/publish), builder schema, approval_action, +governance flags","Present (+0296 flags)"],
 ["erp_workflow_steps","Levels: approver_type, condition, mode, required_approvals, sla_hours, escalate_to","Present (+'permission')"],
 ["erp_workflow_instances","A running approval; has started_by + branch_id (branch scope ready)","Present"],
 ["erp_workflow_tasks","Per-assignee task; decided_by/at, comment","Present"],
 ["erp_workflow_resolve_users / _user_can_act","Resolve & authorize by user/role/company_admin/permission","Present (+permission resolve)"],
 ["erp_within_limit / erp_role_limits","Per-role/user amount authority limits","Present"],
 ["Outcome handlers (workflow-handlers.ts)","Apply the real change on approve/reject","Present (extend per entity)"],
],widths=[2.4,3.0,1.2])

# ---- 2. Gap analysis ----
part("2 · Gap analysis vs current implementation")
tbl(["Capability","Current","After this step","Remaining work"],
[
 ["Permission-based approver resolution","Auth yes; fan-out/notify NO","DONE (0296)","—"],
 ["Role-based approver resolution","Engine yes; inbox hid it","DONE (inbox fix)","—"],
 ["Cannot-approve-own-request","Missing","DONE (flag, 0297)","Set flag per workflow"],
 ["Mandatory reject reason","App-layer only","DONE (flag, 0297)","Set flag per workflow"],
 ["Multi-level / sequential / parallel / quorum","Present","Present","Configure per workflow"],
 ["Amount thresholds","condition gt/lt + authority limits","Present","Configure per workflow"],
 ["Full audit trail","Present (erp_log_audit)","Present","—"],
 ["Branch-scoped approvals","instances.branch_id present; resolver not scoped","Designed","P1b: branch-aware resolve"],
 ["Escalation rules / SLA timers","Columns present; no driver","Designed","P4: scheduled escalation job"],
 ["Unified inbox for ALL approvals","3 surfaces (Field/Workflow/Center)","Inbox surfaces role+perm","P3: route field approvals + retire dupes"],
],widths=[2.3,1.7,1.3,1.5],
fill={(0,2):GREENBG,(1,2):GREENBG,(2,2):GREENBG,(3,2):GREENBG,(7,2):AMBERBG,(8,2):AMBERBG,(9,2):AMBERBG})

# ---- 3. Roadmap ----
part("3 · Migration roadmap")
tbl(["Phase","Scope","Status"],
[
 ["P0 — Foundation","Permission resolution, governance flags+enforcement, inbox role/perm surfacing","DONE (this step, pilot-safe)"],
 ["P1 — Priority workflows","Configure Credit-Limit (perm+threshold), Trade-Spend (perm gate), Price-Change (new) on the engine; wire actions through erp_workflow_start","Next"],
 ["P1b — Branch scope","Branch-aware resolve_users/user_can_act using instances.branch_id","Next"],
 ["P2 — Converge field approvals","Route Van Transfer, Load Request, Customer Transfer, Visit, Day-Close via engine; keep RPCs as outcome handlers","Post-pilot"],
 ["P3 — Unify the inbox","One Approval Queue over engine tasks for ALL types; retire Approval Center + bespoke readers","Post-pilot"],
 ["P4 — Escalation timers","Scheduled job drives sla_hours → escalate_to; status 'escalated'; notifications","Post-pilot"],
 ["P5 — Tenant rule builder","Harden the existing builder for self-serve per-tenant rules + validation guardrails","Post-pilot"],
],widths=[1.5,3.6,1.5],
fill={(0,2):GREENBG})

# ---- 4. Risk ----
part("4 · Risk assessment")
tbl(["Risk","Severity","Mitigation"],
[
 ["Replacing the live decide RPC mid-pilot","High","0297 is a faithful superset; guards default OFF; validate on staging in a transaction before apply; reversible"],
 ["Dual-decision during field-approval cutover","High","Migrate one workflow at a time behind a flag; engine instance becomes the single source of truth; RPC = outcome handler"],
 ["Tenant mis-config locks approvals (no approver resolves)","Medium","Validate-on-publish: every level must resolve ≥1 user; fallback to company_admin; 'no approver' alert"],
 ["Inbox performance (2000-row in-memory filter)","Medium","Move actionable scoping into an indexed query before high volume (P3)"],
 ["Permission resolve set ≠ authorize set","Medium","0296 mirrors erp_user_has_permission exactly; covered by intent + tests"],
 ["Escalation job double-fires","Medium","Idempotent; act only on pending+overdue; audit every escalation (P4)"],
 ["Outcome-handler regressions","Medium","Reuse the existing, tested domain RPCs as handlers; engine only flips state + calls handler"],
],widths=[2.5,0.9,3.0])

# ---- 5. Backward compatibility ----
part("5 · Backward-compatibility plan")
b("Additive-only schema: new columns default to legacy behaviour; new approver_type value is opt-in; no column dropped.")
b("Faithful superset RPC: 0297 reproduces the entire live decide; only adds two flag-gated guards (default off).")
b("Dormant capabilities: no existing definition uses permission-routing or the flags, so live pilot behaviour is identical.")
b("One-workflow-at-a-time convergence behind feature flags; each cutover independently reversible.")
b("Old endpoints kept as outcome handlers during P2 so no caller breaks; bespoke screens stay until P3 retires them.")
b("Reversibility: re-applying the prior decide / dropping the new columns restores the previous engine with no data loss.")

# ---- 6. Phases & 7. Effort ----
part("6 · Execution phases & 7 · Effort estimate")
tbl(["Phase","Key tasks","Effort (eng-days)","Risk class"],
[
 ["P0 Foundation","Done: 0296/0297 + inbox + tests","~ delivered","Pilot-safe"],
 ["P1 Priority workflows","Credit/Trade-Spend/Price-Change definitions + start wiring + outcome handlers + tests","6-9","Pilot-safe→Post-pilot"],
 ["P1b Branch scope","Branch-aware resolve + tests + 1 workflow proof","3-5","Post-pilot"],
 ["P2 Converge field (5 workflows)","Per workflow: start wiring, handler, queue surface, flag, validate (~2-3d each)","12-18","High-risk"],
 ["P3 Unify inbox","One queue over engine tasks; retire Center + bespoke readers; perf query","6-9","High-risk"],
 ["P4 Escalation/SLA","Scheduled job, escalate path, notifications, audit","5-8","Post-pilot"],
 ["P5 Tenant builder hardening","Validation guardrails, self-serve UX, docs","6-10","Post-pilot"],
],widths=[1.5,2.9,1.0,1.2])
para("Indicative total beyond P0: ~6-10 engineer-weeks, sequenced so the pilot keeps running throughout. P1 is the only "
     "near-term track; P2-P5 are a post-pilot programme.", italic=True, color=GREY, size=9)

# ---- 8. Blueprint ----
part("8 · Approval workflow blueprint (all 9)")
para("Target engine configuration for every major FMCG financial + operational approval. (Ready SQL for the first three is "
     "in supabase/blueprints/approval-rules.blueprint.sql.)")
tbl(["Workflow","Entity","Levels & routing","Guards"],
[
 ["Credit Limit Increase","credit_limit_request","L1 <=5k role=branch_manager · L2 >5k perm=credit.request.approve · esc→admin 24h","self-block, reject-reason"],
 ["Trade Spend","trade_promotion","L1 perm=pricing.rule.edit (matches action) 48h","self-block, reject-reason"],
 ["Price Change (new)","price_change","L1 perm=pricing.manage · esc→admin 24h","self-block, reject-reason"],
 ["Customer Transfer","customer_transfer","L1 perm=customer.transfer, branch-scoped","reject-reason"],
 ["Out-of-Route Visit","visit_compliance","L1 perm=visit.approve_out_of_route, branch-scoped","reject-reason"],
 ["Day-Close Exception","work_session","L1 perm=day.approve_close_exception, branch-scoped","add explicit reject"],
 ["Load Request","stock_request","L1 perm=stock_request.approve, branch-scoped","reject-reason"],
 ["Van Transfer","van_transfer","auto < threshold · else L1 perm=stock.transfer.approve","reject-reason"],
 ["Van Reconciliation","van_reconciliation","auto <= variance · else L1 perm=reconciliation.approve","reject-reason"],
],widths=[1.4,1.5,2.6,1.1])
para("Customer Master-Data Change already runs on the engine (change-request workflow) and only needs governance flags + "
     "branch scope applied.", italic=True, color=GREY, size=9)

# ---- 9. Change classification ----
part("9 · Change classification (before destructive work)")
h2("Pilot-safe (additive, dormant — done now or safe to apply now)")
b("0296 permission resolution + approver_type='permission' + governance flag columns.", color=GREEN)
b("0297 governance enforcement (faithful superset; guards default off).", color=GREEN)
b("Inbox role/permission surfacing (TS) — no new tasks exist in the pilot, so no visible change yet.", color=GREEN)
b("Configuring NEW definitions for Trade-Spend / Price-Change (no current workflow to disturb).", color=GREEN)
h2("Post-pilot (functional change; schedule deliberately)")
b("Re-routing Credit-Limit from company_admin to permission/role + thresholds (changes who approves).", color=AMBER)
b("Branch-scoped resolution (P1b) and escalation timers (P4).", color=AMBER)
b("Tenant rule-builder self-serve hardening (P5).", color=AMBER)
h2("High-risk (destructive / behaviour-changing — gated, flagged, reversible)")
b("Converging the five hardcoded field approvals onto the engine (P2) — touches live sell/collect-adjacent flows.", color=RED)
b("Retiring the bespoke queue readers + Approval Center and moving to one inbox (P3).", color=RED)
b("Replacing the live decide RPC on the PILOT database — apply only after staging validation, in a maintenance window.", color=RED)
para("Recommendation: apply P0 + P1 (priority workflows) when ready; hold all High-risk items until after the pilot, each "
     "behind its own flag with a tested rollback. The pilot can run unchanged throughout.", bold=True, color=NAVY)

out="docs/audits/VANTORA-Approval-Convergence-Program.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
