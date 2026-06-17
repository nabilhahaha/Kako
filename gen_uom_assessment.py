#!/usr/bin/env python3
"""VANTORA Multi-UoM Readiness Assessment -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def bl(t,size=9.3,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    for ln in t.split("\n"):
        p=d.add_paragraph(); r=p.add_run(ln); r.font.name="Consolas"; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0x22,0x22,0x22)
def tbl(headers,rows,widths=None,size=8.0,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Multi-UoM Readiness Assessment"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Buy by carton, sell by piece — how ready is VANTORA, and what's the path?"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Assessment only — no implementation · verified against live schema + code · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Headline verdict")
para("VANTORA is NOT starting from zero. It already contains a COMPLETE, working multi-UoM engine — but it lives in the "
     "PHARMACY vertical only. The conversion model, master table, helper library, validation rules, UoM-aware pricing and a "
     "management screen all exist and are proven in pharmacy POS + receiving. The FMCG DISTRIBUTION flows (van-sell, "
     "general POS, purchase orders, stock transfers, returns, reconciliation) are single-UoM (base units) and do NOT use "
     "any of it. So this is a GENERALISE-AN-EXISTING-PATTERN effort, not a greenfield build.", bold=True)
para("Two-speed reality:", bold=True)
bl("Pharmacy = full multi-UoM (feature-flagged pharmacy.multi_unit_support): sell/receive in any UoM, converts qty + "
   "price to base, audits entered vs base. WORKING.", color=GREEN)
bl("FMCG distribution = base-unit only: van-sell uses the old non-UoM price function; POS has no unit picker; no "
   "transaction line records a UoM; the UoM-aware pricing RPC is called only from settings, never from a sale.", color=AMBER)
bl("Pilot data: the conversion master erp_product_uoms is EMPTY (0 rows, all tenants); catalog base/purchase units are "
   "descriptive labels with no linking factor (e.g. base 'strip', purchase 'box', strips-per-box UNKNOWN).", color=RED)

# 1 status
part("1 · Current implementation status")
tbl(["Area","Status","Evidence"],
[
 ["UoM master (erp_product_uoms: uom, factor, barcode, is_case, sort)","Built; UI exists; EMPTY in data","settings/uom + migration 0137"],
 ["Conversion helpers (toBase/fromBase/priceToBase/factorOf)","Built & used (pharmacy)","src/lib/erp/uom.ts, uom-server.ts, uom-pricing.ts, uom-rules.ts"],
 ["Conversion RPCs (erp_uom_to_base, erp_resolve_price uom-overload)","Built","migration 0137"],
 ["Pricing by UoM (erp_prices.uom + min_qty breaks)","Built; called only from settings","3 rows in data"],
 ["Pharmacy POS (sell by box/strip + convert)","WIRED & converting","pharmacy/pos/actions.ts"],
 ["Pharmacy receive (buy UoM → base)","WIRED & converting","pharmacy/receive/actions.ts"],
 ["Van-sell (erp_van_sell)","Base-unit only (no UoM)","uses old 5-arg erp_resolve_price"],
 ["General POS / sales orders","No UoM picker","sales/pos/pos-terminal.tsx"],
 ["Transaction lines (invoice/PO/SO/stock/van/returns)","NO uom column anywhere","schema confirmed"],
 ["Stock balances","Base-unit only (correct invariant)","erp_inventory_stock"],
 ["Reporting","No UoM breakdown","no per-UoM report"],
],widths=[2.6,1.6,2.4],
fill={(0,1):AMBERBG,(4,1):GREENBG,(5,1):GREENBG,(6,1):AMBERBG,(7,1):AMBERBG,(8,1):REDBG,(9,1):GREENBG,(10,1):AMBERBG})

# 2 existing
part("2 · Existing capabilities (the proven foundation)")
bl("A real conversion model: erp_product_uoms holds, per product, each sellable/purchasable UoM with a numeric FACTOR "
   "(base units per 1 of that UoM), a BARCODE (per-UoM scanning), an is_case flag and a sort — enough for a "
   "Case/Inner/Piece hierarchy.")
bl("Stock invariant done right: inventory is ALWAYS stored in base units; conversion happens at the entry points "
   "(sell/receive), so balances are unambiguous. This is the correct architecture and a big head-start.")
bl("Governance rules already modelled: sell_mode (base | sales | all), allow_fractional (e.g. Kg/Gram), validateSell / "
   "validatePurchase — exactly the rules FMCG needs (no half-pieces, allow fractional weight).")
bl("UoM-aware pricing path: erp_prices(uom, min_qty) + erp_resolve_price(uom-overload) resolves a price for a specific "
   "UoM with a base×factor fallback — the pricing-by-UoM requirement is already designed.")
bl("Management UI: settings/uom lets an admin add per-product UoMs with factor + barcode. The data-entry surface exists.")

# 3 missing
part("3 · Missing capabilities (the gap to FMCG multi-UoM)")
bl("Transaction-level UoM capture: NO line table (invoice, PO, sales order, stock movement, van transfer, sales/purchase "
   "return) records which UoM was transacted or the entered qty. Today a line is a bare base quantity — you cannot tell "
   "'2 cartons' from '48 pieces'. This is the single biggest gap.", color=RED)
bl("FMCG transaction wiring: van-sell, general POS, sales orders, PO receiving (non-pharmacy), van/stock transfers, "
   "returns and reconciliation do not call the conversion or UoM-pricing functions — they assume base units.", color=RED)
bl("UoM pickers in the field/desk UI: no carton/piece selector in van-sell or POS; no per-UoM barcode scan-to-sell.", color=AMBER)
bl("Populated factors: erp_product_uoms is empty for FMCG; purchase_uom differs from base on 342 catalog products with NO "
   "factor to convert. pack_size is free-text ('12x1'), not a usable numeric hierarchy.", color=AMBER)
bl("Reporting by UoM: no sell-through, stock or variance report expressed in cartons vs pieces.", color=AMBER)
bl("Cost conversion on purchase: outside pharmacy, receiving doesn't convert purchase-UoM cost to base cost.", color=AMBER)

# 4 architecture
part("4 · Architectural design (recommended)")
para("Reuse the pharmacy engine as the company-wide UoM core; keep base-unit stockkeeping; capture UoM on every "
     "transaction line and convert to base at write time. Concretely:")
h2("Model")
bl("Per product: a base UoM (smallest sellable unit) + N alternate UoMs in erp_product_uoms, each with factor, barcode, "
   "is_case, sort. Case/Inner/Piece = three rows (piece 1, inner 6, case 24). Kg/Gram = base gram, Kg factor 1000 + "
   "allow_fractional.")
bl("Inventory stays in BASE units (unchanged) — the proven invariant. Pricing can be per-UoM (erp_prices.uom) or "
   "base×factor fallback.")
h2("Transaction capture (the new part)")
bl("Add to each line: entered_uom (text), entered_qty (numeric), uom_factor (numeric snapshot), and keep quantity as the "
   "BASE qty (so all existing stock/finance logic is unchanged). The snapshot factor makes the line self-describing and "
   "audit-stable even if the master changes later.")
bl("At write time (sell/receive/transfer/return): base_qty = entered_qty × factor; unit_price resolved per entered_uom; "
   "stock moves in base. This mirrors exactly what pharmacy POS already does — generalised to the shared sell/issue RPCs.")
h2("Surfaces")
bl("Van-sell + POS: a UoM selector per line (defaulting to default_sell_uom), barcode scan picks the matching UoM row.")
bl("PO receiving: enter in purchase_uom, convert to base for stock + cost.")
bl("Returns/transfers/reconciliation: same line shape (entered_uom + factor + base_qty).")
bl("Reporting: aggregate in base, present in a chosen UoM via factor.")
para("Net: it is the pharmacy pattern, lifted into src/lib/erp/uom* (already shared) and wired into the FMCG "
     "transaction RPCs, plus four small additive columns on the line tables.", italic=True, color=GREY, size=9)

# 5 migration impact
part("5 · Migration impact")
tbl(["Change","Type","Risk"],
[
 ["Add entered_uom / entered_qty / uom_factor columns to line tables (quantity stays base)","Additive, nullable","Low (back-compat: null ⇒ base, factor 1)"],
 ["Populate erp_product_uoms with real factors + barcodes (per tenant)","Data","Low (master data; can seed)"],
 ["Wire van-sell / POS / SO to uom-aware price + convert to base","Behaviour (flagged)","Medium-High (core sell path)"],
 ["Wire PO receiving conversion outside pharmacy","Behaviour (flagged)","Medium"],
 ["UoM pickers + barcode-per-uom in van/POS/PO UI","UI","Medium"],
 ["Make pack_size numeric (or derive from product_uoms)","Schema cleanup","Low"],
 ["UoM-aware reporting","Additive","Low"],
],widths=[3.4,1.4,1.8])
para("All of it can be FLAG-GATED (mirroring pharmacy.multi_unit_support) and is backward-compatible: a null entered_uom "
     "means 'base unit', so existing single-UoM tenants and the current pilot are unaffected until enabled. No destructive "
     "change; stock invariant (base units) is preserved throughout.", bold=True, color=GREEN)

# 6 pilot importance
part("6 · FMCG pilot importance")
para("HIGH for a REAL FMCG distributor — but not a blocker for the current single-UoM pilot.", bold=True)
bl("Buy-by-carton / sell-by-piece is the defining FMCG distribution behaviour. Reps sell loose pieces, inners and full "
   "cases off the same van; buyers receive cartons; pricing differs per UoM (piece vs case). Without it, a true "
   "distributor cannot run on the platform.", color=DARK)
bl("Today's pilot WORKS because its products are transacted in a single unit (base) consistently — so the missing "
   "conversion never bites. The moment a pilot product must be 'bought by box, sold by strip', the gap is blocking.", color=AMBER)
bl("It also unlocks correctness: per-UoM pricing, accurate stock ('2 cartons + 3 pieces'), barcode-per-UoM scanning, and "
   "carton/piece reporting that distributors expect.", color=DARK)
para("Recommendation: treat Multi-UoM as the next major FMCG-readiness workstream AFTER the approval program — it is the "
     "biggest remaining functional gap for broad FMCG fit, and the foundation already exists to do it safely.",
     bold=True, color=NAVY)

# 7 roadmap
part("7 · Recommended roadmap")
tbl(["Phase","Scope","Risk"],
[
 ["U1 — Master & factors","Promote uom.ts/uom-server to company-wide core (drop pharmacy-only assumption); seed erp_product_uoms factors + barcodes; numeric pack hierarchy; settings/uom available to FMCG","Low"],
 ["U2 — Transaction capture","Additive line columns (entered_uom/entered_qty/uom_factor); base qty unchanged; write-time conversion helper shared by all sell/issue RPCs","Low-Med"],
 ["U3 — Sell path (flagged)","Van-sell + POS UoM picker + barcode-per-UoM; route pricing through the uom-aware erp_resolve_price; validateSell rules","Med-High"],
 ["U4 — Buy path (flagged)","PO receiving conversion (qty + cost) to base outside pharmacy","Medium"],
 ["U5 — Returns/transfers/recon","Same line shape end-to-end","Medium"],
 ["U6 — Reporting","Per-UoM sell-through / stock / variance views","Low"],
],widths=[1.5,3.7,1.0])
para("Sequencing mirrors the approval program: additive foundation first (U1-U2, pilot-safe), then flag-gated behaviour "
     "(U3-U5) validated on staging before activation, then reporting (U6). Reuse pharmacy.multi_unit_support's pattern as "
     "the per-tenant switch. Estimated ~5-8 engineer-weeks, sequenced so single-UoM tenants are never disrupted.",
     bold=True, color=NAVY)

out="docs/audits/VANTORA-Multi-UoM-Assessment.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
