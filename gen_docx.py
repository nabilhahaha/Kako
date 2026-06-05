#!/usr/bin/env python3
"""Generate VANTORA_Full_Platform_Documentation.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── base styles ──
normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10.5)

BRAND = RGBColor(0x1F, 0x3A, 0x5F)
ACCENT = RGBColor(0x0B, 0xC5, 0xDA)

for i, sz in [(1, 18), (2, 14), (3, 12)]:
    st = doc.styles[f'Heading {i}']
    st.font.color.rgb = BRAND
    st.font.size = Pt(sz)
    st.font.bold = True

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement('w:shd'); sh.set(qn('w:val'), 'clear'); sh.set(qn('w:fill'), hexcolor)
    tcPr.append(sh)

def add_table(headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for j, h in enumerate(headers):
        hdr[j].text = ''
        p = hdr[j].paragraphs[0]; r = p.add_run(h); r.bold = True
        r.font.size = Pt(9.5); r.font.color.rgb = RGBColor(0xFF,0xFF,0xFF)
        shade(hdr[j], '1F3A5F')
    for row in rows:
        cells = t.add_row().cells
        for j, val in enumerate(row):
            cells[j].text = ''
            p = cells[j].paragraphs[0]; rr = p.add_run(str(val)); rr.font.size = Pt(9)
    if widths:
        for j, w in enumerate(widths):
            for row in t.rows:
                row.cells[j].width = Inches(w)
    doc.add_paragraph()
    return t

def h1(txt): doc.add_heading(txt, level=1)
def h2(txt): doc.add_heading(txt, level=2)
def p(txt, bold=False, italic=False):
    par = doc.add_paragraph(); r = par.add_run(txt); r.bold = bold; r.italic = italic
    return par
def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')

# ════════════════════════════ TITLE PAGE ════════════════════════════
ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run('VANTORA'); r.bold = True; r.font.size = Pt(40); r.font.color.rgb = BRAND
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Multi-Tenant SaaS ERP Platform'); r.font.size = Pt(16); r.font.color.rgb = ACCENT
sub2 = doc.add_paragraph(); sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub2.add_run('Full Platform Documentation'); r.font.size = Pt(14); r.bold = True
meta = doc.add_paragraph(); meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Version 1.0  ·  June 2026  ·  Generated for engineering & product handoff').font.size = Pt(10)
note = doc.add_paragraph(); note.alignment = WD_ALIGN_PARAGRAPH.CENTER
rn = note.add_run('Stack: Next.js (App Router, RSC) · Supabase (Postgres + RLS + Auth + Storage) · TypeScript')
rn.italic = True; rn.font.size = Pt(9.5); rn.font.color.rgb = RGBColor(0x66,0x66,0x66)
doc.add_page_break()

# ════════════════════════════ TOC ════════════════════════════
h1('Contents')
toc = [
 '1. Executive Summary','2. Platform Purpose','3. Full System Architecture',
 '4. Modules','5. Platform Owner Area','6. Company Admin Area',
 '7. Roles & Permissions Hierarchy','8. Plans, Modules, Entitlements & Industry Packs',
 '9. Multi-Tenant Structure & Tenant Isolation','10. Database Structure Overview',
 '11. RLS / Security Model','12. Audit Logging System','13. Navigation & UI Structure',
 '14. Existing Pages / Routes','15. Migrations Applied','16. Tests & Verification Status',
 '17. Scalability Review Summary','18. PR / Branch Structure','19. Remaining Technical Debt',
 '20. Recommended Next Development Phase',
]
for t in toc:
    par = doc.add_paragraph(); run = par.add_run(t); run.font.size = Pt(10.5)
doc.add_page_break()

# ════════════════════════════ 1. EXEC SUMMARY ════════════════════════════
h1('1. Executive Summary')
p('VANTORA is a production-grade, multi-tenant SaaS ERP platform that serves many independent '
  'businesses ("companies"/"tenants") of different industries from a single application and database. '
  'Each tenant operates in strict isolation, sees only the modules its subscription plan and business '
  'type unlock, and is governed by a role-and-permission model. A separate vendor ("Platform Owner") '
  'tier administers the entire platform: companies, plans, modules, roles, billing, audit, and support.')
p('The platform is built on Next.js (App Router, React Server Components + Server Actions) and Supabase '
  '(Postgres with Row-Level Security, Auth, and Storage). Authorization is enforced in five consistent '
  'layers — Plan entitlement → Company module → Role permission → UI visibility → Route/API guard — with '
  'the database (RLS) as the hard isolation boundary.')
p('Current status (production-grade checklist):', bold=True)
bullets([
 'Clean platform governance and a unified authorization model.',
 'Tenant isolation verified across all 129 database tables (reads and writes).',
 'All sensitive mutations are audit-logged; tenants and the vendor can each review their audit trail.',
 'Consistent entitlement chain; every licensable module has a navigation surface.',
 'Database performance hygiene: 100% foreign-key index coverage; no per-row RLS auth evaluation.',
 'Continuous-integration guards for architecture, navigation, authorization, and schema health.',
 'A documented, trigger-based scalability plan for growth to 10,000+ companies.',
])

# ════════════════════════════ 2. PURPOSE ════════════════════════════
h1('2. Platform Purpose')
p('VANTORA lets a vendor operate one ERP product that adapts to many industries instead of building a '
  'separate system per vertical. A company signs up, picks a business type (e.g. clothing, clinic, '
  'restaurant, FMCG distribution, pharmacy, supermarket), and immediately gets a tailored experience: '
  'only the relevant modules, navigation, roles, and screens — no irrelevant complexity.')
p('Goals:', bold=True)
bullets([
 'One codebase, many industries — industry "packs" layer on shared commerce/accounting foundations.',
 'Strong multi-tenant isolation so one company can never see or affect another.',
 'A powerful but safe Platform Owner control center for onboarding and administering tenants.',
 'A clean, role-appropriate experience for every end user (owner, manager, cashier, accountant, …).',
 'Scalable foundations ready for large-scale customer onboarding.',
])

# ════════════════════════════ 3. ARCHITECTURE ════════════════════════════
h1('3. Full System Architecture')
p('VANTORA has two authorization tiers and a layered request path:')
p('Two tiers', bold=True)
bullets([
 'Platform (vendor) tier — the operator and its internal staff. Identified by the is_platform_owner '
 'flag on a profile, or membership in erp_platform_staff. They see only the vendor control center '
 '(/platform/*) and never tenant operational screens.',
 'Tenant (company) tier — a customer company’s own users, identified by branch membership '
 '(erp_user_branches). Scoped to exactly one company.',
])
p('Request path', bold=True)
bullets([
 'UI: Next.js App Router — React Server Components, Server Actions, i18n (Arabic/English, RTL).',
 'Authorization core (src/lib/erp): guards, permissions, capabilities, navigation/visibility, '
 'home routing, plan/role admin logic.',
 'Auth context: getUserContext() / getPlatformContext() resolve identity, company, modules, and '
 'permissions — now request-memoized (React cache()) so each request resolves once.',
 'Data: Supabase Postgres with Row-Level Security. Reads use the user’s JWT (RLS-enforced); '
 'mutations go through guarded, audited Server Actions; multi-write operations use SECURITY DEFINER '
 'RPCs (e.g. invoice issue, fashion checkout) for atomicity.',
 'Files: Supabase Storage (object store); the database holds metadata + storage paths only.',
])
p('The five-layer entitlement chain', bold=True)
p('Plan entitlement (erp_plan_modules) → Company module (erp_company_modules) → Role permission '
  '(erp_company_role_permissions / erp_role_permissions) → UI visibility (visibleSections) → '
  'Route/API guard (requireModule / requirePermission). The apex tiers (Platform Owner, Super Admin) '
  'bypass layers 1–3.')

# ════════════════════════════ 4. MODULES ════════════════════════════
h1('4. Modules')
p('A "module" is a licensable capability. Effective modules for a tenant = the modules its company has '
  'enabled ∩ the modules its plan unlocks. The 20 plan-gateable modules split into core capabilities '
  'and industry verticals; three additional "item-level" modules refine screens inside core sections.')
add_table(
 ['Module', 'Type', 'What it does'],
 [
  ['sales','Core','Quotes, orders, invoices, customers, collections.'],
  ['inventory','Core','Products, stock levels, movements, low-stock/expiry.'],
  ['purchasing','Core','Suppliers, purchase orders, supplier returns/payments.'],
  ['accounting','Core','Chart of accounts, journals, double-entry posting, finance reports.'],
  ['pos','Core','Generic counter Point-of-Sale (quick sale).'],
  ['crm','Core','Customer relationship management surface (any-of with sales).'],
  ['workflow','Core','Approvals & workflow engine (approval center, tasks).'],
  ['analytics','Core','Analytics/reporting surface (sales report, dashboards).'],
  ['field_ops','Core','Field/van sales: journeys, GPS visits, day-close, settlements.'],
  ['integrations','Core','Integration hub, API keys, connectors, data import/export, onboarding.'],
  ['hotel','Vertical','Rooms and bookings.'],
  ['clinic','Vertical','Reception, doctor queue, appointments, patients, visits, services.'],
  ['restaurant','Vertical','Tables, orders, kitchen.'],
  ['salon','Vertical','Appointments, tickets, services.'],
  ['pharmacy','Vertical','Dispensing + near-expiry tracking.'],
  ['laundry','Vertical','Laundry orders and services.'],
  ['market','Vertical','Supermarket POS.'],
  ['wholesale','Vertical','Price levels/tiers, wholesale orders.'],
  ['distribution','Vertical','FMCG distribution: routes, journey compliance, MSL, perfect-store.'],
  ['fashion','Vertical','Clothing store pack: variants, POS, installments, cash box, reports.'],
  ['sales_orders','Item-level','Sales-order screen inside Sales (business-type gated).'],
  ['returns','Item-level','Sales-returns screen inside Sales.'],
  ['warehousing','Item-level','Transfers, stock requests, counts, multi-warehouse.'],
 ],
 widths=[1.2,1.0,4.3])
p('Abstract packs: "retail" and "electrical" have no standalone database module — "electrical" is '
  'permission-gated (electrical.rma); "retail" preselects core modules. A continuous-integration test '
  'guarantees every real module has a navigation surface (no "orphan" modules).', italic=True)

# ════════════════════════════ 5. PLATFORM OWNER ════════════════════════════
h1('5. Platform Owner Area')
p('The vendor control center lives under /platform/* and is owner/staff-gated. It is organized into '
  'SaaS-style groups: Overview, Tenants, Catalog, Billing, Team & Access, Reference Data.')
bullets([
 'Overview / Activity / Analytics — platform-wide KPIs, activity feed, cross-tenant analytics.',
 'Companies (Company 360) — per-company profile: information, business type, plan, subscription, '
 'enabled modules, roles & permissions, users, integrations, usage, audit; actions to '
 'suspend/activate, renew, change plan, grant/clear trial, toggle modules, onboard admins, reset '
 'passwords, and View As Company (read-only per-role tenant preview).',
 'Plans & Modules editor — create/edit/clone/archive/reorder plans, set per-plan module '
 'entitlements with a live impact preview, and edit business-type templates.',
 'Global Roles & Permissions — manage the default role catalog: create/rename/clone/delete roles, '
 'grouped permission editor with search, dangerous-permission warnings, and role comparison.',
 'Billing — subscriptions, invoices, unpaid totals, price book.',
 'Team & Access — internal platform staff (roles + granular permissions, escalation-guarded) and '
 'the forensic Audit Log.',
 'Reference Data — e.g. the clinic drug list importer.',
])
p('The Platform Owner never sees an "Out of Plan" / upgrade screen and bypasses tenant gates by design.')

# ════════════════════════════ 6. COMPANY ADMIN ════════════════════════════
h1('6. Company Admin Area')
p('A Company Admin (the tenant "admin" role) administers their own company only. Capabilities include:')
bullets([
 'Staff & organization — invite users, assign roles, manage departments/teams/job titles and '
 'reporting lines (audited).',
 'Branches & regions — manage branches, regions, and areas (within plan limits).',
 'Roles & permissions — per-company role configuration and the Authorization Console '
 '(capabilities, scopes, limits).',
 'Data & fields — custom fields, customer data, units of measure, field governance.',
 'Integrations (plan-gated) — integration hub, API keys, connectors, import/export, onboarding.',
 'Audit Log — a new tenant-facing viewer: the company’s own audit trail (who did what, when).',
 'Operational modules — everything the company’s business type and plan enable (sales, '
 'inventory, the relevant vertical, etc.).',
])

# ════════════════════════════ 7. ROLES & PERMISSIONS ════════════════════════════
h1('7. Roles & Permissions Hierarchy')
p('Two role tiers. Platform roles govern the vendor control center; tenant roles govern a company. '
  'Tenant permissions resolve per company (company-scoped config if present, else global defaults). '
  'A role’s rank is used only to pick the "top role" for navigation display — it does not grant '
  'permissions.')
p('Platform roles & permissions', bold=True)
add_table(['Platform role','Default platform permissions'],
 [['Owner (flag)','All — apex; also bypasses all tenant gates'],
  ['admin','view_companies, create_companies, manage_billing, export_data, manage_users, access_support_tickets, access_audit_logs'],
  ['sales','view_companies, create_companies'],
  ['support','view_companies, access_support_tickets'],
  ['implementation','view_companies, create_companies, export_data, access_support_tickets'],
  ['finance','view_companies, manage_billing, export_data, access_audit_logs']],
 widths=[1.4,5.1])
p('Tenant roles (by rank)', bold=True)
add_table(['Rank','Role','Summary'],
 [['8','admin','Company Owner/Admin — all sales/inventory/purchasing/accounting + every vertical *.manage + settings.users/branches'],
  ['7','manager','Full operational control (no granular fashion split)'],
  ['7','sales_director / national_sales_manager','FMCG commercial leadership'],
  ['6','branch_manager / regional_manager / supervisor / it_admin','Branch ops / scoped management / IT & integrations'],
  ['5','accountant / area_manager / doctor','Finance / scoped commercial / clinical'],
  ['4','warehouse_keeper','Inventory & purchasing'],
  ['3','cashier / technician / stylist','POS & service front-desk'],
  ['2','salesman / driver / receptionist','Field/sales/reception'],
  ['1','staff / housekeeping','Limited operational'],
  ['0','viewer','Read-only']],
 widths=[0.6,2.6,3.3])
p('Permission groups (16): sales, inventory, purchasing, accounting, settings, field_ops, hotel, '
  'clinic, restaurant, salon, pharmacy, market, wholesale, electrical, fashion (plus the fashion.manage '
  'owner umbrella). There are 79 tenant permissions and 7 platform permissions. System roles (21) are '
  'protected from deletion; custom roles are owner-managed.')
p('Representative permissions', bold=True)
add_table(['Group','Example permissions'],
 [['sales','sales.sell, sales.discount, sales.collect, sales.return, customers.manage, customers.approve'],
  ['inventory','inventory.view/adjust/transfer/count, stock_request.create/approve, product.create/import'],
  ['purchasing','purchasing.manage, purchasing.return, suppliers.manage'],
  ['accounting','accounting.view, accounting.post, reports.view, credit.request.approve'],
  ['settings','settings.users, settings.branches, settings.custom_fields, integrations.manage, workflow.manage'],
  ['field_ops','field.sales, visit.override_gps, day.close, target.manage, reconciliation.approve'],
  ['fashion','fashion.manage (umbrella) → fashion.sell/inventory/purchase/installments/cashbox/reports'],
  ['verticals','hotel.manage, clinic.manage/reception/doctor, restaurant.manage, salon.manage, pharmacy.dispense, laundry.manage, market.pos, wholesale.pricing, electrical.rma']],
 widths=[1.3,5.2])

# ════════════════════════════ 8. PLANS / ENTITLEMENTS ════════════════════════════
h1('8. Plans, Modules, Entitlements & Industry Packs')
p('A subscription plan unlocks a set of modules and sets limits (users, branches, products, storage, '
  'trial days). A company’s effective modules are the intersection of its enabled company-modules '
  'and its plan’s modules. Business-type templates seed which modules a new company starts with; '
  'industry packs preselect bundles for a vertical.')
add_table(['Plan','Rank','Max users','Modules unlocked'],
 [['free','0','1','Sales, inventory, and the vertical modules (no accounting/purchasing/analytics/integrations/pos)'],
  ['standard','1','15','+ accounting, purchasing, analytics, integrations, pos'],
  ['pro','2','50','+ workflow, field_ops (full set)'],
  ['unlimited','3','∞','All modules']],
 widths=[1.1,0.7,1.0,3.7])
p('Industry packs: clinic, pharmacy, distribution, retail, electrical, restaurant, hotel, salon, '
  'laundry, fashion. The Plans & Modules editor lets the Platform Owner manage plans, entitlements, and '
  'business-type templates with a live impact preview (which companies gain/lose which modules before '
  'saving). Clothing companies are normalized to a fashion-only module set.')

# ════════════════════════════ 9. MULTI-TENANT ════════════════════════════
h1('9. Multi-Tenant Structure & Tenant Isolation')
p('VANTORA uses a shared-database, shared-schema, row-level-isolation model. Every tenant table carries '
  'a company_id (or scopes via branch_id → company). Row-Level Security policies enforce that a user '
  'reads/writes only rows belonging to their company; the vendor tier (erp_is_platform_owner()) has '
  'cross-tenant access by design.')
p('Isolation was verified across all 129 tables for both reads and writes: a tenant user sees only their '
  'company’s rows, while the platform owner sees all. One historical cross-tenant gap '
  '(erp_supplier_payments) was found and fixed; the table was empty so no data was exposed.')

# ════════════════════════════ 10. DATABASE ════════════════════════════
h1('10. Database Structure Overview')
p('The schema has ~129 erp_* tables. Representative groups:')
add_table(['Area','Key tables'],
 [['Tenancy & identity','erp_companies, erp_branches, erp_profiles, erp_user_branches, erp_departments, erp_teams, erp_job_titles'],
  ['Roles & permissions','erp_roles, erp_role_permissions, erp_company_roles, erp_company_role_permissions'],
  ['Plans & modules','erp_plans, erp_plan_modules, erp_company_modules, erp_business_type_modules, erp_business_type_roles'],
  ['Platform / vendor','erp_platform_staff, erp_platform_staff_permissions, erp_platform_role_permissions'],
  ['Commerce','erp_customers, erp_suppliers, erp_products_catalog, erp_invoices, erp_invoice_lines, erp_sales_orders, erp_sales_returns'],
  ['Inventory','erp_inventory_stock, erp_stock_movements, erp_stock_requests, erp_warehouses'],
  ['Accounting','erp_chart_of_accounts, erp_journal_entries, erp_journal_lines'],
  ['Fashion pack','erp_fashion_styles/variants/colors/sizes, erp_installment_plans/schedule/payments, erp_cash_sessions/movements, erp_expenses'],
  ['Verticals','erp_clinic_*, erp_restaurant_*, erp_salon_*, erp_laundry_*, erp_pharmacy_*, erp_rooms/bookings, erp_wholesale_*'],
  ['Governance','erp_audit_logs, erp_custom_fields, erp_entity_notes/attachments, erp_workflow_instances/tasks, erp_notifications'],
  ['Billing','erp_billing_subscriptions, erp_billing_invoices, erp_billing_payments, erp_billing_plan_prices']],
 widths=[1.7,4.8])

# ════════════════════════════ 11. RLS / SECURITY ════════════════════════════
h1('11. RLS / Security Model')
bullets([
 'Every erp_* table has Row-Level Security enabled with at least one policy (129/129).',
 'Tenant tables: read/write only where company_id = erp_user_company_id() (or via a parent join), '
 'OR erp_is_platform_owner().',
 'Global reference tables (plans, roles, role permissions, business-type templates, clinic reference) '
 'are readable by any authenticated user but writable only by the owner/super-admin.',
 'Sensitive cross-cutting tables are tightly scoped: audit logs (vendor + own-company-admin), billing '
 '(own company for tenants; owner writes), platform staff (owner/self/manage_users).',
 'Apex consistency: the Platform Owner is authoritative on profiles and the global role catalog (RLS '
 'updated for this). Guards (requireModule/requirePermission/capabilities) bypass for owner/super-admin.',
 'Performance hygiene: all 112 unindexed foreign keys were indexed; 6 RLS policies that re-evaluated '
 'auth.uid() per row were fixed to evaluate once per query.',
])

# ════════════════════════════ 12. AUDIT ════════════════════════════
h1('12. Audit Logging System')
p('All sensitive mutations are recorded in erp_audit_logs via the SECURITY DEFINER function '
  'erp_log_audit(action, entity, entity_id, details, company_id), which stamps the actor from the '
  'session (the actor cannot be forged). Columns: actor_id, actor_email, company_id, action, entity, '
  'entity_id, details (JSONB), created_at.')
bullets([
 'Covered actions: create/update/delete, enable/disable, grant/revoke, activate/deactivate, plan '
 'change, role & permission changes, module/plan/business-type changes, user assignment, view-as, etc.',
 'A bilingual describeAuditEvent() renders human-readable sentences.',
 'Vendor forensic viewer at /platform/audit; tenant self-audit viewer at /settings/audit-log '
 '(company-admin, own company only).',
])

# ════════════════════════════ 13. NAVIGATION ════════════════════════════
h1('13. Navigation & UI Structure')
p('Navigation is computed by visibleSections() from the user’s permissions, effective modules, '
  'business type, and tier. A single shared helper (isModuleGateOpen) drives both the sidebar and the '
  'mobile bottom navigation, so gating is consistent.')
bullets([
 'Platform Owner: sees only the grouped vendor panel (Overview, Tenants, Catalog, Billing, Team & '
 'Access, Reference Data).',
 'Tenant users: see only the sections their modules + permissions allow; vertical businesses get a '
 'clean, focused experience (e.g. a clothing store sees only its Fashion store, not generic ERP).',
 'Mobile: a bottom navigation bar with mutually-exclusive groups (e.g. Sell = Fashion POS or generic '
 'Sales, never both).',
 'Home routing: each business type lands on its own home (clothing → /fashion, clinic → role screen, '
 'etc.); the platform owner lands on /platform; no-company users land on /dashboard.',
 'Guards: vendor items never leak to tenants (regression-tested); no dead navigation links '
 '(regression-tested).',
])

# ════════════════════════════ 14. ROUTES ════════════════════════════
h1('14. Existing Pages / Routes')
add_table(['Area','Representative routes'],
 [['Platform','/platform, /platform/activity, /platform/analytics, /platform/audit, /platform/billing, /platform/staff, /platform/companies, /platform/companies/[id] (+ /analytics, /view-as), /platform/plans, /platform/roles, /platform/drugs'],
  ['Dashboard/home','/dashboard, /today, /manager, /territory, /attention, /notifications, /account'],
  ['Sales','/sales/invoices, /sales/pos, /sales/orders, /sales/returns, /sales/pricing, /sales/report, /customers, /customers/[id]'],
  ['Inventory','/products, /inventory, /inventory/low-stock, /inventory/expiry, /inventory/transfers, /warehouses, /suppliers'],
  ['Purchasing/Accounting','/purchases/orders, /purchases/returns, /accounting, /exports'],
  ['Verticals','/fashion (+ sell/products/inventory/customers/installments/suppliers/cashbox/reports), /clinic*, /restaurant*, /salon*, /laundry*, /pharmacy*, /hotel*, /wholesale*, /distribution*'],
  ['Field ops','/rep, /field/route, /field/stock, /sales/journey, /sales/settlement'],
  ['Settings','/settings/staff, /settings/organization, /settings/branches, /settings/users, /settings/permissions, /settings/authz, /settings/audit-log, /settings/integration-hub, /settings/integrations, /settings/import, /settings/export, /settings/custom-fields, /settings/marketplace'],
  ['Upgrade/auth','/upgrade, /login, /register, /setup']],
 widths=[1.6,4.9])

# ════════════════════════════ 15. MIGRATIONS ════════════════════════════
h1('15. Migrations Applied')
p('The schema is managed by sequential SQL migrations (0001–0158). The most recent program (Fashion '
  'pack + Platform Governance + Scalability) is summarized below; all are additive and reversible with '
  'documented rollbacks.')
add_table(['Migration','Purpose'],
 [['0146','Fashion Store pack — tables, RPCs, clothing module/role/permission seeding'],
  ['0147','Fashion clothing role cleanup — curated per-role fashion permissions'],
  ['0148','New-clothing-company trigger also enforces fashion-only modules'],
  ['0149','Platform owner is apex authority on erp_profiles (RLS)'],
  ['0150','Plan storage-limit entitlement (erp_plans.storage_limit_mb)'],
  ['0151','SECURITY: fix cross-tenant leak on erp_supplier_payments (RLS scope)'],
  ['0152','Platform owner may manage the global role catalog (RLS)'],
  ['0153','Tenant Audit Viewer — company-admin read of own audit log (RLS)'],
  ['0154','Integrations becomes a real plan-gated capability + backfill'],
  ['0155','Normalize clothing business-type template to fashion-only'],
  ['0156','Close the POS plan gap — POS grantable + enabled for general retail'],
  ['0157','Scalability: covering indexes for all 112 unindexed foreign keys'],
  ['0158','Scalability: fix 6 per-row auth.uid() RLS policies (initplan)']],
 widths=[1.1,5.4])

# ════════════════════════════ 16. TESTS ════════════════════════════
h1('16. Tests & Verification Status')
bullets([
 '661 unit tests pass; typecheck and production build are green.',
 'Architecture-integrity tests: no dead permissions/modules in nav; catalog completeness; '
 'orphan-module guard (allowlist now empty).',
 'Navigation-routes tests: no dead sidebar links; vendor items never leak to tenants.',
 'Authorization tests: apex-tier bypass; fashion umbrella; home routing (no redirect loops).',
 'Plan-admin / role-admin tests: validation, impact preview, diff, role compare, danger flags.',
 'Schema-health integration test (DB): fails CI if a migration introduces an unindexed FK or a '
 'per-row auth.uid().',
 'Domain integration tests (DB): accounting balance, invoice/payment idempotency, customer scope, etc.',
 'CI on the governance PR applied migrations 0149–0158 cleanly to a staging database.',
])

# ════════════════════════════ 17. SCALABILITY ════════════════════════════
h1('17. Scalability Review Summary')
p('A full capacity review (docs/SCALABILITY-REVIEW.md) covers the multi-tenant model, indexing, '
  'high-volume tables, attachment storage, reporting/analytics patterns, caching, retention, '
  'partitioning, and bottleneck projections.')
add_table(['Scale','Primary bottleneck','Mitigation'],
 [['100 companies','Per-request auth queries','Request memoization (done); transaction pooler'],
  ['1,000 companies','Platform analytics over capped sets; audit/notification growth','Pre-aggregated rollup tables; retention; catalog caching'],
  ['10,000 companies','Total write volume + single-primary ceiling','Monthly range-partitioning; read replica; evaluate sharding']],
 widths=[1.5,2.6,2.4])
p('Applied now: 100% FK index coverage, per-query RLS auth, request-memoized auth context, and a CI '
  'schema-health guard. Deferred (trigger-based, when data justifies): (company_id, created_at) '
  'composite indexes, retention policies, partitioning, analytics rollups, read replica, sharding.')

# ════════════════════════════ 18. PR / BRANCH ════════════════════════════
h1('18. PR / Branch Structure')
p('The work is split into two clean, independently-reviewable pull requests (stacked), preserving all '
  'commit history, migrations, audit trail, and documentation links.')
add_table(['PR','Branch → Base','Scope'],
 [['#121','claude/fashion-store-pack → main','Fashion Store Industry Pack (clothing vertical)'],
  ['#122','claude/platform-governance → claude/fashion-store-pack','Governance, authorization, module architecture, scalability (migrations 0149–0158)']],
 widths=[0.7,3.3,2.5])
p('Land order: review/merge #121, then #122. CI on #122 applied all migrations to staging successfully.')

# ════════════════════════════ 19. TECH DEBT ════════════════════════════
h1('19. Remaining Technical Debt')
add_table(['Item','Severity','Notes'],
 [['Naming divergence (finance↔accounting; market="Supermarket")','Low','Bridged in code; a deliberate reviewed rename pass is recommended (high blast-radius across seeds/tests/i18n).'],
  ['field_ops over-provisioned on many companies','Low','Not a live leak (perm-gated); re-scope as a product decision.'],
  ['Code↔DB role-map divergence (ROLE_PERMISSIONS vs DB)','Low','Runtime uses DB; code map only feeds a hint/tests. Harmless, documented.'],
  ['/platform/copilot-analytics namespace','Low','Shared with company admins but under vendor namespace; move to /insights/*.'],
  ['Multiple permissive RLS policies (advisor)','Low','Minor planner overhead from OR’d SELECT policies; consolidate opportunistically.'],
  ['Trigger-based scalability work','Planned','Composite indexes, retention, partitioning, rollups, replica — when data justifies.']],
 widths=[2.6,0.9,3.0])
p('No correctness or security debt is outstanding: tenant isolation verified, the one real leak fixed, '
  'all sensitive mutations audited, FK coverage complete, and no per-row RLS auth.', italic=True)

# ════════════════════════════ 20. NEXT PHASE ════════════════════════════
h1('20. Recommended Next Development Phase')
add_table(['Phase','Action'],
 [['A. Land & stabilize','Merge #121 → #122 → main; confirm Supabase transaction pooler for serverless.'],
  ['B. Onboarding readiness','Catalog (plans/roles/modules) short-TTL caching; the reviewed naming reconciliation; finalize field_ops/POS-per-vertical defaults; move copilot-analytics out of /platform.'],
  ['C. Scale-readiness (trigger-based)','(company_id, created_at) composite indexes; retention + monthly range-partitioning of high-volume tables; platform-analytics rollup tables; read replica for reporting.'],
  ['D. Customer-facing product','Self-serve plan/upgrade flows on the entitlement model; tenant-facing dashboards & reporting on the rollup foundation.']],
 widths=[1.8,4.7])

doc.add_paragraph()
end = doc.add_paragraph(); end.alignment = WD_ALIGN_PARAGRAPH.CENTER
re = end.add_run('— End of document —'); re.italic = True; re.font.color.rgb = RGBColor(0x88,0x88,0x88)

out = 'VANTORA_Full_Platform_Documentation.docx'
doc.save(out)
print('SAVED', out)
