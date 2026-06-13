#!/usr/bin/env python3
"""VANTORA Pricing & Packaging Proposal -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E); CYAN=RGBColor(0x0E,0x7C,0x86)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.6,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.6); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Pricing & Packaging Proposal"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("FMCG distributors — Small · Medium · Enterprise · suggested SaaS model"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Indicative ranges & rationale (not a quote). Adjust to market + cost-to-serve · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("How to read this")
para("These are SUGGESTED packaging tiers and indicative price ranges to anchor commercial conversations — not a binding "
     "quote. The model is value-based and scales with the customer's field force (the unit that drives both their value and "
     "your cost-to-serve). Currency shown in USD with an EGP guide; localise per market.", bold=True)
para("Pricing levers: (1) a per-tenant platform subscription, (2) a per-active-rep seat fee (the main scaling lever for FMCG), "
     "(3) module/engine add-ons, (4) one-time onboarding. Land with the sell→collect loop; expand into finance + engines.",
     italic=True)

part("1. Packaging tiers")
tbl(["","Small distributor","Medium distributor","Enterprise distributor"],[
 ["Profile","1 branch, 1–10 reps, single warehouse","2–5 branches, 10–50 reps, multi-warehouse","5+ branches, 50+ reps, regional ops"],
 ["Core ERP (sales/inventory/purchasing/accounting)","Included","Included","Included"],
 ["Distribution (routes, van sales, coverage)","Included (basic)","Included (full)","Included (full)"],
 ["Field Force (offline, GPS, journey)","Add-on","Included","Included"],
 ["Approval Queue + governance","Included","Included","Included + custom policies"],
 ["Engines (trade spend, perfect store, alerts, route intel)","—","Choose 1–2","All"],
 ["Integrations / API / e-invoicing","—","Basic","Full + dedicated"],
 ["Support / SLA","Standard (email/WhatsApp)","Priority","Dedicated CSM + SLA"],
 ["Onboarding","Self-serve + guided","Assisted","White-glove + data migration"],
],widths=[1.7,1.6,1.6,1.6],
 fill={(0,1):BLUEBG,(0,2):BLUEBG,(0,3):BLUEBG})

part("2. Suggested SaaS pricing model")
para("Hybrid: platform base + per-rep seat + add-ons. The per-rep seat is the primary scaling lever (mirrors how DMS / "
     "van-sales tools are bought).", bold=True)
tbl(["Tier","Platform base /mo","Per active rep /mo","Engine add-ons /mo","Onboarding (one-time)"],[
 ["Small","$120–250 (≈6–12k EGP)","$8–15","—","$300–800"],
 ["Medium","$400–900 (≈20–45k EGP)","$7–12","$80–200 per engine","$1,500–4,000"],
 ["Enterprise","$1,500–4,000+ (≈75–200k EGP)","$5–10 (volume)","Bundled / custom","$5,000–20,000+"],
],widths=[1.0,1.7,1.3,1.5,1.5])
para("Notes: annual prepay 15–20% discount; per-rep rate decreases with volume (enterprise). Engines (trade spend, perfect "
     "store, route intelligence, alerts) priced as add-ons for Medium, bundled for Enterprise. E-invoicing/integrations "
     "metered or flat per tier.", italic=True, size=9)
h2("Illustrative monthly (rep-driven)")
tbl(["Scenario","Reps","Indicative monthly (USD)","Indicative annual"],[
 ["Small distributor","6","$120 base + 6×$12 = ~$190","~$1,900 (w/ prepay)"],
 ["Medium distributor","25","$600 base + 25×$10 + 2 engines×$150 = ~$1,150","~$11–12k"],
 ["Enterprise distributor","120","$2,500 base + 120×$7 = ~$3,340 (+engines bundled)","~$36–40k"],
],widths=[1.8,0.7,2.5,1.5])

part("3. Why this model")
b("Per-rep seat aligns price with the customer's value (each productive rep generates revenue) and with your cost-to-serve (support/compute scale with reps).",bold=True)
b("Low base + per-rep makes Small easy to say yes to (land), while engines + integrations create natural Medium→Enterprise expansion (expand).")
b("Module/engine add-ons monetise the depth you've already built (trade spend, perfect store, route intel) without discounting the core.")
b("Annual prepay improves cash + retention; onboarding fee covers white-glove setup and qualifies serious buyers.")
h2("Discounting & pilot path")
b("Pilot/founding customers: 1–3 months free or 50% for the first quarter in exchange for a reference + feedback.")
b("Guardrail: never discount the per-rep seat below cost-to-serve; discount the base or onboarding instead.")

part("4. Packaging recommendation")
para("Lead with THREE simple plans (Small / Medium / Enterprise) mapped to the existing Lite / Standard / Enterprise "
     "templates so the product and the price list match 1:1. Sell the sell→collect loop first; let finance depth and engines "
     "be the upsell. Keep the price page to one screen: base + per-rep + a short add-on list.", bold=True, color=NAVY)
b("Small = Lite template (Core ERP + basic distribution).")
b("Medium = Standard template (+ Field Force + 1–2 engines).")
b("Enterprise = Enterprise template (all engines + integrations + SLA).")
para("")
para("Validate the exact numbers against 3–5 real distributor quotes during the pilot, then lock the price list. Treat the "
     "ranges here as the starting anchor.", italic=True, color=AMBER)
out="docs/audits/VANTORA-Pricing-Packaging-Proposal.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
