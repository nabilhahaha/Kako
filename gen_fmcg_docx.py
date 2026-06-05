#!/usr/bin/env python3
"""Generate VANTORA_FMCG_Distribution_Pack_Design.docx"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
n = doc.styles['Normal']; n.font.name='Calibri'; n.font.size=Pt(10.5)
BRAND=RGBColor(0x1F,0x3A,0x5F); ACC=RGBColor(0x0B,0xC5,0xDA); GR=RGBColor(0x66,0x66,0x66)
for i,sz in [(1,17),(2,13.5),(3,11.5)]:
    s=doc.styles[f'Heading {i}']; s.font.color.rgb=BRAND; s.font.size=Pt(sz); s.font.bold=True

def shade(c,h):
    p=c._tc.get_or_add_tcPr(); e=OxmlElement('w:shd'); e.set(qn('w:val'),'clear'); e.set(qn('w:fill'),h); p.append(e)
def table(headers,rows,widths=None,fs=8.7):
    t=doc.add_table(rows=1,cols=len(headers)); t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,h in enumerate(headers):
        cell=t.rows[0].cells[j]; cell.text=''; r=cell.paragraphs[0].add_run(h)
        r.bold=True; r.font.size=Pt(fs); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); shade(cell,'1F3A5F')
    for row in rows:
        cells=t.add_row().cells
        for j,v in enumerate(row):
            cells[j].text=''; rr=cells[j].paragraphs[0].add_run(str(v)); rr.font.size=Pt(fs-0.3)
    if widths:
        for j,w in enumerate(widths):
            for row in t.rows: row.cells[j].width=Inches(w)
    doc.add_paragraph()
def h1(t): doc.add_heading(t,level=1)
def para(t,b=False,it=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=it
def bl(items):
    for i in items: doc.add_paragraph(i,style='List Bullet')
def fig(path,cap,w=6.4):
    doc.add_picture(path,width=Inches(w)); doc.paragraphs[-1].alignment=WD_ALIGN_PARAGRAPH.CENTER
    c=doc.add_paragraph(); c.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=c.add_run(cap); r.italic=True; r.font.size=Pt(8.5); r.font.color.rgb=GR; doc.add_paragraph()

D='docs/diagrams/fmcg'

# ── TITLE ──
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('VANTORA'); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=BRAND
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s.add_run('FMCG Distribution Enterprise Pack'); r.font.size=Pt(17); r.font.color.rgb=ACC; r.bold=True
s2=doc.add_paragraph(); s2.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=s2.add_run('Architecture & Design Document'); r.font.size=Pt(13); r.bold=True
m=doc.add_paragraph(); m.alignment=WD_ALIGN_PARAGRAPH.CENTER
m.add_run('Distribution-first · Saudi Arabia · Route-to-Market / DMS · Design for implementation review').font.size=Pt(10)
m2=doc.add_paragraph(); m2.alignment=WD_ALIGN_PARAGRAPH.CENTER
rr=m2.add_run('Aligned to Nestlé · PepsiCo · Coca-Cola · Mondelez · Unilever · P&G best practices'); rr.italic=True; rr.font.size=Pt(9.5); rr.font.color.rgb=GR
doc.add_page_break()

# ── CONTENTS ──
h1('Contents')
toc=['0. Scope & Reuse','Part II — Enterprise RTM: Distributor Model & DSD',
 'Part II — Executive Dashboards (C-level)',
 '1. Sales Hierarchy (7 roles)','2. Customer Structure','3. Route Execution',
 '4. Perfect Store (RED)','5. Trade Marketing (TPM)','6. Van Sales','7. Distribution KPIs',
 '8. Near-Expiry & Returns','9. Customer Data Governance (KSA)','10. Security & Permissions',
 '11. Database Design','12. Mobile-First Field App','13. Saudi Arabia Localization',
 '14. Enterprise Best-Practice Alignment','15. Governance Compliance','16. Implementation Roadmap',
 '17. Open Decisions & Risks','Appendix — Architecture, ERD & Workflow Diagrams']
for x in toc:
    p=doc.add_paragraph(); p.add_run(x).font.size=Pt(10.5)
doc.add_page_break()

# ── 0 ──
h1('0. Scope & Reuse')
para('A distribution-first FMCG platform (Route-to-Market / Distributor Management System) delivered as '
 'a first-class industry pack on VANTORA — the plan-gated `distribution` module (+ `field_ops` for the '
 'mobile field app) over shared sales, inventory, accounting, pricing, CRM, analytics, and workflow. '
 'Targeted at Saudi Arabia (Arabic-first, VAT 15%, ZATCA e-invoicing, Commercial Registration, National '
 'Address) and aligned to global FMCG best practice.')
fig(f'{D}/fmcg_arch.png','Figure 1 — FMCG pack architecture: pack capabilities over shared VANTORA modules over the platform foundation.', 6.6)
table(['Capability','In VANTORA today','This pack adds'],
 [['Customer master + approval + DFG','yes','channel/sub-channel, classification, GPS, payment terms, frequency, CDG change-requests'],
  ['Routes / regions / scoped roles','partial','route master + territory + PJP + beat + route riding/accompaniment'],
  ['Visits / GPS / day-close / settlement','yes','GPS validation, structured outcomes, Perfect-Store capture'],
  ['Van sales','partial','load sheet, van ledger, mobile sell/return/collect, near/old expiry, reconciled day-close'],
  ['Merchandising / MSL / Perfect-Store / OOS','dashboards','audit capture: OSA, SOS, planogram, competitor, photos'],
  ['Trade marketing','—','NEW: spend, listing fees, displays, promos, claims, ROI/ROTS, activity calendar'],
  ['Near/old expiry & returns','partial','NEW workflows + claims + warehouse processing'],
  ['Mobile-first field apps','responsive web','role-specific, field-optimized, offline-capable apps'],
  ['Saudi localization','bilingual','VAT 15%, ZATCA/Fatoora, CR, VAT no., National Address']],
 widths=[1.9,1.3,3.3])

# ── PART II ──
doc.add_page_break()
h1('Part II — Enterprise Route-to-Market: Distributor Model & DSD')
para('A dual-model RTM platform: the same tenant can run Direct Store Delivery (DSD), the Distributor '
 '(indirect) model, or a Hybrid — across multiple brands, principals, and distributors by region — for '
 'all field-force types (distributor sales force, company sales force, van sales, merchandisers, Key '
 'Accounts).')
fig(f'{D}/fmcg_rtm.png','Figure 2 — Dual operating model: DSD (direct) and Distributor (indirect), with sell-in (primary) vs sell-out (secondary).', 6.7)
para('Company mode (role in the value chain): Manufacturer / Brand-Owner (owns brands → direct and/or '
 'via distributors), Distributor (represents principals → sells to outlets), or Hybrid (both). '
 'Per-region / per-country go-to-market is set independently, so the same company can run direct sales '
 'in one region and distributor sales in another:')
table(['GTM (per region)','Chain','Sell-in (primary)','Sell-out (secondary)','Example'],
 [['Direct (DSD)','Company → Outlet','— (direct)','Company → Outlet (the sale)','Riyadh: own van/sales force; Key Accounts'],
  ['Distributor','Principal → Distributor → Outlet','Principal → Distributor','Distributor → Outlet','Eastern Province: appointed distributor'],
  ['Hybrid','both','both','both','Modern Trade direct + Traditional Trade via distributor']],
 widths=[1.2,1.7,1.3,1.3,1.0], fs=7.7)
para('Principals, Brands & Distributors (multi-brand / multi-principal / multi-distributor):', b=True)
bl(['Principal — a brand owner; the tenant may OWN brands (acts as principal) and/or REPRESENT principals (acts as distributor) → multi-principal.',
 'Brand — belongs to a principal; products map brand → principal → multi-brand. All sales/targets/spend/KPIs slice by brand & principal.',
 'Distributor — a partner with contract, credit, assigned territories/regions/channels, price list, optional linked tenant; multiple distributors by region with exclusivity rules.'])
para('Sell-in vs Sell-out & Secondary Sales (core measurement):', b=True)
bl(['Sell-in (primary): Principal → Distributor — depletion into channel; basis for distributor targets, rebates, trade-spend funding.',
 'Sell-out (secondary): Distributor → Outlet — true market demand; basis for ND/WD, coverage, availability, SOS, Perfect Store.',
 'Stock-in-trade = Σ sell-in − Σ sell-out per distributor/SKU → prevents channel loading; feeds forecasting.',
 'Secondary capture: (a) distributor runs the same pack; (b) secondary-sales ingestion (governed data-share); (c) principal’s own merchandisers capture availability/SOS.'])
para('Distributor management, targets, rebates & trade spend:', b=True)
bl(['Distributor management — onboarding, contracts, price lists, credit, performance scorecards, status (audited).',
 'Distributor targets — primary (sell-in) & secondary (sell-out) per distributor × brand × SKU × period; cascaded; rolled up.',
 'Rebate programs — volume / value / growth / mix / slab / display-linked; accrued on qualifying sales; approval + credit-note settlement; audited.',
 'Trade spend allocated & tracked per distributor (and brand/channel/principal); ROI/ROTS by distributor.',
 'Key Account Management (KAM) — Modern Trade / national chains managed direct (banner→chain→store, JBP, listing, planograms), parallel to distributors for Traditional Trade.'])
para('Multi-tenant linkage & isolation:', b=True)
bl(['Single-tenant — one company runs DSD and/or models distributors as partner entities + ingested secondary data; all rows scoped by company_id (standard RLS).',
 'Connected tenants (enterprise) — principal and distributor are separate VANTORA tenants linked by a consented, scoped, audited data-share over the integrations/sync layer. No cross-tenant RLS is opened; full tenant isolation preserved.'])
para('Distributor inventory, coverage, scorecards & profitability:', b=True)
bl(['Distributor inventory — stock-on-hand per distributor/SKU → days-of-cover, OOS, stock-in-trade.',
 'Distributor coverage — outlets covered ÷ universe in territory; productive coverage.',
 'Distributor scorecard — periodic composite (sell-in vs sell-out, coverage, ND/WD, target %, OTIF/fill-rate, claim & rebate status, Perfect-Store, receivables ageing) → ranked league table.',
 'Distributor profitability — gross-to-net: revenue (sell-in) − COGS − trade spend − rebates − logistics → margin & ROI by distributor/brand/region.'])
para('Strategic planning — JBP, agreements, budgets, growth & market share:', b=True)
bl(['Annual Joint Business Plans (JBP) — per distributor/key account: volume & value targets, growth, trade-spend envelope, activity calendar, agreed KPIs; reviewed quarterly.',
 'Distributor agreements — terms, exclusivity, margins, payment terms, SLAs, period.',
 'Trade-spend budgets — annual by brand/channel/distributor/region, phased to the activity calendar; commitments vs actuals; ROI/ROTS.',
 'Growth targets — YoY by brand/region/distributor. Market-share tracking — share % & trend (panel import or estimated from ND/WD & SOS).',
 'Regional performance — consolidated region/country P&L, coverage, growth, share, target attainment.'])
para('KPIs roll up: outlet (secondary) → distributor → territory → channel → brand → principal → national, with sell-in vs sell-out reconciliation.', it=True)

# ── EXECUTIVE DASHBOARDS ──
h1('Part II — Executive Dashboards (C-level)')
para('Role-scoped, read-only strategic dashboards over the rollups — no live heavy scans; drill-down '
 'respects role scope. Built for enterprise FMCG operating across multiple regions and countries '
 '(region → country → area → territory → route; per-region mode & GTM; per-country localization; '
 'currency-normalized group consolidation).')
fig(f'{D}/fmcg_execdash.png','Figure 3 — C-level executive dashboards fed by the distribution rollup layer.', 6.6)
table(['Persona','Focus KPIs'],
 [['CEO','Revenue & growth (vs target/LY), market share & trend, gross-to-net margin, region/country performance, top risks'],
  ['Commercial Director','Sell-in vs sell-out, distributor scorecards & profitability, JBP attainment, channel & brand mix, receivables'],
  ['Sales Director','Coverage, ND/WD, strike rate, productivity, target attainment by region/area/rep, Perfect-Store'],
  ['Trade Marketing Director','Trade-spend vs budget, ROI/ROTS, promotion & display effectiveness, listing, claims, activity calendar'],
  ['Supply Chain Director','Fill-rate / OTIF, distributor inventory & days-of-cover, stock-in-trade, near/old expiry & write-offs, forecast vs actual']],
 widths=[1.9,4.6])

# ── 1 ──
h1('1. Sales Hierarchy (7 roles)')
para('Reuses VANTORA roles + per-branch reporting lines and region/area/route scope. Targets cascade '
 'down; achievement and approvals roll up. Scope determines data visibility and approval authority.')
fig(f'{D}/fmcg_hierarchy.png','Figure 4 — FMCG sales hierarchy: scope & approval authority.', 6.6)

# ── 2 ──
h1('2. Customer Structure')
para('Extends erp_customers via a 1–1 FMCG profile.')
table(['Attribute','Notes'],
 [['Channel / Sub-channel','Traditional Trade (grocery/kiosk/wholesaler), Modern Trade (hyper/super/mini), HoReCa, Pharmacy, B2B'],
  ['Territory / Route','region → area → territory → route (beat)'],
  ['Classification','outlet grade A/B/C/D, size band, Perfect-Store tier'],
  ['GPS location','validated lat-long + geofence radius'],
  ['Credit limit / Payment terms','credit limit + balance + hold; cash / on-account; terms (days)'],
  ['Visit frequency','F1/F2/F4 + preferred day(s) → PJP'],
  ['KSA compliance','CR number, VAT no. (15-digit), National Address']],
 widths=[1.7,4.8])

# ── 3 ──
h1('3. Route Execution')
bl(['Permanent Journey Plan (PJP): scheduled outlet/date list per rep from route + frequency.',
 'Beat plan: the day’s ordered outlets with time windows.',
 'Route compliance: on-route vs off-route (off-route needs approval).',
 'Coverage = visited ÷ planned outlets; Strike rate = productive ÷ total visits; Call productivity = drop size, lines/SKUs per call.',
 'Route riding & supervisor accompaniment: structured field-coaching evaluation during a ride-along.'])

# ── 4 ──
h1('4. Perfect Store (RED / Picture of Success)')
bl(['MSL compliance — audit vs each outlet’s Must-Stock List.',
 'OSA (On-Shelf Availability) → OOS list.',
 'Shelf Share (SOS) — our facings ÷ category facings.',
 'Display compliance — contracted displays present & correct (photo-backed).',
 'Planogram compliance — shelf layout vs planogram.',
 'Competitor visibility — SKU/price/promo/share-of-shelf.',
 'Photo evidence — before/after shelf photos (object storage).',
 'Perfect-Store score — weighted availability + facings + planogram + display + price.'])

# ── 5 ──
h1('5. Trade Marketing (TPM)')
bl(['Trade spend & budget by customer/channel/promo (commitments + actuals).',
 'Listing fees — per-SKU per-outlet/chain listing payments (Modern Trade).',
 'Displays — types/contracts per outlet + compliance.',
 'Promotions — price-off / BOGO / bundle / slab + eligibility + budget → pricing engine at sale.',
 'Activity calendar — time-phased plan of promos/displays/listings.',
 'Claims — submit (evidence) → validate → tiered approve → settle (credit note) → audited.',
 'ROI / ROTS (Return on Trade Spend) — incremental value vs spend per promo/display/listing.'])

# ── 6 ──
h1('6. Van Sales')
bl(['Van inventory — van warehouse (is_van) stock ledger.',
 'Load sheet — stock issued to the van for a beat day (draft→loaded→reconciled).',
 'Sales — mobile cash/credit invoices vs van stock, priced incl. promotions, ZATCA-compliant.',
 'Collections — on-account vs customer balance + cash session.',
 'Returns — good (restock) vs damaged vs near/old expiry (segregate → claims).',
 'Day close — reconcile van (load ± sales/returns = closing) + cash; variance → supervisor approval.'])

# ── 7 ──
h1('7. Distribution KPIs')
para('Pre-aggregated rollups (no live cross-tenant scans).')
table(['KPI','Definition'],
 [['Numeric Distribution (ND)','outlets stocking SKU ÷ total outlets'],
  ['Weighted Distribution (WD)','ND weighted by outlet sales importance'],
  ['Coverage','visited ÷ planned outlets'],
  ['Productivity / Strike rate','productive visits ÷ total visits'],
  ['SKU per outlet','avg distinct SKUs per productive call'],
  ['Average invoice (drop size)','avg invoice value per productive call'],
  ['MSL compliance / OSA / OOS','must-stock adherence; on-shelf availability; out-of-stock'],
  ['Perfect-Store score','weighted RED score']],
 widths=[2.1,4.4])

# ── 8 ──
h1('8. Near-Expiry & Returns (two workflows)')
bl(['Near-Expiry: batch/expiry tracking → near-expiry threshold → flag → action (sell-through promo / pull-back) before expiry.',
 'Old/Expired: expired stock → segregate (block) → expiry claim → approval → warehouse processing (destroy / return-to-supplier) → credit note + write-off.',
 'Approvals tiered by value (supervisor → area → region); fully audited.'])

# ── 9 ──
h1('9. Customer Data Governance (KSA)')
para('Master-data changes affecting tax/legal/financial integrity are approval-gated change-requests, '
 'built on VANTORA’s customer-approval + Dynamic Field Governance + audit.')
table(['Change request','Why governed','Flow'],
 [['GPS Change Request','outlet location / fencing integrity','field captures new GPS + photo → supervisor/area approve → apply'],
  ['CR Update Request','Commercial Registration (السجل التجاري)','requestor → compliance review → approve → apply'],
  ['VAT Update Request','VAT no. (15-digit) — affects e-invoicing','requestor → finance/compliance approve → apply'],
  ['National Address Update','العنوان الوطني (unified address)','requestor → review → approve → apply']],
 widths=[1.8,2.2,2.5])

# ── 10 ──
h1('10. Security & Permissions')
table(['Group','Permission keys'],
 [['Customers / CDG','customers.manage/approve, customer.classify, credit.request.create/approve, cdg.request.create, cdg.gps/cr/vat/address.approve'],
  ['Route / Journey','route.create/import, journey.create/import, beat.plan, accompaniment.record'],
  ['Van sales','vansales.load, sales.sell/return/collect, day.close, day.approve_close_exception'],
  ['Perfect Store','field.sales, visit.override_gps, visit.approve_out_of_route, audit.capture, assortment.manage, grade.manage, planogram.manage'],
  ['Trade marketing','trade.promo/display/listing/budget.manage, trade.claim.create/approve'],
  ['Near/old expiry','expiry.manage, returns.claim.create/approve, inventory.adjustment.approve'],
  ['Targets / analytics','target.view/manage, reports.view, report.aggregate.view, reconciliation.view/manage/approve'],
  ['Distributor / RTM','principal.manage, brand.manage, distributor.manage, distributor.target.manage, secondary.ingest, rebate.scheme.manage, rebate.accrual.approve, keyaccount.manage'],
  ['Strategy / Executive','exec.dashboard.view, jbp.manage, agreement.manage, growth.target.manage, market.share.manage, distributor.scorecard.view']],
 widths=[1.4,5.1])
para('Role × approval authority (key):', b=True)
table(['Capability','Merch','Sales/Van','Supervisor','Area','Regional','NSM'],
 [['Audit capture','✔','✔','✔','view','view','view'],
  ['Sell/return/collect/day-close submit','—','✔','✔','view','view','view'],
  ['Out-of-route / GPS override / day-close exception','—','—','✔','✔','✔','✔'],
  ['Accompaniment (route riding)','—','—','✔','✔','✔','—'],
  ['GPS change-request approve','—','—','✔','✔','✔','✔'],
  ['CR / VAT / National Address approve','—','—','—','✔','✔','✔'],
  ['Promotions / displays / listing','—','—','—','✔','✔','✔'],
  ['Trade claim approve (tiered)','—','—','≤L1','≤L2','≤L3','all'],
  ['Expiry / write-off approve','—','—','≤L1','≤L2','✔','✔'],
  ['Targets / credit-limit approve','—','—','—','✔','✔','✔']],
 widths=[2.5,0.7,0.85,0.95,0.6,0.7,0.5], fs=7.6)
para('All approvals + sensitive mutations are written to erp_audit_logs via erp_log_audit.', it=True)

# ── 11 ──
h1('11. Database Design')
para('All new tables company_id-scoped; RLS company_id = erp_user_company_id() OR erp_is_platform_owner(); '
 'created_by/updated_by/created_at; audited mutations.')
fig(f'{D}/fmcg_erd.png','Figure 5 — FMCG pack entity-relationship overview.', 6.6)
table(['Table group','Purpose'],
 [['channels / customer_profile / customer_change_requests','channel master, FMCG+KSA attributes, CDG governance'],
  ['journey_plans / journey_stops / beats / accompaniments','PJP, planned visits, day beat, route-riding evals'],
  ['load_sheets / load_sheet_lines / van_reconciliations','van load + day-close variance'],
  ['store_audits / audit_lines / planograms / competitor_obs / visit_photos','Perfect-Store capture + evidence'],
  ['promotions / promo_lines / listing_fees / displays / display_contracts','trade mechanics'],
  ['trade_budgets / trade_spend / activity_calendar / claims / claim_lines','spend + claims lifecycle'],
  ['expiry_batches / return_claims / return_claim_lines','near/old expiry + returns'],
  ['targets / daily_rollup','targets + pre-aggregated KPIs (by brand/principal/distributor/territory)'],
  ['principals / brands / distributors / distributor_territories','RTM master: multi-brand, multi-principal, multi-distributor by region'],
  ['distributor_targets / secondary_sales / channel_stock','sell-in & sell-out targets; secondary transactions; stock-in-trade'],
  ['rebate_schemes / _accruals / _settlements / key_accounts','rebate programs lifecycle; Key Account Management'],
  ['operating_modes / distributor_inventory / distributor_scorecards','company mode + per-region GTM; distributor stock; scorecards'],
  ['distributor_agreements / jbp / growth_targets / market_share','contracts; Joint Business Plans; growth & market share']],
 widths=[3.0,3.5])
para('Index strategy: every FK indexed; composite (company_id, created_at DESC) and (route_id/salesman, '
 'date) on high-volume tables; partial indexes for hot predicates; RLS auth wrapped in (SELECT …). '
 'Scalability: monthly range-partition high-volume tables; analytics via rollups; photos to object storage.')

# ── 12 ──
h1('12. Mobile-First Field App')
para('Field reps work in low-connectivity outlets — mobile-first, fast, offline-capable (queue + sync), '
 'Arabic/RTL, GPS-aware.')
table(['App','User','Core screens'],
 [['Salesman app','Salesman / Van Salesman','My Day/beat; check-in (GPS); order or van-sell; collect; audit; photos; day-close'],
  ['Supervisor app','Supervisor','route riding / accompaniment; approvals (out-of-route, GPS override, day-close exception, CDG GPS); live coverage/strike'],
  ['Merchandiser app','Merchandiser','store audit (MSL/OSA/SOS/planogram); competitor; photos — no selling']],
 widths=[1.4,1.7,3.4])

# ── 13 ──
h1('13. Saudi Arabia Localization')
bl(['VAT 15% standard rate (inclusive/exclusive handling).',
 'ZATCA e-invoicing (Fatoora) Phase 2 — compliant e-invoices, QR, clearance/reporting (via the integrations module; mirrors the existing ETA pattern).',
 'Commercial Registration (CR) + VAT number (15-digit) on customer master, governed via CDG.',
 'National Address (العنوان الوطني): building/street/district/city/postal/additional no.',
 'Arabic-first, RTL; Hijri date optional; currency SAR.'])

# ── 14 ──
h1('14. Enterprise Best-Practice Alignment')
table(['Practice','Source','In this pack'],
 [['Route-to-Market / DMS','all','distribution module; van/pre-sell; routes'],
  ['PJP + beat planning','Nestlé / Unilever','journey plans, beats, coverage'],
  ['Perfect Store / RED / Picture of Success','Coca-Cola / PepsiCo','store audits; MSL/OSA/SOS/planogram; Perfect-Store score'],
  ['Numeric & Weighted Distribution','Nielsen / all','ND/WD KPIs'],
  ['Trade Promotion Management + ROTS','P&G / Mondelez','spend, promos, listing fees, claims, ROI'],
  ['Route riding / field coaching','Unilever / PepsiCo','accompaniments'],
  ['Master-data governance','all','CDG change-requests + approvals + audit']],
 widths=[2.3,1.4,2.8])

# ── 15 ──
h1('15. Governance Compliance')
bl(['Plan-gated distribution + field_ops; distribution/general business-type templates.',
 'Company-scoped RLS on all new tables; sensitive tables tightly scoped; reads/writes verified.',
 'Audit on all approvals + sensitive mutations (erp_log_audit).',
 'Role hierarchy + region/area/route scope reuse platform mechanisms.',
 'Permissions registered (labels + groups + danger flags); managed via Global Roles & Plans editors.',
 'FK index coverage + per-query RLS + rollups/partitioning plan; files in object storage.'])

# ── 16 ──
h1('16. Implementation Roadmap')
table(['Phase','Scope'],
 [['F1 — Foundation','Customer profile/classification/channel/KSA fields; route master; PJP & beats; Customer Data Governance; permissions + RLS + audit'],
  ['F2 — Mobile field & Perfect Store','Salesman/Supervisor/Merchandiser apps; visits + GPS validation; store audits (MSL/OSA/SOS/planogram); competitor; photos; coverage'],
  ['F3 — Van sales','Load sheet; van inventory; mobile sell/return/collect; near-expiry; day-close reconciliation'],
  ['F4 — Trade marketing (TPM)','Promotions (pricing); displays; listing fees; budgets/spend; activity calendar; claims + approvals; ROI/ROTS'],
  ['F5 — Near/old expiry & returns','Batch/expiry; near + expired workflows; claims; warehouse processing + accounting'],
  ['F6 — Distribution analytics','Daily rollups; ND/WD/coverage/strike/SKU-per-outlet/avg-invoice/Perfect-Store; partitioning'],
  ['F7 — KSA compliance & hardening','ZATCA/Fatoora; VAT; approval-threshold matrix; targets cascade; offline-sync hardening; load test; retention']],
 widths=[1.7,4.8])

# ── 17 ──
h1('17. Open Decisions & Risks')
bl(['Pre-sell vs van-sell (or hybrid) default per channel — model supports both.',
 'ZATCA Phase-2 clearance integration scope & certified solution provider.',
 'Trade-spend ↔ accounting posting (accrual vs claim-time) — finance sign-off.',
 'Offline-first sync/conflict strategy for the mobile apps (separate mobile design doc).',
 'Promotion-engine depth (slab/bundle) — extend the existing pricing engine vs a new TPM engine.'])

# ── APPENDIX ──
doc.add_page_break()
h1('Appendix — Workflow Diagrams')
fig(f'{D}/fmcg_workflows.png','Figure 6 — Key FMCG operational workflows (van-sales day cycle, visit & merchandising, trade promotion & claims, near/old expiry, customer data governance).', 6.5)

end=doc.add_paragraph(); end.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=end.add_run('— End of design document —'); r.italic=True; r.font.color.rgb=RGBColor(0x88,0x88,0x88)

out='VANTORA_FMCG_Distribution_Pack_Design.docx'
doc.save(out); print('SAVED', out)
