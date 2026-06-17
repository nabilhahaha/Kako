#!/usr/bin/env python3
"""VANTORA Commercial / Demo Presentation -> .docx (slide-style)"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
CYAN=RGBColor(0x0E,0x7C,0x86)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(11)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def slide(kicker,title):
    d.add_page_break()
    p=d.add_paragraph(); r=p.add_run(kicker.upper()); r.font.color.rgb=CYAN; r.font.size=Pt(10); r.bold=True
    p=d.add_paragraph(); r=p.add_run(title); r.font.color.rgb=NAVY; r.font.size=Pt(22); r.bold=True
    d.add_paragraph()
def para(t,color=DARK,size=11,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=11,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=9.5,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(9.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ── Title slide ──
for _ in range(6): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(52); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("The Business Operating System for FMCG Distribution"); r.font.size=Pt(16); r.font.color.rgb=CYAN; r.bold=True
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("ERP · Distribution · Field Force — one platform, mobile-first, multi-tenant SaaS"); r.font.size=Pt(12); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Commercial & Demo Overview · Arabic + English · June 2026"); r.font.size=Pt(10); r.italic=True; r.font.color.rgb=GREY

slide("The one-liner","What VANTORA is")
para("VANTORA is a single, mobile-first platform that runs an FMCG distributor end-to-end — from the salesman's van on the "
     "street to the general ledger in the back office. One system replaces the usual stack of a desktop ERP + a separate "
     "field-sales app + spreadsheets, with everything tenant-isolated and bilingual (Arabic/English, RTL-ready).", size=12, bold=True)
b("For the rep: sell, collect, and reconcile from a phone — online or offline.")
b("For the manager: real-time coverage, approvals, and KPIs.")
b("For finance: automatic AR/AP, GL posting, and reports — no double entry.")

slide("Why now","The problem we solve")
tbl(["Today's pain","VANTORA"],[
 ["ERP on a desktop + a separate van-sales app that don't reconcile","One system: the van sale IS the invoice IS the GL entry"],
 ["Cash collected in the field, posted days later (or lost)","Collect on the phone; AR updates instantly, idempotent (no double-posting)"],
 ["No real-time view of routes, coverage, or stock on vans","Live routes, coverage, van stock, perfect-store"],
 ["Approvals by phone call / WhatsApp","A single Approval Queue (mobile) for day-close, visits, transfers"],
 ["Spreadsheets for pricing, credit, targets","Pricing engine, credit limits, targets — governed and audited"],
],widths=[3.0,3.4])

slide("Capability 1","ERP capabilities")
b("Sales: invoices, POS / quick-sale, returns, credit notes, ETA e-invoicing readiness.")
b("Inventory: multi-warehouse stock, batches/expiry (FEFO), transfers, counts, valuation (tenant-set costing method).")
b("Purchasing: suppliers, purchase orders, goods receipt, supplier returns & payments.")
b("Accounting: chart of accounts, automatic journals, vouchers, AR/AP, aging, financial reports, exports.")
b("Customers & pricing: full customer master, segments/channels, credit limits + approval workflow, price rules.")
para("Every transaction posts to the GL automatically — the sell-collect loop and finance are one system, not two.", italic=True, color=CYAN)

slide("Capability 2","Distribution capabilities")
b("Routes & territories: route master, customer sequencing, working days, van assignment.")
b("Van sales: load → sell-off-van → collect → reconcile, with stock moving in real time.")
b("Coverage & journey compliance: planned vs visited, out-of-route detection, day-close with exception approval.")
b("Trade spend (foundation): promotions, accruals, claims, ROI — governed and approvable.")
b("Merchandising / perfect store: must-stock lists, surveys, outlet grading, scorecards.")
b("Critical alerts: OOS, credit, collection, route and inventory alerts.")

slide("Capability 3","Field Force capabilities (mobile-first)")
b("The rep's day: a purpose-built 'Today' home — route, visits, sell, collect, close — designed for a phone.")
b("Offline-capable: keep selling and collecting with no signal; sync when back online.")
b("GPS & visit compliance: location-aware visits, out-of-route flags routed to a supervisor.")
b("Van load requests & transfers: request stock to the van; supervisor approves in the queue.")
b("Bottom-nav + 'More' drawer: the core loop is 4 taps; everything else is one tap away.")
para("Role-aware landing: a salesman opens on their day, a supervisor on their approval inbox — no hunting.", italic=True, color=CYAN)

slide("People","Roles")
tbl(["Role","What they do","Opens on"],[
 ["Company Admin","Setup, users, policies, oversight","Dashboard"],
 ["Branch Manager","Branch ops, purchasing, approvals","Manager cockpit"],
 ["Supervisor","Approve day-close / visits / transfers; coverage","Approval Queue"],
 ["Salesman / Van rep","Sell, collect, load, close — the route","Today"],
 ["Warehouse Keeper","Receive, adjust, transfer, approve loads","Inventory requests"],
 ["Accountant","Collections, GL, supplier pay, AR","Collections"],
],widths=[1.6,3.0,1.8])
para("Permissions are granular and enforce separation of duties (a rep can sell & collect but not post the GL or approve "
     "loads). Tenant data is fully isolated by row-level security.", italic=True)

slide("How work flows","Workflows")
b("Sell → collect → auto-post to AR & GL (idempotent — a double-tap never double-charges).")
b("Van load: rep requests → supervisor approves → stock on the van.")
b("Day close: rep closes → coverage exception → supervisor approves.")
b("Out-of-route visit → supervisor approves/rejects with a comment.")
b("Customer / van transfer → request → approve, all in one mobile Approval Queue.")
b("Purchase: PO → receive → stock + AP posted. Credit-limit change → request → approve.")

slide("Business model","SaaS model")
b("Multi-tenant SaaS: one secure platform, each company fully isolated (RLS); spin up a new tenant in minutes.")
b("Subscription tiers by capability (Core ERP → + Distribution → + Field Force / engines), licensed per module.")
b("Per-company feature flags + templates (e.g. Lite / Standard / Enterprise) — powerful backend, simple frontend.")
b("Bilingual (AR/EN), cloud-hosted (Vercel + Supabase Postgres), mobile-first — low IT footprint for the customer.")
b("Pricing levers: per-tenant subscription + per-seat (reps) + module add-ons; land with the sell-collect loop, expand into finance & engines.")

slide("Why us","Competitive advantages")
tbl(["Advantage","Why it matters"],[
 ["One system, not two","The van sale, the invoice and the GL entry are the same record — no reconciliation gap"],
 ["Mobile-first + offline","Reps actually use it on the street; nothing is desktop-only"],
 ["Arabic-first, RTL-native","Built for the region, not translated as an afterthought"],
 ["Governed & auditable","Permissions, approvals, credit, pricing — controlled and logged, validated in pilot"],
 ["Fast to onboard","New tenant + roles + modules seed automatically; role-aware UX from first login"],
 ["Multi-industry core","Same engine also runs pharmacy, retail, clinics, restaurants — one codebase, many packs"],
],widths=[2.2,4.2])

slide("Proof","Pilot-validated — not a prototype")
b("Full platform, role, coverage and workflow audits passed; 6 critical issues fixed and validated on live data.")
b("Separation of duties, tenant isolation, idempotent collections and sequential invoice numbering verified on the live pilot tenant.")
b("Every core FMCG workflow completes end-to-end by the intended role; 1300+ automated tests green; production build green.")
b("A populated pilot tenant (route, sales history, collections, AR, live approvals) is ready to demo today.")
para("")
para("VANTORA is feature-complete and pilot-ready for FMCG distribution — and the same platform extends across industries. "
     "Ready to run a live pilot and convert it into the first paying customers.", size=12, bold=True, color=NAVY)

out="docs/audits/VANTORA-Commercial-Demo-Presentation.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
