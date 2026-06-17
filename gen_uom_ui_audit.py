#!/usr/bin/env python3
"""VANTORA Multi-UoM — UI & Permission Audit / Gap Report -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def bl(t,size=9.3,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=7.8,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(7.9); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Multi-UoM — UI & Permission Audit"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Is every UoM capability exposed (backend · API · UI · nav · permission · mobile)? — gap report before U3"); r.font.size=Pt(10.3); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Read-only · verified against code + live schema/permissions · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Verdict (updated — gap CLOSED)")
para("Most of the IMPLEMENTED multi-UoM surface is correctly exposed end-to-end (backend → API → UI → nav → permission). "
     "The single 'backend-only' configuration gap has now been CLOSED: the product UoM-governance fields are in the FMCG "
     "product form behind uom.manage. The two remaining un-exposed items are intentionally not-yet-built (the U3/U4 "
     "sell/buy flows) plus the U2 line columns (dormant foundation by design).",
     bold=True)
tbl(["Result","Count","Items"],
[
 ["Fully exposed","6","Alt-UoMs (settings/uom) · Base unit (product form) · Product UoM-governance (product form — NOW ADDED) · Price-by-UoM (price-book) · UoM conversion · Enable flag (settings/features)"],
 ["Gap to close before U3","0","— (closed: Default Sell Unit / Purchase Unit / Sell Mode / Allow Fractional added to the product form behind uom.manage)"],
 ["Intentionally pending (U3/U4)","2","Sell-by-UoM picker (van-sell/POS) · Buy-by-UoM (FMCG receiving)"],
 ["Dormant by design (U2)","1","Transaction-line UoM capture columns (no UI until U3)"],
],widths=[1.6,0.6,4.4],
fill={(0,0):GREENBG,(1,0):GREENBG,(2,0):AMBERBG,(3,0):AMBERBG})

part("1 · Six-layer exposure matrix")
tbl(["Capability","Backend","API","UI","Nav","Permission","Mobile","Verdict"],
[
 ["Define alt UoMs (carton/inner/piece + factor + barcode)","erp_product_uoms","list/upsert/deleteProductUom","settings/uom","uom.manage (Settings▸Data)","admin/mgr/branch_mgr/warehouse","More drawer","EXPOSED"],
 ["Base unit (unit → base_uom)","catalog.unit/base_uom","upsertProduct","product form 'unit' field","Products","inventory.* / product.create","Yes","EXPOSED"],
 ["Price by UoM (uom + qty breaks)","erp_prices(uom,min_qty)","upsertPrice","sales/price-book (+preview)","pricing.manage/view","pricing.manage","More drawer","EXPOSED"],
 ["UoM → base conversion","erp_uom_to_base","uomToBaseAction","infra (price-book preview)","—","pricing.view gate","—","EXPOSED (infra)"],
 ["Enable multi-UoM","platform.multi_uom flag","features actions","settings/features","settings.users (admin)","company admin","More drawer","EXPOSED"],
 ["Product UoM governance (purchase/sell UoM, sell_mode, allow_fractional)","catalog columns","upsertProduct (gated)","product form 'Units & Selling' section","Products","uom.manage","Yes","EXPOSED (now)"],
 ["Sell by UoM (picker per line)","backend-ready","pharmacy action only","NONE (even pharmacy POS UI has no picker)","—","—","NO","PENDING U3"],
 ["Buy/receive by UoM","pharmacy only","pharmacy receive","pharmacy receive only","—","—","—","PENDING U4"],
 ["Transaction-line UoM capture","entered_uom/qty/factor cols","lineUomFields (unwired)","NONE (dormant)","—","—","NO","DORMANT (U2)"],
],widths=[1.7,1.0,1.0,1.1,0.95,1.0,0.7,0.85],size=7.0,
fill={(0,7):GREENBG,(1,7):GREENBG,(2,7):GREENBG,(3,7):GREENBG,(4,7):GREENBG,(5,7):GREENBG,(6,7):AMBERBG,(7,7):AMBERBG,(8,7):AMBERBG})

part("2 · The gap — now CLOSED")
h2("Product UoM-governance fields added to the FMCG product form")
para("Previously the standard product form set the base unit ('unit') but NOT purchase_uom / default_sell_uom / sell_mode "
     "/ allow_fractional — those were database-only for FMCG (set only via the pharmacy onboarding wizard). This is now "
     "fixed.", color=GREEN)
h2("What shipped")
bl("A 'Units & Selling' section on the product create/edit form exposing: Default Sell Unit, Purchase Unit, Sell Mode "
   "(base | sales | all) and Allow Fractional. Rendered ONLY for users with uom.manage.")
bl("upsertProduct writes these four columns ONLY when the caller holds uom.manage (otherwise the columns are left "
   "untouched — preserved on edit, DB-default on create), mirroring the U-4 sensitive-field pattern.")
bl("ar/en labels added; the existing columns are reused (no schema change). Demo pilot products were finished with "
   "default-sell/purchase units + sell_mode so U3 has fully-configured products to test.")
para("Why it mattered for U3: sell_mode decides which units a product may be sold in, and allow_fractional governs Kg/Gram "
     "sales — U3's sell-by-UoM picker reads both. The configuration triangle (alternate units in settings/uom, prices in "
     "price-book, governance in the product form) is now complete and UI-exposed.", italic=True, color=GREY, size=9)

part("3 · Intentionally pending / dormant (NOT gaps)")
bl("Sell-by-UoM picker (van-sell / general POS / even pharmacy POS) — this IS U3. The pharmacy CHECKOUT ACTION accepts a "
   "uom and converts, but NO sell UI (pos-fast / pos-terminal / van-sell) currently renders a per-line unit picker. So "
   "sell-by-UoM is backend-ready everywhere and UI-exposed nowhere — exactly the U3 scope.", color=AMBER)
bl("Price-by-UoM is exposed only via the Price-Book screen; the rule-based Sales/Pricing and Wholesale price screens do "
   "NOT carry a uom (base-unit only). Fine for now — price-book is the per-UoM surface — but worth noting for U3 pricing.", color=AMBER)
bl("Buy-by-UoM in FMCG receiving — this IS U4. Pharmacy receive already converts; FMCG PO receiving does not yet.", color=AMBER)
bl("Transaction-line UoM columns (entered_uom/entered_qty/uom_factor) — added in U2 as DORMANT foundation; no UI is "
   "expected until U3 writes them. The shared lineUomFields() helper is ready and unit-tested.", color=AMBER)

part("4 · Permission model — verified")
tbl(["Surface","Permission / gate","Roles that hold it (pilot)"],
[
 ["settings/uom (manage product UoMs)","uom.manage","admin, manager, branch_manager, warehouse_keeper"],
 ["sales/price-book (price by UoM)","pricing.manage / pricing.view","admin, manager, + pricing roles"],
 ["settings/features (enable multi_uom)","company admin (settings.users)","admin"],
 ["Product form (base unit)","product.create / inventory.adjust","admin, manager, branch_manager, warehouse_keeper"],
],widths=[2.3,2.0,2.3])
para("The permission model is coherent: data/inventory managers configure UoMs; pricing roles set per-UoM prices; only "
     "the company admin flips the capability flag. Field roles (salesman) correctly do NOT manage UoM master data. The one "
     "gap (governance fields) would also sit behind uom.manage when added.", italic=True, color=GREY, size=9)

part("5 · Mobile")
para("All UoM CONFIGURATION screens (settings/uom, price-book, settings/features) are reachable on mobile via the sidebar "
     "'More' drawer but are NOT in the bottom-nav — which is correct: these are back-office admin tasks, not field actions. "
     "The field-facing mobile UoM need (a picker when selling) is U3 and is the right place to invest mobile effort.",
     bold=True)

part("6 · Status — cleared for U3")
para("The single configuration gap is CLOSED: product UoM-governance (default sell unit / purchase unit / sell_mode / "
     "allow_fractional) is now in the product form behind uom.manage. Every IMPLEMENTED multi-UoM capability is exposed "
     "across backend → API → UI → nav → permission, with mobile config via the More drawer. The only remaining un-exposed "
     "items are the sell-by-UoM picker (U3) and FMCG buy-by-UoM (U4) — the planned next workstreams. Proceeding to U3.",
     bold=True, color=NAVY)

out="docs/audits/VANTORA-Multi-UoM-UI-Permission-Audit.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
