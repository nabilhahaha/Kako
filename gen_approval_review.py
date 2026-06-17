#!/usr/bin/env python3
"""VANTORA Approval Logic Review -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(8.6); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows,widths=None,size=8.0,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ---- Cover ----
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Approval Logic Review"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Every approval workflow — who requests, who approves, what happens, and is it safe for pilot"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Verified against source + the live pilot tenant's effective permission grants · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---- Executive summary ----
h1("Executive summary (plain English)")
para("VANTORA's approval system works like this: a front-line user (usually a salesman) does something that needs a "
     "second pair of eyes — closing a day with low coverage, visiting a customer off the planned route, moving a customer "
     "to another rep, moving stock between vans, requesting a van load, or asking for a higher credit limit. That action "
     "creates a PENDING request. A senior user (supervisor, branch manager, warehouse keeper, or company admin) then "
     "approves or rejects it. The decision is written to the record, an audit row is logged, and the original action either "
     "completes or is discarded.", bold=True)
para("There are TWO approval surfaces today:")
b("The Field Approval Queue (/approvals/queue) — one mobile-first inbox that covers the five day-to-day FMCG approvals "
  "(day-close, out-of-route visit, customer transfer, van transfer, trade-spend). This is the screen built for the pilot.", color=DARK)
b("The Workflow Inbox / Approval Center (/approvals, /approval-center) — a generic engine for credit-limit, customer "
  "onboarding and change-request approvals. It only appears for users with the workflow.manage permission (in the pilot, "
  "that is the company admin only).", color=DARK)
para("Verdict: the core field approvals are correct and pilot-ready. A few approvals live outside the unified queue "
     "(load requests, van reconciliation, credit-limit) and one (trade-spend) has a view-vs-act permission mismatch — none "
     "are blockers, but they should be tidied so the queue is the single, consistent place approvers look.",
     bold=True, color=GREEN)

# ---- 1. Workflows ----
part("1 · Approval workflows that exist")
para("Eleven genuine approval workflows (request → pending → approve/reject by a different role). Three further operations "
     "— sales-return complete/cancel, purchase-order receipt, inventory-transfer complete/cancel — are status transitions "
     "done by one authorised user, not separate-approver workflows, so they are out of scope here.")
tbl(["#","Workflow","Surface","In unified queue?"],
[
 ["1","Day-Close Exception (coverage below threshold)","Field Queue","Yes"],
 ["2","Out-of-Route / GPS Visit Compliance","Field Queue","Yes"],
 ["3","Customer Transfer (reassign customer to another rep)","Field Queue","Yes"],
 ["4","Van (Stock) Transfer between vans/warehouses","Field Queue","Yes"],
 ["5","Trade-Spend promotion approve/cancel","Field Queue (flag-gated)","Yes*"],
 ["6","Stock / Load Request (van load)","/inventory/requests","No"],
 ["7","Van Reconciliation (stock/cash variance)","/field/van-reconciliation","No"],
 ["8","Credit-Limit Increase Request","Workflow Inbox","No"],
 ["9","Customer Onboarding approval","Workflow Inbox","No"],
 ["10","Customer Change Request (sensitive field edit)","Workflow Inbox","No"],
 ["11","Generic Change Requests (entity edits)","Workflow Inbox","No"],
],widths=[0.4,3.3,1.9,1.0],
fill={(0,3):GREENBG,(1,3):GREENBG,(2,3):GREENBG,(3,3):GREENBG,(4,3):AMBERBG,
      (5,3):AMBERBG,(6,3):AMBERBG,(7,3):AMBERBG,(8,3):AMBERBG,(9,3):AMBERBG,(10,3):AMBERBG})
para("* Trade-spend is gated behind the TRADE_SPEND feature flag (off by default), so it usually does not appear at all.",
     italic=True, color=GREY, size=8.5)

# ---- 2. Per-workflow detail ----
part("2 · Each workflow in detail")
def wf(title, rows):
    h2(title)
    tbl(["Question","Answer"], rows, widths=[1.9,4.6], size=8.6)

wf("1 · Day-Close Exception",[
 ["Who creates","Salesman, when closing a day below the coverage threshold (closeDay → erp_close_day)"],
 ["Created from","/today (end-of-day close)"],
 ["Who approves","Supervisor, Branch Manager, Company Admin"],
 ["Who rejects","No reject path — approve-only (queue returns 'reject_unsupported')"],
 ["Screen","Field Approval Queue"],
 ["After approve","work_session close_status: pending_approval → closed"],
 ["After reject","n/a (stays pending until approved)"],
 ["Audit log","Yes — erp_log_audit('approve_close_day','work_session')"],
 ["Comment","Reason captured at close (close_reason); no comment on approve"],
 ["Permission","day.approve_close_exception"],
])
wf("2 · Out-of-Route Visit Compliance",[
 ["Who creates","Salesman — auto-logged when a visit is off-route / off-GPS (erp_check_in_visit)"],
 ["Created from","Field visit check-in"],
 ["Who approves / rejects","Supervisor, Branch Manager, Company Admin (same permission both)"],
 ["Screen","Field Approval Queue"],
 ["After approve","status pending_approval → approved"],
 ["After reject","status pending_approval → rejected"],
 ["Audit log","Yes — erp_log_audit('decide_compliance','visit_compliance')"],
 ["Comment","Yes — stored in reason"],
 ["Permission","visit.approve_out_of_route"],
])
wf("3 · Customer Transfer",[
 ["Who creates","Salesman / Supervisor (transferCustomer with requireApproval=true)"],
 ["Created from","/customers/transfer"],
 ["Who approves","Supervisor, Branch Manager, Company Admin"],
 ["Who rejects","No reject path — approve-only"],
 ["Screen","Field Approval Queue"],
 ["After approve","status pending → applied (customer reassigned)"],
 ["After reject","n/a (can be cancelled by requester)"],
 ["Audit log","Yes — erp_log_audit('approve_transfer','customer')"],
 ["Comment","Reason captured at request"],
 ["Permission","customer.transfer"],
])
wf("4 · Van (Stock) Transfer",[
 ["Who creates","Salesman / Warehouse keeper (requestVanTransfer)"],
 ["Created from","/inventory/van-transfer"],
 ["Who approves / rejects","Supervisor, Branch Manager, Warehouse Keeper, Company Admin"],
 ["Screen","Field Approval Queue"],
 ["After approve","status pending → completed (stock moves)"],
 ["After reject","status pending → rejected (reason mandatory)"],
 ["Audit log","Yes — request / approve / reject all logged"],
 ["Comment","Yes — reason (mandatory on reject)"],
 ["Auto-approve","If total value < van_transfer_auto_approve_below threshold"],
 ["Permission","stock.transfer.approve"],
])
wf("5 · Trade-Spend (flag-gated)",[
 ["Who creates","Marketing/admin (draft promotion)"],
 ["Who approves / cancels","Whoever holds pricing.rule.edit (typically admin)"],
 ["Screen","Field Approval Queue (only when TRADE_SPEND flag on)"],
 ["After approve","status draft/pending → approved"],
 ["After reject","status → cancelled (reason mandatory)"],
 ["Audit log","Yes — activate / deactivate logged"],
 ["Comment","Yes — reason"],
 ["MISMATCH","Section is shown by reports.view but the action needs pricing.rule.edit (see §4)"],
 ["Permission","view: reports.view · act: pricing.rule.edit"],
])
wf("6 · Stock / Load Request",[
 ["Who creates","Salesman / Driver (createStockRequest, stock_request.create)"],
 ["Created from","/inventory/requests"],
 ["Who approves / rejects","Supervisor, Branch Manager, Warehouse Keeper, Company Admin"],
 ["Screen","/inventory/requests (NOT the unified queue)"],
 ["After approve / reject","status pending → approved | rejected"],
 ["Audit / comment","Via RPC; reject reason via notes"],
 ["Permission","create: stock_request.create · approve: stock_request.approve"],
])
wf("7 · Van Reconciliation",[
 ["Who creates","System/warehouse — computed at day end (erp_compute_van_reconciliation)"],
 ["Who approves / rejects","reconciliation.approve holders (Branch Manager, Admin)"],
 ["Screen","/field/van-reconciliation (NOT the unified queue)"],
 ["After approve / reject","pending_approval → settled | rejected"],
 ["Auto","Variance ≤ threshold auto-drafts (no approval)"],
 ["Permission","settle/manage: reconciliation.manage · approve: reconciliation.approve"],
])
wf("8-11 · Workflow-engine approvals (credit-limit, onboarding, change requests)",[
 ["Who creates","Salesman (credit), system (onboarding/change on sensitive edits)"],
 ["Who approves / rejects","Per workflow step — usually company_admin; credit uses credit.request.approve"],
 ["Screen","Workflow Inbox /approvals + Approval Center (need workflow.manage to SEE the tab)"],
 ["After approve / reject","instance pending → approved | rejected; handler applies/discards the change"],
 ["Audit log","Yes — erp_workflow_decide logs decide on workflow_task"],
 ["Comment","Yes — stored in erp_workflow_tasks.comment"],
 ["Permission","workflow.manage (to see) + the entity permission (e.g. credit.request.approve)"],
])

# ---- 3. Permissions ----
part("3 · Approval permissions (exact strings)")
tbl(["Permission","Governs","Who holds it (pilot tenant)"],
[
 ["day.approve_close_exception","Approve day-close exceptions","Supervisor, Branch Mgr, Admin"],
 ["visit.approve_out_of_route","Approve/reject out-of-route visits","Supervisor, Branch Mgr, Admin"],
 ["customer.transfer","Approve customer transfers","Supervisor, Branch Mgr, Admin"],
 ["stock.transfer.approve","Approve/reject van transfers","Supervisor, Branch Mgr, Warehouse, Admin"],
 ["stock_request.approve","Approve/reject load requests","Supervisor, Branch Mgr, Warehouse, Admin"],
 ["reconciliation.manage","Settle van reconciliation","Supervisor, Branch Mgr, Warehouse, Admin"],
 ["reconciliation.approve","Approve reconciliation variance","Branch Mgr, Admin"],
 ["credit.request.approve","Approve credit-limit increases","Accountant, Branch Mgr, Admin"],
 ["credit.request.create","Raise a credit-limit request","Salesman, Admin"],
 ["stock_request.create","Raise a load request","Salesman, Driver, Admin"],
 ["pricing.rule.edit","Approve/cancel trade-spend","(admin-tier only)"],
 ["workflow.manage","See Workflow Inbox / Approval Center","Admin only"],
 ["reports.view","(used as the trade-spend SECTION gate)","Supervisor, Branch Mgr, Accountant, Viewer, Admin"],
],widths=[2.2,2.3,2.1])

# ---- 4. Consistency ----
part("4 · Is the approval logic consistent?")
h2("All in one queue?")
para("No. The Field Queue holds 5 of the 11 approvals. Load requests and van reconciliation sit on their own screens; "
     "credit-limit, onboarding and change requests sit in the separate Workflow Inbox (admin-only). Approvers therefore "
     "look in up to three places.", color=AMBER)
h2("Approve-only (no reject)?")
para("Day-Close Exception and Customer Transfer are approve-only by design (the queue blocks reject with "
     "'reject_unsupported'). A requester can cancel their own pending transfer, but an approver cannot formally reject — "
     "consider adding an explicit reject for auditability.", color=AMBER)
h2("Automatic approvals?")
b("Van Transfer below the value threshold auto-completes.")
b("Day-Close at/above the coverage threshold closes with no approval.")
b("Customer Transfer with requireApproval=false applies immediately.")
b("Van Reconciliation with variance ≤ threshold auto-drafts.")
para("All four are intentional, threshold-driven, and audit-logged — safe, but worth documenting for pilot operators.",
     italic=True, color=GREY, size=9)
h2("Still backend-only / no unified UI?")
para("Load Request and Van Reconciliation have working screens but are absent from the unified queue. Credit-limit / "
     "onboarding / change-request approvals are only reachable through the Workflow Inbox, whose tab needs workflow.manage "
     "— so in the pilot only the company admin can see them.", color=AMBER)
h2("Two real inconsistencies to fix")
para("A. Trade-spend view-vs-act mismatch.", bold=True, color=RED)
para("The queue shows the trade-spend section (and its approve/reject buttons) to anyone with reports.view, but the action "
     "requires pricing.rule.edit. A supervisor/accountant/viewer would see buttons that always fail 'unauthorized'. "
     "Low impact today (flag is off), but fix the gate to pricing.rule.edit before enabling trade-spend.")
para("B. Credit-limit approvers can't reach the inbox.", bold=True, color=RED)
para("Branch Manager and Accountant hold credit.request.approve, but the Workflow Inbox tab is gated by workflow.manage "
     "(admin only). So they can approve by permission but have no UI to do it. Either grant the inbox to credit approvers "
     "or surface credit-limit requests in the Field Queue.")
h2("Duplicated screens?")
para("No duplicate menu items — the earlier consolidation reduced three approval menu entries to one Approval Queue with "
     "tabs. There are still three underlying list UIs (Field Queue, Workflow Inbox, Approval Center) behind those tabs; "
     "Approval Center largely repeats the inbox and could be retired.", color=AMBER)
h2("Code vs database drift (hygiene)")
para("The app reads permissions from the database (erp_company_role_permissions → erp_role_permissions), NOT the code "
     "ROLE_PERMISSIONS constant. The two have drifted (e.g. the DB grants Branch Manager credit.request.approve and "
     "reconciliation.approve; the code constant does not). Runtime behaviour follows the DB, which is correct for the "
     "pilot — but the code constant should be re-synced so tests and seeds match reality.", color=AMBER)

# ---- 5. Role-by-role ----
part("5 · Role-by-role test (live pilot grants)")
tbl(["Role","Can approve","Cannot / not shown","Sees Approval Queue?"],
[
 ["Salesman","Nothing — requester only (raises day-close, transfers, load & credit requests)","All approvals","No (correct)"],
 ["Supervisor","Day-close, visit, customer transfer, van transfer, load request","Credit-limit (no inbox), trade-spend act","Yes"],
 ["Branch Manager","All supervisor approvals + reconciliation + (credit by permission)","Credit inbox not visible (no workflow.manage)","Yes"],
 ["Warehouse Keeper","Van transfer, load request, reconciliation settle","Day-close, visit, customer transfer (hidden)","Yes"],
 ["Accountant","Credit-limit (by permission only)","Has NO Approval-Queue nav perm → cannot reach any queue UI","No (gap)"],
 ["Company Admin","Everything incl. Workflow Inbox / Approval Center","Trade-spend act (lacks pricing.rule.edit in pilot)","Yes"],
],widths=[1.3,2.5,1.9,0.9],
fill={(0,3):GREENBG,(1,3):GREENBG,(2,3):GREENBG,(3,3):GREENBG,(4,3):REDBG,(5,3):GREENBG})
para("Key gap: the Accountant holds credit.request.approve but none of the five permissions that reveal the Approval "
     "Queue nav item, and lacks workflow.manage for the inbox — so they have an approval right with no screen to use it.",
     bold=True, color=RED)

# ---- 6. Matrix ----
part("6 · Approval Matrix")
tbl(["Workflow","Requester","Approver","Permission","Screen","After approve","After reject","Status"],
[
 ["Day-Close Exception","Salesman","Supervisor / Branch Mgr / Admin","day.approve_close_exception","Field Queue","closed","— (approve-only)","Working"],
 ["Out-of-Route Visit","Salesman (auto)","Supervisor / Branch Mgr / Admin","visit.approve_out_of_route","Field Queue","approved","rejected","Working"],
 ["Customer Transfer","Salesman / Supervisor","Supervisor / Branch Mgr / Admin","customer.transfer","Field Queue","applied","— (approve-only)","Working"],
 ["Van Transfer","Salesman / Warehouse","Supervisor / Branch Mgr / Warehouse / Admin","stock.transfer.approve","Field Queue","completed","rejected","Working"],
 ["Stock / Load Request","Salesman / Driver","Supervisor / Branch Mgr / Warehouse / Admin","stock_request.approve","/inventory/requests","approved","rejected","Partial (not in queue)"],
 ["Van Reconciliation","System","Branch Mgr / Admin","reconciliation.approve","/field/van-reconciliation","settled","rejected","Partial (not in queue)"],
 ["Credit-Limit Request","Salesman","Accountant / Branch Mgr / Admin","credit.request.approve","Workflow Inbox","approved","rejected","Needs UX cleanup"],
 ["Customer Onboarding","System","Company Admin","workflow.manage","Workflow Inbox","approved","rejected","Working (admin)"],
 ["Customer Change Request","System","Company Admin","workflow.manage","Workflow Inbox","applied","discarded","Working (admin)"],
 ["Generic Change Requests","User","Company Admin","workflow.manage","Workflow Inbox","approved","rejected","Working (flag)"],
 ["Trade-Spend","Admin","pricing.rule.edit holder","pricing.rule.edit","Field Queue (flag)","approved","cancelled","Partial (gate mismatch)"],
],widths=[1.25,0.95,1.55,1.35,1.1,0.7,0.75,0.95],size=7.2,
fill={(0,7):GREENBG,(1,7):GREENBG,(2,7):GREENBG,(3,7):GREENBG,(4,7):AMBERBG,(5,7):AMBERBG,(6,7):REDBG,(7,7):GREENBG,(8,7):GREENBG,(9,7):GREENBG,(10,7):AMBERBG})

# ---- 7. Business explanation ----
part("7 · Plain-English business summary")
h2("How the approval system works")
para("Front-line staff raise requests; senior staff approve or reject them; the system records the decision and an audit "
     "trail and then applies or discards the change. The five everyday FMCG approvals are collected in one mobile inbox "
     "(the Approval Queue) so a supervisor can clear them in a few taps. A separate engine handles credit-limit and "
     "customer-data approvals for the company admin.")
h2("Ready for pilot now")
b("Day-Close Exception, Out-of-Route Visit, Customer Transfer, Van Transfer — full request→approve flow, audit, and "
  "(where it makes sense) reject + comment. These are correct and safe.", color=GREEN)
b("Stock / Load Request — works on its own screen with approve/reject; safe to use, just not yet in the unified queue.", color=GREEN)
h2("Needs improvement (not blockers)")
b("Credit-Limit approval — give Branch Manager / Accountant a way to see and action it (today only the admin can).", color=AMBER)
b("Trade-Spend — fix the view-vs-act permission gate before turning the feature on.", color=AMBER)
b("Fold Load Requests + Van Reconciliation into the unified Approval Queue so approvers have ONE place to look.", color=AMBER)
b("Add an explicit Reject (with reason) to Day-Close and Customer Transfer for cleaner audit.", color=AMBER)
b("Re-sync the code permission constant with the database, and retire the duplicate Approval Center tab.", color=AMBER)
h2("Which roles should see the Approval Queue")
para("Supervisor, Branch Manager, Warehouse Keeper and Company Admin — these are the approvers and they correctly see it "
     "today. The Salesman correctly does NOT (requester only). The Accountant SHOULD see it if they are to approve "
     "credit-limit requests — that is the one gap to close.", bold=True, color=NAVY)
h2("Bottom line")
para("The approval logic is correct, permission-safe (every action re-checks its permission server-side, independent of "
     "what the menu shows), and understandable. The field approvals are pilot-ready. The follow-ups are about putting all "
     "approvals in one consistent place and closing two small gate gaps — UX and configuration, not core logic.",
     bold=True, color=GREEN)

out="docs/audits/VANTORA-Approval-Logic-Review.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
