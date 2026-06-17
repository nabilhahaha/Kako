#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
ACCENT=RGBColor(0x1F,0x49,0x7D)
doc=Document(); doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(10.5)
def h(t,l=1):
    p=doc.add_heading(t,level=l)
    for r in p.runs: r.font.color.rgb=ACCENT
    return p
def para(t,b=False,i=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=i; return p
def bullets(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
def table(headers, rows):
    tb=doc.add_table(rows=1, cols=len(headers)); tb.style='Light Grid Accent 1'
    for i,hd in enumerate(headers):
        c=tb.rows[0].cells[i]; c.text=''; c.paragraphs[0].add_run(hd).bold=True
    for row in rows:
        cells=tb.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    return tb
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('FMCG Salesman — Daily Workflow UAT Checklist'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=ACCENT
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=s.add_run('Cleaned 15-permission role model · 2026-06-14'); rs.italic=True; rs.font.size=Pt(11)
para('Validates the complete day of a real FMCG salesman on the cleaned role (no '
     'sales.sell, no customers.manage, no customer.create). No new features — this '
     'pass is to refine the experience before building more.', i=True)
h('Setup')
bullets(['Login: salesman@pilot.test / test.123 (mobile / narrow window).',
         'Demo customers: ★ Demo · Good/Near/Over/Overdue/Cash Only + PILOT-C01…C05.',
         'Bottom-nav for this rep: Home · Today · Sell · Inventory · More (no Customers tab — by design). Hub = /field/van-sales (My Day).'])
h('A. Daily workflow checklist (11 areas)')
table(['#','Area','Where','Expected / Pass'], [
 ['1','Today','bottom-nav Today (/today)','Day flips to open; planned visits/route shown'],
 ['2','Route execution','Hub → Journey (/field/journey) / Route (/field/route)','Visit check-in records GPS + media; sequence advances; End-Day visible'],
 ['3','Customer list','Today/Journey list · Sell/Collect picker','Customers reachable in route + pickers (see Finding F3 — no standalone list)'],
 ['4','Customer Statement','Van-Sell → Statement link (/field/van-sales/statement/[id])','Summary + aging + open invoices + ledger; Collect Now + Print/PDF'],
 ['5','Collect Now','Statement/credit card → Collect (/field/van-sales/collect?customer=)','Outstanding auto-loads; settle; balance drops; receipt'],
 ['6','Van Sale','bottom-nav Sell (/field/van-sales/sell)','Credit card; UoM picker; live totals; Payment step (cash/credit/partial/mixed)'],
 ['7','Invoice Issue','Sell → Issue → confirmation modal → Confirm & Issue','Modal lists customer/lines/UoM/totals/payment; status Paid/Partial/Credit; blocked credit prevented'],
 ['8','Invoice Print','Done → Print Invoice (/print/invoices)','Customer, no./date, lines+UoM, qty, price, discount, VAT, net, paid, remaining, status'],
 ['9','Receipt Print','Done → Print Receipt (/print/receipt)','Collected = paid, Remaining correct, method, signatures (0.00 bug fixed)'],
 ['10','Returns','Hub → Return (/field/van-sales/return)','Return to van; credit note issued + linked; print'],
 ['11','Day Close / Settlement','Hub End Day → /field/van-reconciliation; Journey End Day (day.close)','Reconciliation variance; day closes; settlement consistent'],
])
h('B. Negative checks (the cleanup must hold)')
table(['Check','Expected'], [
 ['Sidebar Sales section (Quick Sale / Orders / Invoices)','Not visible'],
 ['Customers master-data section + Customers bottom-tab','Not visible'],
 ['Open /sales/invoices directly','Redirects to dashboard (page guard)'],
 ['/sales/pos, /sales/orders','Not in nav; gated by sales.sell'],
 ['Collections (/collections)','Visible (rep keeps sales.collect) — acceptable'],
])
h('C. Findings — UX friction / duplication / confusion (refine, not build now)')
table(['ID','Finding','Severity','Recommendation'], [
 ['F1','No Statement entry on the My-Day hub — only reachable inside Van-Sell. The canonical Customer→Statement path isn’t where the day starts.','High','Add a customer-first entry / Statement tile on the hub: pick customer once → Statement → Collect/Sell.'],
 ['F2','Action-first vs customer-first: Sell, Collect, Return each re-pick the customer (up to 3× per visit).','High','One visit/customer context (pick once → Statement·Collect·Sell·Return) — biggest friction removed.'],
 ['F3','No standalone customer list for the rep (customers.manage removed). Reached only via route + pickers.','Medium','Confirm route access is enough; else add a read-only field customer list gated by field.sales.'],
 ['F4','Two End-Day concepts: hub CTA → van-reconciliation (settlement); Journey End Day → work-session close.','Medium','Clarify wording (Settle & reconcile van vs Close visits) and chain them into one finish.'],
 ['F5','Hub tile overload (10), incl. a Coming-soon dead tile (Confirm Load) and two route tiles (Journey + Route).','Medium','Hide coming-soon; merge Journey/Route; permission-filter tiles; order to the daily flow.'],
 ['F6','Inventory bottom-tab → generic /inventory (warehouse), not the rep’s van stock.','Low','Point the rep’s Inventory tab to van stock (/field/stock) or relabel.'],
 ['F7','Collections in two places — hub Collect tile + sidebar Collections menu (kept via sales.collect).','Low','Acceptable; optionally suppress the sidebar link for field-primary roles later.'],
])
h('D. Acceptance')
para('UAT passes when every §A row behaves as expected for the cleaned-role salesman, '
     'every §B negative check holds, and the §C findings are triaged (High first: '
     'F1/F2 — the customer-first context) for the refinement pass before any new '
     'functionality. The end-to-end day — Today → Route → (Customer) Statement → '
     'Collect → Sell → Invoice → Print → Returns → Day Close/Settlement — completes '
     'within the Van Sales workspace with no back-office detours.')
out='docs/audits/FMCG-Salesman-UAT-Checklist.docx'; doc.save(out); print('wrote',out)
