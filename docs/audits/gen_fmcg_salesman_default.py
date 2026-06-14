#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
ACCENT = RGBColor(0x1F,0x49,0x7D)
doc = Document(); doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(10.5)
def h(t,l=1):
    p=doc.add_heading(t,level=l)
    for r in p.runs: r.font.color.rgb=ACCENT
    return p
def para(t,b=False,i=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=i; return p
def bullets(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
def numbered(items):
    for it in items: doc.add_paragraph(it, style='List Number')
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('FMCG Salesman Role — Default + Migration Strategy'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=ACCENT
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=s.add_run('Migration 0307 · applied to staging · 2026-06-14'); rs.i=True; rs.font.size=Pt(11)
para('Goal: New FMCG tenant = correct role model from day one; no manual cleanup. '
     'Additive, reversible, existing tenants unchanged.', i=True)
h('1. What was promoted')
para('The canonical Van-Sales salesman model (15 KEEP permissions) is now the default '
     'template. Removed from the salesman role template (erp_role_permissions): '
     'sales.sell, customers.manage, customer.create. The salesman role is used by the '
     'FMCG business types fmcg / general / wholesale (field-sales role; field.sales is '
     'its core), so this is the FMCG default. Verified: template now has 15 perms.')
h('2. How NEW companies inherit it (automatic)')
para('At company creation, erp_seed_company_roles() (migration 0022) copies '
     'erp_role_permissions → the company own erp_company_role_permissions, filtered to '
     'the business type roles. The template no longer carries the 3 back-office perms, '
     'so every new fmcg/general/wholesale company is seeded with the clean salesman '
     'model (Today → Statement → Collect → Sell → Invoice → Print). No manual cleanup.')
h('3. Why EXISTING companies stay unchanged')
para('Seeding is a snapshot at creation, not a live link. The auth resolver treats a '
     'company own erp_company_role_permissions as authoritative; the template is only a '
     'fallback for a company with NO role config. Verified on staging: all 5 existing '
     'tenants have their own config — none use the fallback — so the template change '
     'touches none of them.')
h('4. Migrating a CHOSEN existing company (explicit opt-in)')
para('Shipped Platform-Owner-only function:')
para('select erp_apply_fmcg_salesman_default(\'<company_id>\');  -- returns rows removed')
para('Removes ONLY the 3 back-office perms from that company salesman role; leaves all '
     'other perms and overrides intact. Reversible. Guarded by erp_is_platform_owner().')
para('Recommended rollout per company:', b=True)
numbered(['Pick the company; snapshot its salesman perms (for revert).',
          'select erp_apply_fmcg_salesman_default(:id);',
          'Re-validate the six field flows for a salesman in that company.',
          'Repeat per company on your schedule — no bulk auto-migration.'])
h('5. Future formal path — role template versioning (optional)')
para('The platform has a versioning system (migration 0226, gated by '
     'KAKO_ROLE_VERSIONING): versioned templates, per-company adopted version, template '
     'edits create a NEW version affecting NEW companies only, existing companies '
     'upgrade explicitly while overrides survive. When enabled, this change becomes '
     'Salesman v2 and per-company adoption is a first-class Upgrade action in the '
     'Platform console, superseding the manual function in §4.')
h('6. Validation checklist (per migrated company)')
para('Salesman (15-perm model) keeps: Van Sales (sell), Collections, Customer '
     'Statement, Returns, Invoice Print, Receipt Print — all gated by field.sales, no '
     'sales.sell dependency. Back-office Quick Sale, Sales Orders, Invoices editor and '
     'Customers-master are hidden / blocked.')
h('7. Rollback')
bullets(['Template (new companies): re-insert the 3 rows into erp_role_permissions for salesman.',
         'A migrated company: re-insert the 3 rows into its erp_company_role_permissions.',
         'Function: drop function erp_apply_fmcg_salesman_default(uuid).',
         'No schema or code change involved — role config only.'])
out='docs/audits/FMCG-Salesman-Default-Migration.docx'; doc.save(out); print('wrote',out)
