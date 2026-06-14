#!/usr/bin/env python3
"""FMCG Pilot Readiness — Refresh (post Collection-in-Sell Phase 1) as a Word doc."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1F, 0x49, 0x7D)
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


def h(t, level=1):
    p = doc.add_heading(t, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = bold
    r.italic = italic
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style='List Number')


def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers))
    tb.style = 'Light Grid Accent 1'
    for i, hd in enumerate(headers):
        c = tb.rows[0].cells[i]
        c.text = ''
        c.paragraphs[0].add_run(hd).bold = True
    for row in rows:
        cells = tb.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return tb


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('FMCG Pilot Readiness — Refresh')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = ACCENT
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('Post Collection-in-Sell Phase 1 · VANTORA FMCG · 2026-06-14')
rs.italic = True
rs.font.size = Pt(11)

para('Re-evaluation of the prior FMCG gap assessment after the sell → invoice → '
     'collect loop and Collection-in-Sell Phase 1 (in-flow payment + full credit '
     'control) landed. No new platform engines proposed — the focus is the '
     'shortest path to a real FMCG pilot company on VANTORA.', italic=True)
para('Prior baseline: FMCG transactional core 95/100 · overall pilot 88/100 '
     '(remaining gap = operational: activation, setup, on-device dry-run, '
     'connectivity).', bold=True)

h('1. Van Selling — status: ✅ Done for pilot')
bullets([
    'One obvious entry: bottom-nav Sell → Van-Sell for van-sales tenants.',
    'Flow: Customer → Products → Review → Payment → Issue, mobile-first.',
    'Multi-UoM sell (Piece / Inner / Carton) with per-line conversion, per-UoM '
    'pricing, base-unit stock invariant (U1–U3; same picker on POS + invoice + '
    'sales-order surfaces).',
    'Server-authoritative pricing, discount cap, van-required, negative-stock '
    'guard, idempotency — atomic in erp_van_sell / erp_van_sell_with_payment.',
    'Post-sale branded receipt + Print / Share / New Sale.',
])
para('Residual: offline van-sell is not wired (currently online-only) — see blockers.', italic=True)

h('2. Collections — status: ✅ Done for pilot (stronger than baseline)')
bullets([
    'Standalone Collect: one receipt across many outstanding invoices '
    '(oldest-first or specified), atomic erp_settle_collection, idempotent, '
    'branded receipt; now auto-loads a deep-linked customer’s outstanding.',
    'Collection-in-Sell Phase 1: take payment before issuing — full cash / credit '
    '/ partial / mixed tenders (cash · card · bank transfer · cheque), reusing the '
    'exact erp_collections posting model (reports/statements/reconciliation '
    'unchanged).',
    'Credit control, server-enforced (the real-world gap NOT closed before): no '
    'overpayment; cash-only must be fully paid; credit limit (unpaid ≤ available); '
    'credit-days / overdue block; Near-limit warning. Validated by a 19/19 matrix '
    '+ the 5,000/4,900/1,000 edge case + invariants on staging.',
    'Blocked customers get status + reason + debt snapshot (outstanding, overdue, '
    'open invoices, oldest age) and a one-tap Collect Now → Collection.',
])

h('3. Remaining FMCG pilot blockers')
para('The transactional loop is no longer the blocker. What remains:')
table(['#', 'Item', 'Severity (real pilot)', 'Notes'], [
    ['B1', 'Offline van-sell / collect (queue + replay)', 'High if weak signal; Low if online-first',
     'Infra exists (offline-sync, idempotency keys = the Phase-6 seam). Not yet wired. #1 real-world gap now.'],
    ['B2', 'Manager day-2 reporting — AR aging, collections by route/rep, van stock & valuation, route KPIs', 'Medium',
     'Assemble from existing data (no new engine). Light version suffices for a pilot.'],
    ['B3', 'Operational activation — real-company setup, enable flags, on-device supervised dry-run', 'Medium (process)',
     'Already packaged (setup wizard, reference-tenant SQL, Readiness Diagnostic, dry-run script).'],
    ['B4', 'Connectivity decision', 'Decision',
     'Online-first pilot is viable today; offline (B1) is the first enhancement after.'],
    ['B5', 'Replenishment in UoM (buy/receive, U4)', 'Low for time-boxed pilot',
     'Pre-stock works; UoM purchasing is additive and deferred.'],
])
para('No item requires a new platform engine.', bold=True)

h('4. Revised pilot readiness')
table(['Track', 'Prior', 'Now', 'Why'], [
    ['FMCG transactional + commercial core', '95', '98',
     'Loop complete + in-flow payment + server-enforced credit control (governance gap closed)'],
    ['Overall pilot — online-first', '88', '92',
     'Engineering essentially done; remaining is activation + light reporting + connectivity decision'],
    ['Overall pilot — offline-required', '—', '~85',
     'Offline van-sell/collect (B1) not yet built'],
])
para('Verdict: GO for a controlled, online-first real FMCG pilot after activation '
     '+ one supervised on-device dry-run. Rollback is one switch (KAKO_VAN_SALES '
     'off, or the per-tenant platform.collect_in_sell / platform.multi_uom toggles).',
     bold=True)

h('5. Minimum remaining work before a real FMCG pilot (online-first)')
para('Shortest path — all reuse, no new engines:')
numbered([
    'Provision the real company (1–2 days): branches, warehouses + vans '
    'assigned/stocked, SKUs with base UoM + factors + price > 0, customers with '
    'credit limit + terms, routes, return reasons (setup wizard + reference-tenant '
    'SQL pattern).',
    'Enable flags for the company: van_sales (+ settings), platform.multi_uom, '
    'platform.collect_in_sell.',
    'Readiness Diagnostic = READY, 0 blockers (/field/van-sales/readiness).',
    'One supervised on-device dry-run of the full loop, including a credit-blocked '
    'customer → Collect Now → settle and a mixed-tender sale.',
    'Light manager reporting (B2): AR aging + collections + van stock from existing '
    'queries.',
    'Rep training (½ day): the Payment step, credit statuses, Collect-Now.',
])
para('Next highest-priority FMCG gaps after the loop (ranked; additive / reuse-only):', bold=True)
numbered([
    'Offline van-sell + collect (queue + replay on the existing offline-sync + '
    'idempotency seam) — biggest real-world lever.',
    'FMCG reporting pack — AR aging, collection performance, van stock/valuation, '
    'route/rep KPIs (from existing tables).',
    'Supervisor credit-override — route through the EXISTING approvals engine (not '
    'a new engine) for one-off limit/overdue overrides.',
    'Buy/receive in UoM (U4) then Returns/Transfers in UoM (U5) — additive, '
    'flag-gated, for steady-state replenishment.',
])
para('Bottom line: the sell → invoice → collect loop with credit governance is '
     'pilot-ready; the shortest path to a live FMCG pilot is operational activation '
     '+ one dry-run on an online-first route, with offline as the first post-pilot '
     'investment.', bold=True)

out = 'docs/audits/FMCG-Readiness-Refresh.docx'
doc.save(out)
print('wrote', out)
