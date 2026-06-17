#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
ACCENT=RGBColor(0x1F,0x49,0x7D)
doc=Document(); doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(10.5)
def h(t,l=1):
    p=doc.add_heading(t,level=l)
    for r in p.runs: r.font.color.rgb=ACCENT
    return p
def para(t,b=False,i=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=i; return p
def bullets(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
def table(headers, rows):
    tb=doc.add_table(rows=1, cols=len(headers)); tb.style='Light Grid Accent 1'
    for i,hd in enumerate(headers):
        c=tb.rows[0].cells[i]; c.text=''; c.paragraphs[0].add_run(hd).bold=True
    for row in rows:
        cells=tb.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    return tb
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('Route Execution → Visit Context — Workflow Review'); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=ACCENT
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=s.add_run('Make the route visit-driven · review only (no implementation) · 2026-06-14'); rs.italic=True; rs.font.size=Pt(11)
para('Should Route Execution become the primary entry into the customer visit '
     'context, so picking the next stop opens the Statement hub with Collect / Sell '
     '/ Return / Print? Recommendation: yes — pure navigation wiring, no new '
     'module/engine. Review only.', i=True)
h('1. Current state — why it feels module-driven')
table(['Piece','Today','Connects to selling?'], [
 ['Route Execution (/field/journey)','Ordered stop list, GPS check-in, coverage KPI, photo, call, blocked-visit reason','No — a stop only offers Check-in · Photo · Call'],
 ['Visit context (/field/van-sales/statement/[id])','Statement + aging + open invoices + ledger, with Collect · Sell · Return · Print','Yes — but reached only via the Customer tile, not the route'],
])
para('The day splits into two disconnected halves: route = compliance (check-in), '
     'visit = selling (statement hub). After check-in nothing carries the rep into '
     'Statement/Collect/Sell — they go back to the hub and re-enter the customer. '
     'That is the "separate module" feel.')
h('2. The visit-driven model (target)')
para('Today → Route → [tap next stop] → Visit Context (Statement hub) → Collect · '
     'Sell · Return · Print → Done → Next customer → back to Route (next stop). One '
     'continuous loop: Route → Customer Visit → Statement → Collect → Sell → Return '
     '→ Print → Next Customer.')
h('3. Recommendation')
para('Make Route Execution the primary entry into the visit context; keep the '
     'Customer tile as the off-route / ad-hoc entry. Both open the SAME visit '
     'context. Concretely (all reuse — destinations already exist):')
bullets([
 'Each route stop → "Open visit" deep-links to /field/van-sales/statement/[customer_id]. Check-in (GPS) becomes the act of starting the visit — one tap records compliance and lands on the statement hub.',
 'Visit context → "Next customer" returns to /field/journey and surfaces the next unvisited stop (highlight, not forced order).',
 'Route banner on the visit context when arrived from the route (Stop 3/12 · Next: …).',
 'Compliance stays intact — GPS check-in / coverage KPI / blocked reason still fire, as a side-effect of starting the visit.',
])
h('4. Why this is low-risk (reuse, no new module)')
bullets([
 'Visit context already built (statement + Collect/Sell/Return/Print).',
 'Sell / Collect / Return already accept ?customer=; statement is /[id].',
 'Route/journey already has the stop list + check-in + GPS.',
 'Change = navigation wiring + a slim route banner + a Next link. No engine, no transaction change; additive + reversible.',
])
h('5. Open design questions (decide before building)')
table(['Question','Options','Lean'], [
 ['Check-in vs open visit','(a) one tap = check-in AND open visit; (b) check-in first, separate Open visit','(a) minimal friction'],
 ['Off-route / walk-in customers','Customer tile opens visit context without a route check-in','Keep the Customer tile as the escape hatch'],
 ['Next customer behaviour','(a) auto-advance to next sequence; (b) return to route, highlight next','(b) routes are not strictly sequential'],
 ['Mandatory check-in to sell?','Enforce GPS check-in before Sell, or advisory','Advisory for the pilot'],
])
h('6. Side benefit — resolves the parked F3')
para('A route-primary model plus the Customer tile cover on-route (planned) and '
     'off-route (ad-hoc) customer access — exactly the F3 gap (no standalone customer '
     'list). Adopting this closes F3 without a separate screen.')
h('7. Suggested phasing (design only)')
bullets([
 'Phase 1 (small, reuse-only): route stop → Open visit deep-link; visit context → Next customer back to route. Delivers the visit-driven loop.',
 'Phase 2 (polish): one-tap check-in+open; route banner (Stop x/y · Next); auto-highlight next stop; optional per-visit summary.',
])
para('Bottom line: Route Execution should be the primary entry into the visit '
     'context — no new module, only wiring the route stops into the existing '
     'Statement hub and a Next-customer return, turning two modules (compliance + '
     'selling) into one visit-driven loop: Route → Customer → Statement → Collect → '
     'Sell → Return → Print → Next.', b=True)
out='docs/audits/FMCG-Route-to-Visit-Context-Review.docx'; doc.save(out); print('wrote',out)
