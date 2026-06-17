#!/usr/bin/env python3
"""Generate the Collection-in-Sell Phase 1 deliverables as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1F, 0x49, 0x7D)
GREEN = RGBColor(0x1E, 0x7A, 0x34)
RED = RGBColor(0xB0, 0x2A, 0x2A)

doc = Document()
styles = doc.styles
styles['Normal'].font.name = 'Calibri'
styles['Normal'].font.size = Pt(10.5)


def h(text, level=1):
    p = doc.add_heading(text, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    if size:
        r.font.size = Pt(size)
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def table(headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    for i, hd in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = ''
        run = c.paragraphs[0].add_run(hd)
        run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Collection-in-Sell — Phase 1')
r.bold = True
r.font.size = Pt(22)
r.font.color.rgb = ACCENT
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('Payment before invoice issuance — Build & Validation report')
rs.italic = True
rs.font.size = Pt(12)
meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('VANTORA FMCG · Van-Sales · 2026-06-14').font.size = Pt(9)

para('Status: BUILT · server-validated on staging · flag-gated (default OFF). '
     'Enabled for the VANTORA Pilot FMCG (DEMO) tenant only for UAT. Additive & '
     'reversible — erp_van_sell, erp_settle_collection and the standalone Collect '
     'screen are untouched.', italic=True)

# ── 1. Architecture ──────────────────────────────────────────────────────────
h('1. Architecture summary')
para('Flow: Customer → Products → Review → Payment → Issue Invoice. The Payment '
     'step appears only when the flag is ON; otherwise the flow is the original '
     'Review → Issue.')
para('Feature flag:', bold=True)
bullets([
    'platform.collect_in_sell (platform pack, domain "pos"). Default OFF — no '
    'template seeds it. collectInSellEnabled(flags) gates the server action, the '
    'page, and the UI step.',
])
para('Atomic RPC — erp_van_sell_with_payment (migration 0306):', bold=True)
bullets([
    'A faithful superset of erp_van_sell (auth, branch, idempotency, customer, '
    'van, discount cap, UoM, stock guard, invoice + issue) plus tender posting.',
    'Tenders (p_tenders = [{method, amount, reference}]) become standard '
    'erp_collections rows (one per method) allocated to the new invoice via '
    'erp_collection_allocations — the EXACT existing posting model, so every '
    'report / statement / reconciliation keeps working unchanged.',
    'Order of operations (key design point): tenders are applied BEFORE '
    'erp_issue_invoice, lowering the customer balance first, so issue’s own credit '
    'check evaluates the true post-payment exposure (balance + unpaid), not the '
    'full net. Invoice status is then set paid / partially_paid from paid_amount.',
    'Idempotent via the invoice idempotency_key (a repeat returns the existing '
    'invoice + its current paid_amount/status without re-issuing or re-charging).',
    'Returns invoice_id, invoice_number, net_amount, paid_amount, status.',
])
para('Tender methods (DB-canonical erp_collections.method codes): cash, '
     'credit_card (Card), bank_transfer, check (Cheque). Cheque / transfer require '
     'a reference.')

para('Credit control (server-enforced; salesman cannot override — Phase 1):', bold=True)
table(['Rule', 'Logic'], [
    ['No overpayment', 'Σ tenders ≤ net (payment_exceeds_total)'],
    ['Cash-only (limit = 0)', 'any unpaid remainder ⇒ over_credit (must be fully paid)'],
    ['Credit limit (limit > 0)', 'unpaid ≤ available = credit_limit − balance; else over_credit'],
    ['Credit days / overdue', 'credit control on + terms set + oldest unpaid invoice age > payment_terms_days ⇒ customer_overdue_blocked for any unpaid remainder; a fully-paid (cash) sale is still allowed'],
])

para('Credit status badge (customer selection + Payment step):', bold=True)
bullets([
    'Good — within credit limit and payment terms.',
    'Near credit limit — available < 10% of limit (warning only, non-blocking).',
    'Over credit limit — balance ≥ limit.',
    'Overdue — oldest unpaid invoice older than allowed credit days.',
    'Cash only — credit limit = 0.',
])
para('Each surfaces a plain reason and the figures behind it (limit, balance, '
     'available, exceeded-by, oldest unpaid age, allowed days). Blocked standings '
     '(over limit / overdue / cash only) show “Credit sales are blocked. '
     'Collection or full cash sale only.”, a debt snapshot (outstanding, overdue '
     'amount, open invoices, oldest invoice age), and a one-tap “Collect Now” that '
     'opens the Collection workflow for the customer (outstanding invoices '
     'auto-load).')

para('Permissions:', bold=True)
bullets([
    'field.sales to sell; sales.collect to enter tenders (a rep without it gets a '
    'credit-only Payment step).',
    'No new permission introduced; no master-data grants. Supervisor override is '
    'explicitly out of Phase-1 scope.',
])

# ── 2. Validation ────────────────────────────────────────────────────────────
h('2. Staging validation report')
para('Project rsjvgehvastmawzwnqcs (staging). All RPC scenarios run as the pilot '
     'salesman inside rolled-back transactions (no pilot data changed). Unit base '
     'price: 1 unit = 79.80 net (70.00 + 14% VAT); 2 units = 159.60.')

para('Payment scenarios', bold=True)
table(['Scenario', 'Tenders', 'Result'], [
    ['Full cash (1u)', 'cash 79.80', 'PASS — paid, paid 79.80'],
    ['Full credit (2u)', '—', 'PASS — issued, paid 0'],
    ['Partial (2u)', 'cash 60', 'PASS — partially_paid, paid 60'],
    ['Mixed (2u)', 'cash 100 + card 59.60', 'PASS — paid, paid 159.60'],
])

para('Credit-control matrix (19/19 PASS)', bold=True)
table(['Scenario', 'Full Cash', 'Full Credit', 'Partial'], [
    ['A · Good', 'Allowed', 'Allowed', 'Allowed'],
    ['B · Near limit (avail < 10%)', 'Allowed', 'Allowed', 'Allowed (warning only)'],
    ['C · Over limit (balance ≥ limit)', 'Allowed', 'Blocked (over_credit)', 'Blocked (over_credit)'],
    ['D · Overdue (age > terms)', 'Allowed', 'Blocked (customer_overdue_blocked)', 'Blocked (customer_overdue_blocked)'],
    ['E · Cash-only (limit 0)', 'Allowed', 'Blocked (over_credit)', 'Blocked (over_credit)'],
])
para('F · Mixed tender (remaining still validated):', bold=True)
table(['Combo', 'Result'], [
    ['cash + credit (good customer)', 'Allowed (partially_paid)'],
    ['cash + transfer + credit (good)', 'Allowed (partially_paid)'],
    ['cash + card + credit (good)', 'Allowed (partially_paid)'],
    ['mixed where remaining > available', 'Blocked (over_credit)'],
])

para('Real-world edge case — limit 5,000 · balance 4,900 · invoice 1,000', bold=True)
para('Available credit = 100. Status = Near credit limit (100 < 10% of 5,000 = 500).')
table(['Test', 'Remaining', 'Expected', 'Result'], [
    ['1. Full credit (paid 0)', '1,000 > 100', 'Blocked', 'PASS — over_credit'],
    ['2a. Partial 900', '100 = avail 100', 'Allowed', 'PASS — partially_paid; final AR 5,000'],
    ['2b. Partial 500', '500 > 100', 'Blocked', 'PASS — over_credit'],
    ['3. Full cash 1,000', '0', 'Allowed', 'PASS — paid; final AR 4,900 (unchanged)'],
])

para('Invariants (partial sale 2u cash 60 + idempotent replay)', bold=True)
table(['Check', 'Expected', 'Got'], [
    ['net / paid / status', '159.60 / 60 / partially_paid', 'PASS'],
    ['stock decrement (base units)', '2', 'PASS — 2 (no double-decrement on replay)'],
    ['customer AR balance delta', '+99.60 (net − paid)', 'PASS — 99.60'],
    ['collection allocation rows', '1', 'PASS — 1'],
    ['invoices for idempotency key', '1', 'PASS — 1'],
    ['idempotent replay → same invoice', 'true', 'PASS — true'],
])
para('Pure-core unit tests: 1386 passed (payment math, credit-limit examples 1–4, '
     'overdue/credit-days, near-limit, status classification, standing-blocked). '
     'tsc clean; build green.')
para('Bugs found & fixed during validation: (1) a non-matching idempotency '
     'SELECT…INTO nulled v_paid → reset before the tender loop; (2) erp_issue_invoice '
     're-checked the full net → tenders now applied before issue; (3) tender method '
     'codes aligned to the erp_collections.method CHECK constraint (credit_card/check).',
     italic=True)

# ── 3. UAT ───────────────────────────────────────────────────────────────────
h('3. UAT guide (per role)')
para('Enabled on the pilot tenant; host = current preview / staging. Demo password test.123.')
para('Salesman (field.sales + sales.collect):', bold=True)
bullets([
    'Bottom-nav Sell → pick a customer; note the credit status badge + reason.',
    'Add products + units; Review; tap Payment.',
    'Try each: Pay full · Cash → Paid; Credit (pay later) → Credit; a partial cash '
    'amount → Partially paid; Add payment twice for a mixed tender (cash + card / '
    'transfer / cheque — reference required for the last two).',
    'Confirm the live status chip, remaining, new balance, and the credit panel.',
    'On an over-limit / overdue / cash-only customer, confirm a remaining balance '
    'disables Issue Invoice with the “Collection only” warning + debt snapshot, '
    'while a full-cash sale still issues; tap Collect Now → outstanding invoices.',
    'Issue → receipt shows the status badge + paid amount; Print / Share / New sale.',
])
para('Supervisor:', bold=True)
para('Review issued invoices + collections in the day’s activity; confirm '
     'blocked-credit customers were not sold to on credit; confirm the standalone '
     'Collect screen still works for debt recovery on those customers.')
para('Accountant:', bold=True)
para('Verify AR — invoice paid_amount/status, the per-tender erp_collections rows '
     '(method + reference), allocations against the invoice, and that the customer '
     'balance moved by exactly net − paid. Confirm cash vs card/transfer/cheque '
     'land in their existing treatments and reconciliation is unaffected.')
para('Company Admin:', bold=True)
para('Toggle platform.collect_in_sell in Company Settings → Features and confirm '
     'the Payment step appears/disappears accordingly (the rest of the sell flow '
     'is unchanged when OFF).')
para('Acceptance: every matrix cell behaves as in §2; AR/stock invariants hold; '
     'standalone Collect remains available for blocked customers.', bold=True)

# ── 4. Rollback ──────────────────────────────────────────────────────────────
h('4. Rollback procedure')
para('All rollbacks are instant and non-destructive (no schema change, no data transform):')
bullets([
    'Per tenant (fastest): Company Settings → Features → turn Collect-in-Sell OFF, '
    "or: update erp_feature_flags set enabled=false where company_id=:co and "
    "feature_key='platform.collect_in_sell'. The sell flow reverts to Review → "
    'Issue; the standalone Collect screen stays.',
    'Platform-wide: leave the flag unset for all tenants (default OFF) — nothing '
    'calls the new RPC.',
    'Remove the engine (optional): DROP FUNCTION '
    'public.erp_van_sell_with_payment(uuid,uuid,jsonb,jsonb,uuid,date,text). No '
    'table/column changes to reverse; invoices and collections already written by '
    'it remain valid (standard model).',
    'Code revert: the UI Payment step, server action, and pure helpers are all '
    'additive; reverting the commit restores prior behaviour with no migration '
    'down-step.',
])
para('Out of Phase-1 scope (future): supervisor credit-override approval workflow, '
     'on-account overpayment/change, offline tender capture.', italic=True)

out = 'docs/audits/Collection-in-Sell-Phase1.docx'
doc.save(out)
print('wrote', out)
