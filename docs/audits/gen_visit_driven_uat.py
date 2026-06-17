#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
ACCENT=RGBColor(0x1F,0x49,0x7D)
doc=Document(); doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(10.5)
def h(t,l=1):
    p=doc.add_heading(t,level=l)
    for r in p.runs: r.font.color.rgb=ACCENT
    return p
def para(t,b=False,i=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=i; return p
def table(headers, rows):
    tb=doc.add_table(rows=1, cols=len(headers)); tb.style='Light Grid Accent 1'
    for i,hd in enumerate(headers):
        c=tb.rows[0].cells[i]; c.text=''; c.paragraphs[0].add_run(hd).bold=True
    for row in rows:
        cells=tb.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    return tb
def bullets(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('Visit-Driven Route (Phase 1) — UAT Scenario'); r.bold=True; r.font.size=Pt(19); r.font.color.rgb=ACCENT
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=s.add_run('Full route-driven day: first customer → end-of-day settlement · 2026-06-14'); rs.italic=True; rs.font.size=Pt(11)
para('Flag platform.visit_driven_route is ON for the pilot. Login salesman@pilot.test '
     '/ test.123 (mobile). Transactional behaviour unchanged — this validates the '
     'navigation loop and the Complete-Visit guard.', i=True)
h('Pre-conditions')
bullets(['Day not yet started (or start it from Today).',
         'A planned route with stops (use the ★ Demo customers to exercise credit states).',
         'Flags ON: multi_uom, collect_in_sell, visit_driven_route; Van Sales active.'])
h('Scenario A — one full visit (the loop)')
table(['#','Step','Expected'], [
 ['1','Today → Start day','Day status = open'],
 ['2','Hub → Route (Journey)','Each stop shows Open visit (primary) + Check in (secondary)'],
 ['3','Tap Open visit on Stop 1','GPS check-in runs (stop visited) AND lands on the visit context (statement)'],
 ['4','Visit context','Route banner "Stop 1 of N · Next: <name>"; statement + actions Collect · Sell · Return · Print'],
 ['5','Statement review','Credit status, available credit, overdue, open invoices correct'],
 ['6','Collect → settle','Collection posts; balance drops; back on the visit context'],
 ['7','Sell → UoM → Payment → Issue → Print','Invoice issued; status correct; Invoice/Receipt prints correct'],
 ['8','Return (optional) → reason → submit','Credit note issued + linked'],
 ['9','Complete Visit','Returns to the route with the next stop highlighted (?focus=<nextId>)'],
 ['10','Repeat 3–9 for next stops','Coverage KPI rises per visited stop'],
 ['11','End Day & settle van → van reconciliation','Variance shown; day closes; settlement consistent'],
])
para('Pass: Route → Customer → Statement → Collect → Sell → Return → Print → '
     'Complete Visit → Next, with no back-office menu and no re-selecting the customer.', b=True)
h('Scenario B — Complete-Visit guard (the new safeguard)')
table(['#','Step','Expected'], [
 ['1','Sell → add product → do NOT issue → Back to statement','Sale left unfinished'],
 ['2','Tap Complete Visit','Blocked — dialog: unfinished sale; Keep working / Discard & complete visit'],
 ['3','Keep working','Dialog closes; stays on the visit context'],
 ['4','Finish the sale (Issue) → Complete Visit','Proceeds to the route (flag cleared on issue)'],
 ['5','Repeat with Collect and Return started-not-finished','Same guard; Discard & complete explicitly clears + completes'],
])
para('Pass: Complete Visit never silently closes a visit while a sale / collection '
     '/ return is open — finish or explicitly discard.', b=True)
h('Scenario C — flag OFF regression')
bullets(['Turn Visit-Driven Route OFF (or test another tenant).',
         'Route stops show only Check in / Photo / Call (no Open visit); statement has no route banner / Complete Visit.',
         'Pass: identical to today — no regression.'])
h('Scenario D — edge cases')
bullets(['Already-visited stop → Open visit opens the context directly (no second check-in).',
         'Off-route / walk-in → Hub Customer tile → same visit context (no banner, no check-in).',
         'Offline → Open visit queues the check-in and stays on the route; no crash.',
         'Blocked customer (★ Demo Over/Overdue/Cash) → credit block + Collect Now as before; full-cash still issues.'])
h('Acceptance')
para('Scenario A completes as one visit-driven loop; Scenario B blocks accidental '
     'completion and requires complete-or-discard; Scenario C confirms the flag gates '
     'the behaviour; Scenario D edge cases behave gracefully. No transaction, schema '
     'or engine change — rollback is the flag OFF.')
out='docs/audits/FMCG-Visit-Driven-Route-UAT.docx'; doc.save(out); print('wrote',out)
