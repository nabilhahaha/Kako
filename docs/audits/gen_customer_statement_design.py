#!/usr/bin/env python3
"""Customer Account Statement — Design (FMCG) as a Word document."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1F, 0x49, 0x7D)
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


def h(t, level=1):
    p = doc.add_heading(t, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.italic = italic
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style='List Number')


def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = 'Light Grid Accent 1'
    for i, hd in enumerate(headers):
        c = tb.rows[0].cells[i]; c.text = ''; c.paragraphs[0].add_run(hd).bold = True
    for row in rows:
        cells = tb.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return tb


title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('Customer Account Statement — Design'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = ACCENT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('FMCG · design-first · reuse-only (no new engine) · 2026-06-14'); rs.italic = True; rs.font.size = Pt(11)

para('Design-first. No code; no new platform engine. Everything reuses existing '
     'tables, helpers, components, print templates and the Collect-Now deep link. '
     'Additive and flag-safe.', italic=True)

h('1. Goal')
para('A full customer account statement (not a flat report): summary header, '
     'running debit/credit/balance ledger, open invoices, AR aging, print/PDF and '
     'a Collect Now action — with role-specific UX for Salesman, Supervisor and '
     'Company Admin, built on existing invoices, collections, allocations, customer '
     'balances and the existing aging logic.')

h('2. The one data-sourcing rule (critical)')
para('The current customer page builds ledger credits from erp_payments — empty '
     'for FMCG/van-sales, where payments live in erp_collections + '
     'erp_collection_allocations (same root cause as the print-receipt bug just '
     'fixed). The statement sources the ledger from both, authoritatively:')
table(['Ledger side', 'Source', 'Amount'], [
    ['Debit (owed)', 'erp_invoices (status ≠ draft/cancelled)', 'net_amount'],
    ['Credit (paid)', 'erp_collections (+ allocations for invoice linkage)', 'amount / applied_amount'],
    ['Credit (legacy)', 'erp_payments', 'amount'],
    ['Credit (returns)', 'credit notes', 'amount'],
])
para('Running balance = Σ debits − Σ credits by date. The closing balance must '
     'reconcile to erp_customers.balance (a built-in self-check surfaced to admins).')

h('3. Server builder (one helper, reused by all roles + print)')
para('A single read-only function (query + existing aging math): '
     'loadCustomerStatement(customerId, {from?, to?}) → { customer, summary '
     '(creditLimit, currentBalance, availableCredit, overdueAmount, openInvoiceCount, '
     'oldestInvoiceDays), aging (current/d30/d60/d90/d90p), openInvoices[], ledger '
     '(StatementEntry[]), openingBalance? }.')
para('Reuses: bucketFor() + bucket labels (accounting/aging); availableCreditFor / '
     'creditStatusOf / overdueDays (van-sales/sell.ts); loadCustomerOutstanding '
     '(collect-server.ts); StatementEntry (statement-table.tsx); '
     'INVOICE_STATUS_LABELS / PAYMENT_METHOD_LABELS (constants.ts). No writes; '
     'RLS-scoped (branch isolation automatic).')

h('4. Sections (the statement screen)')
numbered([
    'Customer summary — name/code/phone, status badge, and the credit panel already '
    'built: Credit limit · Current balance · Available credit · Overdue amount '
    '(+ the Good/Near/Over/Overdue/Cash status + reason from CreditStandingCard).',
    'Aging buckets — Current / 1-30 / 31-60 / 61-90 / 90+ (reuse accounting.aging.* '
    'labels + bucketFor). Each bucket clickable to filter open invoices.',
    'Open invoices — number, date, due date, net, paid, outstanding, status, days '
    'overdue. Oldest-first (matches the collect flow).',
    'Running statement (ledger) — the reusable StatementTable (date · reference · '
    'description · debit · credit · running balance). Optional date-range with '
    'opening balance carried forward.',
    'Actions — Print / PDF and Collect Now.',
])

h('5. Print / PDF (reuse, no PDF engine)')
bullets([
    'Reuse /print/statement/[id] (bilingual, running balance, BrandLogo, '
    'PrintButton) — enhanced to render the FULL statement (summary + aging + open '
    'invoices + ledger) from the SAME builder, so screen and print never diverge. '
    'Fix its credit source to collections (per §2).',
    'PDF = browser “Save as PDF” from that print route (existing pattern; no new '
    'dependency/engine). Optional ?from=&to= for a period statement.',
])

h('6. Collect Now (reuse the deep link)')
para('Collect Now routes to /field/van-sales/collect?customer=<id> (auto-loads '
     'outstanding). For desktop accounting users without a van, it routes to '
     '/collections preselected to the customer. Shown when the user holds '
     'sales.collect and outstanding > 0; on blocked customers it is the primary CTA.')

h('7. Role-specific UX (one builder, three entry points / scopes)')
table(['Aspect', 'Salesman (field.sales)', 'Supervisor', 'Company Admin / Accountant (accounting.view)'], [
    ['Entry point', 'Van-Sell customer card / Collect → Statement; customer row in Today/route',
     'Branch AR list / aging → drill into customer', 'Accounting → Aging / Customers → drill; or /customers/[id]'],
    ['Scope', 'His route/branch only (RLS)', 'His branch(es)', 'All branches'],
    ['Layout', 'Mobile-first single column; big Collect Now', 'Mobile/desktop; “who’s overdue” lens',
     'Desktop dense; period filter, opening balance, Export/PDF, reconciliation check'],
    ['Primary action', 'Collect Now, Print receipt', 'Monitor + Collect Now (if sales.collect)',
     'Print/PDF statement, period statements'],
    ['Write actions', 'none (read + collect)', 'none (read + collect)', 'none from the statement'],
])
para('All three render the SAME component with role props (canCollect, scope, '
     'showPeriodFilter, showReconCheck) — no divergent logic.')

h('8. Permissions & routing')
bullets([
    'New additive route /customers/[id]/statement (or a tab on the customer page). '
    'Gate: customers.manage OR accounting.view OR field.sales. Branch isolation via RLS.',
    'Collect Now gated by sales.collect. Print/PDF available to all who can view. '
    'No new permission string required.',
    'Optional flag platform.customer_statement (default ON) only if a kill-switch is '
    'wanted; otherwise a plain additive screen.',
])

h('9. Reuse map (nothing built from scratch)')
table(['Need', 'Reused asset'], [
    ['Ledger table + running balance', 'src/components/statement-table.tsx'],
    ['Statement print + PDF', 'src/app/(print)/print/statement/[id] (+ BrandLogo, PrintButton)'],
    ['Aging buckets + math', 'src/app/(app)/accounting/aging/page.tsx (bucketFor)'],
    ['Credit summary + status + reason', 'CreditStandingCard / creditStatusOf / availableCreditFor'],
    ['Outstanding invoices', 'loadCustomerOutstanding (collect-server.ts)'],
    ['Collections + allocations', 'erp_collections, erp_collection_allocations, supabase-gateway'],
    ['Collect Now', '/field/van-sales/collect?customer= · /collections'],
    ['Status/method labels', 'INVOICE_STATUS_LABELS, PAYMENT_METHOD_LABELS (constants.ts)'],
    ['i18n', 'customers.stmt*, accounting.aging.*, vanSales.collect.* (+ a few new keys)'],
])

h('10. Build plan (when approved — additive, reversible)')
numbered([
    'loadCustomerStatement builder (read-only; invoices ⊕ collections ⊕ credit-notes '
    '⊕ legacy payments → summary + aging + openInvoices + ledger). Also repoint the '
    'existing customer-statement page to it (fixes the collections-missing gap).',
    'CustomerStatement component (sections §4) with role props; reuse StatementTable '
    '+ CreditStandingCard.',
    'Routes/entry points per role (§7) + Collect Now / Print-PDF actions.',
    'Enhance /print/statement/[id] to the full statement via the same builder.',
    'i18n (new keys, ar/en parity); unit-test pure parts (running balance, aging '
    'bucketing, available-credit, reconciliation check).',
    'Staging validation on CRDEMO-* customers — the ledger closing balance must '
    'equal erp_customers.balance.',
])

h('11. Recommendation')
para('Build it as one read-only statement builder + one component, three role entry '
     'points, and fix the existing statement’s credit source to collections at the '
     'same time (it currently under-reports FMCG payments). This is the highest-value '
     '“light reporting” item from the readiness refresh (B2) and needs no new engine '
     '— only reuse of invoices, collections, allocations, aging and the '
     'print/Collect-Now infrastructure already in place.')

out = 'docs/audits/Customer-Statement-Design.docx'
doc.save(out)
print('wrote', out)
