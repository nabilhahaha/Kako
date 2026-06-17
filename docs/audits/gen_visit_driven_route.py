#!/usr/bin/env python3
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
ACCENT=RGBColor(0x1F,0x49,0x7D)
doc=Document(); doc.styles['Normal'].font.name='Calibri'; doc.styles['Normal'].font.size=Pt(10.5)
def h(t,l=1):
    p=doc.add_heading(t,level=l)
    for r in p.runs: r.font.color.rgb=ACCENT
    return p
def para(t,b=False,i=False):
    p=doc.add_paragraph(); r=p.add_run(t); r.bold=b; r.italic=i; return p
def bullets(items):
    for it in items: doc.add_paragraph(it, style='List Bullet')
def table(headers, rows):
    tb=doc.add_table(rows=1, cols=len(headers)); tb.style='Light Grid Accent 1'
    for i,hd in enumerate(headers):
        c=tb.rows[0].cells[i]; c.text=''; c.paragraphs[0].add_run(hd).bold=True
    for row in rows:
        cells=tb.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    return tb
t=doc.add_paragraph(); t.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=t.add_run('Visit-Driven Route — Design Proposal'); r.bold=True; r.font.size=Pt(20); r.font.color.rgb=ACCENT
s=doc.add_paragraph(); s.alignment=WD_ALIGN_PARAGRAPH.CENTER
rs=s.add_run('Route becomes the primary operational workflow · design-first (no implementation) · 2026-06-14'); rs.italic=True; rs.font.size=Pt(10.5)
para('Visit lifecycle: Route Stop → Check-in → Customer Statement → Collect → Sell '
     '→ Return → Print → Complete Visit → Next Customer. Reuse-only (no new '
     'module/engine). Additive, reversible, flag-safe.', i=True)
h('1. Current workflow')
para('Route (/field/journey): stop list; per stop Check-in (GPS) · Photo · Call · '
     'blocked-reason — and it DEAD-ENDS at "Checked in". The visit context '
     '(/field/van-sales/statement/[id]) is complete (Statement + Collect/Sell/Return/'
     'Print) but is reached only via the Customer tile, not the route. Result: route '
     '= compliance, visit = selling — two modules, one broken seam.')
h('2. Proposed workflow')
para('A single visit-session loop. Route is the spine; tapping a stop runs the '
     'existing GPS check-in AND opens the visit context; Complete Visit marks the '
     'stop visited and returns to the route with the next stop active.')
table(['State','Meaning','Set by'], [
 ['planned','on route, not started','route plan'],
 ['in_progress','checked-in, visit context open','Check-in (existing GPS RPC)'],
 ['completed','visit finished','Complete Visit (= existing visited mark)'],
 ['blocked','could not visit','existing blocked-reason flow'],
])
bullets([
 'Check-in = start visit (one tap: GPS compliance + lands on the visit context).',
 'Inside: Statement first, then Collect / Sell / Return / Print — already customer-scoped.',
 'Complete Visit marks completed (existing visited/KPI) + optional outcome; Next Customer → route, next stop.',
 'Off-route / walk-in via the Customer tile → same visit context (no check-in) — also closes parked F3.',
])
h('3. Reused components')
para('Reused (the bulk — no logic change):', b=True)
table(['Need','Reused asset'], [
 ['Route stops, sequence, coverage KPI','field/journey/journey-screen.tsx'],
 ['GPS check-in + blocked reason + offline queue','checkInVisit + journey flow'],
 ['Visit context (Statement + Collect/Sell/Return/Print)','/field/van-sales/statement/[id] + CustomerStatementView (F1/F2)'],
 ['Statement data (reconciling)','loadCustomerStatement'],
 ['Collect / Sell / Return scoped to customer','existing screens (?customer=)'],
 ['Invoice / receipt / statement print','existing /print/* templates'],
 ['Off-route customer pick','/field/van-sales/customers (F1)'],
])
para('New (small, additive — UI/navigation only):', b=True)
bullets([
 '"Open visit" on each route stop → deep-link to the visit context (check-in first when not visited).',
 'Route banner on the visit context (Stop x/y · Next: …), shown only when arrived from the route (?from=route).',
 '"Complete Visit" + "Next Customer" control → mark visited (existing) + return to route, next stop active.',
 'Visit-session position carried in the URL (?from=route&seq=N) — no new table.',
 'Phase 2 optional: visit outcome (sold / collected / no-order) attached to the existing visit/work-session record.',
])
h('4. Risks')
table(['Risk','Mitigation'], [
 ['Offline — check-in queues offline but the visit context needs data','Keep check-in offline-capable (unchanged); visit context degrades gracefully (online-first today); Phase-2 cache the day’s statements. Never block check-in on the context.'],
 ['Compliance gating — should selling require check-in?','Phase 1 advisory; check-in-before-sell becomes a later policy toggle if needed.'],
 ['GPS denied / far from customer','Reuse existing blocked-reason / distance handling; Open visit still works (recorded per current geofence rules).'],
 ['Partial / resumed visits','in_progress = checked-in, not completed; re-tap reopens the same context; documents persist independently.'],
 ['Double-counting / KPI integrity','Complete Visit reuses the existing visited set / coverage logic — no new counter.'],
 ['Don’t break non-van field reps','Gate Open visit / banner behind Van Sales active so merchandising-only routes are unchanged.'],
 ['Per-stop data load','Statement is one indexed read per customer, on demand (not pre-loaded for all stops).'],
 ['Navigation depth / back button','Define back semantics: context Back → route; Next → route (replace, not push); avoid deep stacks.'],
 ['Blocked customers mid-visit','Credit block + Collect-Now already handle this inside the context — no new path.'],
])
h('5. Migration path')
para('Additive, reversible, flag-gated (platform.visit_driven_route, default OFF). '
     'OFF = today’s behaviour exactly (no regression).')
para('Phase 1 — the loop (small, reuse-only):', b=True)
bullets([
 'Open visit on each route stop (check-in → visit context, ?from=route&seq=N).',
 'Route banner + Complete Visit / Next Customer on the visit context when from=route.',
 'Behind the flag; gated by Van Sales active; existing screens untouched when off.',
 'Validate on the pilot: full route — Stop → Check-in → Statement → Collect → Sell → Return → Print → Complete → Next — confirm coverage KPI, GPS, documents + prints, and the statement reconciles.',
])
para('Phase 2 — polish:', b=True)
bullets([
 'One-tap check-in+open; auto-highlight next stop; per-visit outcome + end-of-visit summary; optional check-in-before-sell policy.',
])
para('Rollout: enable for the pilot first; then promote to the FMCG default (same '
     'template/flag mechanism as the salesman role model), existing tenants opting in '
     'explicitly. Rollback: flag OFF → instant revert; no schema change in Phase 1; '
     'no transaction logic touched.', b=True)
para('Conclusion: turns the route from a compliance screen into the primary '
     'operational workflow — a visit-driven loop built almost entirely from existing '
     'components. Small, UI-only new surface; risk contained by a flag + Van-Sales '
     'gating; clean, reversible, pilot-first migration.', b=True)
out='docs/audits/FMCG-Visit-Driven-Route-Design.docx'; doc.save(out); print('wrote',out)
