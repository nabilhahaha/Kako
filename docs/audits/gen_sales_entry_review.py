#!/usr/bin/env python3
"""FMCG Sales-Entry Point Review — inventory, classification, recommendations (Word)."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ACCENT = RGBColor(0x1F, 0x49, 0x7D)
doc = Document()
doc.styles['Normal'].font.name = 'Calibri'
doc.styles['Normal'].font.size = Pt(10.5)


def h(t, level=1):
    p = doc.add_heading(t, level=level)
    for r in p.runs:
        r.font.color.rgb = ACCENT
    return p


def para(t, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(t); r.bold = bold; r.italic = italic
    return p


def bullets(items):
    for it in items:
        doc.add_paragraph(it, style='List Bullet')


def numbered(items):
    for it in items:
        doc.add_paragraph(it, style='List Number')


def table(headers, rows):
    tb = doc.add_table(rows=1, cols=len(headers)); tb.style = 'Light Grid Accent 1'
    for i, hd in enumerate(headers):
        c = tb.rows[0].cells[i]; c.text = ''; c.paragraphs[0].add_run(hd).bold = True
    for row in rows:
        cells = tb.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = str(v)
    return tb


title = doc.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run('FMCG Sales-Entry Point Review'); r.bold = True; r.font.size = Pt(22); r.font.color.rgb = ACCENT
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
rs = sub.add_run('Inventory · classification · recommendations · 2026-06-14'); rs.italic = True; rs.font.size = Pt(11)

para('Review only. Nothing removed or hidden. Establishes the Van Sales workspace '
     'as the canonical FMCG salesman experience and a safe, reversible path to one '
     'obvious selling flow: Customer → Statement → Collect → Sell → Invoice → Print.',
     italic=True)

h('1. Inventory of all sales-entry points (classified)')
table(['Route', 'Classification', 'Gating (perm · module · flag)'], [
    ['/field/van-sales (+ sell/collect/return/statement/…)', 'Field Sales (Van Sales) — CANONICAL', 'field.sales · van_sales · Van Sales active'],
    ['/sales/invoices', 'Back Office Invoicing', 'sales.sell | sales.collect · sales'],
    ['/sales/orders', 'Back Office Invoicing', 'sales.sell · sales_orders'],
    ['/sales/returns', 'Back Office Invoicing', 'sales.return · returns'],
    ['/collections', 'Back Office Invoicing', 'sales.collect · sales'],
    ['/customers/[id] (statement)', 'Back Office Invoicing', 'customers nav · RLS'],
    ['/accounting/aging · /sales/settlement · /distribution/credit-requests', 'Back Office Invoicing', 'accounting / field.sales / credit.request.*'],
    ['/sales/pos (Quick Sale)', 'POS / Quick Sale', 'sales.sell · pos'],
    ['/pharmacy/pos · /market/pos · /fashion/sell', 'POS / Quick Sale (vertical)', 'vertical perm · vertical module'],
    ['/wholesale/order', 'Back Office (vertical)', 'wholesale.pricing · wholesale'],
    ['/restaurant/orders · /salon/tickets · /laundry/orders · /clinic/visits', 'Vertical POS / Order', 'vertical perms · vertical module'],
    ['/print/receipt · /invoices · /statement · /collection · /credit-note', 'Shared output (all flows)', 'auth · RLS'],
])
para('Legacy / Duplicate (for an FMCG van-sales tenant ONLY): /sales/invoices, '
     '/sales/pos, /sales/orders, /collections, /customers/[id] are not legacy '
     'globally (back office + other verticals need them) but are redundant for the '
     'FMCG salesman — the Van Sales workspace already covers sell, collect, '
     'statement, invoice and print. Hide by role for field reps; do not delete.', bold=True)

h('2. Duplicated workflows (FMCG salesman)')
table(['Job', 'Canonical (Van Sales)', 'Duplicate(s) reachable today'], [
    ['Create a sale', '/field/van-sales/sell', '/sales/invoices, /sales/pos, /sales/orders'],
    ['Collect cash', '/field/van-sales/collect', '/collections'],
    ['Customer statement', '/field/van-sales/statement/[id]', '/customers/[id]'],
    ['Returns', '/field/van-sales/return', '/sales/returns'],
])
para('Root cause (verified on the pilot): the salesman role holds sales.sell AND '
     'customers.manage (role-default and company-override). Those light up the whole '
     'back-office Sales section AND Customers master-data in the sidebar — three '
     'ways to sell plus a master-data editor on top of the Van Sales workspace. This '
     'also contradicts the guardrail "do not give the salesman master-data '
     'permissions."', bold=True)

h('3. Redundant menu entries (sidebar) for the FMCG salesman')
bullets([
    'Sales section: Quick Sale, Sales Orders, Invoices, Collections, Sales Returns.',
    'Customers master-data (create/edit) section.',
    'The bottom-nav already collapses Sell → Van-Sell (good); the sidebar / command '
    'palette still expose the back-office equivalents.',
])

h('4. What can be hidden — by role, entitlement, or flag (all reversible)')
numbered([
    'By role permission (recommended): tighten the pilot salesman to field.sales + '
    'sales.collect (+ customers.view for statements); remove sales.sell + '
    'customers.manage. Effect: back-office Sales + Customers-master sidebar sections '
    'disappear for the rep; Van Sales workspace + Sell bottom-tab (gated by '
    'field.sales) remain; selling still works. Reversible; company-override scope — '
    'other tenants unaffected.',
    'By entitlement (module): if unused, disable pos + sales_orders modules for the '
    'tenant → Quick Sale + Sales Orders nav vanish company-wide. Keep sales (office '
    'invoices/collections).',
    'By feature flag / nav rule (optional, broader): a tenant flag (e.g. '
    'fmcg.field_primary) that suppresses the generic Sales sidebar section for field '
    'roles on van_sales tenants. Most product-ized but largest change; option #1 '
    'meets the pilot goal today with no code.',
])
para('Office roles keep everything: Accountant / Company Admin retain /sales/*, '
     '/collections, /customers, /accounting/aging — Back Office Invoicing is a '
     'legitimate separate experience for them, not a duplicate.')

h('5. Recommended FMCG salesman role model (keep / remove)')
para('The pilot salesman holds 18 permissions today. Recommended posture so the rep '
     'operates exclusively through the Van Sales workspace:')
para('5.1 KEEP — field workspace (15):', bold=True)
table(['Permission', 'Why it stays'], [
    ['field.sales', 'Van-Sell (core selling path)'],
    ['sales.collect', 'Collections + in-sell payment'],
    ['day.close', 'Close the field day'],
    ['reconciliation.view', 'Van reconciliation / settlement'],
    ['stock_request.create', 'Request a van load'],
    ['stock.transfer', 'Van load / return transfers'],
    ['stock.view, inventory.view', 'See van stock (read-only)'],
    ['product.search', 'Find products while selling'],
    ['pricing.view', 'See prices (read-only)'],
    ['target.view, report.aggregate.view', 'Own targets & performance (read-only)'],
    ['field.attach_media', 'Visit photos'],
    ['change_requests.create', 'Request a customer GPS/data change → approval'],
    ['credit.request.create', 'Request a credit-limit change → approval'],
])
para('The rep changes nothing in master data directly — change_requests.* and '
     'credit.request.* only submit requests that a supervisor/admin approves.', italic=True)
para('5.2 REMOVE — back-office / master-data (3):', bold=True)
table(['Permission', 'Effect of removing', 'Breaks field flow?'], [
    ['sales.sell', 'Hides Quick Sale, Invoices, Sales Orders + the generic Sell fallback tab', 'No — Van-Sell uses field.sales'],
    ['customers.manage', 'Hides Customers master-data (create/edit) + generic Customers tab', 'No — field flows load customers by branch (RLS)'],
    ['customer.create', 'No master-data customer creation', 'No — see option below'],
])
para('Decision point — new outlets on the route: if reps must onboard new customers '
     'in the field, do NOT restore customer.create (full master-data). Instead enable '
     'a lightweight quick-create (platform.quick_customer_create) writing a minimal '
     'record pending approval — a separate controlled path, not master-data. Default: '
     'remove; add quick-create only if the pilot needs it.', italic=True)
para('5.3 Menus / screens hidden by this role change (rep only):', bold=True)
bullets([
    'Sales sidebar section: Quick Sale, Sales Orders, Invoices, Collections (back office), Sales Returns (back office).',
    'Customers master-data section + the generic Customers bottom-tab.',
    'POS / Quick Sale.',
    'The generic invoice-entry editor (/sales/invoices).',
    'Command-palette entries for the above (inherit nav visibility).',
])
para('All remain fully available to Accountant / Company Admin — Back Office '
     'Invoicing is their legitimate workspace, untouched.')
para('5.4 What remains available — via the Van Sales workspace:', bold=True)
para('Customer → Statement (/field/van-sales/statement) → Collect '
     '(/field/van-sales/collect) → Sell (/field/van-sales/sell → issue) → Print '
     '(/print/receipt + /print/invoices). Plus returns (/field/van-sales/return), My '
     'Day hub (route, visits, load, settlement), reconciliation, credit status + '
     'Collect Now, and credit/data change requests (submitted for approval).')
para('5.5 Mechanism (reversible, pilot-scoped):', bold=True)
para('Delete the three rows (sales.sell, customers.manage, customer.create) from '
     'erp_company_role_permissions for (pilot company 612af0bd…, role salesman). '
     'Company-scoped — role-default template and other tenants untouched — and '
     'reversible (re-insert). No code/schema change.')

h('6. Target FMCG navigation structure')
para('Principle: Van Sales is the canonical salesman workspace; future FMCG '
     'enhancements extend it rather than adding a parallel sales flow.')
para('Salesman one obvious path: Customer → Statement → Collect → Sell → Invoice → '
     'Print. The canonical flow already exists end-to-end (sell, collect, statement, '
     'credit status, Collect-Now, invoice issue, invoice print, receipt print) — the '
     'only gap is navigation cohesion: surface Statement + Collect on the Van Sales '
     'hub and from the customer context so the rep never needs the back-office menus.')
table(['Group', 'Salesman (field)', 'Accountant / Admin'], [
    ['Van Sales (canonical)', 'Sell · Collect · Returns · Statement · My Day · Settlement', 'visible (oversight)'],
    ['Sales / AR (back office)', 'hidden (no sales.sell)', 'Invoices · Orders · Collections · Returns · Aging · Credit Requests'],
    ['Customers', 'Statement only (customers.view)', 'full master-data (customers.manage)'],
    ['POS / Quick Sale', 'hidden (no pos/perm)', 'optional'],
])
para('Other verticals (Fashion/Pharmacy/Market/Restaurant/…) are unaffected — their '
     'POS/order screens stay gated by their own module + business type.')

h('7. Recommended next steps (each safe + reversible)')
numbered([
    'Tighten the pilot salesman role → field.sales + sales.collect + customers.view; '
    'remove sales.sell + customers.manage (company-override only). Re-test Van-Sell / '
    'Collect / Statement and confirm the back-office sections disappear for the rep.',
    'Add Statement + Collect entries to the Van Sales hub (/field/van-sales) and a '
    'customer-context drilldown, so the canonical flow is reachable without the '
    'sidebar (small additive UI that extends the canonical workspace).',
    'Decide POS / Sales-Orders entitlement for the pilot: disable pos + sales_orders '
    'modules if unused (hides Quick Sale + Orders company-wide).',
    'Optional, later: productize an fmcg.field_primary nav flag so this is a '
    'one-switch posture for any FMCG van-sales tenant.',
    'Only after sign-off, consider deprecating truly unused generic screens — nothing '
    'is deleted now.',
])
para('Net: for the FMCG pilot salesman, a single canonical Van Sales workspace '
     'delivering Customer → Statement → Collect → Sell → Invoice → Print, with the '
     'duplicate back-office and Quick-Sale entries hidden by role/entitlement — all '
     'reversible, no impact on other tenants or verticals.', bold=True)

out = 'docs/audits/FMCG-Sales-Entry-Point-Review.docx'
doc.save(out)
print('wrote', out)
