#!/usr/bin/env python3
"""VANTORA Full Platform Audit (read-only) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"; GREYBG="ECECEC"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows,widths=None,size=8.2,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# Cover
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Full Platform Audit"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("End-to-end read-only audit · Blockers, security, data integrity, RLS, permissions, navigation · Pilot-readiness"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Branch claude/fmcg-sell-collect-loop · Supabase rsjvgehvastmawzwnqcs (vantora-staging) · No code changed · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# EXEC
h1("Executive summary")
para("This is a complete, evidence-based read-only audit of the VANTORA platform after the Edge Functions fix. Nothing was "
     "changed. Every finding cites the exact file:line or migration. The two reported live errors are root-caused, and the "
     "broader sweep surfaced four additional Blockers — including two cross-tenant security issues — plus eight Major and "
     "seven Minor findings.", bold=True)
para("Headline: the platform is NOT ready for an external pilot until the 6 Blockers and the top separation-of-duties Major "
     "(MJ-1) are fixed. The good news: every P0 fix is small, idempotent, additive, and loses no data — the SECURITY DEFINER "
     "fixes share one guard, and the numbering/idempotency fixes mirror patterns already present in the codebase. After a "
     "single focused remediation batch the platform is pilot-conditional.", color=NAVY, bold=True)
h2("Counts")
tbl(["Severity","Count","Meaning"],[
 ["Blocker","6","Must fix before pilot — functional outage or cross-tenant security"],
 ["Major","8","Should fix soon — privilege, integrity, or security gaps"],
 ["Minor","7","Nice to have — disclosure, fragility, policy"],
],widths=[1.2,0.8,5.0],fill={(0,0):REDBG,(1,0):AMBERBG,(2,0):GREENBG})
para("Verified CLEAN: all 277 erp_ tables have RLS; service-role confined to /api/v1 + cron; money math matches RPCs; "
     "i18n parity test-enforced (69 namespaces); mobile bottom-nav gated; role permissions TS-typed (no orphans).", italic=True)

# 1 BLOCKERS
part("1. Full Blocker list (6)")
tbl(["#","Module","Issue","Root cause (evidence)","Data risk","Fix"],[
 ["BL-1","Pharmacy · Medicine catalog","Catalog refresh fails: FK violation erp_products_catalog_medicine_ref_id_fkey","importEgyptianDrugs() DELETE-then-INSERT on erp_clinic_reference; FK medicine_ref_id has no ON DELETE (reference-actions.ts:112-123; 0274:37)","None (FK protects); re-insert would orphan links","Code+SQL: UPSERT on stable key + soft-delete"],
 ["BL-2","Sales · Invoices","Invoice create fails: 'duplicate code' on (branch_id, invoice_number)","erp_sequences.current_val lags max existing number; erp_next_number re-issues a used number (0267; invoices/actions.ts:99)","None; sequence repair fixes","SQL: sequence repair + seed hardening"],
 ["BL-3","Security · RBAC","Any authenticated user can WIPE another tenant's role->permission set","erp_seed_fashion_role_perms(p_company_id) SECURITY DEFINER, unconditional DELETE, no caller guard, granted to authenticated (0147:20-47)","DESTRUCTIVE cross-tenant — victim RBAC wiped","SQL: caller-vs-tenant guard"],
 ["BL-4","Security · RBAC/modules","Any authenticated user can re-seed/enable another tenant's roles, permissions, modules","erp_seed_company_roles / erp_seed_company_modules SECURITY DEFINER, no caller guard, granted to authenticated (0021/0036/0071:29-46)","Cross-tenant RBAC/module pollution (additive)","SQL: same guard"],
 ["BL-5","FMCG · Van sales","Van stock-request & load-manifest numbers can SILENTLY duplicate","VLR-${Date.now()}-${Math.random()} / VDL- with NO UNIQUE constraint (request-actions.ts:55,229; 0011:17; 0194:21)","Duplicate numbers -> ambiguous van reconciliation","Code+SQL: erp_next_number + UNIQUE"],
 ["BL-6","Collections","Web collection double-submit -> duplicate receipt + DOUBLE balance reduction","recordCollection passes p_idempotency_key:null though RPC supports it (collections/actions.ts:39-48; 0267:100-109)","AR corruption (payment applied twice)","Code: generate+pass UUID key"],
],widths=[0.45,1.2,1.5,2.2,1.2,1.1],
 fill={(i,0):REDBG for i in range(6)})

# 2 MAJOR
part("2. Major issues (8)")
tbl(["#","Area","Issue","Evidence","Fix"],[
 ["MJ-1","Permissions (SoD)","Financial/stock posting actions check only requireAuth() — no specific permission. RPCs guard branch access, not the permission, so a viewer/staff on a branch can issue invoices, record payments, post vouchers, adjust stock, approve stock requests","issueInvoice, recordPayment, quickSale, postVoucher, completeReturn, receivePurchaseOrder, adjustStock, finalizeStockCount, approveStockRequest","Code: add hasPermission gates (accounting.post, sales.collect, stock.adjust, stock_request.approve, purchasing.manage)"],
 ["MJ-2","Navigation","/settings/workflows/templates show-then-404 on KAKO_WORKFLOW_BUILDER (nav lacks flag) — same class as A3","templates/page.tsx:20 vs navigation.ts:479","Code: add flag:'workflow_builder' token"],
 ["MJ-3","Data integrity","Non-RPC create paths (orders, POs, returns, transfers, stock-requests) have NO idempotency key -> double-submit duplicate documents","sales/orders; purchases/orders; pharmacy/purchasing; sales/returns; inventory/transfers|requests","Code+SQL: idempotency key + partial unique index"],
 ["MJ-4","Security","Legacy ts_* trade-spend tables have USING(true) open RLS on tenant data incl. ts_users.password plaintext","0004:185-195,28","SQL: scope or DROP the legacy tables"],
 ["MJ-5","Security","erp_today_journey / erp_customer_in_today_plan allow cross-tenant read of salesman journey PII + GPS","0129:76-108 (SECURITY DEFINER, no tenant scope)","SQL: scope by erp_user_company_id()"],
 ["MJ-6","Data integrity","nextProductCode SELECT-MAX-then-INSERT race -> second concurrent create fails","products/actions.ts:82-94","Code: atomic sequence"],
 ["MJ-7","Data integrity","Bulk drug import builds codes from Date.now()+index -> collision under concurrency","products/actions.ts:24-26","Code: per-row atomic code"],
 ["MJ-8","Permissions","Master-data writes auth-only: arbitrary credit limits & product prices settable by any user","customers/actions.ts upsertCustomer; products/actions.ts upsertProduct","Code: customer.create/edit, pricing.manage gates"],
],widths=[0.45,1.1,2.4,1.7,1.6],
 fill={(i,0):AMBERBG for i in range(8)})

# 3 MINOR
part("3. Minor issues (7)")
tbl(["#","Area","Issue","Fix"],[
 ["MN-1","Error handling","~25 admin/staff/platform actions return raw Postgres error.message (schema disclosure) instead of friendlyDbError (settings/*, platform/companies/actions.ts, customers/actions.ts:494)","Code"],
 ["MN-2","Module visibility","Electrical nav section has no module gate — relies only on electrical.rma perm; an admin in a non-electronics tenant could see RMA screens (navigation.ts:423-428)","Code/config"],
 ["MN-3","RLS","erp_workflow_steps cross-tenant READ of approver chains/conditions (0088:101 using auth.uid() is not null)","SQL"],
 ["MN-4","RLS","erp_product_cost(p_product) cross-tenant read of cost basis, one scalar/call (0288:47-62)","SQL"],
 ["MN-5","RLS","erp_seed_action_policies(p_company) cross-tenant default-row insert (nuisance) (0272:93)","SQL (same guard as BL-3/4)"],
 ["MN-6","Data","ON DELETE RESTRICT on product_id everywhere is correct — ensure product delete is soft-delete only (it is: toggleProductActive)","Config/none"],
 ["MN-7","Data","erp_companies cascade chain: deleting one company row wipes the tenant history — intentional teardown, restrict who can hard-delete","Config/policy"],
],widths=[0.45,1.1,4.0,0.9],
 fill={(i,0):GREENBG for i in range(7)})

# 4 SECURITY
part("4. Security findings")
b("Cross-tenant WRITE (BL-3, BL-4, MN-5): unguarded SECURITY DEFINER erp_seed_*(p_company_id) granted to authenticated. BL-3 is DESTRUCTIVE. All fixed by one guard: erp_is_platform_owner() OR p_company_id = erp_user_company_id().",bold=True,color=RED)
b("Cross-tenant READ (MJ-5, MN-3, MN-4): salesman journey PII/GPS, workflow config, product cost basis.")
b("Open legacy tables (MJ-4): ts_* USING(true) + ts_users.password plaintext; never dropped.")
b("Latent watch-item: erp_post_revenue has no internal isolation but is NOT granted to authenticated (0070:54) — safe today; add a defensive in-body guard so a future GRANT can't open it.",color=AMBER)
b("VERIFIED CLEAN: all 277 erp_ tables have RLS enabled; createServiceClient confined to /api/v1 (API-key auth) + CRON_SECRET-gated crons + HMAC callback; non-super-admins' client-supplied company_id is ignored and pinned to ctx.companyId.",color=GREEN)

# 5 DATA INTEGRITY
part("5. Data-integrity findings")
para("BL-1, BL-2, BL-5, BL-6, MJ-3, MJ-6, MJ-7, MN-6, MN-7 (see those sections).")
b("VERIFIED SOUND: money math (per-line round-then-sum) matches the erp_van_sell / erp_settle_collection RPCs byte-for-byte; entity codes tenant-scoped (0019); document numbers branch-scoped (0268); idempotency SCHEMA present — the gaps are in the CALLERS (BL-6, MJ-3), not the schema.",color=GREEN)

# 6 PERMISSIONS
part("6. Permissions & role findings")
para("MJ-1 (separation of duties on posting — highest-priority Major), MJ-8 (master-data credit/price), MN-2 (electrical).")
b("VERIFIED CLEAN: ROLE_PERMISSIONS is TS-typed Record<BranchRole, Permission[] | '*'> — no orphan permission strings possible; salesman holds sales.sell; engine modules excluded from ALL_MODULES.",color=GREEN)
b("Smell (not a finding): accountant holds customers.change_status; cashier holds cross-vertical perms (pharmacy.dispense, restaurant.manage, fashion.*) — harmless because each is module-gated.",color=GREY)

# 7 NAV/UI
part("7. Navigation & UI findings")
para("MJ-2 (workflow-templates show-then-404).")
b("VERIFIED CLEAN: all other page guards are SUPERSETS of their nav gates (no show-then-redirect); pharmacy nav items carry tenant-feature flag tokens fed by getFeatureFlags; mobile bottom-nav is perm+module gated and targets no flag-gated routes; i18n parity is test-enforced (i18n.test.ts, keys-usage.test.ts).",color=GREEN)

# 8 MODULE VISIBILITY
part("8. Module-visibility findings")
para("MN-2 (electrical section has no module gate — permission-only boundary, fragile if electrical.rma is granted broadly).")
b("Otherwise sound: MODULE_LABELS fully covers the Module union; engine modules stay erp_company_modules-driven; pharmacy/clinic/hotel/fashion sections carry module gates.",color=GREEN)

# 9 RLS
part("9. Multi-tenant / RLS findings")
para("BL-3, BL-4, MJ-4, MJ-5, MN-3, MN-4, MN-5, plus the erp_post_revenue latent watch-item (Section 4). The four "
     "Blocker/Major SECURITY DEFINER issues share one root cause and one fix shape (the caller-vs-tenant guard).")

# 10 READINESS
part("10. Pilot-readiness score")
para("Overall: 5.5 / 10 — NOT READY for an external pilot until the 6 Blockers + MJ-1 are fixed. Conditionally ready after "
     "a short, well-scoped remediation sprint (fixes are small and mostly one-line / idempotent).", bold=True, color=NAVY)
tbl(["Dimension","Score","Note"],[
 ["Auth / login / edge functions","9/10","Solid; edge functions now deployed"],
 ["Tenant isolation (RLS)","4/10","2 cross-tenant write Blockers (one destructive)"],
 ["Data integrity / financial flows","5/10","4 Blockers/Majors in numbering & idempotency"],
 ["Permissions / separation of duties","5/10","Posting actions under-gated (MJ-1)"],
 ["Navigation / module visibility","8/10","One show-then-404 + electrical gate"],
 ["i18n / mobile","9/10","Test-guarded, clean"],
 ["Error handling","7/10","Raw-message leaks (minor)"],
],widths=[2.6,0.8,3.0],
 fill={(1,1):REDBG,(2,1):AMBERBG,(3,1):AMBERBG,(0,1):GREENBG,(4,1):GREENBG,(5,1):GREENBG,(6,1):BLUEBG})

# REMEDIATION
part("Consolidated remediation plan (prioritized)")
h2("P0 — Must fix before pilot")
b("BL-3, BL-4, MN-5 — add caller-vs-tenant guard to the erp_seed_* functions (one SQL migration, ~3 functions). DO FIRST — destructive.",bold=True,color=RED)
b("BL-6 — recordCollection: generate + pass a UUID idempotency key (code, tiny).")
b("BL-1 — medicine catalog UPSERT + soft-delete + unique key (code + SQL).")
b("BL-2 — sequence-repair migration + seed hardening (SQL).")
b("BL-5 — van request/manifest -> erp_next_number + UNIQUE (code + SQL).")
b("MJ-1 — add permission gates to financial/stock posting actions (code, mechanical).")
h2("P1 — Should fix soon (remaining Majors)")
b("MJ-4 (drop/scope ts_*), MJ-5 (journey PII scope), MJ-2 (workflow-templates flag), MJ-3 (idempotency on create paths), MJ-8 (credit/price gates), MJ-6/MJ-7 (atomic product codes).")
h2("P2 — Nice to have (Minors)")
b("MN-1 (friendlyDbError everywhere), MN-2 (electrical module), MN-3/MN-4 (cross-tenant reads), erp_post_revenue defensive guard.")
h2("P3 — Documentation / policy")
b("MN-7 (restrict company hard-delete), MN-6 (confirm soft-delete only), production/main vs pilot-preview divergence (the live pilot is the feature-branch preview, not main).")
para("")
para("Effort: P0 is realistically a single focused remediation batch — the SECURITY DEFINER fixes are one shared guard, and "
     "the idempotency/numbering fixes mirror patterns already in the codebase. All P0 fixes are idempotent, additive, and "
     "lose no data (no demo-data deletion required). This audit changed nothing; fixes proceed only on approval, diff by diff.",
     bold=True, color=NAVY)

out="docs/audits/VANTORA-Full-Platform-Audit.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
