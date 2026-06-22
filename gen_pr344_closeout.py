#!/usr/bin/env python3
"""VANTORA — PR #344 Close-Out Confirmation (report-only) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"; GREYBG="ECECEC"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)

def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.6,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.6); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ---------------- Cover ----------------
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Pull Request #344 — Close-Out Confirmation"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Report-only PR closed without merging · Guardrails verified · Production step held"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("nabilhahaha/Kako · closed 2026-06-22 16:52 UTC · main unchanged (1ba58c5) · No production writes"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---------------- Outcome ----------------
h1("Outcome")
para("Pull Request #344 (the Route Planner migration-repair dry-run handoff report) was closed WITHOUT merging, at "
     "your direction. The report remains available on its branch; nothing landed on main; and the production "
     "migration-repair step stays blocked pending your explicit approval.", bold=True)
tbl(["Field","Value"],
    [["PR","#344 — docs: Route Planner migration-repair dry-run handoff report"],
     ["State","CLOSED (merged = false)"],
     ["Closed at","2026-06-22 16:52 UTC"],
     ["Branch (preserved)","claude/migration-dryrun-report @ dcf3c13"],
     ["Base branch","main @ 1ba58c5 — unchanged"],
     ["Subscription","Unsubscribed (auto-resubscribe only if reopened)"]],
    widths=[1.9,4.6],
    fill={(1,1):GREENBG})

# ---------------- 9-point confirmation ----------------
h1("Confirmation — all 9 points")
tbl(["#","Item","Status"],
    [["1","PR #344 closed","Closed (not merged), 16:52 UTC"],
     ["2","Report remains on claude/migration-dryrun-report","Branch + commit dcf3c13 intact; .docx reports preserved"],
     ["3","No changes merged to main","main still at 1ba58c5; merged = false"],
     ["4","No production DB writes","Read-only throughout"],
     ["5","No migration repair","Not performed"],
     ["6","No Route Planner product code","Only report artifacts"],
     ["7","No PR #325","Untouched"],
     ["8","No field-insights changes","Untouched"],
     ["9","No route-guard / permissions changes","Untouched"]],
    widths=[0.4,3.3,2.8],
    fill={(i,2):GREENBG for i in range(9)})

# ---------------- Hold ----------------
h1("Production step — HELD")
para("The production migration-repair step remains BLOCKED pending explicit approval. No migration repair and no "
     "Route Planner product-code merge will be started until you clearly authorize it.", bold=True, color=RED)
h2("Safe next step when you are ready (from the report)")
b("Pipeline owner runs the dry-run via the Supabase CLI on a clone — not via this MCP session.")
b("Match this project's scheme: name = file stem (e.g. 0358_rp_connector_admin), version = a fresh timestamp like 0353-0357.")
b("Require supabase db diff to come back clean on the clone before anything touches production.")
b("Only after a clean clone dry-run, and only with your explicit approval, reconcile 0358-0364 on production.")

h2("Constraints honored")
para("No production writes. No migration repair on production. No Route Planner product-code merge. No PR #325. "
     "No field-insights changes. No route-guard or permissions changes. Read-only throughout; no branch provisioned ($0).",
     italic=True, color=GREY, size=9)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"VANTORA_PR344_CloseOut_Confirmation.docx")
d.save(out)
print("WROTE", out, os.path.getsize(out), "bytes")
