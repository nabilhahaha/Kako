#!/usr/bin/env python3
"""VANTORA Approval Convergence — P1b/P4 delivery + P3/P2 plans -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def bl(t,size=9.3,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    for ln in t.split("\n"):
        p=d.add_paragraph(); r=p.add_run(ln); r.font.name="Consolas"; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0x22,0x22,0x22)
def tbl(headers,rows,widths=None,size=8.0,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Approval Convergence — Phase 2"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Delivered: branch scope + SLA escalation · Planned: unified inbox + 6 field-workflow migrations"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Additive engine primitives validated on staging · field migrations flagged + reversible · pilot untouched · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Status")
tbl(["Item","Phase","Status"],
[
 ["Permission resolution + governance + inbox fix","P0","Delivered (committed)"],
 ["Credit-Limit / Trade-Spend / Price-Change on engine","P1","Delivered (flagged)"],
 ["Branch-scoped approvers","P1b","Delivered this phase (additive, staging-validated)"],
 ["SLA escalation routing","P4 foundation","Delivered this phase (additive, staging-validated)"],
 ["Unified approval inbox","P3","Designed (this doc) — build next"],
 ["6 field-workflow migrations","P2","Planned (this doc) — flagged, post-pilot activation"],
],widths=[3.2,1.0,2.4],
fill={(0,2):GREENBG,(1,2):GREENBG,(2,2):GREENBG,(3,2):GREENBG,(4,2):AMBERBG,(5,2):AMBERBG})

part("1 · Branch-scoped approvers (P1b) — DELIVERED")
h2("Design")
para("A step may set branch_scoped=true. When set, its approvers are resolved to the users holding the role/permission IN "
     "THE REQUEST'S BRANCH (erp_workflow_instances.branch_id) and assigned as plain 'user' tasks. Because the resulting "
     "tasks are 'user' tasks, the decide RPC is UNCHANGED. erp_workflow_start now records branch_id from context.branch_id "
     "(additive — null when not provided, so legacy behaviour is identical).")
h2("Implementation (0299)")
bl("erp_workflow_steps.branch_scoped column (default false).")
bl("erp_workflow_resolve_users_in_branch(company, type, ref, branch) — branch-restricted resolver.")
bl("erp_workflow_make_tasks — faithful superset + branch-scoped fan-out to in-branch users.")
bl("erp_workflow_start — faithful superset that records instance.branch_id from context.")
h2("Tests / staging validation")
tbl(["Check","Result"],
[
 ["resolve_users_in_branch(branch1)","3 in-branch approvers"],
 ["resolve_users_in_branch(random branch)","0 (correctly excluded)"],
 ["branch-scoped start → make_tasks","3 'user' tasks for exactly the in-branch approvers"],
 ["decide path","unchanged ('user' tasks already authorised)"],
],widths=[3.6,3.0],fill={(i,1):GREENBG for i in range(4)})
h2("Rollback / risk")
para("Rollback: drop column branch_scoped; restore prior make_tasks/start. No step sets branch_scoped, so DORMANT — zero "
     "current effect. Risk: LOW (additive, behaviour-preserving).", color=GREEN)

part("2 · SLA escalation routing (P4 foundation) — DELIVERED")
h2("Design")
para("A new function escalates overdue pending tasks: for any pending task past its due_at whose step defines escalate_to, "
     "it ADDS the escalation target (company_admin or a role) as a 'user' task and stamps the original's escalated_at. It "
     "does NOT change instance status or the decide RPC, so existing decisions keep working and either approver can act. "
     "Idempotent (escalated_at guards re-escalation).")
h2("Implementation (0300)")
bl("erp_workflow_escalate_overdue() → integer (count escalated). Additive, new function, execute revoked from public/anon.")
bl("Wiring (infra, not in this migration): schedule it on pg_cron (e.g. every 15 min) — `select cron.schedule('wf-escalate','*/15 * * * *', $$select erp_workflow_escalate_overdue()$$);`")
h2("Tests / staging validation")
tbl(["Check","Result"],
[
 ["Overdue task → escalate","Escalation 'user' task created for company_admin; original escalated_at set"],
 ["Audit","erp_log_audit('escalate', 'workflow_task') row written"],
 ["Re-run (idempotent)","No duplicate escalation task"],
],widths=[3.0,3.6],fill={(i,1):GREENBG for i in range(3)})
h2("Rollback / risk")
para("Rollback: drop function (and unschedule the cron job). DORMANT until scheduled. Risk: LOW (new function, no live path "
     "touched, idempotent).", color=GREEN)

part("3 · Unified approval inbox (P3) — DESIGN")
h2("Goal")
para("One Approval Queue for approvers: field approvals (day-close, visit, transfers…) AND engine workflow tasks "
     "(credit, trade-spend, price-change, change-requests) in a single, mobile-first list. Retire the duplicate Approval "
     "Center; remove the in-memory 2000-row scan.")
h2("Design")
bl("New SQL function erp_workflow_my_tasks() (SECURITY DEFINER): returns the pending engine tasks the caller can act on, "
   "pushing the company_admin/user/role/permission predicate INTO SQL (indexed) instead of the current in-memory filter — "
   "the scalability fix.")
bl("The unified queue page reads (a) the existing field-queue items and (b) erp_workflow_my_tasks(), maps both to one "
   "ApprovalItem shape, and dispatches: field types → their existing actions; engine types → decideTask.")
bl("ApprovalsTabs + /approval-center retired behind a flag (KAKO_UNIFIED_INBOX); the field queue page becomes the single "
   "surface. /approvals (workflow inbox) kept as a redirect during transition.")
h2("Tests / rollback / risk")
para("Tests: erp_my_tasks returns exactly actionable tasks per assignee type (unit on the SQL + a page-mapping test). "
     "Rollback: flag off restores the three-tab layout. Risk: MEDIUM (UX + a read-path change); no decision logic changes.",
     color=AMBER)

part("4 · Field-workflow migration plans (P2) — PLANNED")
para("All six follow ONE pattern and stay HIGH-RISK: each is feature-flagged (default OFF), keeps its existing RPC as the "
     "OUTCOME HANDLER (engine decides who/when; the proven RPC does what), and is validated on staging before any "
     "activation. The operational FMCG flows are NOT activated until after pilot validation. The engine instance becomes "
     "the single source of truth; the legacy direct path remains for flag-off.", bold=True)
h2("Common pattern")
code("create action  ──▶  erp_workflow_start(key, entity, record_id, {branch_id, amount})\n"
     "decision (inbox) ─▶  decideTask ─▶ erp_workflow_decide ─▶ applyWorkflowOutcome(entity)\n"
     "                                            └─ handler calls the existing RPC (do the thing)")

def plan(title, rows, risk):
    h2(title)
    tbl(["Aspect","Plan"], rows, widths=[1.4,5.1], size=8.4)
    para("Risk: "+risk, italic=True, color=AMBER, size=9)

plan("4.1 Day-Close Exception",[
 ["Design","Definition day_close_approval / entity work_session; approver permission day.approve_close_exception; branch_scoped; flags on. Adds an explicit REJECT (today approve-only)."],
 ["Implementation","closeDay (coverage < threshold) → erp_workflow_start with {branch_id}. Handler work_session: approve→close_status='closed'; reject→'open' (re-openable). Flag KAKO_APPROVAL_DAYCLOSE; legacy approveDayClose kept."],
 ["Tests","Handler unit; staging: approve closes, reject reopens, self-approval block, branch-scoped to the rep's branch."],
 ["Rollback","Flag off → legacy approveDayClose; delete definition."],
],"Medium — daily operational; flag-gated + reversible; pilot unaffected while off.")

plan("4.2 Out-of-Route Visit",[
 ["Design","visit_compliance_approval / visit_compliance; permission visit.approve_out_of_route; branch_scoped; already approve+reject."],
 ["Implementation","erp_check_in_visit (off-route) → start. Handler sets status approved/rejected. Flag KAKO_APPROVAL_VISIT; legacy decideVisitCompliance kept."],
 ["Tests","Staging: off-route visit creates task to in-branch supervisor; approve/reject; reject-reason enforced."],
 ["Rollback","Flag off → legacy path; delete definition."],
],"Medium — high volume; flag-gated; validate inbox load with erp_my_tasks first.")

plan("4.3 Customer Transfer",[
 ["Design","customer_transfer_approval / customer_transfer; permission customer.transfer; branch_scoped (note: cross-branch transfer may route to BOTH branches — model as 2 levels). Adds explicit reject."],
 ["Implementation","transferCustomer(requireApproval) → start. Handler applies the transfer on approve / marks rejected. Flag KAKO_APPROVAL_CUSTTRANSFER; legacy approveCustomerTransfer kept."],
 ["Tests","Staging: pending→applied on approve; reject path; SoD (requester ≠ approver)."],
 ["Rollback","Flag off → legacy; delete definition."],
],"Medium — touches customer ownership; flag-gated + reversible.")

plan("4.4 Load Request",[
 ["Design","stock_request_approval / stock_request; permission stock_request.approve; branch_scoped."],
 ["Implementation","createStockRequest → start. Handler calls existing erp_approve_stock_request / reject. Flag KAKO_APPROVAL_LOADREQ; legacy approve/reject kept."],
 ["Tests","Staging: request→approved/rejected via engine; in-branch warehouse approver."],
 ["Rollback","Flag off → legacy; delete definition."],
],"Medium — stock to van; flag-gated.")

plan("4.5 Van Transfer",[
 ["Design","van_transfer_approval / van_transfer; permission stock.transfer.approve; threshold (auto-approve below value KEPT; above → workflow)."],
 ["Implementation","requestVanTransfer (above threshold) → start. Handler calls erp_approve_van_transfer / erp_reject_van_transfer. Flag KAKO_APPROVAL_VANTRANSFER; auto + legacy paths kept."],
 ["Tests","Staging: below-threshold auto still works; above routes to engine; approve completes stock move; reject reason."],
 ["Rollback","Flag off → legacy; delete definition."],
],"Medium-High — moves stock; preserve the auto-approve threshold; staging-validate stock conservation.")

plan("4.6 Van Reconciliation",[
 ["Design","van_reconciliation_approval / van_reconciliation; permission reconciliation.approve; threshold on variance."],
 ["Implementation","erp_compute_van_reconciliation (variance > threshold) → start. Handler calls erp_settle / erp_reject_van_reconciliation. Flag KAKO_APPROVAL_VANRECON; legacy kept."],
 ["Tests","Staging: under-variance auto-settle unchanged; over-variance routes; settle/reject via engine; audit + balance checks."],
 ["Rollback","Flag off → legacy; delete definition."],
],"Medium-High — financial variance; flag-gated; validate settlement amounts on staging.")

part("5 · Sequencing & guardrails")
bl("Build order next: P3 unified inbox (so all migrated approvals have one home), then P2 one workflow at a time, "
   "lowest-risk first (Load Request, Day-Close → Visit → Customer Transfer → Van Transfer → Van Reconciliation).", color=DARK)
bl("Each P2 workflow: ship flagged OFF + staging-validated; activate only after pilot sign-off on the prior one.", color=DARK)
bl("Pilot stability: every operational flow keeps its legacy path while its flag is OFF. Nothing activates without "
   "staging validation. All steps reversible by flag.", color=GREEN)
para("No destructive change, production risk, or architectural conflict was encountered in this phase — branch scope and "
     "escalation are additive supersets, validated and dormant. Proceeding to P3 next.", bold=True, color=NAVY)

out="docs/audits/VANTORA-Approval-Convergence-Phase2.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
