#!/usr/bin/env python3
"""VANTORA Product Positioning recommendation -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E); CYAN=RGBColor(0x0E,0x7C,0x86)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.6,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.6); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Product Positioning"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("How to present VANTORA: ERP · Distribution ERP · DMS · Field Force Automation · Business OS"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Recommendation + rationale · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Recommendation (one line)")
para("Lead with “FMCG Distribution Management System (DMS) with built-in ERP” — i.e., sell VANTORA primarily as a DMS / "
     "Field-Force-enabled distribution platform, with ERP as the included foundation, and reserve “Business OS” as the "
     "long-term vision narrative for larger/multi-industry accounts.", bold=True, color=NAVY, size=11)

part("The five candidate positions")
tbl(["Position","What the buyer hears","Fit for VANTORA"],[
 ["ERP","Heavy back-office finance/inventory system","True but commoditised; sounds slow/expensive; not the differentiator"],
 ["Distribution ERP","ERP specialised for distributors","Accurate; better than plain ERP; still leads with 'ERP'"],
 ["DMS (Distributor Mgmt System)","Routes, van sales, coverage, secondary sales","Exactly what FMCG distributors search/buy; matches the differentiator"],
 ["Field Force Automation (FFA)","Rep app: visits, orders, GPS","A capability, not the whole product — too narrow to lead with alone"],
 ["Business OS","One system to run the whole company","Aspirational; right vision, too abstract for the first sale"],
],widths=[1.5,2.4,2.5],
 fill={(2,0):GREENBG})

part("Recommended positioning architecture")
h2("Primary (how you open every conversation)")
para("“VANTORA — the Distribution Management System that runs your whole distributor, from the van to the ledger.”", bold=True, color=NAVY)
b("Leads with DMS (the category FMCG distributors recognise and budget for) and immediately signals the ERP depth ('to the ledger') that point DMS/van-sales apps lack.")
h2("Supporting pillars (the proof under the headline)")
b("Field Force Automation — the mobile rep app (sell/collect/visit, offline) — your wedge and daily-use hook.")
b("Built-in ERP — sales, inventory, purchasing, accounting — 'no separate ERP, no reconciliation gap.'")
b("Governance — approvals, permissions, credit, pricing, audit — 'control, not chaos.'")
h2("Vision (for enterprise / investors)")
para("“…and the same platform is a Business OS — one codebase that also runs pharmacy, retail, clinics and more.” Use this to "
     "raise the ceiling on large accounts and the fundraising story, AFTER the DMS value lands.", italic=True)

part("Rationale")
b("Category clarity: FMCG distributors search for 'DMS', 'van sales', 'secondary sales', 'SFA/FFA' — not 'ERP'. Leading with DMS meets demand where it already exists.",bold=True)
b("Differentiation: pure DMS/van-sales tools have no real ERP/GL; pure ERPs have no real field force. VANTORA's wedge is being BOTH — 'DMS with built-in ERP' captures that in one phrase.")
b("Sales motion: DMS is a concrete, budgeted line item with a fast ROI story (collections, coverage). 'ERP' triggers long, IT-led evaluations; 'Business OS' triggers 'what is that?'.")
b("Expansion: land as the DMS for the field force, expand into finance and engines — the per-rep pricing and the modular product both support this.")
b("Don't over-claim: 'Business OS' is true at the platform level but dilutes the FMCG sale; keep it as vision, not the lead.")

part("Messaging cheat-sheet")
tbl(["Audience","Lead with","Avoid"],[
 ["FMCG distributor owner","DMS that also does your accounting; collect cash & see coverage from a phone","'ERP migration', jargon"],
 ["Operations / sales manager","Field Force Automation + Approval Queue + coverage","Back-office finance detail"],
 ["Finance","Built-in ERP: van sale = invoice = GL, automatic AR/AP","'Just a sales app'"],
 ["Enterprise / investor","Multi-tenant Business OS, FMCG-first, multi-industry","Narrow 'rep app' framing"],
],widths=[1.7,2.7,2.0])

part("Naming & tagline options")
b("Tagline (primary): 'Run your distributor from the van to the ledger.'")
b("Category line: 'FMCG Distribution Management System — with built-in ERP.'")
b("Vision line (enterprise): 'The Business OS for distribution — and beyond.'")
para("")
para("Bottom line: position as a DMS-with-built-in-ERP for FMCG distribution. It matches buyer demand, states the "
     "differentiator, shortens the sale, and still leaves room to grow the story into a multi-industry Business OS once the "
     "first distributors are won.", bold=True, color=NAVY)
out="docs/audits/VANTORA-Product-Positioning.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
