#!/usr/bin/env python3
"""VANTORA FMCG Adoption Improvements Report (U1-U8) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.4,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.5); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("FMCG Adoption Improvements"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("U1–U8 implemented — role-aware landing, approval consolidation, quick actions, cheat sheets"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Branch claude/fmcg-sell-collect-loop · commit ee9b133 · tsc clean · 1321 tests · build green · no new features · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Summary")
para("This patch transitions VANTORA from 'feature-complete' to 'feels like a mature FMCG product' — entirely through "
     "usability, navigation and onboarding changes. No new business functionality was added; every change reduces clicks, "
     "improves discoverability, or guides the user. The single biggest lever — role-aware landing — means each role now "
     "opens directly on the screen they use first thing every day.", bold=True)
tbl(["Item","Improvement","Type","Status"],[
 ["U1","Role-aware landing pages (each FMCG role opens on its work screen)","Code (1 fn) + tests","Done"],
 ["U2","Consolidate 3 approval menu items into one 'Approvals' with tabs","Code (nav + tabs)","Done"],
 ["U3","Getting-started step CTAs on the dashboard","Already present (kept)","Done"],
 ["U4","Role-gated one-tap quick-action bar on the dashboard","Code (UI) + i18n","Done"],
 ["U5","Reduce nav clutter (3→1 approvals); keep the rest permission-gated","Code (via U2)","Done (safe)"],
 ["U8","Per-role cheat sheets (login · landing · daily flow)","Docs","Done"],
],widths=[0.5,3.5,1.5,0.9],fill={(i,3):GREENBG for i in range(6)})

part("U1 — Role-aware landing  (the biggest adoption lever)")
para("Before: every FMCG role landed on the generic /dashboard. Now resolveHomePath (src/lib/erp/home.ts) opens each role on "
     "the screen they work from first thing every day. The dashboard already redirects to the resolved home, so field roles "
     "never linger on a KPI page.", bold=True)
tbl(["Role","Lands on now","Why"],[
 ["Salesman / Van rep","/today","Their day + route, not KPIs"],
 ["Supervisor (+ area/regional/national/director)","/approvals/queue","Their inbox is their #1 daily job"],
 ["Accountant","/collections","Chase AR — the daily task"],
 ["Warehouse Keeper","/inventory/requests","Load requests waiting for them"],
 ["Branch Manager","/manager","Branch cockpit"],
 ["Company Admin / Manager","/dashboard","The overview IS their view"],
],widths=[2.6,1.8,2.2],fill={(0,1):GREENBG,(1,1):GREENBG,(2,1):GREENBG,(3,1):GREENBG,(4,1):GREENBG})
b("Falls back to /dashboard when roles aren't known (view-as preview / legacy callers) — no regression.")
b("Most-senior role wins for multi-role users; unit tests added for every role + fallback + precedence.")

part("U2 — Approval navigation consolidated")
para("Before: three similarly-named items in one menu — Approval Center, Approvals, Approval Queue — left users unsure where "
     "to act. Now there is ONE 'Approvals' entry → the unified queue, permission-gated (incl. workflow.manage). A shared "
     "tab bar (Field / Workflow / Center) sits on all three approval pages, so every surface stays reachable without the "
     "menu clutter.", bold=True)
b("Field = the unified queue (day-close, visit, customer/van transfer, trade-spend).")
b("Workflow = the workflow-task inbox; Center = the approval center — shown as tabs only for users who run the workflow engine.")
b("Nav tests updated to the new single-entry, permission-gated model.")

part("U3 + U4 — Onboarding & shortcuts")
h2("U3 — Getting started")
b("The dashboard already shows a getting-started checklist (Add first product → customer → invoice) with done-state ticks; "
  "kept as the onboarding surface.")
h2("U4 — Quick actions")
b("A role-gated one-tap quick-action bar on the dashboard: New Invoice · Record Collection · New Customer · Receive PO · "
  "Reports — each shown only to roles that hold the permission. Cuts the 'hunt through menus' clicks for admins/managers.")

part("U5 — Navigation clutter")
para("The heaviest clutter (the three approval items) is removed by U2. The remaining nav items stay permission-gated, so "
     "each role already sees only its tools. Deliberately NO further hiding was applied — that would risk removing access; "
     "deeper per-role curation is a safe optional follow-up.", bold=True)

part("U8 — Per-role cheat sheets")
para("docs/pilot/ROLE_CHEAT_SHEETS.md — one page per role: login, where you now land, and your daily flow (Company Admin, "
     "Branch Manager, Supervisor, Salesman/Van rep, Warehouse Keeper, Accountant). Ready to hand to pilot users.", bold=True)

part("Verification & guardrails")
tbl(["Gate","Result"],[
 ["Type check (tsc --noEmit)","PASS (exit 0)"],
 ["Unit/integration suite","PASS — 1321 passed, 0 failed (+ new home + nav tests)"],
 ["Production build","PASS (exit 0)"],
 ["New business functionality","NONE — usability/navigation/onboarding only"],
 ["Access removed from any role","NONE — only landing, naming, shortcuts changed"],
],widths=[2.6,3.8],fill={(0,1):GREENBG,(1,1):GREENBG,(2,1):GREENBG,(3,1):GREENBG,(4,1):GREENBG})

part("Net effect")
para("VANTORA now opens like a mature FMCG platform: a salesman sees their day, a supervisor sees their approval inbox, an "
     "accountant sees who owes money — immediately, on login. Combined with the enriched live demo data (route, 12 invoices, "
     "collections, ~4,906 EGP open AR, 4 pending approvals), every role meets a populated, working system rather than empty "
     "screens.", bold=True, color=NAVY)
h2("Optional next (safe, no new features)")
b("Role-tuned mobile quick tabs (manager/accountant).")
b("Light per-role desktop nav curation.")
para("")
para("Operational-pilot focus maintained — reduce clicks, improve discoverability, make daily FMCG operations faster, "
     "without new functionality.", bold=True, color=NAVY)

out="docs/audits/VANTORA-FMCG-Adoption-Improvements.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
