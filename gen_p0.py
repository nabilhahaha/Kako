#!/usr/bin/env python3
"""VANTORA P0 Remediation Report -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows,widths=None,size=8.2,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("P0 Remediation Report"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("6 Blockers + separation-of-duties · additive, idempotent, reversible · no data deletion"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Branch claude/fmcg-sell-collect-loop · commit 4facc3e · tsc clean · 1317 tests · build green · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("1. Summary")
para("All seven approved P0 items are implemented, verified, and pushed. Every change is additive, idempotent and reversible; "
     "NO data was deleted and no tenant data was cleaned up. The two reported live errors (medicine catalog, invoice numbering) "
     "and the four other Blockers are fixed, and the separation-of-duties gap (MJ-1) is closed. P1 items are intentionally "
     "untouched (separate phase).", bold=True)
tbl(["Item","Fix","Change type","Status"],[
 ["BL-1","Medicine catalog UPSERT-in-place (no DELETE)","Code","Done"],
 ["BL-2","Sequence repair migration (advance counters)","SQL (0294)","Done"],
 ["BL-3","Revoke EXECUTE on destructive cross-tenant seed fn","SQL (0293)","Done"],
 ["BL-4","Revoke EXECUTE on RBAC/module seed fns","SQL (0293)","Done"],
 ["BL-5","Van numbering -> erp_next_number + UNIQUE index","Code + SQL (0295)","Done"],
 ["BL-6","Collection idempotency key threaded from client","Code","Done"],
 ["MJ-1","Permission gates on financial/stock posting actions","Code + test","Done"],
],widths=[0.5,3.4,1.5,0.9],fill={(i,3):GREENBG for i in range(7)})

part("2. What changed (per item)")
h2("BL-1 — Medicine catalog (clinic/reference-actions.ts)")
b("Replaced DELETE-then-INSERT with an in-place UPSERT keyed on the PRIMARY KEY: existing drugs keep their id (so "
  "erp_products_catalog.medicine_ref_id links survive), new drugs are inserted, drugs absent from the feed are left "
  "untouched. The DELETE that violated the FK is gone entirely — no FK can be hit, no row is deleted.")
h2("BL-2 — Invoice/document numbering (migration 0294)")
b("Advances erp_sequences.current_val to GREATEST(current, max trailing-integer of existing numbers) per (branch, type) for "
  "invoices, sales orders, POs, sales returns, journals, payment/receipt vouchers, collections. Over-advance is safe; the "
  "duplicate-code collision is eliminated. Touches only the counter rows — no business data read-for-write or deleted. "
  "Idempotent (GREATEST).")
h2("BL-3 / BL-4 — Cross-tenant seed functions (migration 0293)")
b("REVOKE EXECUTE ... FROM authenticated, public, anon on erp_seed_fashion_role_perms / erp_seed_company_roles / "
  "erp_seed_company_modules. Verified safe: these are only ever called by AFTER-INSERT triggers on erp_companies, and every "
  "such trigger function is SECURITY DEFINER, so onboarding runs as the owner and is unaffected; no app code calls them "
  "directly (zero .rpc() sites). service_role keeps EXECUTE. Reversible by re-GRANT (documented in the migration).")
h2("BL-5 — Van document numbers (request-actions.ts + migration 0295)")
b("request_number / manifest_number now come from the atomic erp_next_number(branch_id, 'stock_request'|'van_load') counter "
  "instead of Date.now()+Math.random(); migration 0295 adds branch-scoped UNIQUE indexes so any residual collision FAILS "
  "LOUDLY instead of silently duplicating.")
h2("BL-6 — Collection idempotency (collections actions + manager)")
b("recordCollection now accepts and forwards a client-generated UUID idempotency key (was hardcoded null); the client holds "
  "one stable key per submit attempt (reused across a rapid double-click, cleared on success), so erp_settle_collection "
  "dedupes a retry instead of creating a second receipt and reducing the balance twice.")
h2("MJ-1 — Separation of duties (10 posting actions + test)")
b("Added a specific hasPermission() gate to each financial/stock posting action on top of the RPC's branch-access check: "
  "createInvoice/issueInvoice -> sales.sell; recordPayment/quickSale -> sales.collect; createVoucher/postVoucher -> "
  "accounting.post; completeReturn -> sales.return|sales.sell; receivePurchaseOrder -> purchasing.manage; "
  "recordSupplierPayment -> accounting.post|suppliers.manage; adjustStock -> inventory.adjust|stock.adjust; "
  "finalizeStockCount -> inventory.count; completeTransfer -> inventory.transfer|stock.transfer.approve; "
  "approveStockRequest -> stock_request.approve. Gates reuse existing i18n keys (no new keys).")

part("3. Files changed")
tbl(["File / migration","Item"],[
 ["supabase/migrations/0293_security_revoke_cross_tenant_seed_fns.sql","BL-3/BL-4"],
 ["supabase/migrations/0294_sequence_repair.sql","BL-2"],
 ["supabase/migrations/0295_van_document_number_unique.sql","BL-5"],
 ["src/app/(app)/clinic/reference-actions.ts","BL-1"],
 ["src/app/(app)/field/van-sales/request-actions.ts","BL-5"],
 ["src/app/(app)/collections/actions.ts + collections-manager.tsx","BL-6"],
 ["sales/invoices, sales/pos, sales/returns, accounting/vouchers, purchases/orders, suppliers, inventory, inventory/count, inventory/transfers, inventory/requests (actions.ts)","MJ-1"],
 ["src/lib/erp/mj1-posting-permissions.test.ts (new)","MJ-1 regression test"],
],widths=[5.0,1.5])

part("4. Verification")
tbl(["Gate","Result"],[
 ["Type check (tsc --noEmit)","PASS (exit 0)"],
 ["Unit/integration suite (vitest run)","PASS — 1317 passed, 181 skipped, 0 failed (+3 new MJ-1 tests)"],
 ["Production build (next build)","PASS (exit 0)"],
 ["New MJ-1 regression test","PASS — viewer/staff blocked; salesman/cashier/accountant/warehouse_keeper allowed; rep cannot approve own load"],
],widths=[2.6,3.8],fill={(0,1):GREENBG,(1,1):GREENBG,(2,1):GREENBG,(3,1):GREENBG})
h2("Manual validation steps (post-deploy on vantora-staging)")
b("Apply migrations 0293-0295 to staging (migrate-staging workflow or psql). Then:")
b("BL-1: as platform owner, run the drug-catalog refresh on a pharmacy that has products linked to drugs — succeeds, no FK error; product links intact.")
b("BL-2: create an invoice on a branch that previously errored — a fresh, non-colliding number is issued.")
b("BL-3/4: as a normal tenant user, attempt to RPC erp_seed_company_roles with another company's id — permission denied.")
b("BL-5: create two van load requests; numbers are STO-/VAN- sequential and unique.")
b("BL-6: double-click 'record collection' — exactly one receipt; balance reduced once.")
b("MJ-1: sign in as a viewer/staff and attempt to issue an invoice / post a voucher / adjust stock — blocked; salesman/accountant/warehouse_keeper still work.")

part("5. Safety & guardrails honoured")
tbl(["Rule","Honoured?","How"],[
 ["No destructive changes","Yes","No DROP/DELETE of data; only index/grant/counter changes + code"],
 ["No data deletion","Yes","BL-1 stopped deleting; no migration deletes rows"],
 ["No tenant data cleanup","Yes","Sequence repair touches counters only; no tenant rows altered"],
 ["Additive / idempotent","Yes","IF NOT EXISTS, REVOKE no-op, GREATEST upsert, UPSERT-on-PK"],
 ["Reversible","Yes","0293/0295 document exact reversal; 0294 is a forward-only counter bump (by nature)"],
 ["No P1 mixed in","Yes","Only BL-1..6 + MJ-1; all P1 items untouched"],
],widths=[2.0,1.0,3.4],fill={(i,1):GREENBG for i in range(6)})
para("Note: migration 0294 advances counters and is therefore forward-only (you cannot un-issue a number); it is "
     "non-destructive and idempotent, but unlike 0293/0295 it has no clean down-migration. This is the standard, safe "
     "sequence-repair approach.", italic=True, color=AMBER)

part("6. Re-scored pilot readiness")
para("With all six Blockers and the top separation-of-duties Major resolved, the platform clears the bar that previously "
     "blocked an external pilot. Remaining items are P1 (Should-fix-soon) and below — none is a functional outage or a "
     "cross-tenant breach.", bold=True)
tbl(["Dimension","Before","After","Note"],[
 ["Tenant isolation (RLS)","4/10","7/10","BL-3/BL-4 closed; M-tier cross-tenant reads remain (P1/P2)"],
 ["Data integrity / financial flows","5/10","8/10","BL-1/2/5/6 fixed; M-tier idempotency on create paths remains (P1)"],
 ["Permissions / separation of duties","5/10","8/10","MJ-1 gates added + tested; master-data gates remain (P1)"],
 ["Navigation / module visibility","8/10","8/10","Workflow-templates 404 + electrical gate are P1"],
 ["Auth / login / edge functions","9/10","9/10","Unchanged (solid)"],
 ["i18n / mobile","9/10","9/10","Unchanged (solid)"],
 ["Error handling","7/10","7/10","Raw-message leaks are P2"],
],widths=[2.4,0.7,0.7,2.6],
 fill={(0,2):GREENBG,(1,2):GREENBG,(2,2):GREENBG})
para("Overall: 5.5 / 10  ->  7.8 / 10 — PILOT-CONDITIONAL. The blockers are cleared; proceeding to a controlled pilot is "
     "reasonable once migrations 0293-0295 are applied to staging and the manual validation above passes. The remaining "
     "Majors (P1) should be scheduled as the next phase but do not block a controlled pilot.", bold=True, color=NAVY)
h2("Next phase (P1 — separate batch, on approval)")
b("MJ-4 drop/scope ts_* legacy tables; MJ-5 scope journey PII/GPS; MJ-2 workflow-templates flag; MJ-3 idempotency on "
  "order/PO/return/transfer create paths; MJ-8 credit/price master-data gates; MJ-6/MJ-7 atomic product codes.")
para("")
para("P0 remediation complete. No P1 work was started; it will be implemented as a separate, independently-reviewed phase "
     "on your approval.", bold=True, color=NAVY)

out="docs/audits/VANTORA-P0-Remediation-Report.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
