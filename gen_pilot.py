#!/usr/bin/env python3
"""VANTORA FMCG Pilot Readiness Audit (final) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.3,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("FMCG Pilot Readiness Audit"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Final — after P0 blockers, U-4 gating, Approval Queue & transfer request screens"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Branch claude/fmcg-sell-collect-loop · tsc clean · 1318 tests · build green · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Verdict")
para("VANTORA FMCG is PILOT-READY for a controlled pilot. All six P0 blockers are fixed and validated on staging; the "
     "sensitive-write gaps (pricing, credit) are closed (U-4); separation of duties is enforced and tested; and the "
     "field-management approval workflows are now completable END-TO-END through the unified Approval Queue plus the new "
     "request screens. No functional pilot blocker remains. Two P1 SECURITY items (legacy ts_* tables, journey PII scope) "
     "are recommended to fix early in the pilot window but do not block a controlled internal pilot.", bold=True, color=NAVY)
tbl(["Dimension","Score","Note"],[
 ["Core sell→collect→finance loop","9/10","Complete & validated"],
 ["Field-management approvals","9/10","End-to-end via Approval Queue + request screens"],
 ["Tenant isolation (RLS)","7/10","P0 cross-tenant writes closed; 2 P1 reads remain"],
 ["Permissions / separation of duties","9/10","MJ-1 + U-4 gates, tested"],
 ["Navigation / mobile","9/10","0 dead links; mobile==desktop; approver bottom tab"],
 ["i18n / usability","9/10","ar+en parity test-enforced"],
],widths=[2.6,0.8,3.0])
para("Overall readiness: 8.6 / 10 — GO for a controlled FMCG pilot.", bold=True, color=GREEN)

part("1. Completed workflows  (end-to-end by intended role)")
tbl(["Workflow","Request (UI)","Approve (UI)","Role(s)","Status"],[
 ["Sell → collect (invoice/POS → payment)","✓","n/a","salesman, cashier","Complete"],
 ["Van load: request → approve → load","✓ /field/van-sales","✓ /inventory/requests","salesman → whse/supervisor","Complete"],
 ["Purchase: PO → receive → AP post","✓","✓","procurement/whse → accountant","Complete"],
 ["Sales return: create → complete","✓","✓","cashier/salesman → mgr","Complete"],
 ["Collections → reconciliation","✓","✓","salesman → supervisor","Complete"],
 ["GL: voucher → post","✓","✓ accountant","accountant","Complete"],
 ["Day-close exception → approve","✓ field screen","✓ Approval Queue","salesman → supervisor","Complete"],
 ["Out-of-route visit → approve/reject","auto-logged","✓ Approval Queue","salesman → supervisor","Complete"],
 ["Customer transfer → approve","✓ /customers/transfer (NEW)","✓ Approval Queue","mgr → mgr","Complete"],
 ["Van transfer → approve/reject","✓ /inventory/van-transfer (NEW)","✓ Approval Queue","rep → supervisor","Complete"],
],widths=[2.1,1.5,1.4,1.4,0.9],
 fill={(i,4):GREENBG for i in range(10)})
para("All ten core FMCG workflows are now completable end-to-end in the UI by the intended role.", bold=True, color=GREEN)

part("2. Remaining gaps  (non-blocking)")
tbl(["#","Gap","Severity","Recommendation"],[
 ["G1","Legacy ts_* tables: open RLS + plaintext password (MJ-4)","Major (security)","Drop/scope before EXTERNAL exposure; safe to monitor in a controlled internal pilot"],
 ["G2","Journey PII/GPS cross-tenant read via SECURITY DEFINER (MJ-5)","Major (security)","Scope by tenant early in the pilot window"],
 ["G3","Trade-spend promo CREATION has no screen (approve/cancel exist)","Minor","trade_spend is flag-gated/foundation; add a create screen if enabling promotions"],
 ["G4","Electrical section has no module gate (MN-2)","Minor","Not relevant to FMCG pilot; fix in P1"],
 ["G5","Some admin actions leak raw DB errors (MN-1)","Minor","Route through friendlyDbError (P2)"],
 ["G6","Customer transfer = approve-only (no reject action)","Minor","By design today; add reject action if needed (would be new backend)"],
 ["G7","Company-admin user/permission model (M-2)","Minor","Keep boundary; optional 'Staff'→'Users & Roles' label"],
],widths=[0.4,3.3,1.0,2.1],
 fill={(0,2):AMBERBG,(1,2):AMBERBG})

part("3. Pilot blockers")
para("FUNCTIONAL BLOCKERS: NONE. All six P0 blockers (medicine catalog FK, invoice numbering, cross-tenant RBAC ×2, van "
     "numbering, collection idempotency) are fixed and validated on staging; the financial/stock separation-of-duties gap "
     "(MJ-1) and the price/credit gaps (U-4) are closed and tested.", bold=True, color=GREEN)
para("SECURITY (gate before EXTERNAL/multi-customer exposure, not a controlled-pilot blocker): G1 (ts_* legacy tables) and "
     "G2 (journey PII scope). Recommend fixing in the first pilot iteration.", color=AMBER)

part("4. Recommended pilot roles")
tbl(["Pilot role","Code role","Why"],[
 ["Company Admin","admin","Tenant setup, staff onboarding, oversight"],
 ["Branch Manager","branch_manager","Branch ops, approvals, purchasing"],
 ["Supervisor","supervisor","Field approvals (day-close, visit, transfers) via Approval Queue"],
 ["Salesman / Van rep","salesman","Sell, collect, request load/transfer, day close (Cash Van ≡ salesman)"],
 ["Accountant","accountant","Collections, GL posting, supplier payments, reports"],
 ["Warehouse","warehouse_keeper","Stock adjust/transfer/count, approve loads, receive POs"],
],widths=[1.6,1.4,3.6])
para("Note: 'Cash Van' is the salesman role (verified — not a separate role in code or DB). Viewer can be added as a "
     "read-only observer.", italic=True)

part("5. Recommended pilot test scenarios")
b("S1 — Sell & collect: salesman creates an invoice (price honored), issues it, records a payment; double-click the "
  "collection to confirm NO duplicate (BL-6).")
b("S2 — Invoice numbering: create invoices on a previously-failing branch; confirm sequential, no 'duplicate code' (BL-2).")
b("S3 — Medicine catalog: platform owner refreshes the drug list on a pharmacy with linked products; confirm success, "
  "links intact (BL-1).")
b("S4 — Day-close approval: salesman closes a day below coverage → supervisor approves it from the Approval Queue (mobile).")
b("S5 — Out-of-route visit: rep logs an out-of-route visit → supervisor approves/rejects with a comment in the queue.")
b("S6 — Customer transfer: manager submits /customers/transfer → it appears pending in the queue → another manager approves.")
b("S7 — Van transfer: rep submits /inventory/van-transfer (from/to + lines) → supervisor approves/rejects in the queue.")
b("S8 — Separation of duties: a Viewer attempts to issue an invoice / post a voucher / change a price or credit limit — "
  "all correctly BLOCKED (MJ-1 + U-4).")
b("S9 — Mobile: a supervisor uses ONLY the phone — bottom-nav Approvals tab → approve a day-close + a transfer end-to-end.")
b("S10 — Tenant isolation: a user in Tenant A cannot see Tenant B's customers/invoices/approvals (RLS).")
b("S11 — Arabic/English: switch locale; verify the queue, transfer forms, and approvals render correctly RTL/LTR.")

part("6. Overall readiness score")
para("8.6 / 10 — GO for a controlled FMCG pilot.", bold=True, color=NAVY)
b("Functionally complete: all ten core workflows are end-to-end; the field-management approval loop is fully exposed and "
  "mobile-accessible.")
b("Protected: P0 blockers fixed, separation of duties + price/credit gates enforced and regression-tested.")
b("Caveat to lift before broad/external rollout: the two P1 security items (G1 ts_* tables, G2 journey PII) — schedule in "
  "the first pilot iteration.")
para("")
para("Recommendation: proceed with the controlled FMCG pilot using the six roles above and scenarios S1–S11. Track feedback "
     "with the existing pilot templates; fix G1/G2 early. No new feature work is required to start the pilot.", bold=True, color=NAVY)

out="docs/audits/VANTORA-FMCG-Pilot-Readiness-Final.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
