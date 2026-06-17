#!/usr/bin/env python3
"""VANTORA Approval Architecture Assessment -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    for ln in t.split("\n"):
        p=d.add_paragraph(); r=p.add_run(ln); r.font.name="Consolas"; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0x22,0x22,0x22)
def tbl(headers,rows,widths=None,size=8.0,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ---- Cover ----
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Approval Architecture Assessment"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Should approvals move to a configurable rule engine? — a strategy decision before pilot expansion"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Type · Permission · Role · Branch scope · Amount threshold · Escalation — assessed against the live code · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---- Executive recommendation ----
h1("Executive recommendation")
para("Yes — VANTORA should standardise on a configurable, rule-driven approval engine, and it is closer to that than it "
     "looks. The platform ALREADY contains a configurable approval engine (the Phase-2 workflow engine) that supports "
     "conditional/threshold routing, role/admin/user assignees, parallel approvals with quorum, and SLA + escalation "
     "columns. The problem is not a missing engine — it is BIFURCATION: only 3 of the 9 workflows use it; the other 6 are "
     "hardcoded as one-off RPCs with a single permission check.", bold=True)
para("Recommendation: DO NOT build a new engine. CONVERGE the six hardcoded approvals onto the existing configurable "
     "engine, and extend that engine with a few missing primitives (permission-based assignee, branch scope, "
     "self-approval prohibition, and live escalation timers). This turns approvals into tenant-configurable rules — the "
     "single biggest lever for multi-tenant SaaS scalability, because new customers' governance becomes configuration, "
     "not code.", bold=True, color=GREEN)
h2("What approvals should be based on (the headline answer)")
para("A layered model, not a single dimension:")
b("PERMISSION is the capability gate — always enforced server-side, independent of UI. Keep this; it is already correct.", color=DARK)
b("RULE-BASED CONDITIONS (request type + amount threshold + branch scope) decide WHICH approval LEVEL is required.", color=DARK)
b("ROLE (resolved within branch scope) is the default way to find the approver for a level; PERMISSION-holder is the "
  "broader fallback. NAMED USERS only as an explicit override or escalation target — never as the primary mechanism.", color=DARK)
b("APPROVAL LEVELS (sequential steps with quorum) express multi-tier sign-off (e.g. Branch Manager → Company Admin).", color=DARK)
para("In one line: permission-gated, rule-routed, role-resolved, branch-scoped, level-stepped — with named users only as "
     "an escape hatch. Avoid hardcoded approver roles in code entirely.", bold=True, color=NAVY)

# ---- What exists ----
part("1 · What VANTORA already has (and what it lacks)")
para("The Phase-2 workflow engine (entity-agnostic, per-tenant definitions) already provides most of the target model:")
tbl(["Capability","Status today","Where"],
[
 ["Per-tenant workflow definitions (by key + entity)","Present","erp_workflow_definitions"],
 ["Multi-step / approval levels","Present","erp_workflow_steps.step_no"],
 ["Conditional routing (amount thresholds, etc.)","Present (when/op/value: eq/neq/gt/lt/in)","erp_workflow_steps.condition + erp_workflow_condition_met"],
 ["Assignee by role / company_admin / user","Present","approver_type, erp_workflow_resolve_users / _user_can_act"],
 ["Parallel approvals + quorum","Present","mode='parallel', required_approvals"],
 ["Audit trail on every decision","Present","erp_log_audit('decide','workflow_task')"],
 ["Mandatory comment / reject reason","Partial (app-layer per action)","—"],
 ["Assignee by PERMISSION (any holder of perm X)","MISSING","resolve_users handles role, not permission"],
 ["BRANCH SCOPE (approver in same branch as request)","MISSING","resolve_users scopes to company, not branch"],
 ["Cannot-approve-own-request guard","MISSING","not enforced anywhere"],
 ["Live escalation timers (escalate after N h)","MISSING (columns exist)","sla_hours / escalate_to seeded; no pg_cron driver"],
],widths=[3.0,2.4,1.2],
fill={(0,1):GREENBG,(1,1):GREENBG,(2,1):GREENBG,(3,1):GREENBG,(4,1):GREENBG,(5,1):GREENBG,
      (6,1):AMBERBG,(7,1):REDBG,(8,1):REDBG,(9,1):REDBG,(10,1):AMBERBG})
para("So the gap to a full configurable approval-rule engine is FOUR primitives — not a new engine.", bold=True, color=GREEN)

# ---- Per-workflow ----
part("2 · The nine workflows assessed")
para("For each: current implementation · current approver logic · hardcoded or configurable · recommended future state · "
     "migration complexity · risk · SaaS impact.")
tbl(["Workflow","Current implementation","Approver logic (today)","H/C","Migration"],
[
 ["Credit Limit Increase","Phase-2 engine workflow (2 steps, >50k senior review)","company_admin (config); permission credit.request.approve held but engine routes to admin","Configurable","Low (re-config rules)"],
 ["Customer Master Data Change","Engine change-request workflow + outcome handler","company_admin","Configurable","Low"],
 ["Customer Transfer","Bespoke RPC erp_approve_customer_transfer (approve-only)","Single perm customer.transfer","Hardcoded","Medium"],
 ["Out-of-Route Visit","Bespoke RPC erp_decide_visit_compliance","Single perm visit.approve_out_of_route","Hardcoded","Medium"],
 ["Day-Close Exception","Bespoke RPC erp_approve_day_close (approve-only)","Single perm day.approve_close_exception","Hardcoded","Medium"],
 ["Load Request","Bespoke RPC erp_approve_stock_request","Single perm stock_request.approve","Hardcoded","Medium"],
 ["Van Transfer","Bespoke RPC erp_approve/reject_van_transfer; auto-approve under value threshold","Single perm stock.transfer.approve + fmcg_settings threshold","Hardcoded","Medium"],
 ["Trade Spend","Status update on erp_trade_promotions (flag-gated)","Single perm pricing.rule.edit","Hardcoded","Medium"],
 ["Price Change","No approval — direct edit","Permission-gated only (pricing.manage)","None (n/a)","New (optional)"],
],widths=[1.35,2.0,2.0,0.9,1.0],
fill={(0,3):GREENBG,(1,3):GREENBG,(2,3):AMBERBG,(3,3):AMBERBG,(4,3):AMBERBG,(5,3):AMBERBG,(6,3):AMBERBG,(7,3):AMBERBG,(8,3):REDBG})
para("H/C = Hardcoded / Configurable. Six of nine are hardcoded single-permission gates with no thresholds, levels, branch "
     "scope or escalation — the strategic debt this assessment targets.", italic=True, color=GREY, size=8.6)

h2("Recommended future state — same for all")
para("Every workflow becomes a tenant-configurable definition on the unified engine: a request creates an instance "
     "(erp_workflow_start) that is routed by RULES (type → conditions on amount/branch → level → assignee by "
     "permission/role within branch scope), decided through ONE inbox, audited, and escalated on SLA. Domain logic "
     "(apply the credit limit, move the stock, close the day) stays in per-entity OUTCOME HANDLERS — the engine decides "
     "WHO and WHEN; the handler does WHAT. This is the existing handler pattern, generalised.")

# ---- Target architecture ----
part("3 · Target architecture — the rule model")
h2("Layered resolution (request → approver)")
b("1. CAPABILITY GATE — the action always re-checks a permission server-side (e.g. credit.request.approve). Unchanged; "
  "this is the security floor and must never depend on config.")
b("2. RULE MATCH — by request type + context: amount thresholds (gt/lt), branch scope, customer segment, channel. "
  "Decides the required approval LEVEL(s).")
b("3. ASSIGNEE RESOLUTION per level — permission-holder OR role, resolved WITHIN the request's branch (branch scope), "
  "with an optional named-user override and an escalation target.")
b("4. GUARDS — cannot approve own request; mandatory reject reason; quorum (N-of-M) for parallel levels.")
b("5. ESCALATION — if undecided after the SLA, auto-route to escalate_to (a higher level / named user).")
h2("Why permission+role+rule, not named users")
tbl(["Basis","Use it for","Why / caveat"],
[
 ["Permission","The capability floor + 'any holder can act' assignment","Survives org change; multi-tenant safe; already enforced"],
 ["Role","Default assignee resolution within a branch","Human-readable; maps to org chart; tenant-editable"],
 ["Approval levels","Multi-tier sign-off (threshold escalation)","Expresses governance (SoD) without code"],
 ["Rule conditions","Choosing the level (amount/branch/type)","The configurable heart; no redeploys per customer"],
 ["Named users","Override / escalation target ONLY","Brittle at scale (turnover); never the primary rule"],
],widths=[1.3,2.6,2.6])

h2("Proposed rule shape (illustrative, matches your example)")
code('{\n'
     '  "request_type": "credit_limit_increase",\n'
     '  "requester_scope": "same_branch",\n'
     '  "capability_permission": "credit.request.approve",\n'
     '  "can_approve_own": false,\n'
     '  "requires_reject_reason": true,\n'
     '  "levels": [\n'
     '    { "when": "amount <= 5000", "approver": { "role": "branch_manager", "scope": "branch" },\n'
     '      "escalate_after_h": 24, "escalate_to": { "role": "company_admin" } },\n'
     '    { "when": "amount > 5000",  "approver": { "permission": "credit.request.approve", "scope": "company" },\n'
     '      "quorum": 1, "escalate_after_h": 24, "escalate_to": { "user": "<finance_director_id>" } }\n'
     '  ]\n'
     '}')
para("This maps 1:1 onto the existing tables: each levels[] entry is an erp_workflow_steps row (condition = the `when`, "
     "approver_type/ref = the approver, sla_hours = escalate_after_h, escalate_to = the target). The only NEW pieces are "
     "approver_type='permission', a branch-scope flag, can_approve_own, and the escalation timer driver.", italic=True, color=GREY, size=9)

# ---- Migration ----
part("4 · Migration complexity & roadmap")
tbl(["Phase","Scope","Complexity","Risk"],
[
 ["P1 — Engine primitives","Add approver_type='permission'; branch-scope resolution; can_approve_own guard; mandatory reject reason in engine","Medium","Low (additive, behind config)"],
 ["P2 — Converge field approvals","Route Van Transfer, Load Request, Customer Transfer, Visit, Day-Close through erp_workflow_start + outcome handlers; keep RPCs as handlers","Medium-High","Medium (dual-path during cutover)"],
 ["P3 — Unify the inbox","One Approval Queue reads engine tasks for ALL types; retire bespoke queue readers + Approval Center","Medium","Low-Medium (UX)"],
 ["P4 — Escalation timers","pg_cron (or scheduled function) drives sla_hours → escalate_to; status 'escalated'","Medium","Medium (infra; idempotency)"],
 ["P5 — Tenant rule builder","Admin UI to edit definitions/levels/thresholds per tenant (builder-lite already exists)","Medium","Medium (mis-config guardrails)"],
],widths=[1.6,3.0,1.0,1.0])
para("Sequencing note: P1 is a prerequisite. P2 can be done one workflow at a time (lowest-risk first: Van Transfer / Load "
     "Request, which already have clean pending→decided status models). Credit-Limit and Customer-Data are already on the "
     "engine, so they only need P1 rule enrichment.", italic=True, color=GREY, size=9)

# ---- Risks ----
part("5 · Risks & mitigations")
tbl(["Risk","Mitigation"],
[
 ["Dual-write / double-decision during cutover","Migrate one workflow at a time; feature-flag each; the engine instance is the single source of truth, RPC becomes the outcome handler only"],
 ["Tenant mis-configuration locks approvals (no valid approver)","Validate definitions on save (every level must resolve ≥1 user); fall back to company_admin; 'no approver' alert"],
 ["Inbox performance at scale","Today's inbox fetches up to 2000 tasks and filters in memory — replace with an indexed assignee-resolution query before high volume"],
 ["Self-approval / segregation-of-duties gaps","Add can_approve_own=false guard in the engine (compliance requirement for finance approvals)"],
 ["Escalation infra (timers) reliability","Idempotent scheduled job; escalate only pending+overdue; audit every escalation"],
 ["Outcome-handler correctness (apply the real change)","Reuse the existing, tested domain RPCs as handlers; engine only flips state + calls handler"],
],widths=[2.7,3.9])

# ---- SaaS ----
part("6 · SaaS scalability impact")
b("Onboarding becomes configuration, not code: a new distributor's approval policy (thresholds, who signs off, "
  "escalation) is set per tenant — no per-customer branches or deploys.", color=GREEN)
b("Enterprise governance unlocks deals: multi-level sign-off, segregation of duties, branch-scoped approval and "
  "audit are exactly what larger FMCG buyers require — and become a sales feature.", color=GREEN)
b("One inbox, one audit model across all approvals → lower support load, consistent UX, simpler training.", color=GREEN)
b("Caution — multi-tenant blast radius: a shared engine means a bug or a bad global default hits every tenant. "
  "Mitigate with per-tenant definitions, validation-on-save, and the capability-permission floor that no config can "
  "bypass.", color=AMBER)
b("Caution — performance: assignee resolution and the inbox must be query-driven and indexed before volume grows; the "
  "current in-memory filter is fine for a pilot, not for scale.", color=AMBER)

# ---- Decision ----
part("7 · The decision & immediate (pilot) stance")
para("Strategic decision: ADOPT the configurable rule-engine direction by CONVERGING onto the engine VANTORA already "
     "has, extended with four primitives (permission assignee, branch scope, self-approval guard, escalation timers). Do "
     "NOT build a parallel engine, and do NOT keep adding hardcoded single-permission approvals — every new one is future "
     "migration debt.", bold=True, color=NAVY)
para("For the pilot NOW (no rebuild): the existing field approvals are safe and correct as-is. The only near-term, "
     "low-risk cleanups identified separately remain valid and are compatible with this direction:")
b("Align the Trade-Spend visibility gate with its action permission (pricing.rule.edit).")
b("Give credit.request.approve holders a working UI path (interim: configure the credit workflow's first level to a "
  "role they hold + surface role-tasks in the inbox; this is the engine doing exactly what P1 generalises).")
b("Fold Load Request + Van Reconciliation into the one Approval Queue.")
para("These are stepping-stones toward the target architecture, not throwaway work — each is a small move in the same "
     "direction. The full convergence is a post-pilot programme (Phases P1-P5), not a pre-pilot blocker.",
     bold=True, color=GREEN)

out="docs/audits/VANTORA-Approval-Architecture-Assessment.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
