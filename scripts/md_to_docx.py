#!/usr/bin/env python3
"""Consolidate the VANTORA handover markdown into a single Word document.

Lightweight Markdown -> .docx: headings, tables (with shaded header), bullet/
numbered lists, code fences, blockquotes, horizontal rules, and inline
**bold** / `code`. Not a full CommonMark engine — tuned for these docs.
"""
import re
import sys
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BRAND = RGBColor(0x1F, 0x49, 0x7D)
HEADER_FILL = "1F497D"


def shade(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    sh = OxmlElement("w:shd")
    sh.set(qn("w:val"), "clear")
    sh.set(qn("w:fill"), fill)
    tcPr.append(sh)


def add_inline(par, text):
    """Render inline **bold** and `code` into runs."""
    # Split on bold / code while keeping delimiters.
    for chunk in re.split(r"(\*\*.+?\*\*|`[^`]+`)", text):
        if not chunk:
            continue
        if chunk.startswith("**") and chunk.endswith("**"):
            r = par.add_run(chunk[2:-2]); r.bold = True
        elif chunk.startswith("`") and chunk.endswith("`"):
            r = par.add_run(chunk[1:-1]); r.font.name = "Consolas"; r.font.size = Pt(9.5)
        else:
            par.add_run(chunk)


def emit_table(doc, rows):
    cols = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=cols)
    t.style = "Light Grid Accent 1"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row in enumerate(rows):
        cells = t.add_row().cells
        for j in range(cols):
            val = row[j] if j < len(row) else ""
            p = cells[j].paragraphs[0]
            add_inline(p, val.strip())
            if i == 0:
                shade(cells[j], HEADER_FILL)
                for rr in p.runs:
                    rr.bold = True
                    rr.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    doc.add_paragraph()


def convert(md, doc):
    lines = md.split("\n")
    i = 0
    in_code = False
    code_buf = []
    while i < len(lines):
        line = lines[i]
        # code fence
        if line.strip().startswith("```"):
            if in_code:
                p = doc.add_paragraph()
                r = p.add_run("\n".join(code_buf))
                r.font.name = "Consolas"; r.font.size = Pt(9)
                p.paragraph_format.left_indent = Pt(12)
                code_buf = []; in_code = False
            else:
                in_code = True
            i += 1; continue
        if in_code:
            code_buf.append(line); i += 1; continue

        # table: a | line followed by a |---| separator
        if line.lstrip().startswith("|") and i + 1 < len(lines) and re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i + 1]):
            rows = []
            while i < len(lines) and lines[i].lstrip().startswith("|"):
                if re.match(r"^\s*\|[\s:|-]+\|\s*$", lines[i]):
                    i += 1; continue
                cells = [c for c in lines[i].strip().strip("|").split("|")]
                rows.append(cells); i += 1
            emit_table(doc, rows); continue

        stripped = line.strip()
        if not stripped:
            i += 1; continue
        if re.match(r"^---+$", stripped):
            i += 1; continue
        # headings
        m = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if m:
            level = len(m.group(1)); txt = m.group(2)
            h = doc.add_heading(level=min(level, 4))
            add_inline(h, txt)
            for r in h.runs:
                if level <= 2:
                    r.font.color.rgb = BRAND
            i += 1; continue
        # blockquote
        if stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            r = p.add_run("")
            add_inline(p, stripped.lstrip("> ").strip())
            for rr in p.runs:
                rr.italic = True
            i += 1; continue
        # bullets / numbered
        mb = re.match(r"^(\s*)[-*]\s+(.*)$", line)
        mn = re.match(r"^(\s*)\d+\.\s+(.*)$", line)
        if mb:
            p = doc.add_paragraph(style="List Bullet")
            add_inline(p, mb.group(2)); i += 1; continue
        if mn:
            p = doc.add_paragraph(style="List Number")
            add_inline(p, mn.group(2)); i += 1; continue
        # plain paragraph
        p = doc.add_paragraph()
        add_inline(p, stripped)
        i += 1
    return doc


def main():
    out = sys.argv[1]
    title = sys.argv[2]
    files = sys.argv[3:]
    doc = Document()
    # Title page
    doc.add_paragraph()
    t = doc.add_paragraph(); t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run(title); r.bold = True; r.font.size = Pt(26); r.font.color.rgb = BRAND
    s = doc.add_paragraph(); s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = s.add_run("FMCG Van-Sales · Pilot, Reference Tenant & Certification Package")
    rs.font.size = Pt(13); rs.italic = True
    d = doc.add_paragraph(); d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d.add_run("Generated 2026-06-10 · PR #311").font.size = Pt(11)
    doc.add_page_break()
    for idx, f in enumerate(files):
        with open(f, encoding="utf-8") as fh:
            convert(fh.read(), doc)
        if idx != len(files) - 1:
            doc.add_page_break()
    doc.save(out)
    print("wrote", out)


if __name__ == "__main__":
    main()
