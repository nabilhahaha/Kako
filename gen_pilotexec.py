#!/usr/bin/env python3
"""VANTORA Pilot Execution Report -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.0,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Pilot Execution Report"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Pilot tenant provisioned + scenarios S1–S11 executed on the live staging database"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("vantora-staging (rsjvgehvastmawzwnqcs) · company VANTORA Pilot FMCG (DEMO) · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Method & honesty note")
para("The pilot tenant, demo data and user accounts were CREATED on the live vantora-staging database. Scenarios were "
     "EXECUTED at the data/RPC layer (calling the real RPCs and queries as the intended user via JWT context, so RLS, "
     "tenant isolation and in-RPC authorization are genuinely exercised). Eight scenarios were run live with real results "
     "below. Three scenarios (S3 pharmacy catalog, S9 mobile gesture, S11 RTL rendering) are pure-UI or out-of-FMCG-scope; "
     "they are marked CODE-VERIFIED with a recommended operator click-through — NOT fabricated as executed.", bold=True)

part("1. Pilot tenant & data created  (items 1–9)")
tbl(["#","Item","Created","Detail"],[
 ["1","Pilot Tenant","✓","VANTORA Pilot FMCG (DEMO) · id 612af0bd… · business_type=fmcg · 21 modules + 22 roles auto-seeded"],
 ["2","Demo FMCG data","✓","1 branch (PILOT) + 2 warehouses (main + van) + stock loaded (1000/product)"],
 ["3","Sample Products","✓","8 FMCG SKUs (oil, sugar, rice, tea, soap, paste, pasta, detergent) with cost/sell/tax"],
 ["4","Sample Customers","✓","5 grocery/market customers, approved/active, credit limits, assigned to the salesman"],
 ["5","Sample Warehouse","✓","Pilot Main Warehouse (PILOT-WH) + Pilot Van (PILOT-VAN)"],
 ["6","Sample Salesman","✓","salesman@pilot.test (role salesman)"],
 ["7","Sample Supervisor","✓","supervisor@pilot.test (role supervisor)"],
 ["8","Sample Accountant","✓","accountant@pilot.test (role accountant)"],
 ["9","Sample Warehouse Keeper","✓","warehouse@pilot.test (role warehouse_keeper)"],
],widths=[0.3,1.7,0.7,4.3],fill={(i,2):GREENBG for i in range(9)})
para("Plus a Company Admin (admin@pilot.test, role admin) so the tenant has an owner. All accounts: password test.123, email "
     "confirmed, one identity each, bcrypt-verified.", italic=True)

part("2. Scenario execution  (S1–S11)")
para("Legend: PASS = executed live, expected = actual · VERIFIED = code-verified, operator UI walkthrough recommended.", italic=True)
def srow(s,scn,exp,act,scr,iss,sev,rec):
    return [s,scn,exp,act,scr,iss,sev,rec]
rows=[
 srow("S1","Collection double-submit (idempotency, BL-6)","One receipt; no double balance","PASS — 2 calls, same key → 1 collection created","Collections / POS payment","None","—","Ship as-is"),
 srow("S2","Invoice numbering (BL-2)","Sequential, no duplicate code","PASS — INV-PILOT-000001/2/3","Sales · Invoices","None","—","Ship as-is"),
 srow("S3","Medicine catalog refresh (BL-1)","UPSERT, links intact","VERIFIED — pharmacy feature; not in FMCG pilot; UPSERT fix validated in P0","Pharmacy · Catalog","Out of FMCG scope","—","Exercise on a pharmacy tenant"),
 srow("S4","Day-close exception approval","pending → closed by supervisor","PASS — close_status 'pending_approval' → 'closed'","Field close + Approval Queue","None","—","Ship as-is"),
 srow("S5","Out-of-route visit approval","pending → approved","PASS — status 'pending_approval' → 'approved'","Approval Queue","None","—","Ship as-is"),
 srow("S6","Customer transfer → approve","pending → applied","PASS — request (supervisor) → approve (admin) → 'applied'","/customers/transfer + Approval Queue","None","—","Ship as-is"),
 srow("S7","Van transfer → approve","pending → completed; stock moved","PASS — request (salesman) → approve (supervisor) → 'completed'","/inventory/van-transfer + Approval Queue","None","—","Ship as-is"),
 srow("S8","Separation of duties","Viewer blocked; salesman sell/collect only","PASS — viewer/staff: none; salesman: sell/collect only (no GL/price/credit/approve)","All write screens","None","—","Ship as-is"),
 srow("S9","Mobile approver flow","Approver acts from phone","VERIFIED — bottom-nav Approvals tab + More drawer mirror sidebar","Mobile bottom nav","None (UI)","—","Operator device walkthrough"),
 srow("S10","Tenant isolation (RLS)","Sees only own tenant","PASS — salesman sees 5/190 customers, 0 of Nile's 154, 0 cross-tenant invoices","All list screens","None","—","Ship as-is"),
 srow("S11","Arabic / English","Correct RTL/LTR","VERIFIED — ar+en parity test-enforced; queue/forms have both","All screens","None (UI)","—","Operator locale walkthrough"),
]
tbl(["#","Scenario","Expected","Actual (live)","Screens","Issues","Sev","Recommendation"],
    [r for r in rows],widths=[0.3,1.5,1.2,1.7,1.1,0.7,0.4,1.0],size=7.4,
    fill={(i,3):(GREENBG if rows[i][3].startswith("PASS") else BLUEBG) for i in range(len(rows))})

part("3. Results summary")
tbl(["Outcome","Count","Scenarios"],[
 ["PASS (executed live)","8","S1, S2, S4, S5, S6, S7, S8, S10"],
 ["VERIFIED (code; operator UI walkthrough)","3","S3 (pharmacy), S9 (mobile), S11 (i18n)"],
 ["FAIL","0","—"],
 ["Issues found","0","No defects surfaced during execution"],
],widths=[2.6,0.8,3.0],fill={(0,0):GREENBG,(2,0):GREENBG,(3,0):GREENBG})
para("Every executable workflow passed on the live pilot tenant with expected results, and tenant isolation + separation of "
     "duties were genuinely enforced (not just asserted). No defect was found.", bold=True, color=GREEN)

part("4. What the live execution proves")
b("The P0 fixes hold on real data: invoice numbering is sequential (no duplicate-code), and the collection RPC is idempotent "
  "(double-submit = one receipt).")
b("The approval workflows transition correctly end-to-end through the real RPCs, performed AS the intended approver role.")
b("RLS tenant isolation is real: the pilot salesman cannot see any of the 185 other-tenant customers or 1,106 other-tenant "
  "invoices.")
b("Separation of duties is enforced at the permission layer: viewer/staff hold no sensitive capability; the salesman can "
  "sell/collect but cannot post the GL, set prices/credit, or approve loads.")

part("5. Recommendation")
para("PROCEED with the controlled FMCG pilot. The pilot tenant is provisioned and operationally validated. Remaining steps "
     "are operator-side UI confirmation (S9 mobile, S11 locale) and exercising the pharmacy catalog (S3) only if piloting the "
     "pharmacy pack. The two P1 security items (legacy ts_* tables, journey PII scope) should be scheduled in the first "
     "pilot iteration but did not affect any executed FMCG scenario.", bold=True, color=NAVY)
h2("Pilot login credentials (staging)")
tbl(["Role","Email","Password"],[
 ["Company Admin","admin@pilot.test","test.123"],
 ["Salesman","salesman@pilot.test","test.123"],
 ["Supervisor","supervisor@pilot.test","test.123"],
 ["Accountant","accountant@pilot.test","test.123"],
 ["Warehouse Keeper","warehouse@pilot.test","test.123"],
],widths=[2.0,2.6,1.4])
para("")
para("Operational validation complete — no new feature development was performed; this was tenant provisioning + live "
     "scenario execution only.", bold=True, color=NAVY)

out="docs/audits/VANTORA-Pilot-Execution-Report.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
