#!/usr/bin/env python3
"""VANTORA Multi-UoM U1+U2 — Foundation Delivered -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def bl(t,size=9.3,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.2,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Multi-UoM — U1 & U2 Foundation"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Shared company-wide engine + transaction-level UoM capture — additive, reversible, flagged"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("No sell/buy/returns/reporting wiring (U3-U6 not started) · base-unit invariant preserved · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Summary")
para("U1 (Master & Factors) and U2 (Transaction Capture) are delivered as the additive, reversible, flag-gated foundation "
     "for company-wide Multi-UoM. The pharmacy multi-UoM engine is now promoted to a shared company-wide capability, FMCG "
     "products can hold Case/Inner/Piece hierarchies, and every transaction line can record which UoM was used — WITHOUT "
     "changing the base-unit inventory invariant and WITHOUT wiring any sell or buy flow (U3-U6 deliberately not started).",
     bold=True)
para("Nothing is activated: the company-wide flag platform.multi_uom is OFF for the pilot, the new line columns are "
     "nullable and dormant, and no transaction reads them yet. Current pilot behaviour is unchanged.", bold=True, color=GREEN)

h1("U1 — Master & Factors (promote the engine company-wide)")
h2("What shipped")
bl("Company-wide feature flag platform.multi_uom (platform pack, 'inventory' domain) added to the feature catalog with "
   "ar/en labels — the vertical-agnostic twin of pharmacy.multi_unit_support. The engine now turns on for FMCG / "
   "distribution / any pack, not just pharmacy.")
bl("multiUomEnabled(flags) helper (src/lib/erp/uom.ts): true when platform.multi_uom OR the legacy "
   "pharmacy.multi_unit_support is on — so pharmacy keeps working unchanged and FMCG gains the same capability.")
bl("The engine core (uom.ts / uom-server.ts / uom-rules.ts) was already pure + platform-wide, and the management screen "
   "/settings/uom is gated by the generic uom.manage permission — so the master-data surface is already company-wide.")
bl("FMCG products prepared: example Case/Inner/Piece hierarchies seeded on pilot products (Inner ×6, Carton ×24, base "
   "Piece) to validate end-to-end.")
h2("Validation (staging, real data)")
tbl(["Check","Result"],
[
 ["erp_uom_to_base(carton, 2)","48 base (2 × 24)"],
 ["erp_uom_to_base(inner, 3)","18 base (3 × 6)"],
 ["erp_uom_to_base(piece, 5)","5 base (factor 1)"],
 ["erp_uom_to_base(unknown unit, 2)","2 (safe factor-1 fallback)"],
 ["erp_resolve_price(carton) fallback","672 = base 28 × 24 (UoM-aware price works)"],
 ["multiUomEnabled (unit tests)","off by default; on via platform.multi_uom or pharmacy flag"],
],widths=[3.4,3.0],fill={(i,1):GREENBG for i in range(6)})

h1("U2 — Transaction Capture (additive line columns)")
h2("What shipped")
bl("Migration 0304 adds three NULLABLE columns to the seven transaction line tables (invoice, PO, sales order, stock "
   "movement, van transfer, sales return, purchase return): entered_uom, entered_qty, uom_factor.")
bl("The existing quantity / received_qty columns REMAIN the BASE quantity — so all stock, finance and reporting logic is "
   "unchanged. NULL entered_uom = base unit (legacy), so existing rows and single-UoM tenants are unaffected.")
bl("Shared capture helper lineUomFields(units, enteredUom, enteredQty) (src/lib/erp/uom-capture.ts): returns the snapshot "
   "columns + the BASE qty (baseQty = enteredQty × factor). Pure + unit-tested; DORMANT until U3/U4 adopt it.")
h2("Validation")
tbl(["Check","Result"],
[
 ["3 UoM columns added to all 7 line tables","Confirmed via information_schema"],
 ["lineUomFields(carton, 2) on a 24-factor product","{entered_uom:carton, entered_qty:2, uom_factor:24, baseQty:48}"],
 ["lineUomFields(base / empty / unknown)","base passthrough (snapshot null, baseQty = entered)"],
 ["Full suite / build","1358 tests passed · build green · tsc clean"],
],widths=[3.4,3.0],fill={(i,1):GREENBG for i in range(4)})

h1("Rollback")
bl("U1 flag: remove the platform.multi_uom catalog entry + helper; pharmacy path untouched. Seeded product UoMs are "
   "deletable (additive demo data).")
bl("U2 columns: drop entered_uom / entered_qty / uom_factor from the seven tables. No data is transformed; quantity was "
   "never changed.")
bl("Everything is reversible with zero impact on the base-unit invariant or existing transactions.")

h1("Pilot impact")
para("None functionally. platform.multi_uom is OFF, the line columns are dormant, and no transaction reads or writes them. "
     "The only pilot-data change is additive + inert: three demo products now carry example Carton/Inner/Piece UoMs (visible "
     "only in /settings/uom), seeded to validate the engine. Documented per the recoverability guardrail.", bold=True)

h1("Scope honored / next")
bl("Delivered: U1 (master + factors + company-wide promotion) and U2 (transaction capture). Reviewed-ready.", color=GREEN)
bl("NOT started (as instructed): U3 Sell Path, U4 Buy Path, U5 Returns/Transfers, U6 Reporting.", color=AMBER)
para("The foundation is in place: when you approve, U3 wires van-sell/POS to lineUomFields + the UoM-aware price (flagged, "
     "staging-validated) — no schema churn needed, just adoption of the dormant helper + columns.", bold=True, color=NAVY)

out="docs/audits/VANTORA-Multi-UoM-U1-U2-Foundation.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
