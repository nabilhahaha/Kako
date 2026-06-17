#!/usr/bin/env python3
"""VANTORA Pilot FMCG — UAT Access & Execution Guide -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def step(n,t,size=9.6):
    p=d.add_paragraph(); r=p.add_run(f"{n}. "); r.bold=True; r.font.color.rgb=NAVY; r.font.size=Pt(size); r2=p.add_run(t); r2.font.size=Pt(size)
def bl(t,size=9.3,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.4,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Pilot FMCG — UAT Access & Execution Guide"); r.bold=True; r.font.size=Pt(21); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("A complete, ready-to-test FMCG distribution environment with Multi-UoM"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Tenant: VANTORA Pilot FMCG (DEMO) · Van Sales + Multi-UoM active · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Environment summary")
tbl(["Item","Configured"],
[
 ["Company","VANTORA Pilot FMCG (DEMO)"],
 ["Branch","Pilot Branch"],
 ["Main warehouse","PILOT-WH (Pilot Main Warehouse)"],
 ["Van","PILOT-VAN — assigned to the Salesman, stocked 20 cartons (480 base) of each of 8 SKUs"],
 ["Customers","6 (5 approved) — e.g. Al Nour Grocery (PILOT-C01), City Mini Market (PILOT-C03), Corner Shop (PILOT-C05)"],
 ["Products","8 FMCG SKUs, each with Piece / Inner (×6) / Carton (×24) units — e.g. Sunflower Oil 1L (PILOT-P01), White Sugar 1kg (PILOT-P02), Egyptian Rice 1kg (PILOT-P03)"],
 ["Van Sales","ENABLED (company toggle)"],
 ["Multi-UoM (platform.multi_uom)","ENABLED for this tenant"],
],widths=[1.8,4.7],fill={(3,1):GREENBG,(6,1):GREENBG,(7,1):GREENBG})
para("One deployment setting to confirm: Van Sales also requires the env flag KAKO_VAN_SALES=1 in the hosting (Vercel) "
     "environment. If the Van-Sell screen shows 'not found', set KAKO_VAN_SALES=1 and redeploy — everything else is ready.",
     bold=True, color=AMBER, size=9.5)

part("1 · Login credentials")
para("All users share the demo password below. URL: your pilot preview (…vercel.app). If a password was rotated, an admin "
     "can reset it from Settings → Staff/Users.")
tbl(["Role","Email","Password","Lands on","Use it to test"],
[
 ["Company Admin","admin@pilot.test","test.123","Dashboard","Configure UoMs, pricing, enable features, see everything"],
 ["Warehouse Keeper","warehouse@pilot.test","test.123","Inventory / Load Requests","Product UoM setup, stock, van loading, approvals"],
 ["Supervisor","supervisor@pilot.test","test.123","Approval Queue","Approve day-close / visits / transfers; coverage"],
 ["Salesman (Van Rep)","salesman@pilot.test","test.123","Today","VAN-SELL by carton/piece, collect cash"],
 ["Branch Manager","branchmgr@pilot.test","test.123","Branch cockpit","Approvals, purchasing, reports"],
 ["Accountant","accountant@pilot.test","test.123","Collections","Collections, AR aging, vouchers"],
 ["Viewer","viewer@pilot.test","test.123","Dashboard","Read-only reports"],
],widths=[1.3,1.7,0.9,1.3,1.6],
fill={(3,0):GREENBG,(1,0):GREENBG,(2,0):GREENBG})

part("2 · Pilot execution guide (the FMCG loop)")
para("Run the chain end-to-end. The headline test is the Multi-UoM Van-Sell.", bold=True)

h2("A · Product & UoM setup (Warehouse Keeper or Admin)")
step(1,"Login as warehouse@pilot.test. Go to Settings → Units of Measure. Pick a product (e.g. White Sugar 1kg) — confirm it has Piece, Inner (×6), Carton (×24).")
step(2,"Go to Products → edit a product → the 'Units & Selling' section: confirm Default Sell Unit / Purchase Unit / Sell Mode / Allow Fractional are editable (this section is visible only to uom.manage roles).")
step(3,"Optional: Sales → Price Book — set a per-UoM price (e.g. a Carton price). Leave blank to price as piece × 24.")

h2("B · Inventory & Van load")
step(4,"As Warehouse Keeper, open Inventory — confirm Main Warehouse stock and that PILOT-VAN already holds 20 cartons of each SKU (pre-loaded). (In a fresh run, the rep raises a Load Request and the keeper approves it.)")

h2("C · Van-Sell with UoM (Salesman — the main test)")
step(5,"Login as salesman@pilot.test. Open Van Sales → Sell (the field sell screen).")
step(6,"Pick a customer (e.g. Al Nour Grocery). Add a product (e.g. White Sugar). In the line, use the UNIT selector and choose Carton (×24).")
step(7,"Enter quantity 2 (cartons). Confirm the price shows the carton price (piece price × 24, or the Price-Book carton price).")
step(8,"Review → Issue. EXPECT: invoice created; the line records 2 cartons but stock decrements 48 (base); van stock drops by 48 for that SKU.")
step(9,"Sell another line in Pieces (base) on the same or a new sale — confirm it behaves normally.")

h2("D · Collection (Salesman or Accountant)")
step(10,"On the same visit/customer, collect cash (full or partial) against the invoice. Confirm the customer balance and invoice status update (issued → partially_paid / paid).")

h2("E · Reporting")
step(11,"As Admin/Accountant, open Reports / Sales — confirm the sale appears with the correct value; quantities are reported in BASE units (e.g. 48), which is expected (per-UoM reporting is a later phase).")

h2("F · Approvals (Supervisor) — optional")
step(12,"Trigger an exception (e.g. close the day below coverage, or an out-of-route visit) and approve it as supervisor@pilot.test from the Approval Queue.")

part("3 · What to verify (acceptance checklist)")
tbl(["Area","Pass criteria"],
[
 ["UoM picker","Carton/Inner/Piece selectable on van-sell (and POS) lines; only for uom.manage-configured products"],
 ["Conversion","2 cartons → stock −48 base; line shows entered unit + base quantity"],
 ["Pricing","Carton price = piece × 24 (or the Price-Book carton special)"],
 ["Stock invariant","All balances stay in base units; no negative van stock unless allowed"],
 ["Collection","Balance + invoice status update correctly"],
 ["Reporting","Sale value correct; quantities in base units"],
 ["Permissions","Salesman cannot reach Settings → Units of Measure (no master-data access)"],
 ["Mobile","Van-Sell (incl. the unit picker) usable on a phone"],
],widths=[1.7,4.8])

part("4 · Notes & rollback")
bl("Multi-UoM is enabled for THIS tenant only (platform.multi_uom). Turn it OFF in Settings → Features to revert to "
   "base-only selling instantly.")
bl("Disabled by design (not part of this UAT): buy-by-UoM (receiving), returns/transfers UoM, and reporting-by-UoM.")
bl("If Van-Sell 404s: set KAKO_VAN_SALES=1 in the deployment env (Vercel) and redeploy.")
bl("The van is pre-stocked for convenience; to test the full load workflow, empty the van and have the rep raise a Load "
   "Request for the keeper to approve.")
para("Everything here is reversible: feature flags toggle off, the van/stock are demo data, and no production system is "
     "involved. Happy testing.", italic=True, color=GREY, size=9)

out="docs/audits/VANTORA-Pilot-UAT-Access-Guide.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
