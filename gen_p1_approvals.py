#!/usr/bin/env python3
"""VANTORA P1 Approval Activation — Implementation & Rollback -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def bl(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    for ln in t.split("\n"):
        p=d.add_paragraph(); r=p.add_run(ln); r.font.name="Consolas"; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0x22,0x22,0x22)
def tbl(headers,rows,widths=None,size=8.2,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("P1 Approval Activation"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Credit-Limit · Trade-Spend · Price-Change on the configurable engine — implementation, validation & rollback"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Feature-flagged · backward-compatible · validated on staging · operational FMCG workflows untouched · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Summary")
para("P1 puts the three low-risk commercial approvals — Credit-Limit, Trade-Spend, Price-Change — on the configurable "
     "workflow engine with the new governance foundation (permission-based approvers, amount thresholds, self-approval "
     "block, mandatory reject reason, full audit). Every operational FMCG workflow (day-close, visit, customer transfer, "
     "load request, van transfer, van reconciliation) is UNTOUCHED, as instructed. P2-P5 were not started.", bold=True)
para("Activation is per-workflow feature-flagged and DEFAULT OFF, so applying this release changes nothing until a flag is "
     "switched on. The new engine paths were validated end-to-end on staging.", bold=True, color=GREEN)

h2("Requirements coverage")
tbl(["Requirement","Status","How"],
[
 ["Self-approval blocked","DONE","block_self_approval flag on the 3 definitions; enforced in erp_workflow_decide"],
 ["Reject reason required","DONE","require_reject_reason flag; enforced in engine + decideTask"],
 ["Full audit trail","DONE","erp_log_audit('decide',...) per decision; request/submit audited"],
 ["Threshold support","DONE","Credit: <=5k permission approver, >5k adds company-admin co-sign (condition gt 5000)"],
 ["Permission-based approvers","DONE","approver_type='permission' (credit.request.approve / pricing.manage)"],
 ["Backward-compatible rollout","DONE","flags default OFF preserve legacy paths exactly"],
 ["Feature-flagged activation","DONE","KAKO_APPROVAL_CREDIT / _TRADE_SPEND / _PRICE_CHANGE"],
 ["Rollback path documented","DONE","see section 'Rollback'"],
],widths=[2.1,0.8,3.5],
fill={(i,1):GREENBG for i in range(8)})

part("1 · What was implemented")
h2("Foundation already in place (Phase 0)")
bl("0296 permission resolution + governance flag columns; 0297 governance enforcement (faithful superset of decide); "
   "inbox surfaces role/permission tasks. All additive + dormant.")
h2("P1 changes")
tbl(["Area","Change"],
[
 ["supabase/migrations/0298","Seeds 3 global definitions: credit_limit_approval_v2 (2 levels), trade_spend_approval, price_change_approval — all with governance flags on. Idempotent; dormant until flagged."],
 ["src/lib/erp/approval-flags.ts","3 env flags (default OFF) + creditWorkflowKey() selector."],
 ["customers/actions.ts","requestCreditLimitChange routes to v2 key when KAKO_APPROVAL_CREDIT on, else legacy."],
 ["distribution/trade-spend/actions.ts","submitTradeSpendForApproval() starts the engine (flag-gated); legacy approve/cancel kept for flag-off."],
 ["products/actions.ts","requestPriceChange() stages erp_price_change_requests + starts engine (flag-gated, additive; direct edits unchanged)."],
 ["workflow-handlers.ts","Outcome handlers for trade_promotion (approve→approved / reject→cancelled) and price_change_request (apply sell_price on approve)."],
 ["approvals-access.ts + queue/page.tsx + navigation.ts","Workflow Inbox reachable by P1 approvers (credit.request.approve / pricing.manage) — UI only."],
 ["Tests","approval-flags + approvals-access + workflow-inbox unit tests."],
],widths=[2.1,4.5])

h2("Credit-limit routing (the threshold model)")
code('L1  (always)        approver_type=permission  ref=credit.request.approve   SLA 24h  esc company_admin\n'
     'L2  (amount > 5000)  approver_type=company_admin                              SLA 24h\n'
     'flags: block_self_approval=true, require_reject_reason=true')
para("<=5,000 → one permission approver. >5,000 → permission approver THEN a company-admin co-sign. Tenant-editable.",
     italic=True, color=GREY, size=9)

part("2 · Staging validation (end-to-end)")
para("Exercised directly against the engine on staging as the real pilot users, then cleaned up. All passed:", bold=True)
tbl(["Test","Result"],
[
 ["Permission resolution: resolve approvers of credit.request.approve","accountant + branch_manager + admin (exact)"],
 ["Self-approval (starter approves own)","BLOCKED — 'cannot approve your own request'"],
 ["Unauthorized (salesman decides)","BLOCKED — 'not authorized to decide this task'"],
 ["Reject with empty reason","BLOCKED — 'reject reason required'"],
 ["Approve 3,000 by a different permission holder","Instance + task approved (single level)"],
 ["Threshold: 8,000 request","After L1 approve → L2 company_admin task created; instance pending"],
 ["L2 company-admin co-sign","Instance approved (completed)"],
 ["Audit trail","erp_log_audit decide row per decision"],
],widths=[3.6,3.0],
fill={(i,1):GREENBG for i in range(8)})
para("Test records were removed afterwards; the three global definitions remain. Existing legacy workflows are unaffected "
     "(their flags default to false, so the new guards never fire).", italic=True, color=GREY, size=9)

part("3 · Activation runbook")
para("Apply in the normal migration pipeline (order matters): 0296 → 0297 → 0298. All additive; nothing changes until a "
     "flag is set. Then enable per workflow by setting the env var to 1 and redeploying:")
code('KAKO_APPROVAL_CREDIT=1        # credit-limit → v2 (permission + threshold + governance)\n'
     'KAKO_APPROVAL_TRADE_SPEND=1   # trade-spend → engine (requires KAKO_TRADE_SPEND=1)\n'
     'KAKO_APPROVAL_PRICE_CHANGE=1  # price-change request approval')
para("Recommended pilot order: enable Credit-Limit first (already engine-based, lowest blast radius), validate with real "
     "approvers, then Trade-Spend, then Price-Change. Approvers reach pending items in the Workflow Inbox tab (now visible "
     "to credit.request.approve / pricing.manage holders).")

part("4 · Rollback")
h2("Instant (no deploy) — preferred")
para("Set the relevant flag back to OFF (unset or =0) and redeploy/restart. The workflow immediately reverts to its legacy "
     "behaviour: credit-limit → company_admin workflow; trade-spend → direct approve/cancel; price-change → direct edit "
     "only. In-flight engine instances can be left to complete or cancelled in the inbox.", color=GREEN)
h2("Full (remove the definitions)")
code("delete from erp_workflow_definitions\n where company_id is null\n   and key in ('credit_limit_approval_v2','trade_spend_approval','price_change_approval');\n-- steps cascade. Legacy 'credit_limit_approval' is untouched.")
h2("Foundation rollback (only if required)")
para("Re-apply the previous erp_workflow_decide (restores pre-0297 behaviour) and drop the two flag columns. The "
     "permission-resolution branch in resolve_users is additive and safe to leave. No business data is affected by any "
     "rollback step.")

part("5 · Scope honored")
bl("Implemented + activated (flag-gated): Credit-Limit, Trade-Spend, Price-Change.", color=GREEN)
bl("NOT started: P2, P3, P4, P5.", color=AMBER)
bl("NOT migrated (left stable for post-pilot): Day-Close, Out-of-Route Visit, Customer Transfer, Load Request, "
   "Van Transfer, Van Reconciliation.", color=AMBER)
para("The new approval architecture is now validated on low-risk commercial workflows while the operational FMCG flows "
     "remain exactly as they were.", bold=True, color=NAVY)

out="docs/audits/VANTORA-P1-Approval-Activation.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
