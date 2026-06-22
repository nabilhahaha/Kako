#!/usr/bin/env python3
"""VANTORA — PR #344 CI / Status Report (read-only) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"; GREYBG="ECECEC"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)

def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.6,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.6); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ---------------- Cover ----------------
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Pull Request #344 — CI & Status"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Migration-repair dry-run report PR · Vercel deployment status · Pre-existing failure triage"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("nabilhahaha/Kako · branch claude/migration-dryrun-report · No production writes · June 22, 2026 16:29 UTC"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---------------- Summary ----------------
h1("Summary")
para("This PR carries a documentation/report artifact only. The relevant Vercel deployment (kako) is green and the PR "
     "is ready for review. The single failing check belongs to the field-insights project, a pre-existing Vercel "
     "root-directory misconfiguration that is unrelated to this change and out of scope.", bold=True)

h1("CI status (as of Jun 22, 2026 16:29 UTC)")
tbl(["Check","State","Detail / Action"],
    [["Vercel — kako","READY","Deployed cleanly. Preview live. No action."],
     ["Vercel — field-insights","ERROR","Pre-existing root-dir misconfig, unrelated to this PR. Out of scope — not actioning."]],
    widths=[2.3,1.3,3.0],
    fill={(0,1):GREENBG,(1,1):REDBG})

h1("Why field-insights is not a blocker")
b("The failure exists independently of this PR — it is a Vercel project-settings (root directory) issue, not a code or "
  "migration defect introduced here.")
b("This PR changes only a report artifact; it does not touch field-insights, product code, route guards, permissions, "
  "or any database state.")
b("Fixing it would require changing the field-insights Vercel project configuration — explicitly out of the agreed scope.")

h1("Recommendation")
b("Treat the kako READY check as the gate for this PR — it is green.", bold=True)
b("Optionally add a one-line PR comment noting the field-insights red is a pre-existing Vercel settings issue, so it is "
  "not mistaken for a blocker (only if desired).")
b("PR remains under watch until merged or closed; any new CI failure or review comment will be triaged automatically.")

h2("Constraints honored")
para("No production writes. No migration repair on production. No Route Planner product-code merge. No field-insights "
     "changes. No route-guard or permissions changes. Read-only throughout.", italic=True, color=GREY, size=9)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"VANTORA_PR344_Status_Report.docx")
d.save(out)
print("WROTE", out, os.path.getsize(out), "bytes")
