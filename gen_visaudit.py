#!/usr/bin/env python3
"""VANTORA Role Visibility & UX Simplification Audit -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"; GREYBG="ECECEC"
# classification colors
C_DAILY=GREENBG; C_OCC=BLUEBG; C_RARE=GREYBG; C_WRONG=AMBERBG; C_HIDE=REDBG
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
CLSFILL={"Daily":C_DAILY,"Occasional":C_OCC,"Rarely":C_RARE,"Wrong role":C_WRONG,"Hide":C_HIDE}
def tbl(headers,rows,widths=None,size=8.3,clscol=None,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if clscol is not None and j==clscol and v in CLSFILL: _sh(c,CLSFILL[v])
            elif fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Role Visibility & UX Simplification Audit"); r.bold=True; r.font.size=Pt(21); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Not security — relevance. What each role SHOULD see to do its daily job."); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER; r=p.add_run("Computed from the live nav gates + role permissions · Salesman in depth · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Method & scale")
para("This audit ignores whether a role is technically blocked. It asks: of the items a role actually SEES (after permission "
     "gating), which help their daily job and which are noise? Every visible item is classified on one relevance scale.", bold=True)
tbl(["Class","Meaning"],[
 ["Daily","Used every day — must be front-and-centre"],
 ["Occasional","Useful sometimes — keep, but secondary"],
 ["Rarely","Edge cases — fine in an 'All/More' drawer, not primary"],
 ["Wrong role","Belongs to a different role's job — noise here"],
 ["Hide","Should not be in this role's primary view at all"],
],widths=[1.4,5.0],clscol=0)
para("Recommendation model (no access removed): give each role a curated PRIMARY navigation (Daily + key Occasional), and "
     "move everything else into a single 'All / More' expander. Permissions are unchanged — only RELEVANCE-based ordering and "
     "default visibility change. This is the cleanest way to make every role see only what helps them.", italic=True, color=NAVY)

# ===================== SALESMAN (deep) =====================
part("ROLE 1 — Salesman / Van Rep  (the focus)")
para("The most-used role and the worst-served by the current dense nav. After login the rep lands on /today (good), but the "
     "sidebar exposes ~26 items spanning sell, field, inventory and catalog — most irrelevant to a rep whose entire day is: "
     "work the route → sell → collect → request stock → close the day.", bold=True)
h2("Current visible items — classified")
tbl(["Section","Item","Class","Why"],[
 ["Main","Today (route home)","Daily","Their day starts here"],
 ["Main","Dashboard","Wrong role","Finance KPIs; redundant now they land on Today"],
 ["Main","Attention Center","Occasional","Useful nudges, not core"],
 ["Main","Coaching","Rarely","Training, not the daily loop"],
 ["Main","Route Execution","Daily","Working the route"],
 ["Main","Van Stock","Daily","What's on my van"],
 ["Main","Notifications","Occasional","Alerts, secondary"],
 ["Sales","Quick Sale (POS)","Daily","Primary sell action"],
 ["Sales","Invoices","Daily","Sell / view invoices"],
 ["Sales","Collections","Daily","Collect cash"],
 ["Sales","Rep App","Occasional","Overlaps Today/Route — duplicate entry"],
 ["Sales","Rep Settlement","Occasional","End-of-day money"],
 ["Sales","Sales Orders","Rarely","Pre-sales orders, not van cash sales"],
 ["Sales","Cashbox","Occasional","Cash drawer"],
 ["Sales","Journey","Daily","Plan/work visits"],
 ["Sales","Today's Journey","Daily","Duplicate of Today/Journey"],
 ["Sales","Field Offline","Occasional","Only when no signal"],
 ["Sales","Customers","Daily","Customer info on the route"],
 ["Inventory","Load Requests","Daily","Request stock to the van"],
 ["Inventory","Van Reconciliation","Daily","Reconcile the van"],
 ["Inventory","Van Transfer","Occasional","Move stock"],
 ["Inventory","Stock / Products","Wrong role","Catalog & warehouse stock — not a rep's job"],
 ["Inventory","Low-Stock Alerts","Wrong role","Warehouse/manager concern"],
 ["Inventory","Warehouses","Wrong role","Back-office"],
 ["Inventory","Near-Expiry","Rarely","Occasional check"],
],widths=[1.0,2.0,1.0,2.4],clscol=2)
para("Tally: ~9 Daily · ~6 Occasional · ~3 Rarely · ~4 Wrong-role. Nearly HALF the rep's sidebar is non-daily, and several "
     "items duplicate each other (Today / Today's Journey / Journey / Rep App all open the field loop; POS / Invoices / Sales "
     "Orders all sell).", bold=True, color=AMBER)
h2("Recommended PRIMARY navigation for the Salesman (6 items + More)")
tbl(["Keep as primary","Maps to"],[
 ["Today","the route + day"],
 ["Sell","POS / quick-sale (one 'Sell' entry, not POS + Invoices + Orders + Rep App)"],
 ["Collect","Collections + cashbox under one 'Collect'"],
 ["Customers","route customers"],
 ["Van","van stock + load request + reconciliation under one 'Van'"],
 ["More","everything else (offline, coaching, attention, near-expiry, …)"],
],widths=[1.8,4.6],fill={(i,0):GREENBG for i in range(6)})
h2("Items to hide from the rep's primary view")
b("Dashboard (finance KPIs), Stock/Products, Low-Stock Alerts, Warehouses — warehouse/back-office, wrong audience.")
b("Collapse duplicates: Today's Journey, Rep App, Sales Orders, Cashbox → fold into Today / Sell / Collect.")
b("Move to 'More': Coaching, Attention Center, Field Offline, Near-Expiry, Rep Settlement.")
h2("Reasoning & expected improvement")
para("A van rep on a phone needs five fingers' worth of choices, not a 26-item menu. Collapsing duplicates and hiding "
     "warehouse/finance items turns the rep app from 'an ERP a rep must navigate' into 'a sell-collect tool'. Expected: "
     "faster task start (1 tap to sell/collect), far less scrolling, lower training time, higher field adoption — the single "
     "biggest UX win because the rep is the highest-volume, lowest-patience user.", bold=True, color=NAVY)

# ===================== Other roles (condensed) =====================
def rolepage(title, intro, rows, primary, hide, improve):
    part(title)
    para(intro, bold=True)
    h2("Current visible items — classified (key items)")
    tbl(["Item","Class","Why"],[[r[0],r[1],r[2]] for r in rows],widths=[2.4,1.0,3.0],clscol=1)
    h2("Recommended primary")
    para(primary, color=GREEN)
    h2("Hide / move to 'More'")
    for x in hide: b(x, color=AMBER)
    h2("Expected improvement")
    para(improve, color=NAVY)

rolepage("ROLE 2 — Supervisor",
 "Lands on /approvals/queue (good). Their job is approvals + coverage; they also see sell/field items they rarely action.",
 [["Approval Queue","Daily","Their inbox"],["Supervisor Home","Daily","Coverage/team"],["Manager Home","Occasional","Overlaps supervisor home"],
  ["Reports Center","Occasional","Review"],["Territory / Distribution dashboards","Occasional","Coverage intel"],
  ["Van Reconciliation","Daily","Approve/verify reps"],["Invoices / POS","Wrong role","Supervisors don't sell daily"],
  ["Customers","Occasional","Reassignments"],["Inventory catalog/warehouses","Wrong role","Not their job"],
  ["Coaching","Occasional","Team development"]],
 "Approvals · Supervisor Home · Coverage/Territory · Van Reconciliation · Reports.",
 ["Hide Invoices/POS and inventory-catalog/warehouse items from primary.","Merge Supervisor Home + Manager Home (duplicate cockpits).","Move sell/cashbox to 'More'."],
 "A supervisor opens on their queue and sees only approve + monitor tools — no selling clutter. Faster daily approvals, less noise.")

rolepage("ROLE 3 — Branch Manager",
 "Lands on /manager. Broadest ops role — sees almost everything; benefits most from grouping, not hiding.",
 [["Manager Home","Daily","Branch cockpit"],["Approvals","Daily","Clear branch requests"],["Purchasing (POs/Suppliers)","Daily","Branch supply"],
  ["Reports","Daily","Branch performance"],["Customers","Occasional","Onboard/edit"],["Inventory (stock/transfers/counts)","Occasional","Branch stock"],
  ["Invoices / POS","Occasional","Spot transactions"],["Distribution dashboards","Occasional","Coverage"],["Settings (most)","Wrong role","Company-level admin, not branch"],
  ["Pricing","Occasional","If delegated"]],
 "Manager Home · Approvals · Purchasing · Reports · Customers · Inventory.",
 ["Hide company-level Settings (Users/Permissions/Integrations/Features) — admin territory.","Group the many distribution dashboards under one 'Distribution' entry."],
 "A branch manager gets a branch-ops cockpit, not a company-admin console — clearer scope, faster oversight.")

rolepage("ROLE 4 — Warehouse Keeper",
 "Lands on /inventory/requests (good). Sees the full Inventory section (right) plus some sell/field items (wrong).",
 [["Load Requests (approve)","Daily","Approve van loads"],["Stock / Products","Daily","Manage stock & catalog"],["Transfers","Daily","Move stock"],
  ["Stock Count","Daily","Counts"],["Receive PO (Purchasing)","Daily","Book in deliveries"],["Warehouses","Occasional","Setup"],
  ["Near-Expiry / Low-Stock","Daily","Stock health"],["Invoices / POS / Collections","Wrong role","Not their job (they see via sell perms? no — keep hidden)"],
  ["Approval Queue","Occasional","Load approvals also here"],["Reports","Occasional","Stock reports"]],
 "Inventory (Requests · Stock · Transfers · Counts · Expiry) · Receive PO · Approval Queue.",
 ["Hide any sell/collect entries from primary.","Consolidate Load Requests vs Approval Queue (load approvals appear in two places) into one path."],
 "A warehouse keeper sees a stock console — receive, move, count, approve — with no sales noise and one clear approval path.")

rolepage("ROLE 5 — Accountant",
 "Lands on /collections (good). Sees the Accounting section (right) but Collections sits under Sales, splitting their day.",
 [["Collections","Daily","Record payments"],["Accounting (Vouchers/Journal/Reports/Aging)","Daily","Core finance"],["Supplier Payments","Daily","AP"],
  ["Exports","Occasional","Period close"],["Cashbox","Occasional","Cash"],["Invoices (view)","Occasional","Reference"],
  ["Inventory (view)","Rarely","Stock valuation reference"],["Distribution / field items","Wrong role","Not finance"],
  ["POS / Sell","Wrong role","Accountants don't sell"],["Customers","Occasional","Credit/AR"]],
 "Collections · Accounting (Vouchers/Journal/Reports/Aging) · Supplier Payments · Exports.",
 ["Surface Collections + AR aging UNDER Accounting (one finance home).","Hide POS/sell and distribution/field items from primary."],
 "An accountant gets a single finance cockpit — collect, post, report — instead of hunting between Sales and Accounting.")

rolepage("ROLE 6 — Viewer",
 "Read-only observer. Sees dashboards/reports/inventory-view — but also many list screens with disabled actions, which look broken.",
 [["Dashboard / Reports","Daily","Their whole purpose"],["Inventory (view)","Occasional","Read stock"],["Accounting view","Occasional","Read finance"],
  ["Lists with greyed-out Add/Edit","Wrong role","Looks broken — they can see but every action is disabled"],
  ["Settings","Hide","No reason for a viewer to see admin menus"]],
 "Dashboards · Reports · read-only Inventory/Accounting views.",
 ["Give the Viewer a reports-and-dashboards-only profile.","Hide action-heavy CRUD list screens (they only frustrate)."],
 "A viewer gets a clean reporting portal instead of an editing UI where every button is disabled.")

# ===================== Cross-cutting =====================
part("Cross-cutting findings (your 5 questions)")
tbl(["#","Question","Finding"],[
 ["1","Screens visible but not useful to a role","Salesman: Stock/Products/Warehouses/Low-Stock + Dashboard. Supervisor/Accountant: POS/sell. Warehouse: collections/invoices."],
 ["2","Dashboard widgets that create noise","Payables/Overdue/Receivables on a field/ops user's dashboard; the dashboard itself is noise for roles now landing elsewhere (rep→Today, supervisor→Queue, accountant→Collections)."],
 ["3","Menus that should be hidden entirely","Per role: Settings for Branch Mgr/Warehouse/Accountant/Viewer; Inventory-catalog for Salesman/Supervisor/Accountant; Distribution-dashboards collapsed for everyone but managers."],
 ["4","Right feature, wrong audience","Inventory catalog & warehouses (warehouse role) shown to reps; sell/POS shown to supervisors/accountants; company Settings shown to branch-level roles."],
 ["5","Cleaner role-specific experience","Adopt PER-ROLE NAVIGATION PROFILES: a short curated primary set + one 'All/More' expander. No permission change — pure relevance curation."],
],widths=[0.3,2.2,3.9])

part("Recommendation")
para("Introduce per-role NAVIGATION PROFILES (relevance-based), not permission changes:", bold=True, color=NAVY)
b("Salesman (priority): 5-item primary — Today · Sell · Collect · Customers · Van — everything else under 'More'. Fold the "
  "duplicate sell/field entries (POS/Invoices/Orders/Rep App/Today's-Journey) into single actions.",bold=True)
b("Supervisor: Approvals · Supervisor Home · Coverage · Van Reconciliation · Reports.")
b("Branch Manager: Manager Home · Approvals · Purchasing · Reports · Customers · Inventory (group distribution dashboards).")
b("Warehouse Keeper: Inventory console · Receive PO · Approvals (single load-approval path).")
b("Accountant: one Finance home — Collections + Accounting + Supplier Pay + AR aging together.")
b("Viewer: reports-and-dashboards-only profile.")
para("Implementation is a NAV-CONFIG change (curated primary lists + a 'More' group per role) — no new features, no security "
     "change, fully reversible. The Salesman profile is the highest-impact single change and should ship first.", bold=True, color=NAVY)
para("")
para("Goal achieved: every role sees a short, relevant primary menu for its daily job; nothing is removed — only "
     "de-cluttered. Expected platform-wide effect: faster task starts, less scrolling, lower training, higher adoption — "
     "especially for the salesman.", bold=True, color=NAVY)
out="docs/audits/VANTORA-Role-Visibility-UX-Audit.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
