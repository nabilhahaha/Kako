#!/usr/bin/env python3
"""VANTORA FMCG UX & Adoption Review (role-by-role) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def tbl(headers,rows,widths=None,size=8.3,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("FMCG UX & Adoption Review"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Role-by-role walkthrough · friction, shortcuts, navigation, speed, onboarding · operational-pilot focus"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Against the live pilot tenant (VANTORA Pilot FMCG DEMO) · usability over new features · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

h1("Executive summary")
para("VANTORA is functionally production-ready for an FMCG pilot (all audits + live scenarios passed). This review shifts to "
     "the ADOPTION layer: does each role's daily flow feel fast, obvious and low-friction? Verdict: the foundations are "
     "strong (mobile == desktop, clean RBAC, working approval loop), but a handful of small, high-leverage UX gaps stand "
     "between 'works' and 'feels like a product people want to use daily'. None require major features — they are landing-"
     "page routing, naming, shortcuts, and onboarding polish.", bold=True)
h2("The single biggest adoption lever")
para("Every FMCG role lands on the generic /dashboard. resolveHomePath (src/lib/erp/home.ts:52) routes by VERTICAL module "
     "(fashion/clinic/restaurant/…) but has no branch for FMCG field ROLES, so a salesman, supervisor, accountant and "
     "warehouse keeper all open the same KPI dashboard instead of their work screen. Fixing this one routing function is the "
     "highest-impact, lowest-effort adoption win.", bold=True, color=RED)

part("1. Demo environment enriched (live)")
para("The pilot tenant now contains realistic operating history so every role sees a populated, lifelike system on first "
     "login — not empty screens.", bold=True)
tbl(["Data","Detail"],[
 ["Tenant","VANTORA Pilot FMCG (DEMO) · 21 modules + 22 roles"],
 ["Users","admin / salesman / supervisor / accountant / warehouse @pilot.test (test.123)"],
 ["Warehouses","Pilot Main Warehouse + Pilot Van"],
 ["Products","8 FMCG SKUs with cost/sell/tax + 1000 stock each"],
 ["Customers","5 grocery/market customers, credit limits, on a route"],
 ["Route","Cairo East Route (rep = salesman, 5 customers in sequence)"],
 ["Sales history","12 issued invoices, backdated ~24 days (real AR + stock deduction)"],
 ["Collections","7 collections (partial payments) → ~4,906 EGP open AR"],
 ["Transfers","1 customer transfer (applied) + 1 van transfer (completed)"],
 ["Pending approvals","4 LIVE items waiting: day-close, out-of-route visit, customer transfer, van transfer"],
],widths=[1.5,5.0])
para("Effect: a supervisor logging in now has a real approval queue to work; an accountant sees real AR; a salesman sees a "
     "real route and history.", italic=True)

part("2. Cross-cutting UX findings")
tbl(["#","Finding","Why it hurts adoption","Evidence"],[
 ["X1","FMCG roles all land on /dashboard","A rep wants 'my day', an accountant wants AR — not a generic KPI page; daily extra clicks","home.ts:52 (no FMCG-role branch)"],
 ["X2","Three similarly-named approval items in one menu","'Approval Center' + 'Approvals' + 'Approval Queue' confuse users about where to act","navigation.ts:220,226,227"],
 ["X3","Dense 'main' nav section (~16 items)","Reps/clerks scroll past manager/territory/coaching items they never use","navigation.ts main section"],
 ["X4","No first-run onboarding / empty-state CTAs","A fresh tenant lands on empty screens with no 'do this next' guidance","no getting-started surface"],
 ["X5","Few desktop quick-actions","Mobile has a bottom-nav; desktop users hunt through menus for New Invoice / Collect","sidebar nav only"],
 ["X6","Collections lives under Sales, not Accounting","Accountants look in the Accounting section first and don't find collections","nav main/sales placement"],
],widths=[0.4,2.0,2.6,1.4],
 fill={(0,0):REDBG})

# Role template
def role(name,login,journey,works,confusing,blocker):
    part(name)
    para(f"Pilot login: {login}", italic=True, color=GREY)
    h2("Typical daily journey")
    para(journey)
    h2("What works well");
    for x in works: b(x, color=GREEN)
    h2("What feels confusing / friction");
    for x in confusing: b(x, color=AMBER)
    h2("What would prevent daily adoption (fix first)")
    para(blocker, bold=True, color=RED)

role("3. Company Admin","admin@pilot.test",
 "Sets up branches, warehouses, users and the catalog; configures policies; watches company KPIs; manages customers and pricing.",
 ["Full, well-grouped Settings (Organization / Data & Fields / Integrations / Governance / Personal).",
  "Staff management, marketplace (modules), action policies and feature flags are all reachable.",
  "Dashboard KPIs are a reasonable admin landing page (unlike for field roles)."],
 ["User & permission management is split: the admin manages people via Settings · Staff, while Settings · Users / "
  "Permissions are super-admin-only (M-2) — admins may look for 'Users' and not find the management screen.",
  "The three 'approval' items (X2) are ambiguous for the person who configures them.",
  "Many advanced Settings screens with no 'recommended setup order' for a first-time admin (X4)."],
 "Onboarding: a fresh admin has no guided 'set up your company' checklist. Add a role-based getting-started panel + rename "
 "Staff → 'Users & Roles' so the admin immediately finds user management.")

role("4. Branch Manager","(role: branch_manager)",
 "Runs one branch: approves van loads and day-close exceptions, oversees reps, manages purchasing and local customers.",
 ["Broad branch-ops permissions without company settings/billing (correct scoping).",
  "Approval Queue exposes the day-close / transfer approvals they own.",
  "Purchasing (suppliers, POs, receive) is complete."],
 ["Lands on the generic /dashboard rather than a branch cockpit (X1).",
  "Sees the full dense main nav (X3) including company-level items that aren't theirs to action."],
 "Landing page: route the branch manager to a branch/manager home (e.g., /manager) so their day starts on branch KPIs + "
 "their approval queue, not a generic dashboard.")

role("5. Supervisor","supervisor@pilot.test",
 "Supervises the field team: approves out-of-route visits, day-close exceptions, customer/van transfers; monitors coverage and reconciliations.",
 ["The unified Approval Queue (now populated with 4 live items) + the mobile bottom-nav Approvals tab make approvals fast.",
  "Dedicated /supervisor home and reconciliation screens exist.",
  "Permissions are exactly right: approve, but not post GL or set prices."],
 ["Lands on /dashboard, then must navigate to /approvals/queue or /supervisor (X1) — two of the three 'approval' names "
  "compete for attention (X2).",
  "Overlap between /supervisor, /approval-center, /approvals and /approvals/queue is unclear at a glance."],
 "Landing page + naming: send the supervisor straight to the Approval Queue (their primary daily job) and collapse the "
 "three approval entries into one clear 'Approvals' with tabs.")

role("6. Salesman","salesman@pilot.test",
 "Runs a daily route: starts the day, visits customers in sequence, sells (POS/invoice), collects cash, requests van load, closes the day.",
 ["/today is a purpose-built field home; the mobile bottom-nav (Home/Today/Customers/Sell/Inventory) covers the core loop.",
  "POS / quick-sale is fast; collections and van-sales flows exist; the 'More' drawer mirrors the full sidebar so nothing is unreachable on mobile.",
  "Field permissions are tightly scoped (sell/collect, no GL/price/credit)."],
 ["BIGGEST rep friction: lands on /dashboard (a KPI page irrelevant to a rep) instead of /today (X1).",
  "The desktop sidebar shows many management items a rep never uses (X3).",
  "Creating an invoice line-by-line is slower than scanning; reps will prefer POS — make it the default sell path."],
 "Landing page: a salesman MUST open on /today. Today they open a generic dashboard — the most-used role gets the least-"
 "tailored landing. This single fix most improves rep daily adoption.")

role("7. Warehouse Keeper","warehouse@pilot.test",
 "Receives purchase orders, adjusts stock, approves van load requests, runs transfers and stock counts.",
 ["Inventory section is complete: products, stock, transfers, load requests, counts, warehouses.",
  "Approves stock requests / van transfers (separation of duties from the requesting rep is enforced).",
  "Receiving POs posts stock + Inventory/AP correctly."],
 ["Lands on /dashboard rather than a warehouse view (X1).",
  "The load-request / transfer APPROVAL lives in two places (Inventory · Load Requests and the Approval Queue) — not "
  "obvious which to use (X2).",
  "Inventory actions are spread across several screens; no single 'warehouse today' surface."],
 "Landing page + approval clarity: route to an inventory/requests home and make the load-approval entry singular so the "
 "keeper isn't hunting between Inventory and the Approval Queue.")

role("8. Accountant","accountant@pilot.test",
 "Records collections, posts the GL, pays suppliers, reviews AR aging and financial reports.",
 ["Accounting section is well-formed: chart, vouchers, journal, financial reports, aging, exports.",
  "Posting + supplier payments are permission-gated to the accountant correctly.",
  "Real AR now exists in the demo (~4,906 EGP) so aging/reports are meaningful."],
 ["Lands on /dashboard rather than an AR/collections cockpit (X1).",
  "Collections sits under the Sales/main area, not Accounting (X6) — accountants look in Accounting first.",
  "'Aging' vs 'Financial Reports' overlap; the daily AR view isn't a single obvious entry."],
 "Discoverability: surface Collections + AR aging under the Accounting section (or route the accountant's landing there). "
 "The accountant's daily job (collect + chase AR) should be one click from login.")

part("9. Prioritized usability backlog  (usability/speed/adoption — no major features)")
tbl(["#","Improvement","Impact","Effort","Type"],[
 ["U1","Role-aware landing in resolveHomePath for FMCG roles (salesman→/today, supervisor→/approvals/queue, accountant→/collections or /accounting, warehouse→/inventory/requests, branch_manager→/manager)","HIGH","S","Code (1 fn)"],
 ["U2","Consolidate/rename the 3 approval nav items into one 'Approvals' with tabs (Field / Workflow)","HIGH","S–M","Code (nav)"],
 ["U3","Role-based 'Getting Started' panel + empty-state CTAs (Add product / customer / load)","Med","M","Code (UI)"],
 ["U4","Desktop quick-action bar on the dashboard (New Invoice, Record Collection, New Customer)","Med","S–M","Code (UI)"],
 ["U5","Curate the dense 'main' nav per role (hide manager/territory/coaching from reps who lack the perms — tighten gates)","Med","S","Code (nav)"],
 ["U6","Rename Settings · Staff → 'Users & Roles' (M-2 clarity)","Low","S","i18n"],
 ["U7","Cross-link Collections + AR aging under the Accounting section","Low","S","Code (nav)"],
 ["U8","Ship the pilot demo data + a one-page per-role cheat-sheet (login + 'your daily flow')","Med","S","Docs/data"],
],widths=[0.4,3.4,0.7,0.6,1.2],
 fill={(0,2):REDBG,(1,2):REDBG})
para("Effort: S=hours · M=1–2 days. Everything here is usability polish; none adds a new business capability.", italic=True, size=8.5)

part("10. Recommendation")
para("VANTORA already WORKS for FMCG. To make it FEEL like a production platform people adopt daily, do U1 and U2 first — "
     "role-aware landing and approval-naming clarity remove the friction every role hits on every login, at minimal effort "
     "and zero new functionality. U3–U8 are fast follow-ons. With the enriched demo data already live, the platform now also "
     "DEMOS like a real, operating FMCG business.", bold=True, color=NAVY)
b("Do now (highest leverage, smallest change): U1 role-aware landing, U2 approval consolidation.")
b("Fast follow: U3 onboarding/empty-states, U4 quick-actions, U5 nav curation, U8 role cheat-sheets.")
b("Polish: U6 label, U7 accounting cross-links.")
para("")
para("No major new features are needed for daily operations — the daily FMCG loop is complete and validated. These changes "
     "are purely about speed, clarity and first-run guidance.", bold=True, color=NAVY)

out="docs/audits/VANTORA-FMCG-UX-Adoption-Review.docx"
os.makedirs("docs/audits",exist_ok=True); d.save(out); print("saved",out)
