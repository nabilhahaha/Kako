#!/usr/bin/env python3
"""VANTORA Route Planner — Migration-Repair Dry-Run Report (read-only) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"; GREYBG="ECECEC"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)

def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(8.3); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows,widths=None,size=8.2,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ---------------- Cover ----------------
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Route Planner — Migration-Repair Dry-Run"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Tracking reconciliation for migrations 0358–0364 · Isolated dry-run attempt · Read-only evidence · Recommendation"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Supabase rsjvgehvastmawzwnqcs (vantora-staging) · No production writes · No product-code merge · June 22, 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---------------- Outcome banner ----------------
h1("Outcome")
para("The dry-run could NOT be executed as specified. Two environment-level blockers prevented running it, and the work "
     "was stopped rather than worked around in any way that could touch production. Production was read-only throughout; "
     "no branch was created (so $0 cost was incurred), no migration repair was run on production, and no product code was merged.",
     bold=True)
tbl(["Item","Status"],
    [["Isolated branch/clone created","NO — create_branch harness-gated (see Blocker 1)"],
     ["Migration repair run on clone","NO — Supabase MCP cannot run the CLI (see Blocker 2)"],
     ["Production database written","NO — read-only SELECT statements only"],
     ["Migration repair run on production","NO"],
     ["Route Planner product code merged","NO"],
     ["Branch cost incurred","$0.00 — no branch was provisioned"]],
    widths=[3.0,3.6],
    fill={(0,1):REDBG,(1,1):REDBG,(2,1):GREENBG,(3,1):GREENBG,(4,1):GREENBG,(5,1):GREENBG})

# ---------------- Blockers ----------------
h1("Why the dry-run could not run")
h2("Blocker 1 — create_branch is gated by the execution environment")
para("The request to provision an isolated Supabase branch returned “MCP tool call requires approval” on both "
     "attempts, despite explicit in-chat approval. Branch provisioning is disallowed by this environment's permission "
     "policy, and an in-chat approval does not satisfy that harness-level gate. No branch was created and no cost was incurred.")
h2("Blocker 2 — the Supabase MCP cannot run the CLI")
para("The available MCP toolset exposes only execute_sql, list_migrations and apply_migration. There is no migration "
     "repair, migration list, or db diff CLI command available in this session. Even on a branch, the repair could only "
     "be SIMULATED with a raw INSERT into the tracking table — it could not be exercised through the real Supabase CLI "
     "the way the production reconciliation eventually will be.")

# ---------------- Read-only evidence ----------------
h1("Read-only evidence collected (zero writes)")
para("The following was confirmed against production using SELECT statements only. Production state was read, never modified.")
tbl(["Check","Result","Verdict"],
    [["Route Planner tables present","15 of 15","INTACT"],
     ["Pilot data rows (7 populated tables)","28 rows","INTACT / unchanged"],
     ["Migrations recorded in tracking","0353–0357 only","Drift confirmed"],
     ["Migrations 0358–0364 objects exist but unrecorded","Yes","Drift confirmed"],
     ["Repair touches objects or data","No — INSERT into schema_migrations only","Metadata-only"]],
    widths=[3.4,2.3,1.5],
    fill={(0,2):GREENBG,(1,2):GREENBG,(2,2):AMBERBG,(3,2):AMBERBG,(4,2):GREENBG})
para("Because the 0358–0364 objects already exist, marking them applied re-runs nothing: it is purely a metadata "
     "insert into supabase_migrations.schema_migrations. No object DDL and no data DML are involved.")

# ---------------- Critical finding ----------------
h1("Critical finding — the tracking scheme is non-standard")
para("The recorded rows do not key migrations the way the Supabase CLI default expects. The version column holds a "
     "14-digit apply-timestamp, while the migration files are named 0353_… etc. The name column carries the file stem:")
tbl(["version (recorded)","name (file stem)"],
    [["20260620142235","0353_route_planner_access"],
     ["20260620142253","0354_rp_reporting_graph"],
     ["20260620142324","0355_rp_integration"],
     ["20260620142400","0356_rp_request_center"],
     ["20260620142410","0357_rp_schema_health"]],
    widths=[2.6,3.8])
para("Implication: a naive “supabase migration repair --status applied 0358” keys on the version (expecting a "
     "timestamp) and may NOT map correctly to this project's 0XXX-named files. The reconciliation must follow this "
     "project's own convention — insert name = the file stem (e.g. 0358_rp_connector_admin) with a version that "
     "matches the existing timestamp pattern. This is exactly the operational risk a real CLI dry-run on a clone needs "
     "to confirm, and exactly why running it blindly on production would be unwise.", bold=True)

# ---------------- Recommendation ----------------
h1("Recommendation")
b("Do not run any reconciliation on production yet — consistent with the stated constraints; nothing was done to production.", bold=True)
b("Run the dry-run via the Supabase CLI, by the owner of the deploy pipeline, in an environment where (a) branch/clone "
  "creation is permitted and (b) the supabase CLI is available — neither is possible from this MCP session.")
para("Suggested sequence on the clone:", bold=True, size=9.5)
code("supabase migration list                 # baseline: 0353-0357 recorded")
code("# reconcile 0358-0364 (see note on version-keying below)")
code("supabase migration list                 # confirm: no pending migrations")
code("supabase db diff                         # confirm: no schema drift")
code("# verify 15/15 RP objects + 28 pilot rows intact, then delete the branch")
b("Match the existing scheme when reconciling — name = file stem, version = a fresh timestamp like 0353–0357 — "
  "rather than relying on migration repair's default version-keying, because of the timestamp / 0XXX mismatch above.")
b("The underlying change is low-risk (metadata-only; objects and pilot data already present and verified intact) — but "
  "the CLI dry-run's db diff should come back clean on a clone before anyone marks 0358–0364 applied on production.")

h2("Constraints honored")
para("No production writes. No migration repair on production. No Route Planner product-code merge. No PR #325. "
     "No field-insights. No route-guard or permissions changes. Read-only throughout; no branch provisioned ($0 cost).",
     italic=True, color=GREY, size=9)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"VANTORA_RP_Migration_DryRun_Report.docx")
d.save(out)
print("WROTE", out, os.path.getsize(out), "bytes")
