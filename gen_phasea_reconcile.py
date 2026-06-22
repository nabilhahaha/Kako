#!/usr/bin/env python3
"""VANTORA — Phase A: Production Migration Tracking Reconciliation (0358-0364) -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"; GREYBG="ECECEC"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)

def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(8.2); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows,widths=None,size=8.6,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.6); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ---------------- Cover ----------------
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Phase A — Migration Tracking Reconciliation"); r.bold=True; r.font.size=Pt(21); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Route Planner migrations 0358–0364 recorded · Metadata-only · Production verified stable"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Supabase rsjvgehvastmawzwnqcs (vantora-staging) · ACTIVE_HEALTHY · Executed 2026-06-22"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---------------- Result ----------------
h1("Result — COMPLETED & VERIFIED")
para("The Route Planner migration tracking records for 0358–0364 were reconciled on the production database "
     "(vantora-staging) as a metadata-only, duplicate-safe transaction. All seven rows were inserted; tracking now "
     "records the full 0353–0364 set. No migration SQL was re-run, no objects were created/dropped, and no product or "
     "pilot data was modified. Production remained ACTIVE_HEALTHY throughout.", bold=True)

# ---------------- Preconditions ----------------
h1("Pre-write confirmation (read-only)")
tbl(["Check","Expected","Actual","Result"],
    [["Target DB","vantora-staging (rsjvgehvastmawzwnqcs)","vantora-staging, ACTIVE_HEALTHY","MATCH"],
     ["0353–0357 recorded","5","5","MATCH"],
     ["0358–0364 recorded","0","0","MATCH"],
     ["0358–0364 objects exist (10 checks)","all present","all present","MATCH"],
     ["RP tables present","15","15","MATCH"],
     ["Pilot fingerprint (7 tables)","—","6 / 6 / 7 / 2 / 2 / 4 / 1 = 28","BASELINE"],
     ["Convention","version=timestamp, name=stem","confirmed from 0353–0357","MATCH"]],
    widths=[1.9,1.9,1.9,0.9],
    fill={(i,3):GREENBG for i in range(7)})

# ---------------- Write ----------------
h1("Reconciliation performed")
para("Single transaction, duplicate-safe (guarded by NOT EXISTS on name); versions continue the existing timestamp "
     "sequence after 0357 (20260620142410), preserving 0358 < … < 0364 order:")
code("begin;")
code("insert into supabase_migrations.schema_migrations (version, name)")
code("select v, n from (values")
code("  ('20260620142420','0358_rp_connector_admin'),")
code("  ('20260620142430','0359_rp_planning_persistence_a'),")
code("  ('20260620142440','0360_rp_dataset_persistence'),")
code("  ('20260620142450','0361_rp_saved_plans'),")
code("  ('20260620142500','0362_rp_mission_perms'),")
code("  ('20260620142510','0363_rp_missions'),")
code("  ('20260620142520','0364_rp_plan_approvals')")
code(") as t(v,n)")
code("where not exists (select 1 from supabase_migrations.schema_migrations m where m.name = t.n);")
code("commit;")

# ---------------- Post-write ----------------
h1("Post-write verification")
tbl(["Check","Expected","Actual","Result"],
    [["0353–0364 recorded","12","12","PASS"],
     ["0358–0364 newly recorded","7","7","PASS"],
     ["Duplicate names","0","0","PASS"],
     ["RP tables intact","15","15","PASS"],
     ["Pilot fingerprint","6/6/7/2/2/4/1 (28)","6/6/7/2/2/4/1 (28)","IDENTICAL"],
     ["Schema/object changes","none","none","PASS"],
     ["Product data changes","none","none","PASS"],
     ["Production status","stable","ACTIVE_HEALTHY","PASS"]],
    widths=[1.9,1.7,1.9,1.0],
    fill={(i,3):GREENBG for i in range(8)})

# ---------------- Rollback ----------------
h1("Rollback note")
para("If ever required, the reconciliation is reversible by deleting exactly the seven inserted tracking rows "
     "(objects/data untouched):", size=9)
code("delete from supabase_migrations.schema_migrations")
code("where name in ('0358_rp_connector_admin','0359_rp_planning_persistence_a','0360_rp_dataset_persistence',")
code("  '0361_rp_saved_plans','0362_rp_mission_perms','0363_rp_missions','0364_rp_plan_approvals');")

# ---------------- Scope ----------------
h2("Scope guardrails honored")
para("Metadata/bookkeeping only. Did not re-run migration SQL. Did not create/drop tables. Did not modify product "
     "data. Did not delete pilot data. Did not change route-guard or permissions logic. No PR #325. No field-insights. "
     "Phase B (product merge) proceeds separately in small, gated PRs.", italic=True, color=GREY, size=9)

out=os.path.join(os.path.dirname(os.path.abspath(__file__)),"VANTORA_PhaseA_Migration_Reconciliation.docx")
d.save(out)
print("WROTE", out, os.path.getsize(out), "bytes")
