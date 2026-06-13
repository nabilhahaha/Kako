#!/usr/bin/env python3
"""VANTORA Role-Based Navigation Profiles — Implementation Report -> .docx"""
import os
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

NAVY=RGBColor(0x1F,0x2A,0x44); GREY=RGBColor(0x55,0x55,0x55); DARK=RGBColor(0x1A,0x1A,0x1A)
RED=RGBColor(0xB3,0x1B,0x1B); GREEN=RGBColor(0x1B,0x7F,0x3B); AMBER=RGBColor(0xB4,0x6A,0x00); BLUE=RGBColor(0x1B,0x4F,0x9E)
HEADER_BG="1F2A44"; ZEBRA="F2F4F8"; GREENBG="E5F2E9"; AMBERBG="FBF1DD"; REDBG="F8E3E3"; BLUEBG="E5ECF7"
d=Document(); d.styles["Normal"].font.name="Calibri"; d.styles["Normal"].font.size=Pt(10)
def _sh(c,h):
    tcPr=c._tc.get_or_add_tcPr(); s=OxmlElement("w:shd"); s.set(qn("w:val"),"clear"); s.set(qn("w:color"),"auto"); s.set(qn("w:fill"),h); tcPr.append(s)
def part(t):
    d.add_page_break(); p=d.add_heading(level=0); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(18); r.bold=True
def h1(t):
    p=d.add_heading(level=1); r=p.add_run(t); r.font.color.rgb=NAVY; r.font.size=Pt(14); r.bold=True
def h2(t):
    p=d.add_heading(level=2); r=p.add_run(t); r.font.color.rgb=BLUE; r.font.size=Pt(11); r.bold=True
def para(t,color=DARK,size=10,italic=False,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.color.rgb=color; r.font.size=Pt(size); r.italic=italic; r.bold=bold
def b(t,size=9.5,color=DARK,bold=False):
    p=d.add_paragraph(style="List Bullet"); r=p.add_run(t); r.font.size=Pt(size); r.font.color.rgb=color; r.bold=bold
def code(t):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name="Consolas"; r.font.size=Pt(8.6); r.font.color.rgb=RGBColor(0x33,0x33,0x33)
def tbl(headers,rows,widths=None,size=8.3,fill=None):
    t=d.add_table(rows=1,cols=len(headers)); t.style="Table Grid"; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for j,hh in enumerate(headers):
        c=t.rows[0].cells[j]; c.text=""; r=c.paragraphs[0].add_run(hh); r.bold=True; r.font.size=Pt(8.4); r.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); _sh(c,HEADER_BG)
    for i,row in enumerate(rows):
        cells=t.add_row().cells
        for j,v in enumerate(row):
            c=cells[j]; c.text=""; r=c.paragraphs[0].add_run(v); r.font.size=Pt(size)
            if fill and (i,j) in fill: _sh(c,fill[(i,j)])
            elif i%2==1: _sh(c,ZEBRA)
    if widths:
        for j,w in enumerate(widths):
            for r in t.rows: r.cells[j].width=Inches(w)

# ---- Cover ----
for _ in range(4): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("VANTORA"); r.bold=True; r.font.size=Pt(38); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Role-Based Navigation Profiles"); r.bold=True; r.font.size=Pt(22); r.font.color.rgb=NAVY
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Implementation Report — focused field-sales UX without touching permissions or backend"); r.font.size=Pt(10.5); r.font.color.rgb=GREY
for _ in range(2): d.add_paragraph()
p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
r=p.add_run("Priority 1 Salesman + Supervisor · Warehouse Keeper · Accountant · Branch Manager · Viewer · June 2026"); r.font.size=Pt(9); r.italic=True; r.font.color.rgb=GREY
d.add_page_break()

# ---- Exec summary ----
h1("Executive summary")
para("VANTORA's sidebar previously showed every role the full ERP tree, gated only by permissions. The Salesman — the "
     "highest-volume daily user — saw 20+ items across Dashboard, Catalog, Warehouses, Alerts and Coaching, with the five "
     "things he actually does each day scattered among them. This change adds a thin RELEVANCE layer on top of the existing "
     "permission gating: each role gets a short, curated PRIMARY menu of its daily tools, and everything else folds into a "
     "single \"More\" group.", bold=True)
para("Nothing about access changed. The profile can only surface items the user could already reach — it reorders the "
     "already-permission-filtered navigation, it never adds an entry the permission layer withheld. No backend, no RPC, no "
     "RLS, no migration. Pure navigation/UX. Admin, Manager, Super-Admin and Platform-Owner keep the full, un-profiled tree.")

h2("What shipped")
tbl(["Role","Primary menu (in order)","Everything else"],
[
 ["Salesman / Driver","Today · Sell · Collect · Customers · Van","→ More"],
 ["Supervisor","Approvals · Team · Coverage · Van Reconciliation · Reports","→ More"],
 ["Branch Manager","Branch · Approvals · Purchasing · Reports · Customers · Inventory","→ More"],
 ["Warehouse Keeper","Load Requests · Stock · Receive · Transfers · Approvals","→ More"],
 ["Accountant","Collect · Accounting · Vouchers · AR Aging · Suppliers","→ More"],
 ["Viewer","Home · Reports · Inventory","→ More"],
],widths=[1.5,3.6,1.4],
fill={(0,1):GREENBG})
para("Salesman primary highlighted: this was the explicit Priority-1 target — turn the field rep's sidebar into a focused "
     "field-sales app (Today / Sell / Collect / Customers / Van + More).", italic=True, color=GREY, size=9)

# ---- The salesman problem ----
part("1 · The Salesman problem (Priority 1)")
h2("Before")
para("The salesman's sidebar mirrored the full system. Daily actions competed for attention with admin-only and "
     "occasional screens. From the UX audit, the items that did NOT belong in a field rep's primary view:")
b("Dashboard (generic ERP landing — rep lands on /today instead)")
b("Stock / Products catalog, Warehouses, Low Stock Alerts (inventory-keeper concerns)")
b("Coaching, Attention Center, Near Expiry (supervisor / occasional)")
b("Duplicate entry points: Today + Today's Journey + Journey + Rep App; POS + Invoices + Sales Orders; Collections + Cashbox")

h2("After")
para("Five primary items — the entire daily loop — then one \"More\" group for the long tail. The duplicates collapse "
     "naturally: the primary slot owns the canonical entry (Today→/today, Sell→/sales/pos, Collect→/collections), and the "
     "variants (Journey, Rep App, Invoices, Orders, Cashbox) drop into More instead of competing at the top level.")
tbl(["Primary","Label","Route","Replaces / absorbs"],
[
 ["1","Today","/today","Today's Journey, Journey, Rep App"],
 ["2","Sell","/sales/pos","Invoices, Sales Orders"],
 ["3","Collect","/collections","Cashbox"],
 ["4","Customers","/customers","—"],
 ["5","Van","/field/stock","Van stock / load"],
],widths=[0.7,1.1,1.5,3.0])

# ---- How it works ----
part("2 · How it works (mechanics)")
h2("A relevance layer, not a permission layer")
para("The whole feature is one pure function, applyNavProfile, applied AFTER visibleSections has already done permission, "
     "module and feature-flag gating. It cannot widen access:")
code("sections = applyNavProfile(")
code("  visibleSections(permissions, isSuperAdmin, …, enabledFlags),  // already gated")
code("  roles,                                                        // user's branch roles")
code("  { isSuperAdmin, isPlatformOwner },")
code(");")
para("Guarantees that keep this safe:", bold=True)
b("A profile item is promoted to Primary ONLY if its href is in the already-visible set. If the rep lacks the permission "
  "for /collections, the Collect slot silently drops — the profile can never reveal a hidden screen.")
b("\"More\" is built from the remaining visible items in their original order, de-duplicated by href — so no screen the "
  "user could previously reach disappears; it just moves down.")
b("Elevated users (Super-Admin, Platform-Owner) and the admin / manager roles return the sections UNCHANGED — they keep "
  "the full ERP tree.")
b("Role seniority mirrors home.ts: branch_manager → supervisor → accountant → warehouse_keeper → salesman → driver. "
  "driver reuses the salesman profile.")

h2("Files changed")
tbl(["File","Change"],
[
 ["src/lib/erp/nav-profiles.ts","NEW — NAV_PROFILES config (6 roles), profileRoleFor(), applyNavProfile()."],
 ["src/components/layout/sidebar.tsx","Added roles prop; wrap visibleSections() in applyNavProfile(). Renders for desktop sidebar + mobile \"More\" drawer (shared content)."],
 ["src/app/(app)/layout.tsx","Pass roles={ctx.memberships.map(m => m.role)} to <Sidebar>."],
 ["src/lib/i18n/messages/core.ts","Added nav.sections.primary / .more and the nav.profile.* label namespace (ar + en, full parity)."],
 ["src/lib/erp/nav-profiles.test.ts","NEW — 11 unit tests (profileRoleFor precedence, salesman/supervisor primary, visibility gating, More de-dup, elevated/admin pass-through)."],
],widths=[2.6,4.0])

# ---- Validation ----
part("3 · Validation")
para("All gates green. No permission, RPC, RLS or migration touched.", bold=True, color=GREEN)
tbl(["Gate","Result"],
[
 ["npx tsc --noEmit","Clean — no type errors"],
 ["nav-profiles.test.ts","11 passed"],
 ["i18n.test.ts (ar/en parity)","Passed — new keys present in both locales"],
 ["keys-usage.test.ts","Passed — nav.profile.* + nav.sections.* keys resolve"],
 ["navigation.test.ts","Passed"],
 ["Full suite (vitest run)","1332 passed · 181 skipped"],
 ["npm run build","Compiled successfully"],
],widths=[3.0,3.4],
fill={(i,1):GREENBG for i in range(7)})

h2("Scope discipline")
b("No permission changes — ROLE_PERMISSIONS and erp_role_permissions untouched.")
b("No backend changes — no actions, RPCs, migrations or edge functions.")
b("Navigation/UX only — a render-time reorganisation of an already-filtered list.")
b("Fully reversible — removing the applyNavProfile wrapper restores the prior full-tree sidebar with zero data impact.")

h2("Reviewer quick-check")
para("To confirm no access changed: the set of hrefs rendered for any role is identical before and after — only their "
     "grouping/order differs. Primary ⊆ (visible ∩ profile); More = visible − Primary. Union(Primary, More) == visible.",
     italic=True, color=GREY, size=9)

out="docs/audits/VANTORA-Navigation-Profiles-Implementation.docx"
os.makedirs("docs/audits",exist_ok=True)
d.save(out)
print("WROTE",out)
